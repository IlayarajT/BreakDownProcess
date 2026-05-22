import os
import re
import shutil
import time

import win32com.client as win32
import yaml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import (
    WD_COLOR_INDEX,
    WD_PARAGRAPH_ALIGNMENT,
    WD_LINE_SPACING,
)
from docx.shared import Pt, RGBColor
from docxcompose.composer import Composer
from win32com.client import constants
import zipfile
import tempfile
import getAppPath
from dbprocess import DataBase
from loadconfig import getconfig
from com_manager import WordApplicationManager, COMManager
import threading

class DocxMerger:
    # ==============================================================
    # INIT
    # ==============================================================

    def __init__(self):
        getAppPath.getapppath()  # preserved side-effect
        self.configFolder, self.breakDownConfig = getconfig()
        self.db_process = DataBase()

        merger_yaml = os.path.join(
            self.configFolder, "config", "mMerger.yaml"
        )
        with open(merger_yaml, "r") as stream:
            self.merger_config = yaml.safe_load(stream)

    # ==============================================================
    # FILE MOVE / CLEANUP
    # ==============================================================

    def move_files_to_docs(self, input_list, unique_id):
        source_files_option = self.merger_config["SOURCE"]
        selected_docs = input_list["selected"]
        removed_docs = input_list.get("removed", {})

        files = list(selected_docs.values())
        file_path = os.path.dirname(files[0])
        backup_path = os.path.join(file_path, "docs")

        os.makedirs(backup_path, exist_ok=True)

        def move_to_backup(doc_map):
            for name, path in doc_map.items():
                shutil.move(path, os.path.join(backup_path, name))

        def remove_files(doc_map):
            for path in doc_map.values():
                os.remove(path)

        if source_files_option == "Retain":
            move_to_backup(selected_docs)
            move_to_backup(removed_docs)
        elif source_files_option == "Remove":
            remove_files(selected_docs)
            remove_files(removed_docs)

    # ==============================================================
    # MERGE USING WORD (DOC)
    # ==============================================================

    def merge_in_doc_old(self, input_list, unique_id):
        merger_result = False
        customer = input_list["customer"]
        folder = input_list["folder"]

        cln_doc_name = self.merger_config["FILENAME"][customer]
        cln_doc_name = re.sub("FOLDER", folder, cln_doc_name)

        files = list(input_list["selected"].values())
        file_path = os.path.dirname(files[0])
        composed_file = os.path.join(file_path, cln_doc_name)

        file1 = files.pop(0)
        macro_body = ""

        for file in files:
            file = os.path.abspath(file)
            macro_body += (
                f'\t\t.InsertFile FileName := "{file}", '
                "ConfirmConversions := False, "
                "Link := False, Attachment := False\n"
            )

        word = win32.gencache.EnsureDispatch("Word.Application")
        word.Visible = True
        doc = word.Documents.Open(file1)
        time.sleep(3)

        try:
            macro = doc.VBProject.VBComponents.Add(1)
        except Exception as e:
            print(e)
            return False

        code = f"""
        sub FileMerger()
            Selection.EndKey Unit:=wdStory
            With Selection
            {macro_body}
            End With
        end sub
        """

        macro.CodeModule.AddFromString(code)

        try:
            doc.Application.Run("FileMerger")
            merger_result = True
        except Exception as e:
            print(e)

        if merger_result:
            try:
                word.ActiveDocument.SaveAs(
                    composed_file,
                    FileFormat=constants.wdFormatXMLDocument,
                )
                word.ActiveDocument.Close()
                self.find_duplicates_new(composed_file)
            except Exception as e:
                print(e)
                merger_result = False

        time.sleep(3)
        return merger_result

    def merge_in_doc(self, input_list, unique_id):
        """Merge documents using Word automation with proper COM handling"""
        merger_result = False
        customer = input_list["customer"]
        folder = input_list["folder"]

        cln_doc_name = self.merger_config["FILENAME"][customer]
        cln_doc_name = re.sub("FOLDER", folder, cln_doc_name)

        files = list(input_list["selected"].values())
        file_path = os.path.dirname(files[0])
        composed_file = os.path.join(file_path, cln_doc_name)

        file1 = files.pop(0)
        macro_body = ""

        for file in files:
            file = os.path.abspath(file)
            macro_body += (
                f'\t\t.InsertFile FileName := "{file}", '
                "ConfirmConversions := False, "
                "Link := False, Attachment := False\n"
            )

        try:
            # Use WordApplicationManager for proper COM handling
            with WordApplicationManager(visible=True) as word:
                # Open the first document
                doc = word.Documents.Open(file1)
                time.sleep(2)  # Give Word time to open

                # Try to add macro
                try:
                    macro = doc.VBProject.VBComponents.Add(1)
                except Exception as e:
                    print(f"Could not add VBProject macro: {e}")
                    # Fallback: use Selection.InsertFile directly
                    word.Selection.EndKey(Unit=constants.wdStory)

                    for file in files:
                        try:
                            word.Selection.InsertFile(
                                FileName=file,
                                ConfirmConversions=False,
                                Link=False,
                                Attachment=False
                            )
                        except Exception as insert_error:
                            print(f"Error inserting file {file}: {insert_error}")
                            continue

                    merger_result = True
                else:
                    # Add macro code
                    code = f"""
                    sub FileMerger()
                        Selection.EndKey Unit:=wdStory
                        With Selection
                        {macro_body}
                        End With
                    end sub
                    """

                    try:
                        macro.CodeModule.AddFromString(code)
                        doc.Application.Run("FileMerger")
                        merger_result = True
                    except Exception as macro_error:
                        print(f"Macro execution failed: {macro_error}")
                        # Try alternative method
                        word.Selection.EndKey(Unit=constants.wdStory)
                        for file in files:
                            word.Selection.InsertFile(
                                FileName=file,
                                ConfirmConversions=False,
                                Link=False,
                                Attachment=False
                            )
                        merger_result = True

                # Save the merged document
                if merger_result:
                    try:
                        doc.SaveAs(
                            FileName=composed_file,
                            FileFormat=constants.wdFormatXMLDocument
                        )
                        doc.Close(SaveChanges=0)

                        # Process duplicates
                        self.find_duplicates_new(composed_file)

                    except Exception as save_error:
                        print(f"Error saving document: {save_error}")
                        merger_result = False

        except Exception as e:
            print(f"Word automation failed: {e}")
            merger_result = False

        return merger_result

    # ==============================================================
    # THREAD-SAFE WORD OPERATIONS
    # ==============================================================

    def merge_in_doc_threadsafe(self, input_list, unique_id):
        """Thread-safe version for use in worker threads"""
        result = [False]  # Use list to capture result from nested function

        def merge_worker():
            try:
                # Use com_context for worker threads
                from com_manager import COMManager
                with COMManager.com_context():
                    result[0] = self.merge_in_doc(input_list, unique_id)
            except Exception as e:
                print(f"Thread-safe merge failed: {e}")
                result[0] = False

        # Run in a separate thread
        thread = threading.Thread(target=merge_worker)
        thread.start()
        thread.join(timeout=300)  # 5 minute timeout

        if thread.is_alive():
            print("Word automation timeout")
            return False

        return result[0]

    # ==============================================================
    # ENHANCED MERGE WITH FALLBACK STRATEGY
    # ==============================================================

    def merge_docx_robust(self, input_list, unique_id):
        """Robust merging with multiple fallback strategies"""
        customer = input_list["customer"]
        folder = input_list["folder"]

        cln_doc_name = self.merger_config["FILENAME"][customer]
        cln_doc_name = re.sub("FOLDER", folder, cln_doc_name)

        files = list(input_list["selected"].values())
        file_path = os.path.dirname(files[0])
        composed_file = os.path.join(file_path, cln_doc_name)

        # Strategy 1: Try python-docx for simple documents
        if not self._has_diagrams_or_complex_structure(files):
            try:
                if len(files) == 1:
                    shutil.copy(files[0], composed_file)
                    self.find_duplicates_new(composed_file)
                    return True

                result = Document(files[0])
                composer = Composer(result)

                for i, fname in enumerate(files[1:], start=1):
                    doc = Document(fname)
                    if i != len(files) - 1:
                        doc.add_page_break()
                    composer.append(doc)

                composer.save(composed_file)
                self.find_duplicates_new(composed_file)
                return True

            except Exception as e:
                print(f"Python-docx merge failed: {e}, falling back to Word...")

        # Strategy 2: Try Word automation
        try:
            result = self.merge_in_doc(input_list, unique_id)
            if result:
                return True
        except Exception as e:
            print(f"Word automation failed: {e}")

        # Strategy 3: Thread-safe Word automation as last resort
        try:
            print("Attempting thread-safe Word merge...")
            result = self.merge_in_doc_threadsafe(input_list, unique_id)
            return result
        except Exception as e:
            print(f"All merge strategies failed: {e}")

        return False

    # ==============================================================
    # ERROR HANDLING AND LOGGING
    # ==============================================================

    def log_merge_attempt(self, files, method, success, error=None):
        """Log merge attempts for debugging"""
        log_entry = {
            'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
            'files': [os.path.basename(f) for f in files],
            'method': method,
            'success': success,
            'error': str(error) if error else None
        }

        # Log to file
        log_file = os.path.join(self.configFolder, 'merge_log.json')
        try:
            import json
            logs = []
            if os.path.exists(log_file):
                with open(log_file, 'r') as f:
                    logs = json.load(f)
            logs.append(log_entry)
            with open(log_file, 'w') as f:
                json.dump(logs, f, indent=2)
        except:
            pass

        # Update database if available
        if not success and hasattr(self, 'db_process'):
            error_msg = f"{method}: {error}" if error else method
            self.db_process.update_db(
                unique_id, "mMerger", "ERROR", "", error_msg
            )

    # ==============================================================
    # MERGE DOCX (COMPOSER)
    # ==============================================================
    def merge_docx(self, input_list, unique_id):
        customer = input_list["customer"]
        folder = input_list["folder"]

        cln_doc_name = self.merger_config["FILENAME"][customer]
        cln_doc_name = re.sub("FOLDER", folder, cln_doc_name)

        files = list(input_list["selected"].values())
        file_path = os.path.dirname(files[0])
        composed_file = os.path.join(file_path, cln_doc_name)

        # For documents with diagrams or complex structure, use Word
        if self._has_diagrams_or_complex_structure(files):
            print("Document has diagrams, using Word automation...")
            return self.merge_in_doc(input_list, unique_id)

        # Simple documents can use python-docx
        try:
            if len(files) > 1:
                result = Document(files[0])
                composer = Composer(result)

                for i, fname in enumerate(files[1:], start=1):
                    doc = Document(fname)
                    if i != len(files) - 1:
                        doc.add_page_break()
                    composer.append(doc)

                composer.save(composed_file)
                self.find_duplicates_new(composed_file)
                return True
            else:
                shutil.copy(files[0], composed_file)
                self.find_duplicates_new(composed_file)
                return True

        except Exception as e:
            print(f"Python-docx merge failed: {e}, falling back to Word...")
            return self.merge_in_doc(input_list, unique_id)

    def _has_diagrams_or_complex_structure(self, files):
        """Check if any document has diagrams or might cause merge issues"""
        for file in files:
            try:
                with zipfile.ZipFile(file, 'r') as doc_zip:
                    file_list = doc_zip.namelist()
                    # Check for diagrams or other complex elements
                    if any('diagram' in f.lower() for f in file_list):
                        return True
                    if any('chart' in f.lower() for f in file_list):
                        return True
                    if any('drawing' in f.lower() for f in file_list):
                        return True
            except:
                pass
        return False

    def merge_docx_old(self, input_list, unique_id):
        customer = input_list["customer"]
        folder = input_list["folder"]

        cln_doc_name = self.merger_config["FILENAME"][customer]
        cln_doc_name = re.sub("FOLDER", folder, cln_doc_name)

        files = list(input_list["selected"].values())
        file_path = os.path.dirname(files[0])
        composed_file = os.path.join(file_path, cln_doc_name)

        if len(files) > 1:
            result = Document(files[0])
            result.add_page_break()
            composer = Composer(result)

            for i, fname in enumerate(files[1:], start=1):
                doc = Document(fname)
                if i != len(files) - 1:
                    doc.add_page_break()
                try:
                    composer.append(doc)
                except Exception as e:
                    self.db_process.update_db(
                        unique_id,
                        "mMerger",
                        "ERROR",
                        "",
                        f"Unable to append document {fname}",
                    )
                    print(e)
                    return False

            composer.save(composed_file)
            self.find_duplicates_new(composed_file)
            return True

        # Single file case
        shutil.copy(files[0], composed_file)
        self.find_duplicates_new(composed_file)
        return True

    # ==============================================================
    # DUPLICATE DETECTION (ACTIVE VERSION)
    # ==============================================================

    def find_duplicates_new(self, docxfile):
        document = Document(docxfile)

        try:
            document.styles["Duplicate"].delete()
        except KeyError:
            pass

        dup_style = document.styles.add_style(
            "Duplicate", WD_STYLE_TYPE.PARAGRAPH
        )
        dup_style.base_style = document.styles["Normal"]

        font = dup_style.font
        font.name = "Times New Roman"
        font.size = Pt(12)
        font.color.rgb = RGBColor(255, 0, 0)
        font.highlight_color = WD_COLOR_INDEX.YELLOW

        para_fmt = dup_style.paragraph_format
        para_fmt.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para_fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para_fmt.left_indent = Pt(0)
        para_fmt.right_indent = Pt(0)
        para_fmt.space_before = Pt(0)
        para_fmt.space_after = Pt(0)

        para_map = {}
        for idx, para in enumerate(document.paragraphs):
            if para.text.strip():
                para_map.setdefault(para.text, []).append(idx)

        for indexes in para_map.values():
            if len(indexes) > 1:
                for dup_idx in indexes[1:]:
                    document.paragraphs[dup_idx].style = dup_style

        document.save(docxfile)

    # ==============================================================
    # PARAGRAPH DELETE (UTILITY)
    # ==============================================================

    def delete_paragraph(self, paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
