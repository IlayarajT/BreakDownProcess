import re
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
import langid
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

class DocxPreClean:
    def __init__(self):
        pass

    def preCleanDocx(self, docxfile):
        document = Document(docxfile)
        document = self.remove_header_footer(document)
        document = self.remove_multi_cols(document)
        document = self.change_other_lang_styles(document)
        document = self.clean_tab_para(document)
        document = self.clean_eqn_table(document)
        document = self.remove_section_break(document)
        document = self.remove_page_break(document)
        document = self.remove_tab_after_math(document)
        document = self.remove_run_properties(document)
        document = self.remove_alignments(document)
        document.save(docxfile)

    def remove_run_properties(self, document):
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                for bookmark in run.element.findall(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkStart'):
                    bookmark.getparent().remove(bookmark)
                for bookmark in run.element.findall(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkEnd'):
                    bookmark.getparent().remove(bookmark)
            pPr = paragraph._element.pPr
            if pPr is not None:
                for child in pPr:
                    # Specify the tags of the child elements you want to remove
                    if child.tag.endswith('rPr'):
                        pPr.remove(child)
                    # if child.tag.endswith('spacing') or \
                    #         child.tag.endswith('rPr') or \
                    #         child.tag.endswith('ind'):
                    #     pPr.remove(child)
            for run in paragraph.runs:
                # print(run.text)
                rPr = run._element.rPr
                if rPr is not None:
                    for child in rPr:
                        if child.tag.endswith('sz') or child.tag.endswith('szCs') or \
                                child.tag.endswith('shd') or \
                                child.tag.endswith('rStyle') or \
                                child.tag.endswith('color') or \
                                child.tag.endswith('u') or \
                                child.tag.endswith('spacing') or \
                                child.tag.endswith('textOutline') or \
                                child.tag.endswith('position'):
                            rPr.remove(child)
        #                                 child.tag.endswith('rFonts') or \
        return document

    def clean_eqn_table(self, document):
        styles = document.styles
        if "EqnTable" not in styles:
            style = document.styles.add_style("EqnTable", WD_STYLE_TYPE.TABLE)
        for table in document.tables:
            row_count = len(table.rows)
            row_removed = False
            if row_count == 2:
                for row in table.rows:
                    cell_count = 1
                    cell_content = {}
                    cell_content.clear()
                    col_count = len(table.rows[0].cells)
                    if col_count == 2:
                        del cell_count
                        cell_count = 1
                        cell_content.clear()
                        cell_content[1] = None
                        cell_content[2] = None
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                para_text = paragraph.text
                                para_text = re.sub(r"(\s+|\t+|\;|\.)", '', para_text)
                                para_len = len(para_text)
                                if para_len == 0:
                                    para_xml = paragraph._element.xml
                                    if 'w:object' in para_xml or 'm:oMathPara' in para_xml or 'w:drawing' in para_xml:
                                        cell_content[cell_count] = False
                                    else:
                                        if cell_content[cell_count] is None:
                                            cell_content[cell_count] = True
                                else:
                                    cell_content[cell_count] = False
                            cell_count = cell_count + 1
                        if cell_content[1] is True and cell_content[2] is True:
                            self.remove_row(table, row)
                            row_removed = True
                    elif col_count == 3:
                        del cell_count
                        cell_count = 1
                        cell_content.clear()
                        cell_content[1] = None
                        cell_content[2] = None
                        cell_content[3] = None
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                para_text = paragraph.text
                                para_text = re.sub(r"(\s+|\t+|\;|\.)", '', para_text)
                                para_len = len(para_text)
                                if para_len == 0:
                                    para_xml = paragraph._element.xml
                                    if 'w:object' in para_xml or 'm:oMathPara' in para_xml or 'w:drawing' in para_xml:
                                        cell_content[cell_count] = False
                                    else:
                                        if cell_content[cell_count] is None:
                                            cell_content[cell_count] = True
                                else:
                                    cell_content[cell_count] = False
                            cell_count = cell_count + 1
                        if cell_content[1] is True and cell_content[2] is True and cell_content[3] is True:
                            self.remove_row(table, row)
                            row_removed = True
                if row_removed is True:
                    self.check_cols(table)
            elif row_count == 1:
                self.check_cols(table)
        return document

    def remove_row(self, table, row):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)

    def check_cols(self, table):
        col_count = len(table.rows[0].cells)
        cell_count = 1
        cell_content = {}
        cell_content.clear()
        if col_count == 2:
            cells = table.rows[0].cells
            del cell_count
            cell_count = 1
            cell_content.clear()
            nested_table = False
            for cell in cells:
                for nst_table in cell.tables:
                    nested_table = True
                    break
                if nested_table is True:
                    break
                if nested_table is False:
                    for paragraph in cell.paragraphs:
                        para_text = paragraph.text
                        # print(paragraph._p.xml)
                        re.sub(r"(\s+|\t+)", "", para_text)
                        para_len = len(para_text)
                        if para_len == 0:
                            para_xml = paragraph._element.xml
                            if 'ProgID="Equation' in para_xml:
                                cell_content[cell_count] = "Math"
                                break
                            elif 'm:oMathPara' in para_xml:
                                cell_content[cell_count] = "Math"
                                break
                            elif "w:drawing" in para_xml:
                                cell_content[cell_count] = "Pic"
                                break
                        else:
                            cell_content[cell_count] = para_text
                    cell_count = cell_count + 1
            if nested_table is False:
                if cell_content[1] == "Math" and bool(re.search(r"\([0-9a-z]+\)", cell_content[2])) is True:
                    table.style = "EqnTable"
                elif cell_content[2] == "Math" and bool(re.search(r"\([0-9a-z]+\)", cell_content[1])) is True:
                    table.style = "EqnTable"
        if col_count == 1:
            cells = table.rows[0].cells
            del cell_count
            cell_count = 1
            cell_content.clear()
            nested_table = False
            for cell in cells:
                for nst_table in cell.tables:
                    nested_table = True
                    break
                if nested_table is True:
                    break
                if nested_table is False:
                    for paragraph in cell.paragraphs:
                        para_text = paragraph.text
                        if para_text is None:
                            objects = paragraph._element.xpath('.//w:object')
                            para_xml = paragraph._element.xml
                            if 'ProgID="Equation' in para_xml:
                                table.style = "EqnTable"
        if col_count == 3:
            cells = table.rows[0].cells
            del cell_count
            cell_count = 1
            cell_content.clear()
            nested_table = False
            for cell in cells:
                for nst_table in cell.tables:
                    nested_table = True
                    break
                if nested_table is True:
                    break
                if nested_table is False:
                    for paragraph in cell.paragraphs:
                        para_text = paragraph.text
                        para_text = re.sub(r"(\s+|\t+|\;|\.)", '', para_text)
                        para_len = len(para_text)
                        if para_len == 0:
                            para_xml = paragraph._element.xml
                            if 'ProgID="Equation' in para_xml:
                                cell_content[cell_count] = "Math"
                                break
                            elif 'm:oMathPara' in para_xml:
                                cell_content[cell_count] = "Math"
                                break
                            elif "w:drawing" in para_xml:
                                cell_content[cell_count] = "Pic"
                                break
                            else:
                                cell_content[cell_count] = "Empty"
                        else:
                            cell_content[cell_count] = para_text
                    cell_count = cell_count + 1
            if cell_content[1] == "Empty" and cell_content[2] == "Math" and bool(
                    re.search(r"\([0-9a-z]+\)", cell_content[3])) is True:
                table.style = "EqnTable"
            elif cell_content[1] == "Empty" and cell_content[3] == "Math" and bool(
                    re.search(r"\([0-9a-z]+\)", cell_content[2])) is True:
                table.style = "EqnTable"
            elif cell_content[1] == "Math" and cell_content[2] == "Empty" and bool(
                    re.search(r"\([0-9a-z]+\)", cell_content[3])) is True:
                table.style = "EqnTable"
            elif cell_content[1] == "Math" and bool(re.search(r"\([0-9a-z]+\)", cell_content[2])) is True and \
                    cell_content[3] == "Empty":
                table.style = "EqnTable"
            elif bool(re.search(r"\([0-9a-z]+\)", cell_content[1])) is True and cell_content[2] == "Empty" and \
                    cell_content[3] == "Math":
                table.style = "EqnTable"
            elif bool(re.search(r"\([0-9a-z]+\)", cell_content[1])) is True and cell_content[2] == "Math" and \
                    cell_content[3] == "Empty":
                table.style = "EqnTable"

    def clean_tab_para(self, document):
        for paragraph in document.paragraphs:
            run_cnt = len(paragraph.runs)
            if run_cnt > 0:
                f_run = paragraph.runs[0]
                f_text = f_run.text
                f_text = re.sub("\t", "", f_text)
                para_len = len(f_text)
                if para_len == 0:
                    para_xml = f_run._r.xml
                    if 'w:object' in para_xml or 'm:oMathPara' in para_xml or 'w:drawing' in para_xml:
                        pass
                    elif 'w:tab' in para_xml:
                        first_run = f_run._r
                        p = paragraph._p
                        p.remove(first_run)
        return document

    def remove_tab_after_math(self, document):
        for paragraph in document.paragraphs:
            run_cnt = len(paragraph.runs)
            if run_cnt > 0:
                run_count = 0
                math_found = False
                remove_tab = False
                for run in paragraph.runs:
                    run_xml = run._r.xml
                    if 'ProgID="Equation' in run_xml or 'm:oMathPara' in run_xml:
                        math_found = True
                        break
                    run_count = run_count + 1
                if math_found is True:
                    run_count = run_count + 1
                    run_lenth = len(paragraph.runs)
                    if run_count < run_lenth:
                        next_run = paragraph.runs[run_count]
                        run_text = next_run.text
                        if re.search("\t", run_text):
                            next_run.clear()
                            run_text = re.sub("\t", " ", run_text)
                            next_run.text = run_text
                            paragraph.runs[run_count] = next_run
        return document

    def remove_all_bookmarks(self, document):
        bookmartStart = document._element.xpath(".//w:bookmarkStart")
        bookmartEnd = document._element.xpath(".//w:bookmarkEnd")
        if len(bookmartStart) != 0:
            if len(bookmartStart) == len(bookmartEnd):
                for start in bookmartStart:
                    start.getparent().remove(start)
                for end in bookmartEnd:
                    end.getparent().remove(end)
        return document

    def remove_page_break(self, document):
        for paragraph in document.paragraphs:
            p = paragraph._p
            sectPrs = p.xpath("./w:pPr/w:sectPr")
            if not sectPrs:
                continue
            sectPr = sectPrs[0]
            sectPr.getparent().remove(sectPr)
        return document

    def remove_section_break(self, document):
        for paragraph in document.paragraphs:
            p = paragraph._p
            sectPrs = p.xpath("./w:pPr/w:sectPr")
            if not sectPrs:
                continue
            sectPr = sectPrs[0]
            sectPr.getparent().remove(sectPr)
        return document

    def remove_multi_cols(self, document):
        section = document.sections[0]
        sectPr = section._sectPr
        sec_cols = sectPr.xpath('./w:cols')
        if sec_cols:
            cols = sectPr.xpath('./w:cols')[0]
            if len(cols) > 0:
                cols.set(qn('w:num'), '1')
        return document

    def change_other_lang_styles(self, document):
        normal_style = document.styles["Normal"]
        font = normal_style.font
        font.name = 'Times New Roman'
        font.color.rgb = RGBColor(0, 0, 0)
        para = normal_style.paragraph_format
        para.first_line_indent = None
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para.left_indent = Pt(0)
        para.right_indent = Pt(0)
        para.space_after = Pt(0)
        para.space_before = Pt(0)
        non_english = False
        for paragraph in document.paragraphs:
            style_name = paragraph.style.name
            lang_tuple = langid.classify(style_name)
            lang_name = lang_tuple[0]
            if not lang_name == 'en':
                paragraph.style.name = "Normal"
        return document

    def remove_page_break(self, document):
        for paragraph in document.paragraphs:
            p = paragraph._p
            breakPrs = p.xpath(".//w:br")
            if not breakPrs:
                continue
            breakPr = breakPrs[0]
            if breakPr.type == "page":
                breakPr.getparent().remove(breakPr)
        return document

    def remove_header_footer(self, document):
        section = document.sections[0]
        header = section.header
        footer = section.footer
        header_text = None
        footer_text = None
        for paragraph in header.paragraphs:
            header_text = paragraph.text
            paragraph.text = ""
        for paragraph in footer.paragraphs:
            footer_text = paragraph.text
            paragraph.text = ""
        if header_text is not None and re.search("[a-z]", header_text, re.IGNORECASE):
            header_text = re.sub("\t+", "\t", header_text)
            header_text = re.sub("\t([0-9]+)", "", header_text)
            document.add_paragraph(f"Header: {header_text}")
            document.sections[0].header.text = ""
        if footer_text is not None and re.search("[a-z]", footer_text, re.IGNORECASE):
            footer_text = re.sub("\t+", "\t", footer_text)
            footer_text = re.sub("\t([0-9]+)", "", footer_text)
            document.add_paragraph(f"Footer: {footer_text}")
        return document

    def is_paragraph_in_table(self, paragraph):
        return paragraph._element.xpath('ancestor::w:tbl')

    def remove_alignments(self, document):
        for paragraph in document.paragraphs:
            if self.is_paragraph_in_table(paragraph):
                continue
            pPr = paragraph._p.pPr
            if pPr is None or pPr.jc is None:
                continue
            jc = pPr.jc
            jc_val = jc.get(qn("w:val"))
            if jc_val in ("center", "start"):
                jc.set(qn("w:val"), "left")
        return document

# preCln = DocxPreClean()
# preCln.preCleanDocx("V:\\FOR_BREAKDOWN\\MERGER_INPUT\\SAGE\\TAB_1412835\\HPOS_ProtocolPaper_09.12.2025_clean.docx")
# document = Document('V:\FOR_BREAKDOWN\ParaStyler_INPUT\SAGE\JIV_1260004\JIV_1260004_CLN.docx')
# document = preCln.remove_alignments(document)
# document.save('V:\FOR_BREAKDOWN\ParaStyler_INPUT\SAGE\JIV_1260004\JIV_1260004_CLN.docx')

