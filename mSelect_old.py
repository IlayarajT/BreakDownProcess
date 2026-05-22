import re
import yaml
import glob
import os
from PyQt6.QtCore import Qt
from PyQt6 import QtCore, QtGui, QtWidgets
from PyQt6.QtWidgets import QMessageBox, QApplication, QMainWindow, QPushButton, QLabel, QVBoxLayout, QWidget
import sys
import getAppPath
from mMerger import DocxMerger
import shutil
from docx2pdf import convert
from PyQt6.QtCore import QUrl
from PyQt6.QtWebEngineWidgets import QWebEngineView
from subprocess import *
from applyStyles import ApplyStyles
from breakDownProcess import BreakDownProcess
import win32com.client as win32
from dbprocess import DataBase
from docx import Document
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from openDocFile_new import OpenDocFile
from CreateArticleInfo import GetArticleId
from loadconfig import getconfig
import subprocess
from docxManipulator import DocxManipulator
from TransformXml import XmlTransform
from createQuoteInfo import CreateParaInfo
from docx.oxml import parse_xml
from lxml import etree
from italic_bookmark import ItalicBookmarkProcessor
import pythoncom
from com_manager import COMManager


class ChildWindow(QMainWindow):
    def __init__(self):
        self.processDoc = OpenDocFile()
        self.app_path = getAppPath.getapppath()
        self.configFolder, self.breakDownConfig = getconfig()
        self.db_process = DataBase()
        super(ChildWindow, self).__init__()
        self.setWindowTitle("File Viewer [C&M]")
        self.setGeometry(200, 50, 800, 600)
        self.webView = QWebEngineView()
        self.webView.settings().setAttribute(self.webView.settings().WebAttribute.PluginsEnabled, True)
        self.webView.settings().setAttribute(self.webView.settings().WebAttribute.PdfViewerEnabled, True)
        self.setCentralWidget(self.webView)

    def closeEvent(self, event):
        widgetList = QApplication.topLevelWidgets()
        numWindows = len(widgetList)
        os.remove(pdf_file)
        if numWindows > 1:
            event.accept()
        else:
            event.ignore()

    def url_changed(self):
        self.setWindowTitle(self.webView.title())

    def go_back(self):
        self.webView.back()


class Ui_Dialog(QWidget):
    def __init__(self):
        super(Ui_Dialog, self).__init__()
        self.normalizer_input = None
        self.removed_list = None
        self.selected_list = None
        self.file_names = None
        self.processCombo = None
        self.file_path = None
        self.return_list = ['Select']
        self.input_list = []
        self.listofjournals = []
        self.folder_details = {}
        self.file_details = {}
        self.app_path = getAppPath.getapppath()
        self.configFolder, self.breakDownConfig = getconfig()
        dialog_yml = os.path.join(self.configFolder, 'config\\dialogConfig.yaml')
        with open(dialog_yml, "r") as stream:
            self.dialog_config = yaml.safe_load(stream)
        breakdown_yaml = os.path.join(self.configFolder, 'config\\breakDown.yaml')
        with open(breakdown_yaml, "r") as stream:
            self.breakDownConfig = yaml.safe_load(stream)
        self.mergerInput = self.breakDownConfig['FOLDERS']['MERGER_INPUT']
        self.mergerError = self.breakDownConfig['FOLDERS']['MERGER_ERROR']
        self.ParaStylerInput = self.breakDownConfig['FOLDERS']['ParaStyler_INPUT']
        self.ParaStylerError = self.breakDownConfig['FOLDERS']['ParaStyler_ERROR']
        self.BreakDownInput = self.breakDownConfig['FOLDERS']['BreakDown_INPUT']
        self.BreakDownError = self.breakDownConfig['FOLDERS']['BreakDown_ERROR']
        self.BreakDownDone = self.breakDownConfig['FOLDERS']['BreakDown_DONE']
        self.firstprocess = self.dialog_config['PROCESS'][0]
        self.firstcustomer = self.dialog_config['CUSTOMERS'][0]
        self.jid_aid_list, self.folder_details = self.get_aid_jid_list(self.firstprocess, self.firstcustomer,
                                                                       self.dialog_config)

    def setupUi(self, Dialog):
        Dialog.setObjectName("Dialog")
        Dialog.resize(700, 430)

        # Set the overall stylesheet for the dialog
        Dialog.setStyleSheet("""
            QDialog {
                background-color: #2E3440; /* Dark gray background */
                color: #ECEFF4; /* Light gray text */
            }
            QFrame {
                background-color: #3B4252; /* Slightly lighter gray for frames */
                border-radius: 5px;
            }
            QLabel {
                color: #ECEFF4; /* Light gray text */
            }
            QComboBox {
                background-color: #4C566A; /* Dark gray background */
                color: #ECEFF4; /* Light gray text */
                border: 1px solid #5E81AC; /* Blue border */
                border-radius: 3px;
                padding: 5px;
            }
            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 15px;
                border-left-width: 1px;
                border-left-color: #5E81AC; /* Blue border */
                border-left-style: solid;
                border-top-right-radius: 3px;
                border-bottom-right-radius: 3px;
            }
            QComboBox::down-arrow {
                width: 0;
                height: 0;
                border-style: solid;
                border-width: 6px 5px 0 3px; /* Triangle shape */
                border-color: #ECEFF4 transparent transparent transparent; /* White triangle */
                margin-right: 4px; /* Align it neatly within the drop-down */
                margin-left: 6px; /* Align it neatly within the drop-down */
            }
            QComboBox QAbstractItemView {
                background-color: #4C566A; /* Dark gray background */
                color: #ECEFF4; /* Light gray text */
                selection-background-color: #5E81AC; /* Blue selection */
            }
            QPushButton {
                background-color: #5E81AC; /* Blue background */
                color: #ECEFF4; /* Light gray text */
                font-weight: bold;
                border: none;
                border-radius: 3px;
                padding: 5px 10px;
            }
            QPushButton:hover {
                background-color: #88C0D0; /* Lighter blue on hover */
                color: #2E3440; /* Dark text for improved readability */
            }
            QPushButton:pressed {
                background-color: #BF616A; /* Dark gray when pressed */
                color: #ECEFF4; /* Keeping text light for contrast */
            }
            QListWidget {
                background-color: #4C566A; /* Dark gray background */
                color: #ECEFF4; /* Light gray text */
                border: 1px solid #5E81AC; /* Blue border */
                border-radius: 3px;
            }
            QListWidget::item:selected {
                background-color: #FFFFFF;
                color: #000000;
                font-weight: bold;
                }
            QListWidget::item:hover {
                background-color: #434C5E;
                }
            QProgressBar {
                background-color: #4C566A; /* Dark gray background */
                color: #ECEFF4; /* Light gray text */
                border: 1px solid #5E81AC; /* Blue border */
                border-radius: 3px;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #5E81AC; /* Blue progress */
                border-radius: 3px;
            }
        """)

        font = QtGui.QFont()
        font.setFamily("Calibri")
        font.setPointSize(9)
        font.setWeight(75)
        font1 = QtGui.QFont()
        font1.setFamily("Calibri")
        font1.setPointSize(10)
        font1.setWeight(600)
        Dialog.setFont(font)

        self.frame = QtWidgets.QFrame(Dialog)
        self.frame.setGeometry(QtCore.QRect(20, 10, 671, 50))
        self.frame.setFrameShape(QtWidgets.QFrame.Shape.StyledPanel)
        self.frame.setFrameShadow(QtWidgets.QFrame.Shadow.Raised)
        self.frame.setObjectName("frame")

        self.processLabel = QtWidgets.QLabel(self.frame)
        self.processLabel.setGeometry(QtCore.QRect(15, 10, 60, 28))
        self.processLabel.setFont(font)
        self.processLabel.setObjectName("label")

        self.processCombo = QtWidgets.QComboBox(self.frame)
        self.processCombo.setGeometry(QtCore.QRect(75, 10, 125, 28))
        self.processCombo.setObjectName("comboBox")
        self.processCombo.addItems(self.dialog_config['PROCESS'])
        self.processCombo.currentTextChanged.connect(self.process_combobox_changed)

        self.customerCombo = QtWidgets.QComboBox(self.frame)
        self.customerCombo.setGeometry(QtCore.QRect(310, 10, 80, 28))
        self.customerCombo.setObjectName("comboBox_2")
        self.customerCombo.addItems(self.dialog_config['CUSTOMERS'])
        self.customerCombo.currentTextChanged.connect(self.customer_combobox_changed)

        self.jidAidCombo = QtWidgets.QComboBox(self.frame)
        self.jidAidCombo.setGeometry(QtCore.QRect(480, 10, 150, 28))
        self.jidAidCombo.setObjectName("comboBox_3")
        self.jidAidCombo.addItems(self.jid_aid_list)
        self.jidAidCombo.currentTextChanged.connect(self.get_file_list)

        self.customerLabel = QtWidgets.QLabel(self.frame)
        self.customerLabel.setGeometry(QtCore.QRect(235, 10, 65, 28))
        self.customerLabel.setFont(font)
        self.customerLabel.setObjectName("label_2")

        self.jidAidLabel = QtWidgets.QLabel(self.frame)
        self.jidAidLabel.setGeometry(QtCore.QRect(420, 10, 50, 28))
        self.jidAidLabel.setFont(font)
        self.jidAidLabel.setObjectName("label_4")

        self.frame_2 = QtWidgets.QFrame(Dialog)
        self.frame_2.setGeometry(QtCore.QRect(20, 70, 670, 310))
        self.frame_2.setFrameShape(QtWidgets.QFrame.Shape.StyledPanel)
        self.frame_2.setFrameShadow(QtWidgets.QFrame.Shadow.Raised)
        self.frame_2.setObjectName("frame_2")

        self.selectedList = QtWidgets.QListWidget(self.frame_2)
        self.selectedList.setGeometry(QtCore.QRect(20, 30, 530, 145))
        self.selectedList.setObjectName("listWidget")

        self.removeList = QtWidgets.QListWidget(self.frame_2)
        self.removeList.setGeometry(QtCore.QRect(20, 200, 530, 95))
        self.removeList.setObjectName("listWidget_2")

        self.btnIns = QtWidgets.QPushButton(self.frame_2)
        self.btnIns.setGeometry(QtCore.QRect(575, 30, 85, 20))
        self.btnIns.setFont(font)
        self.btnIns.setObjectName("btnIns")

        self.btnOpen = QtWidgets.QPushButton(self.frame_2)
        self.btnOpen.setGeometry(QtCore.QRect(575, 60, 85, 20))
        self.btnOpen.setFont(font)
        self.btnOpen.setObjectName("btnOpen")

        self.btnMoveUp = QtWidgets.QPushButton(self.frame_2)
        self.btnMoveUp.setGeometry(QtCore.QRect(575, 90, 85, 20))
        self.btnMoveUp.setFont(font)
        self.btnMoveUp.setObjectName("btnMoveUp")

        self.btnMoveDown = QtWidgets.QPushButton(self.frame_2)
        self.btnMoveDown.setGeometry(QtCore.QRect(575, 120, 85, 20))
        self.btnMoveDown.setFont(font)
        self.btnMoveDown.setObjectName("btnMoveDown")

        self.btnRemove = QtWidgets.QPushButton(self.frame_2)
        self.btnRemove.setGeometry(QtCore.QRect(575, 150, 85, 20))
        self.btnRemove.setFont(font)
        self.btnRemove.setObjectName("btnRemove")

        self.btnAdd = QtWidgets.QPushButton(self.frame_2)
        self.btnAdd.setGeometry(QtCore.QRect(575, 230, 85, 20))
        self.btnAdd.setFont(font)
        self.btnAdd.setObjectName("btnAdd")

        self.selFilesLabel = QtWidgets.QLabel(self.frame_2)
        self.selFilesLabel.setGeometry(QtCore.QRect(20, 10, 130, 20))
        self.selFilesLabel.setFont(font)
        self.selFilesLabel.setObjectName("label_5")

        self.rmvFilesLabel = QtWidgets.QLabel(self.frame_2)
        self.rmvFilesLabel.setGeometry(QtCore.QRect(20, 180, 110, 20))
        self.rmvFilesLabel.setFont(font)
        self.rmvFilesLabel.setObjectName("label_6")

        self.btnOK = QtWidgets.QPushButton(Dialog)
        self.btnOK.setGeometry(QtCore.QRect(160, 390, 80, 25))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Fixed, QtWidgets.QSizePolicy.Policy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.btnOK.sizePolicy().hasHeightForWidth())
        self.btnOK.setSizePolicy(sizePolicy)
        self.btnOK.setSizeIncrement(QtCore.QSize(0, 0))
        self.btnOK.setFont(font1)
        self.btnOK.setObjectName("pushButton")

        self.btnCancel = QtWidgets.QPushButton(Dialog)
        self.btnCancel.setGeometry(QtCore.QRect(290, 390, 80, 25))
        self.btnCancel.setFont(font1)
        self.btnCancel.setObjectName("btnCancel")

        self.btnLog = QtWidgets.QPushButton(Dialog)
        self.btnLog.setGeometry(QtCore.QRect(412, 390, 80, 25))
        self.btnLog.setFont(font1)
        self.btnLog.setObjectName("btnLog")

        self.pbar = QtWidgets.QProgressBar(self.frame_2)
        self.pbar.setGeometry(175, 75, 250, 25)
        self.pbar.setHidden(True)

        self.retranslateUi(Dialog)
        self.update_buttons_status()
        self.connections()
        QtCore.QMetaObject.connectSlotsByName(Dialog)

    def retranslateUi(self, Dialog):
        _translate = QtCore.QCoreApplication.translate
        Dialog.setWindowTitle(_translate("Dialog", "BreakDown [C&M]"))
        breakdown_logo = os.path.join(self.configFolder, 'SupportingFiles', 'BD.ico')
        if os.path.exists(breakdown_logo):
            Dialog.setWindowIcon(QtGui.QIcon(breakdown_logo))
            if sys.platform == "win32":
                import ctypes
                myappid = 'com.companyname.BreakDown'  # Unique ID for your application
                ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
        # if os.path.exists(breakdown_logo):
        #     Dialog.setWindowIcon(QtGui.QIcon(breakdown_logo))
        self.processLabel.setText(_translate("Dialog",
                                             "<html><head/><body><p><span style=\"  font-weight:600; font-size:10pt;\">Process:</span></p></body></html>"))
        self.customerLabel.setText(_translate("Dialog",
                                              "<html><head/><body><p><span style=\" font-weight:600; font-size:10pt;\">Customer:</span></p></body></html>"))
        self.jidAidLabel.setText(_translate("Dialog",
                                            "<html><head/><body><p><span style=\"  font-weight:600; font-size:10pt;\">JID_AID:</span></p></body></html>"))
        self.btnIns.setText(_translate("Dialog", "Instruction(s)"))
        self.btnOpen.setText(_translate("Dialog", "Open"))
        self.btnMoveUp.setText(_translate("Dialog", "Move UP"))
        self.btnMoveDown.setText(_translate("Dialog", "Move Down"))
        self.btnRemove.setText(_translate("Dialog", "Remove"))
        self.btnAdd.setText(_translate("Dialog", "Add"))
        self.selFilesLabel.setText(_translate("Dialog",
                                              "<html><head/><body><p><span style=\" font-weight:600; color:#FFFFFF;\">Selected Files:</span></p></body></html>"))
        self.rmvFilesLabel.setText(_translate("Dialog",
                                              "<html><head/><body><p><span style=\" font-weight:600; color:#FFFFFF;\">Removed Files:</span></p></body></html>"))
        self.btnOK.setText(_translate("Dialog", "OK"))
        self.btnCancel.setText(_translate("Dialog", "Cancel"))
        self.btnLog.setText(_translate("Dialog", "Log"))

    def process_combobox_changed(self):
        self.customerCombo.clear()
        self.selectedList.clear()
        self.removeList.clear()
        self.customerCombo.addItems(self.dialog_config['CUSTOMERS'])

    def customer_combobox_changed(self):
        self.jidAidCombo.clear()
        self.selectedList.clear()
        self.removeList.clear()
        self.jid_aid_list = []
        self.jid_aid_list, self.folder_details = self.aid_jid_list(self.processCombo.currentText(),
                                                                   self.customerCombo.currentText(), self.dialog_config)
        self.jidAidCombo.addItems(self.jid_aid_list)
        return self.jid_aid_list

    def aid_jid_list(self, process, customer, config_details):
        input_path = config_details['FOLDERS'][process]['INPUT_PATH']
        input_path = re.sub(r"\[CUSTOMER\]", customer, input_path, re.IGNORECASE)
        self.input_list = glob.glob(input_path + r"\*")
        self.return_list = ['Select']
        self.selectedList.clear()
        self.removeList.clear()
        if len(self.input_list) == 0:
            self.return_list = ['None']
            self.folder_details = {}
        for path_name in self.input_list:
            if os.path.isdir(path_name):
                folder_name = os.path.split(path_name)[1]
                self.folder_details[folder_name] = path_name
                self.return_list.append(folder_name)
        return self.return_list, self.folder_details

    def get_aid_jid_list(self, process, customer, config_details):
        input_path = config_details['FOLDERS'][process]['INPUT_PATH']
        input_path = re.sub(r"\[CUSTOMER\]", customer, input_path, re.IGNORECASE)
        self.input_list = glob.glob(input_path + r"\*")
        self.return_list = ['Select']
        if len(self.input_list) == 0:
            self.return_list = ['None']
            self.folder_details = {}
        for path_name in self.input_list:
            if os.path.isdir(path_name):
                folder_name = os.path.split(path_name)[1]
                self.folder_details[folder_name] = path_name
                self.return_list.append(folder_name)
        return self.return_list, self.folder_details

    def get_file_list(self, value):
        self.file_details = {}
        self.file_names = []
        self.selectedList.clear()
        self.removeList.clear()
        if value == "Select" or value == "" or value == "None":
            self.selectedList.clear()
            self.btnOK.setDisabled(True)
        else:
            self.file_path = self.folder_details[value]
            file_fullpath = glob.glob(self.file_path + u"/*.doc")
            file_fullpath.extend(glob.glob(self.file_path + u"/*.docx"))
            file_fullpath.extend(glob.glob(self.file_path + u"/*.xls"))
            file_fullpath.extend(glob.glob(self.file_path + u"/*.xlsx"))
            for file in file_fullpath:
                file_name = os.path.split(file)[1]
                self.file_details[file_name] = file
                self.file_names.append(file_name)
            if len(self.file_names) > 0:
                self.btnOK.setEnabled(True)
            self.selectedList.addItems(self.file_names)
        return self.file_details, self.file_names

    @QtCore.pyqtSlot()
    def update_buttons_status(self):
        self.btnMoveUp.setDisabled(not bool(self.selectedList.selectedItems()) or self.selectedList.currentRow() == 0)
        self.btnRemove.setDisabled(not bool(self.selectedList.selectedItems()))
        # self.btnView.setDisabled(not bool(self.selectedList.selectedItems()))
        # self.btnIns.setDisabled(not bool(self.selectedList.selectedItems()))
        self.btnIns.setDisabled(not bool(self.selectedList.currentRow() == 0))
        self.btnOpen.setDisabled(not bool(self.selectedList.selectedItems()))
        self.btnOpen.setDisabled(not bool(self.selectedList.selectedItems()))
        self.btnAdd.setDisabled(not bool(self.selectedList.selectedItems()))
        self.btnOK.setDisabled(self.selectedList.count() == 0)
        self.btnMoveDown.setDisabled(
            not bool(self.selectedList.selectedItems()) or self.selectedList.currentRow() == (
                        self.selectedList.count() - 1))

    def connections(self):
        self.selectedList.itemSelectionChanged.connect(self.update_buttons_status)
        self.removeList.itemSelectionChanged.connect(self.update_buttons_status)
        self.btnOpen.clicked.connect(self.on_btnOpen_clicked)
        # self.btnView.clicked.connect(self.on_btnView_clicked)
        self.btnIns.clicked.connect(self.on_btnInstruction_clicked)
        self.btnMoveUp.clicked.connect(self.on_btnMoveUP_clicked)
        self.btnMoveDown.clicked.connect(self.on_btnMoveDown_clicked)
        self.btnRemove.clicked.connect(self.on_btnRemove_clicked)
        self.btnAdd.clicked.connect(self.on_btnAdd_clicked)
        self.btnCancel.clicked.connect(self.on_btnCancel_clicked)
        self.btnOK.clicked.connect(self.on_btnOK_clicked)

    @QtCore.pyqtSlot()
    def on_btnOpen_clicked(self):
        row = self.selectedList.currentRow()
        currentItem = self.selectedList.item(row).text()
        selectedFile = self.file_details[currentItem]
        # self.open_doc_file(selectedFile)
        try:
            os.startfile(selectedFile, 'open')
        except Exception as e:
            QMessageBox.information(self, "mSelect [C&M]",
                                    f"File: {selectedFile}\nUnable to Open file.\n"
                                    f"Please open, Check, and Save manually..")
            os.startfile(selectedFile)

    @QtCore.pyqtSlot()
    def on_btnView_clicked(self):
        row = self.selectedList.currentRow()
        currentItem = self.selectedList.item(row).text()
        selectedFile = self.file_details[currentItem]
        self.open_pdf_file(selectedFile)

    @QtCore.pyqtSlot()
    def on_btnInstruction_clicked(self):
        configPath, breakDownConfig = getconfig()
        jrnJson = os.path.join(configPath, "SupportingFiles/BreakDown.json")
        with open(jrnJson, 'r') as file:
            jsonDetails = yaml.safe_load(file)
        jid_aid = self.jidAidCombo.currentText()
        jidAid = jid_aid.split("_")
        jid = jidAid[0]
        jrnInstructions = jsonDetails['journal_details'][jid]['Instruction']
        jrnIns = jid + ": Instructions\n--------------------\n\n" + jrnInstructions
        self.open_instructions(jrnIns)
        # self.open_pdf_file(selectedFile)

    @QtCore.pyqtSlot()
    def on_btnMoveUP_clicked(self):
        row = self.selectedList.currentRow()
        currentItem = self.selectedList.takeItem(row)
        self.selectedList.insertItem(row - 1, currentItem)
        self.selectedList.setCurrentRow(row - 1)

    @QtCore.pyqtSlot()
    def on_btnMoveDown_clicked(self):
        row = self.selectedList.currentRow()
        currentItem = self.selectedList.takeItem(row)
        self.selectedList.insertItem(row + 1, currentItem)
        self.selectedList.setCurrentRow(row + 1)

    @QtCore.pyqtSlot()
    def on_btnRemove_clicked(self):
        row = self.selectedList.currentRow()
        currentItem = self.selectedList.takeItem(row)
        self.removeList.insertItem(0, currentItem)

    @QtCore.pyqtSlot()
    def on_btnAdd_clicked(self):
        row = self.removeList.currentRow()
        currentItem = self.removeList.takeItem(row)
        self.selectedList.insertItem(0, currentItem)

    @QtCore.pyqtSlot()
    def on_btnCancel_clicked(self):
        sys.exit(0)

    def open_doc_file(self, file):
        file_name = os.path.split(file)[1]
        QMessageBox.information(self, "mSelect [C&M]",
                                f"File: {file_name}\nIf any updates, Please save and close...")
        if file.lower().endswith(".doc") or file.lower().endswith(".docx"):
            word = win32.gencache.EnsureDispatch('Word.Application')
            word.Visible = True
            try:
                doc = word.Documents.Open(file)
            except Exception as e:
                QMessageBox.information(self, "mSelect [C&M]",
                                        f"File: {file}\nUnable to Open file.\n"
                                        f"Please open, Check, and Save manually..")
                os.startfile(file)

    def open_pdf_file(self, file):
        global pdf_file
        pdf_file = re.sub(".docx", ".pdf", file)
        convert(file, pdf_file)
        self.childWin = ChildWindow()
        if self.childWin.isVisible():
            self.childWin.close()
        else:
            self.childWin.show()
            url = QUrl.fromLocalFile(pdf_file)
            self.childWin.webView.load(url)

    def open_instructions(self, instructions):
        temp_file = "temp.txt"
        if os.path.exists("temp.txt"):
            subprocess.run(["del", temp_file], shell=True)
        with open(temp_file, "w") as file:
            file.write(instructions)
        subprocess.run(["notepad.exe", temp_file])

    @QtCore.pyqtSlot()
    def on_btnOK_clicked(self):
        items = []
        ritems = []
        self.selected_list = {}
        self.removed_list = {}
        self.normalizer_input = {}
        close_dialog = False
        for x in range(self.selectedList.count()):
            items.append(self.selectedList.item(x).text().encode("utf8").decode("utf8"))
        for x in range(self.removeList.count()):
            ritems.append(self.removeList.item(x).text().encode("utf8").decode("utf8"))
        for ritem in ritems:
            self.removed_list[ritem] = self.file_details[ritem]
        for item in items:
            self.selected_list[item] = self.file_details[item]
        if len(self.selected_list) > 0:
            self.normalizer_input['selected'] = self.selected_list
            self.normalizer_input['customer'] = self.customerCombo.currentText()
            self.normalizer_input['folder'] = self.jidAidCombo.currentText()
            self.normalizer_input['process'] = self.processCombo.currentText()
        if len(self.removed_list) > 0:
            self.normalizer_input['removed'] = self.removed_list
        if len(self.selected_list) == 0:
            QMessageBox.information(self, "mSelect [C&M]",
                                    "Please Select atleast one document to proceed...")
        else:
            process_doc = OpenDocFile()
            selected_process = self.processCombo.currentText()
            self.btnOK.setDisabled(True)
            self.btnCancel.setDisabled(True)
            self.jidAidCombo.setDisabled(True)
            jid_aid = self.jidAidCombo.currentText()
            jidAid = jid_aid.split("_")
            jrn_id = jidAid[0]
            art_id = jidAid[1]
            mydb = DataBase()
            uniq_id = mydb.get_uniqueid(jrn_id, art_id)
            if selected_process == "mMerger":
                self.pbar.setHidden(False)
                self.pbar.setValue(50)
                merger = DocxMerger()
                # merger_result = merger.merge_docx(self.normalizer_input, uniq_id)
                # merger_result = False
                # if merger_result is False:
                #     merger_result = merger.merge_in_doc(self.normalizer_input, uniq_id)
                # if merger_result is True:
                #     merger.move_files_to_docs(self.normalizer_input, uniq_id)
                merger_result = merger.merge_docx_robust(self.normalizer_input, uniq_id)
                if merger_result:
                    merger.move_files_to_docs(self.normalizer_input, uniq_id)
                    mydb.update_db(uniq_id, "mMerger", "COMPLETED", "", "")
                else:
                    QMessageBox.warning(self, "mSelect [C&M]",
                                        "Merging failed. Please check logs and try manual merge.")
                    mydb.update_db(uniq_id, "mMerger", "ERROR", "",
                                   "All merge strategies failed")
                self.pbar.setValue(100)
                if merger_result is True:
                    if merger_result is True:
                        merger_folder = self.normalizer_input['folder']
                        customer = self.normalizer_input['customer']
                        source_folder = re.sub(r"\[CUSTOMER\]", customer, self.mergerInput)
                        dest_folder = re.sub(r"\[CUSTOMER\]", customer, self.ParaStylerInput)
                        source_folder = os.path.join(source_folder, merger_folder)
                        dest_folder = os.path.join(dest_folder, merger_folder)
                        try:
                            shutil.move(source_folder, dest_folder)
                        except Exception as e:
                            print(e)
                    QMessageBox.information(self, "mSelect [C&M]",
                                            "Merging Process Completed...")
                    #                jidAidCombo_index = self.jidAidCombo.currentIndex()
                    self.jidAidCombo.removeItem(self.jidAidCombo.currentIndex())
                    mydb.update_db(uniq_id, "mMerger", "COMPLETED", "", "")
                elif merger_result is False:
                    QMessageBox.information(self, "mSelect [C&M]",
                                            "Some of the documents may contain shapes please merger manually")
                    mydb.update_db(uniq_id, "mMerger", "ERROR", "", f"Some of the documents may contain shapes please merger manually")
                self.pbar.setHidden(True)
            elif selected_process == "mParaStyler":
                if len(self.selected_list) > 1:
                    QMessageBox.information(self, "mSelect [C&M]", "For Para Styler input must be single file\nPlease remove unwanted files and run again.")
                elif len(self.selected_list) == 1:
                    if self.selectedList.currentItem() is None:
                        first_item = self.selectedList.item(0)
                        self.selectedList.setCurrentItem(first_item)
                    currentItem = self.selectedList.currentItem().text()
                    QMessageBox.information(self, "mSelect [C&M]", f"{currentItem} para styler")
                    article_folder = self.normalizer_input['folder']
                    customer = self.normalizer_input['customer']
                    parastyler_folder = re.sub(r"\[CUSTOMER\]", customer, self.ParaStylerInput)
                    breakdown_folder = re.sub(r"\[CUSTOMER\]", customer, self.BreakDownInput)
                    breakdown_done = re.sub(r"\[CUSTOMER\]", customer, self.BreakDownDone)
                    file_full_path = os.path.join(parastyler_folder, article_folder, currentItem)
                    file_folder = os.path.split(file_full_path)[0]
                    file_name = os.path.split(file_full_path)[1]
                    backup_folder = os.path.join(file_folder, "docs")
                    if not os.path.exists(backup_folder):
                        os.makedirs(backup_folder)
                    backup_file = os.path.join(file_folder, "docs", file_name)
                    shutil.copy(file_full_path, backup_file)
                    my_path = os.path.dirname(file_full_path)
                    info_path = os.path.join(my_path, "para_info.xml")
                    if os.path.exists(info_path):
                        os.remove(info_path)
                    quot_xml = None
                    qoteInfo = CreateParaInfo()
                    ItalicBookmarkProcessor(file_full_path, mode="apply_bookmark").process()
                    # try:
                    #     quot_xml = qoteInfo.create_info_xml(file_full_path)
                    # except Exception as e:
                    #     pass
                    # if quot_xml is not None and os.path.exists(quot_xml):
                    #         qoteInfo.add_quot_prefix_to_paragraphs(file_full_path, quot_xml)
                    #         qoteInfo.add_list_prefix_to_paragraphs(file_full_path, quot_xml)
                    docxprocess = DocxManipulator()
                    processresult, as_file = docxprocess.docx_processor(file_full_path)
                    if processresult is False:
                        as_file = self.run_para_styler(file_full_path)
                    if os.path.exists(as_file):
                        # as_file_copy = re.sub("\.docx", "_copy.docx", as_file)
                        # if os.path.exists(as_file_copy):
                        #     os.remove(as_file_copy)
                        # shutil.copy(as_file, as_file_copy)
                        ItalicBookmarkProcessor(as_file, mode="apply_italic").process()
                        applySty = ApplyStyles()
                        applySty.as_post_clean(as_file)
                        # if quot_xml is not None and os.path.exists(quot_xml):
                        #     qoteInfo.apply_styles_to_paragraphs(as_file, quot_xml)
                        #     os.remove(quot_xml)
                        # self.apply_quot_to_paragraphs(as_file_copy, quot_xml)
                        document = Document(as_file)
                        document = self.add_math_style(document)
                        document = self.remove_box(document)
                        style_count = 0
                        for paragraph in document.paragraphs:
                            if paragraph.style.name == "articletitle":
                                style_count += 1
                                if style_count > 1:  # If more than one "AT" style found
                                    paragraph.style = "sectiona"
                        if style_count == 0:
                            if document.paragraphs[0].style.name == "sectiona":
                                document.paragraphs[0].style = "articletitle"
                        # document = self.update_table_styles(document)
                        # document = self.clean_document_xml(document)
                        for style in document.styles:
                            if style.name == "dummy":
                                # Apply highlighting to the style
                                style.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                break
                        document.save(as_file)
                        list_macros = ["AutoFitTable", "LoadSageStyles", "RemoveLineNumbers"]
                        # list_macros = ["copyStylesFromTemplate", "AutoFitTable", "TableCellStyle"]
                        process_error, error_log, word = process_doc.processDocFile(as_file, True, True, True, list_macros)
                        if os.path.exists(as_file):
                            applyStyles = ApplyStyles()
                            applyStyles.apply_styles(as_file)
                            # file_folder = os.path.split(file_full_path)[0]
                            # file_name = os.path.split(file_full_path)[1]
                            # backup_folder = os.path.join(file_folder, "docs")
                            # if not os.path.exists(backup_folder):
                            #     os.makedirs(backup_folder)
                            # backup_file = os.path.join(file_folder, "docs", file_name)
                            # shutil.move(file_full_path, backup_file)
                            if os.path.exists(as_file) and os.path.exists(file_full_path):
                                os.remove(file_full_path)
                            os.rename(as_file, file_full_path)
                            print(f"Running Check Normal styles: {file_full_path}")
                            applyStyles.check_normal_styles(file_full_path)
                            print(f"Running Cross Check styles: {file_full_path}")
                            applyStyles.cross_check_styles(file_full_path)
                            table_updates = XmlTransform()
                            table_updates.udpate_table_cells(file_full_path)
                            # if quot_xml is not None and os.path.exists(quot_xml):
                            #     self.apply_quot_to_paragraphs(file_full_path, quot_xml)
                            breakdown_input = os.path.join(breakdown_folder, article_folder)
                            breakdown_file = os.path.join(breakdown_input, currentItem)
                            process_completed = True
                            if os.path.exists(breakdown_input):
                                try:
                                    os.remove(breakdown_input)
                                except Exception as e:
                                    remove_error = 1
                                    process_completed = False
                                    QMessageBox.information(self, "mSelect [C&M]",
                                                            f"Unable to remove old folder {breakdown_input}\nFolder May be in use\nPlease move it manually")
                            if process_completed is True:
                                try:
                                    shutil.move(file_folder, breakdown_input)
                                except Exception as e:
                                    process_completed = False
                                    QMessageBox.information(self, "mSelect [C&M]",
                                                            f"Unable to move folder {breakdown_input}\nFolder Maybe in use\nMove it manually")
                        if process_completed is True:
                            QMessageBox.information(self, "mSelect [C&M]", "Para Styler Completed")
                            # process_error, error_log, word = process_doc.processDocFile(breakdown_file, True, False, False, "")
                    else:
                        QMessageBox.information(self, "mSelect [C&M]", "Problem In running Styler,\nCould you please proceed Mannuallay")
            elif selected_process == "mBreakDown":
                applyStyles = ApplyStyles()
                if len(self.selected_list) > 1:
                    QMessageBox.information(self, "mSelect [C&M]", "For Breakdown Process input must be Single file\nPlease remove unwanted files and run again.")
                elif len(self.selected_list) == 1:
                    if self.selectedList.currentItem() is None:
                        first_item = self.selectedList.item(0)
                        self.selectedList.setCurrentItem(first_item)
                    currentItem = self.selectedList.currentItem().text()
                    QMessageBox.information(self, "mBreakDown [C&M]", f"{currentItem} BreakDown Process")
                    article_folder = self.normalizer_input['folder']
                    customer = self.normalizer_input['customer']
                    breakdown_folder = re.sub(r"\[CUSTOMER\]", customer, self.BreakDownInput)
                    breakdown_done = re.sub(r"\[CUSTOMER\]", customer, self.BreakDownDone)
                    article_input = os.path.join(breakdown_folder, article_folder)
                    article_done = os.path.join(breakdown_done, article_folder)
                    backup_file = currentItem
                    backup_file = re.sub("_CLN", "_CLN_AS", backup_file)
                    file_full_path = os.path.join(breakdown_folder, article_folder, currentItem)
                    doc_folder = os.path.join(breakdown_folder, article_folder, "docs")
                    if not os.path.exists(doc_folder):
                        os.mkdir(doc_folder)
                    back_up_full_path = os.path.join(breakdown_folder, article_folder, "docs", backup_file)
                    shutil.copy(file_full_path, back_up_full_path)
                    done_file_path = os.path.join(article_done, currentItem)
                    breakdown_process = BreakDownProcess()
                    art_folder = article_folder.split("_")
                    jid = art_folder[0]
                    aid = art_folder[1]
                    json_name = jid + "_" + aid + ".json"
                    json_path = os.path.join(article_input, json_name)
                    if os.path.exists(json_path):
                        breakdown_process.create_breakdown_docx(jid, aid, file_full_path)
                    else:
                        creat_info = GetArticleId()
                        info_generated, article_id, jrn_tla = creat_info.smart_login(aid, None, jid, article_input)
                        breakdown_process.create_breakdown_docx(jid, aid, file_full_path)
                    process_completed = True
                    if os.path.exists(article_done):
                        try:
                            shutil.rmtree(article_done)
                        except Exception as e:
                            process_completed = False
                    if process_completed is True:
                        try:
                            shutil.move(article_input, article_done)
                        except Exception as e:
                            process_completed = False
                    if process_completed is True:
                        QMessageBox.information(self, "mBreakDown [C&M]", f"{currentItem} BreakDown Process Completed")
                        mydb.update_db(uniq_id, "mBreakDown", "COMPLETED", "", "")
                        # process_error, error_log, word = process_doc.processDocFile(done_file_path, True, False, False, "")
                        # word_open = process_doc.openWordDocumentVisible(done_file_path)
                        word_instance, doc_instance = process_doc.openWordDocumentVisible(done_file_path)
                        if word_instance is not None and doc_instance is not None:
                            print("Document opened successfully and should be visible.")
                            print(f"Opened: {doc_instance.Name}")
                            close_dialog = True
                        else:
                            print("Failed to open document.")
                            close_dialog = False
            self.selectedList.clear()
            self.removeList.clear()
            self.btnOK.setEnabled(True)
            self.btnCancel.setEnabled(True)
            self.jidAidCombo.setEnabled(True)
            if close_dialog is True:
                sys.exit(0)

    def run_para_styler(self, doc_file):
        styler_path = os.path.join(self.configFolder, "ParaStyler\\run.bat")
        #QApplication.setOverrideCursor(Qt.WaitCursor)
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        p = Popen([styler_path, doc_file], stdout=PIPE, stderr=PIPE)
        as_file = re.sub(".docx", "_AS.docx", doc_file)
        output, errors = p.communicate()
        QApplication.restoreOverrideCursor()
        # print(output)
        # print(errors)
        p.wait()
        return as_file

    def remove_box(self, document):
        style = document.styles['abstract']
        prValue = style._element.pPr
        style._element.remove(prValue)
        return document

    def header_count(self, table):
        span_count = 1
        for row in table.rows:
            row_xml = row._tr.xml
            if re.search(r"(w\:gridSpan)", row_xml, re.IGNORECASE):
                span_count = span_count + 1
            else:
                break
        return span_count

    def update_table_styles(self, document):
        styles = document.styles
        if "tablehead" not in styles:
            style = document.styles.add_style("tablehead", WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = document.styles["Normal"]
        if "tablebody" not in styles:
            style = document.styles.add_style("tablebody", WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = document.styles["Normal"]
        styles = document.styles
        for table in document.tables:
            thead_count = self.header_count(table)
            row_count = 1
            for row in table.rows:
                if row_count <= thead_count:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.style = styles['tablehead']
                else:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.style = styles['tablebody']
                    row_count = row_count + 1
                row_count = row_count + 1
        return document

    def add_math_style(self, doc):
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                 'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
        for prefix, uri in nsmap.items():
            etree.register_namespace(prefix, uri)
        for paragraph in doc.paragraphs:
            math_elements = paragraph._element.xpath('.//m:oMathPara')
            if math_elements:
                p = paragraph._element
                ppr = p.find('.//w:pPr', namespaces=nsmap)
                if ppr is None:
                    ppr = parse_xml('<w:pPr xmlns:w="{}"/>'.format(nsmap['w']))
                    p.insert(0, ppr)
                pstyle_elem = ppr.find('.//w:pStyle', namespaces=nsmap)
                if pstyle_elem is not None and pstyle_elem.get(
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') != "displaymath":
                    ppr.remove(pstyle_elem)
                xml_str = '<w:pStyle w:val="displaymath" xmlns:w="{}"/>'.format(nsmap['w'])
                pstyle = parse_xml(xml_str)
                ppr.append(pstyle)
        return doc

    def clean_document_xml(self, document):
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        p_style = paragraph.style
                        para_style = paragraph.style.name
                        # if para_style != "Normal":
                        stId = document.styles[para_style]._element.styleId
                        pprElement = paragraph._p.pPr
                        new_ppr = OxmlElement("w:pPr")
                        new_ppr.style = stId
                        paragraph._p.remove(pprElement)
                        paragraph._p.insert(0, new_ppr)
                        for run in paragraph.runs:
                            rprValue = run._r.rPr
                            if rprValue is not None:
                                szVal = run._r.rPr.sz
                                colVal = run._r.rPr.color
                                # szcValue = run._r.rPr.szCs
                                rFontsValue = run._element.rPr.rFonts
                                asciiValue = run._element.rPr.rFonts_ascii
                                hAnsiValue = run._element.rPr.rFonts_hAnsi
                                if asciiValue is not None and asciiValue == "Arial":
                                    run._r.rPr.remove(rFontsValue)
                                # if szcValue is not None:
                                #     run._r.rPr.remove(szcValue)
                                if szVal is not None:
                                    run._r.rPr.remove(szVal)
                                if colVal is not None:
                                    colXml = run._r.rPr.color.xml
                                    if bool(re.search('w:val="000000"', colXml)) is True:
                                        run._r.rPr.remove(colVal)
        for paragraph in document.paragraphs:
            p_style = paragraph.style
            para_style = paragraph.style.name
            para_text = paragraph.text
            if para_style != "Normal":
                stId = document.styles[para_style]._element.styleId
                pprElement = paragraph._p.pPr
                new_ppr = OxmlElement("w:pPr")
                new_ppr.style = stId
                paragraph._p.remove(pprElement)
                paragraph._p.insert(0, new_ppr)
                for run in paragraph.runs:
                    rprValue = run._r.rPr
                    if rprValue is not None:
                        szVal = run._r.rPr.sz
                        colVal = run._r.rPr.color
                        boldVal = run._r.rPr.b
                        # szcValue = run._r.rPr.szCs
                        rFontsValue = run._element.rPr.rFonts
                        asciiValue = run._element.rPr.rFonts_ascii
                        hAnsiValue = run._element.rPr.rFonts_hAnsi
                        if asciiValue is not None and asciiValue == "Arial":
                            run._r.rPr.remove(rFontsValue)
                        if asciiValue is not None and asciiValue == "Calibri":
                            run._r.rPr.remove(rFontsValue)
                        if asciiValue is not None and asciiValue == "majorHAnsi":
                            run._r.rPr.remove(rFontsValue)
                        if boldVal is not None:
                            if re.search("articletitle|abstracttitle|sectiona|sectionb|sectionc|sectiond|conflictofinterest|acknowledgementstitle|references", para_style, re.IGNORECASE):
                                run._r.rPr.remove(boldVal)
                        # if szcValue is not None:
                        #     run._r.rPr.remove(szcValue)
                        if szVal is not None:
                            run._r.rPr.remove(szVal)
                        if colVal is not None:
                            colXml = run._r.rPr.color.xml
                            if bool(re.search('w:val="000000"', colXml)) is True:
                                # if bool(re.search('w:themeColor="text1"', colXml)):
                                run._r.rPr.remove(colVal)
        return document



# app = QtWidgets.QApplication(sys.argv)
# dialog = QtWidgets.QDialog()
# ui = Ui_Dialog()
# ui.setupUi(dialog)
# dialog.show()
# sys.exit(app.exec())
