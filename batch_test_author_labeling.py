"""
batch_test_author_labeling.py — Batch test runner for labelAuthorsNER.

Scans a folder for <jid>_<aid>_CLN_AS.docx + <jid>_<aid>.json pairs,
runs the full labeling pipeline, produces labeled docx files, and
generates an HTML report.

Usage:
    python batch_test_author_labeling.py "V:\\AU_TEST"

    Optional: specify output folder for labeled docx files and report
    python batch_test_author_labeling.py "V:\\AU_TEST" --output "V:\\AU_TEST\\results"

Output:
    - <jid>_<aid>_CLN_AS_LABELED.docx  (in same folder or --output folder)
    - test_report.html                  (in same folder or --output folder)

Report columns:
    No. | Input File | Validation Result | Input Author Line |
    Output Author Line | Json Count | Output Count
"""

import sys
import os
import re
import json
import glob
import traceback
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================================
# Import from labelAuthorsNER
# ============================================================================
from labelAuthorsNER import (
    AuthorLabeler,
    AuthorEntity,
    LabelResult,
    get_docx_authors_ner,
    extract_au_runs,
    load_json_authors,
    is_degree_token,
)


# ============================================================================
# Docx writing helpers (same as test_author_labeling.py)
# ============================================================================

def ensure_aulabel_style(doc: Document):
    if "aulabel" not in {s.name for s in doc.styles}:
        style = doc.styles.add_style("aulabel", WD_STYLE_TYPE.CHARACTER)
        style.font.color.rgb = RGBColor(0, 0, 255)
        style.font.size = Pt(10)
    return doc


def rebuild_au_paragraph(para, authors):
    for child in list(para._p):
        if child.tag != qn('w:pPr'):
            para._p.remove(child)

    def add_run(text, superscript=False, style_name=None, color_rgb=None):
        r_elem = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        if style_name:
            rs = OxmlElement('w:rStyle')
            rs.set(qn('w:val'), style_name)
            rpr.append(rs)
        if superscript:
            v = OxmlElement('w:vertAlign')
            v.set(qn('w:val'), 'superscript')
            rpr.append(v)
        if color_rgb:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color_rgb)
            rpr.append(c)
        r_elem.append(rpr)
        t = OxmlElement('w:t')
        t.text = text
        if text and (text[0] == ' ' or text[-1] == ' '):
            t.set(qn('xml:space'), 'preserve')
        r_elem.append(t)
        para._p.append(r_elem)

    total = len(authors)
    for i, a in enumerate(authors):
        add_run(a.full_name)
        if a.superscripts:
            sup = ",".join(a.superscripts)
            if a.is_corresponding:
                sup += ",*"
            add_run(sup, superscript=True)
        elif a.is_corresponding:
            add_run("*", superscript=True)
        if a.degrees:
            add_run(" " + ", ".join(a.degrees))
        add_run(a.label, style_name="aulabel", color_rgb="0000FF")
        if i < total - 2:
            add_run(", ")
        elif i == total - 2:
            add_run(", and ")
    add_run(".")


def get_au_text_from_docx(docx_path):
    """Read the AU paragraph text from a docx file."""
    try:
        doc = Document(docx_path)
        for p in doc.paragraphs:
            if p.style.name in ("AU", "AU0"):
                return p.text
    except Exception:
        pass
    return ""


# ============================================================================
# Find docx+json pairs in folder
# ============================================================================

def find_test_pairs(folder):
    """
    Scan folder for <jid>_<aid>_CLN_AS.docx + <jid>_<aid>.json pairs.
    Returns list of (docx_path, json_path, jid_aid) tuples.
    """
    pairs = []
    # Find all _CLN_AS.docx files
    # Find all _AS.docx or _CLN_AS.docx files
    for pattern_suffix in ["*_CLN_AS.docx", "*_AS.docx"]:
        pat = os.path.join(folder, pattern_suffix)
        for docx_path in sorted(glob.glob(pat)):
            basename = os.path.basename(docx_path)
            # Extract <jid>_<aid> from <jid>_<aid>_CLN_AS.docx or <jid>_<aid>_AS.docx
            match = re.match(r'^(.+?)(?:_CLN)?_AS\.docx$', basename, re.IGNORECASE)
            if not match:
                continue
            jid_aid = match.group(1)

            # Skip if we already have this jid_aid (prefer _CLN_AS over _AS)
            if any(j == jid_aid for _, _, j in pairs):
                continue

            # Look for matching JSON
            json_name = f"{jid_aid}.json"
            json_path = os.path.join(folder, json_name)

            if os.path.exists(json_path):
                pairs.append((docx_path, json_path, jid_aid))
            else:
                pairs.append((docx_path, None, jid_aid))

    return pairs


# ============================================================================
# Process a single file pair
# ============================================================================

def process_one(docx_path, json_path, output_folder):
    """
    Process one docx+json pair:
      1. Read original AU line
      2. Run labelAuthorsNER pipeline
      3. Write labeled docx
      4. Read back the labeled AU line
      5. Return result dict for the report

    Returns dict with keys:
        input_file, validation, input_au_line, output_au_line,
        json_count, output_count, details, error
    """
    result = {
        "input_file": os.path.basename(docx_path),
        "validation": "FAIL",
        "input_au_line": "",
        "output_au_line": "",
        "json_count": 0,
        "output_count": 0,
        "details": [],
        "error": "",
    }

    try:
        # 1. Read original AU line
        input_au = get_au_text_from_docx(docx_path)
        result["input_au_line"] = input_au

        # 2. Get JSON author count
        if json_path and os.path.exists(json_path):
            with open(json_path) as f:
                jdata = json.load(f)
            json_count = len(jdata.get("authors_info", {}))
            result["json_count"] = json_count
        else:
            result["validation"] = "NO JSON"
            result["error"] = "JSON file not found"
            return result

        # 3. Run labelAuthorsNER
        labeler = AuthorLabeler()
        label_result = labeler.label_from_docx_and_json(docx_path, json_path)
        result["output_count"] = len(label_result.authors)
        result["details"] = label_result.match_details

        if not label_result.authors:
            result["validation"] = "FAIL"
            result["error"] = "No authors parsed"
            return result

        # 4. Write labeled docx
        doc = Document(docx_path)
        doc = ensure_aulabel_style(doc)

        au_para = None
        for p in doc.paragraphs:
            if p.style.name in ("AU", "AU0"):
                au_para = p
                break

        if au_para is None:
            result["validation"] = "FAIL"
            result["error"] = "No AU paragraph in docx"
            return result

        rebuild_au_paragraph(au_para, label_result.authors)

        # Save labeled docx
        base = os.path.basename(docx_path)
        name, ext = os.path.splitext(base)
        labeled_name = f"{name}_LABELED{ext}"
        labeled_path = os.path.join(output_folder, labeled_name)
        doc.save(labeled_path)

        # 5. Read back the labeled AU line
        output_au = get_au_text_from_docx(labeled_path)
        result["output_au_line"] = output_au

        # 6. Validate
        output_label_count = len(re.findall(r'\[AU\d+\]', output_au))
        # Check if input already had [AU] labels (pre-labeled file)
        input_pre_labels = len(re.findall(r'\[AU\d+\]', input_au))
        errors = []
        warnings = []

        if len(label_result.authors) != json_count:
            errors.append(f"Count mismatch: parsed {len(label_result.authors)} vs JSON {json_count}")

        if input_pre_labels > 0:
            # Input already had labels — label count check is unreliable
            warnings.append(f"Input had {input_pre_labels} pre-existing [AU] labels")
        elif output_label_count != len(label_result.authors):
            errors.append(f"Label count mismatch: {output_label_count} labels in docx vs {len(label_result.authors)} authors")

        if not label_result.json_matched:
            unmatched = [a for a in label_result.authors if not a.json_key]
            if unmatched:
                names = ", ".join(a.full_name for a in unmatched)
                # If count matches but names don't, it's likely a cross-ref issue
                if len(label_result.authors) == json_count:
                    warnings.append(f"Unmatched authors (count OK): {names}")
                else:
                    errors.append(f"Unmatched authors: {names}")

        # Check no degree leaked as author name
        for a in label_result.authors:
            if is_degree_token(a.full_name):
                errors.append(f"Degree '{a.full_name}' misidentified as author")

        if errors:
            result["validation"] = "FAIL"
            result["error"] = "; ".join(errors)
        elif warnings:
            result["validation"] = "WARN"
            result["error"] = "; ".join(warnings)
        else:
            result["validation"] = "PASS"

    except Exception as e:
        result["validation"] = "ERROR"
        result["error"] = f"{type(e).__name__}: {str(e)}"
        result["details"].append(traceback.format_exc())

    return result


# ============================================================================
# HTML Report Generator
# ============================================================================

def escape_html(text):
    """Escape HTML special characters."""
    if not text:
        return ""
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;"))


def highlight_au_labels(text):
    """Highlight [AU1], [AU2], etc. in blue within the AU line."""
    if not text:
        return ""
    escaped = escape_html(text)
    # Highlight [AUn] labels
    highlighted = re.sub(
        r'\[AU(\d+)\]',
        r'<span style="color:#0000FF;font-weight:bold;">[AU\1]</span>',
        escaped
    )
    return highlighted


def generate_html_report(results, input_folder, output_path):
    """Generate the HTML test report."""

    pass_count = sum(1 for r in results if r["validation"] == "PASS")
    fail_count = sum(1 for r in results if r["validation"] == "FAIL")
    warn_count = sum(1 for r in results if r["validation"] == "WARN")
    error_count = sum(1 for r in results if r["validation"] == "ERROR")
    total = len(results)

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Author Labeling Test Report</title>
<style>
    * {{ margin: 0; padding: 0; box-sizing: border-box; }}
    body {{
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        background: #f5f7fa;
        color: #333;
        padding: 20px;
    }}
    .header {{
        background: linear-gradient(135deg, #1a237e, #283593);
        color: white;
        padding: 24px 32px;
        border-radius: 10px;
        margin-bottom: 20px;
    }}
    .header h1 {{ font-size: 22px; margin-bottom: 8px; }}
    .header .meta {{ font-size: 13px; opacity: 0.85; }}

    .summary {{
        display: flex;
        gap: 16px;
        margin-bottom: 20px;
    }}
    .summary .card {{
        flex: 1;
        padding: 16px 20px;
        border-radius: 8px;
        color: white;
        font-weight: 600;
        text-align: center;
    }}
    .card.total {{ background: #455a64; }}
    .card.pass {{ background: #2e7d32; }}
    .card.fail {{ background: #c62828; }}
    .card.error {{ background: #e65100; }}
    .card .num {{ font-size: 28px; display: block; }}
    .card .label {{ font-size: 12px; text-transform: uppercase; letter-spacing: 1px; }}

    table {{
        width: 100%;
        border-collapse: collapse;
        background: white;
        border-radius: 8px;
        overflow: hidden;
        box-shadow: 0 1px 4px rgba(0,0,0,0.1);
    }}
    thead {{ background: #263238; color: white; }}
    th {{
        padding: 12px 14px;
        text-align: left;
        font-size: 12px;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        white-space: nowrap;
    }}
    td {{
        padding: 10px 14px;
        border-bottom: 1px solid #e0e0e0;
        font-size: 13px;
        vertical-align: top;
    }}
    tr:hover {{ background: #f5f5f5; }}

    .au-line {{
        max-width: 500px;
        max-height: 80px;
        overflow: auto;
        font-size: 12px;
        line-height: 1.5;
        word-break: break-word;
        white-space: pre-wrap;
    }}

    .badge {{
        display: inline-block;
        padding: 3px 10px;
        border-radius: 4px;
        font-size: 11px;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }}
    .badge.pass {{ background: #e8f5e9; color: #2e7d32; }}
    .badge.fail {{ background: #ffebee; color: #c62828; }}
    .badge.error {{ background: #fff3e0; color: #e65100; }}
    .badge.warn {{ background: #fff8e1; color: #f57f17; }}
    .badge.nojson {{ background: #fce4ec; color: #880e4f; }}

    .count-match {{ color: #2e7d32; font-weight: 600; }}
    .count-mismatch {{ color: #c62828; font-weight: 600; }}

    .error-text {{
        color: #c62828;
        font-size: 11px;
        margin-top: 4px;
    }}

    .details-toggle {{
        cursor: pointer;
        color: #1565c0;
        font-size: 11px;
        text-decoration: underline;
    }}
    .details-content {{
        display: none;
        margin-top: 6px;
        padding: 8px;
        background: #f5f5f5;
        border-radius: 4px;
        font-size: 11px;
        font-family: monospace;
        max-height: 200px;
        overflow: auto;
        white-space: pre-wrap;
    }}
</style>
<script>
function toggleDetails(id) {{
    var el = document.getElementById(id);
    el.style.display = el.style.display === 'none' ? 'block' : 'none';
}}
</script>
</head>
<body>

<div class="header">
    <h1>Author Labeling Test Report</h1>
    <div class="meta">
        Module: <strong>labelAuthorsNER</strong> &nbsp;|&nbsp;
        Input Folder: <strong>{escape_html(input_folder)}</strong> &nbsp;|&nbsp;
        Generated: <strong>{timestamp}</strong>
    </div>
</div>

<div class="summary">
    <div class="card total"><span class="num">{total}</span><span class="label">Total Files</span></div>
    <div class="card pass"><span class="num">{pass_count}</span><span class="label">Passed</span></div>
    <div class="card" style="background:#f57f17;"><span class="num">{warn_count}</span><span class="label">Warnings</span></div>
    <div class="card fail"><span class="num">{fail_count}</span><span class="label">Failed</span></div>
    <div class="card error"><span class="num">{error_count}</span><span class="label">Errors</span></div>
</div>

<table>
<thead>
<tr>
    <th>No.</th>
    <th>Input File</th>
    <th>Validation</th>
    <th>Input Author Line</th>
    <th>Output Author Line</th>
    <th>Json Count</th>
    <th>Output Count</th>
</tr>
</thead>
<tbody>
"""

    for idx, r in enumerate(results, 1):
        # Validation badge
        v = r["validation"]
        if v == "PASS":
            badge = '<span class="badge pass">PASS</span>'
        elif v == "WARN":
            badge = '<span class="badge warn">WARN</span>'
        elif v == "FAIL":
            badge = '<span class="badge fail">FAIL</span>'
        elif v == "ERROR":
            badge = '<span class="badge error">ERROR</span>'
        else:
            badge = f'<span class="badge nojson">{escape_html(v)}</span>'

        # Error text
        error_html = ""
        if r["error"]:
            error_html = f'<div class="error-text">{escape_html(r["error"])}</div>'

        # Details toggle
        details_html = ""
        if r["details"]:
            details_id = f"details_{idx}"
            details_text = "\n".join(r["details"])
            details_html = (
                f'<span class="details-toggle" onclick="toggleDetails(\'{details_id}\')">details</span>'
                f'<div class="details-content" id="{details_id}">{escape_html(details_text)}</div>'
            )

        # Count comparison
        jc = r["json_count"]
        oc = r["output_count"]
        if jc == oc and jc > 0:
            count_class = "count-match"
        elif jc > 0:
            count_class = "count-mismatch"
        else:
            count_class = ""

        # Input/output AU lines with highlighted labels
        input_line_html = f'<div class="au-line">{escape_html(r["input_au_line"])}</div>'
        output_line_html = f'<div class="au-line">{highlight_au_labels(r["output_au_line"])}</div>'

        html += f"""<tr>
    <td>{idx}</td>
    <td><strong>{escape_html(r["input_file"])}</strong>{error_html}{details_html}</td>
    <td>{badge}</td>
    <td>{input_line_html}</td>
    <td>{output_line_html}</td>
    <td style="text-align:center;">{jc}</td>
    <td style="text-align:center;" class="{count_class}">{oc}</td>
</tr>
"""

    html += """</tbody>
</table>

</body>
</html>
"""

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)

    return output_path


# ============================================================================
# Main
# ============================================================================

def main():
    if len(sys.argv) < 2:
        print("Usage: python batch_test_author_labeling.py <input_folder> [--output <output_folder>]")
        print()
        print('Example: python batch_test_author_labeling.py "V:\\AU_TEST"')
        print()
        print("Scans for <jid>_<aid>_CLN_AS.docx + <jid>_<aid>.json pairs,")
        print("produces labeled docx files and test_report.html")
        sys.exit(1)

    input_folder = sys.argv[1]

    # Parse --output flag
    output_folder = input_folder
    if "--output" in sys.argv:
        oi = sys.argv.index("--output")
        if oi + 1 < len(sys.argv):
            output_folder = sys.argv[oi + 1]

    if not os.path.isdir(input_folder):
        print(f"ERROR: '{input_folder}' is not a directory")
        sys.exit(1)

    os.makedirs(output_folder, exist_ok=True)

    # Find test pairs
    pairs = find_test_pairs(input_folder)
    if not pairs:
        print(f"No *_CLN_AS.docx + *.json pairs found in '{input_folder}'")
        sys.exit(1)

    print(f"Found {len(pairs)} test file(s) in '{input_folder}'")
    print(f"Output folder: '{output_folder}'")
    print()

    # Process each pair
    results = []
    for i, (docx_path, json_path, jid_aid) in enumerate(pairs, 1):
        tag = jid_aid.replace("_", " ")
        json_status = "+" if json_path else "NO JSON"
        print(f"[{i}/{len(pairs)}] {tag} ... ", end="", flush=True)

        r = process_one(docx_path, json_path, output_folder)
        results.append(r)

        v = r["validation"]
        oc = r["output_count"]
        jc = r["json_count"]
        count_str = f"{oc}/{jc}" if jc else "?"
        print(f"{v} ({count_str} authors)")

    # Generate report
    report_path = os.path.join(output_folder, "test_report.html")
    generate_html_report(results, input_folder, report_path)

    # Summary
    pass_count = sum(1 for r in results if r["validation"] == "PASS")
    warn_count = sum(1 for r in results if r["validation"] == "WARN")
    fail_count = sum(1 for r in results if r["validation"] == "FAIL")
    error_count = sum(1 for r in results if r["validation"] == "ERROR")

    print()
    print("=" * 60)
    print(f"  TOTAL: {len(results)}  |  PASS: {pass_count}  |  WARN: {warn_count}  |  FAIL: {fail_count}  |  ERROR: {error_count}")
    print(f"  Report: {report_path}")
    print("=" * 60)


if __name__ == "__main__":
    main()
