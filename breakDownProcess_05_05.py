import re
import os
import json
import datetime
import calendar
import ctypes
import string
import shutil
import time
import traceback
from docx import Document
from docx.enum.text import WD_BREAK
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_COLOR_INDEX
from applyStyles import ApplyStyles
from sys import exit
from docx.text.run import Run
from titlecase import titlecase
from openDocFile_new import OpenDocFile
from loadconfig import getconfig
from lxml import etree
import zipfile
import subprocess
# import saxonc  # This is commented in original
from TransformXml import XmlTransform
# ── NEW: HTML token-stream pipeline replaces run-level XML manipulation ──────
# labelAuthorsNER provides label_au_paragraph_in_docx (full DOCX pipeline)
# and LabelAuthor (drop-in for the HTML-string path used by get_docx_authors).
from labelAuthorsNER import (
    LabelAuthor,
    label_au_paragraph_in_docx,
    load_json_authors as _ner_load_json,
)
from TransformXmlJar import XmlTransformJar
from fuzzywuzzy import process, fuzz
# from saxonc import PySaxonProcessor  # This is commented in original
from bs4 import BeautifulSoup

try:
    import saxonc
    from saxonc import PySaxonProcessor
    SAXON_AVAILABLE = True
except ImportError:
    SAXON_AVAILABLE = False
    print("Warning: saxonc module not available. Using Java Saxon fallback.")

class BreakDownProcessor:
    def __init__(self):
        self.config_folder, self.breakdown_config = getconfig()
        breakdown_sequence = os.path.join(self.config_folder, 'config\\breakdownSequence.json')
        with open(breakdown_sequence, "r") as breakdown_styles:
            self.breakdown_json = json.loads(breakdown_styles.read())
            self.breakdown_styles = self.breakdown_json['breakdownStyles']
            self.breakdown_mapping_tags = self.breakdown_json['breakdownMappingTags']
        breakdown_json = os.path.join(self.config_folder, 'SupportingFiles\\BreakDown.json')
        with open(breakdown_json, "r") as f:
            self.breakdown_data = json.load(f)
            self.journal_data = self.breakdown_data['journal_details']
            self.default_data = self.breakdown_data['default_values']
            self.abbr_to_num = {name: num for num, name in enumerate(calendar.month_abbr) if num}

    # ====================
    # Paragraph Operations
    # ====================

    def insert_paragraph_after(self, paragraph, text=None, style=None):
        """Insert a new paragraph after the given paragraph."""
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            self._format_and_add_text(new_para, text, style)
        if style is not None:
            new_para.style = style
        return new_para

    def insert_paragraph_before(self, paragraph, text=None, style=None):
        """Insert a new paragraph before the given paragraph."""
        new_p = OxmlElement("w:p")
        paragraph._p.addprevious(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        if style is not None:
            new_para.style = style
        return new_para

    def delete_paragraph(self, paragraph):
        """Remove a paragraph from the document."""
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    # ====================
    # Content Analysis Methods
    # ====================

    def check_declaration_section(self, paragraphs):
        """Check for declaration of conflicting interests section."""
        para_count = 0
        declaration_found = False
        declaration_text = ''
        for paragraph in paragraphs:
            para_count = para_count + 1
            para_text = paragraph.text
            if re.search(
                    "Declaration of Conflicting Interests|Conflict of Interests statement|Conflict of Interests|Conflict of interest|Competing interest|Disclosure|Competing interests|Conflicts of Interest|Conflicting interests",
                    para_text, re.IGNORECASE):
                paragraph.style = "Duplicate"

                # Case 1: heading and content are in the SAME paragraph (e.g. "Conflicting interests: No conflicts...")
                colon_split = re.split(r':\s*', para_text, maxsplit=1)
                if len(colon_split) == 2 and colon_split[1].strip():
                    inline_text = colon_split[1].strip()
                    if not re.search(
                            "no conflicts of interest|no conflict of interest|no potential conflicts of interest|no potential conflict of interest|no competing interests|no competing interest",
                            inline_text, re.IGNORECASE):
                        declaration_text = inline_text
                        declaration_found = True

                # Case 2: content is in the NEXT paragraph (guard against last paragraph)
                elif para_count < len(paragraphs):
                    next_para = paragraphs[para_count]
                    next_para_text = next_para.text
                    paragraphs[para_count].style = "Duplicate"
                    if not re.search(
                            "no conflicts of interest|no conflict of interest|no potential conflicts of interest|no potential conflict of interest|no competing interests|no competing interest",
                            next_para_text, re.IGNORECASE):
                        declaration_text = next_para_text
                        declaration_found = True
                break
        return paragraphs, declaration_found, declaration_text


    def check_supplementary_material(self, paragraphs):
        """Check for supplementary material sections."""
        supp_found = False
        for paragraph in paragraphs:
            para_text = paragraph.text
            if re.match("Supplementary|Supplemental", para_text, re.IGNORECASE):
                supp_found = True
                break
        return supp_found

    def check_corresponding_author(self, authors):
        """Check if any author is marked as corresponding author."""
        corresp_found = False
        for author_key in authors:
            author = authors[author_key]
            if author['corresponding-author'] is True:
                corresp_found = True
                break
        return corresp_found

    def check_funding_section(self, paragraphs):
        """Check for funding information sections."""
        para_count = 0
        funder_found = False
        funder_text = ''
        for paragraph in paragraphs:
            para_count = para_count + 1
            para_text = paragraph.text
            if re.search(r"^((Funding:|Fundings:|Funding|Fundings|Funding\(s\):|Funding\(s\))( |))$", para_text,
                         re.IGNORECASE):
                paragraph.style = "Duplicate"
                next_para = paragraphs[para_count]
                if re.search("study|supported|funding|funded|fund", next_para.text, re.IGNORECASE):
                    funder_text = next_para.text
                    funder_found = True
                    next_para.style = "Duplicate"
                    break
        return paragraphs, funder_found, funder_text

    def check_funder_paragraph(self, paragraphs):
        """Alternative check for funding paragraphs."""
        para_count = 0
        funder_found = False
        funder_text = ''
        for paragraph in paragraphs:
            para_count = para_count + 1
            para_text = paragraph.text
            if re.search("^(This work was supported)", para_text, re.IGNORECASE):
                paragraph.style = "Duplicate"
                funder_found = True
                funder_text = para_text
                break
        return paragraphs, funder_found, funder_text

    def check_keywords_section(self, paragraphs):
        """Check for keywords section."""
        key_para_found = False
        for paragraph in paragraphs:
            para_text = paragraph.text
            if re.search(r"(^(Keywords|Keyword|Keyword\(s\))$|^(Keyword\:|Keywords\:|Keyword\(s\))$)", para_text,
                         re.IGNORECASE):
                key_para_found = True
        return key_para_found

    # ====================
    # Query Processing
    # ====================

    def process_query_count(self, query_content, query_count):
        """Process and update query counts in query content."""
        query_xcount_found = False
        if query_content is not None:
            if re.search(r"(AQ|GQ)(([0-9]+)|XXX|XX|X|)\:", query_content, re.IGNORECASE):
                query_xcount_found = True
                query_prefix = re.search(r"(AQ|GQ)(([0-9]+)|XXX|XX|X|)\:", query_content, re.IGNORECASE).group(1)
                new_qu = query_prefix + str(query_count) + ":"
                query_content = re.sub(r"(AQ|GQ)(([0-9]+)|XXX|XX|X|)\:", new_qu, query_content, re.IGNORECASE)
                query_count = query_count + 1
        return query_xcount_found, query_content, query_count

    # ====================
    # Paragraph Location Methods
    # ====================

    def find_keyword_paragraph(self, paragraphs):
        """Find the keyword paragraph in the document."""
        key_para_found = False
        para_index = 0
        key_para = None
        for paragraph in paragraphs:
            para_style = paragraph.style.name
            if para_style == "Keyword_Para":
                key_para = paragraph
                key_para_found = True
                final_index = para_index
            else:
                if re.match("(H1|H2|H3|H4|H5|H6|TEXT IND|TEXT|nomenclature)", para_style):
                    break
            para_index = para_index + 1
        if key_para_found is False:
            para_index = 0
            for paragraph in paragraphs:
                para_style = paragraph.style.name
                if para_style == "ABKW":
                    key_para = paragraph
                    key_para_found = True
                    final_index = para_index
                else:
                    if re.match("(H1|H2|H3|H4|H5|H6|TEXT IND|TEXT|nomenclature)", para_style):
                        break
                para_index = para_index + 1
        if key_para_found is False:
            para_index = 0
            for paragraph in paragraphs:
                para_style = paragraph.style.name
                if re.match("(H1|H2|H3|H4|H5|H6|TEXT IND|TEXT|nomenclature)", para_style):
                    final_index = para_index - 1
                    key_para = paragraphs[final_index]
                    break
                para_index = para_index + 1
        return key_para_found, key_para, final_index

    def find_reference_paragraph(self, paragraphs):
        """Find the reference section in the document."""
        ref_head_found = False
        ref_para_found = False
        ref_para = None
        para_index = 0
        final_index = -1
        for paragraph in paragraphs:
            para_style = paragraph.style.name
            if para_style == "Reference_Title" or para_style == "Reference_Head":
                ref_para = paragraph
                ref_head_found = True
                final_index = para_index
            para_index = para_index + 1
        if ref_head_found is False:
            para_index = 0
            for paragraph in paragraphs:
                para_style = paragraph.style.name
                if para_style == "REF":
                    ref_para = paragraph
                    ref_para_found = True
                    final_index = para_index
                    break
                para_index = para_index + 1
        return ref_head_found, ref_para_found, ref_para, final_index

    # ====================
    # Author Processing Methods
    # ====================

    def get_running_head_authors(self, authors_data):
        """Generate running head author string."""
        authors_count = len(authors_data.keys())
        if authors_count == 1:
            author_rhead = authors_data['1']['last-name']
        elif authors_count == 2:
            author_rhead = authors_data['1']['last-name'] + " and " + authors_data['2']['last-name']
        elif authors_count > 2:
            author_rhead = authors_data['1']['last-name'] + " et al."
        return author_rhead

    def get_author_rhead(self, authors_data):
        """Generate author running head (same as get_running_head_authors)."""
        authors_count = len(authors_data.keys())
        if authors_count == 1:
            author_rhead = authors_data['1']['last-name']
        elif authors_count == 2:
            author_rhead = authors_data['1']['last-name'] + " and " + authors_data['2']['last-name']
        elif authors_count > 2:
            author_rhead = authors_data['1']['last-name'] + " et al."
        return author_rhead

    # ====================
    # Main Processing Method
    # ====================

    def create_breakdown_document(self, jid, aid, docxfile):
        new_styles = ApplyStyles()
        document = Document(docxfile)
        labeled_author = ""
        aut_dict = {}

        # ── Resolve article JSON path ──────────────────────────────────────
        article_json_path = os.path.join(
            os.path.split(docxfile)[0], f"{jid}_{aid}.json"
        )

        if self.journal_data[jid]["BreakDown"] is True:
            styles = document.styles
            if 'aulabel' not in styles:
                document = new_styles.create_character_style(document, "aulabel")
            au_found = False
            for paragraph in document.paragraphs:
                if paragraph.style.name in ["AU", "AU0"]:
                    au_found = True
                    break
                for run in paragraph.runs:
                    if run.style.name in ["AU", "AU0"]:
                        au_found = True
                        break
            document.save(docxfile)
            time.sleep(0.5)  # Allow AV/indexer to release file handle before NER reads it

            if au_found is True:
                # ── Author labeling with retry and fallback ───────────────
                # Strategy:
                #   1. Try NER pipeline (get_docx_authors) — has internal retry
                #   2. If NER fails → fallback to XSLT pipeline (label_docx_authors)
                #   3. If both fail → continue with empty author data and log warning
                # ──────────────────────────────────────────────────────────
                json_path_for_ner = (
                    article_json_path
                    if os.path.exists(article_json_path)
                    else None
                )

                # ── Step 1: NER pipeline (primary) ────────────────────────
                author_labeled, labeled_author, aut_dict = self.get_docx_authors(
                    docxfile, json_path=json_path_for_ner
                )
                if labeled_author is None:
                    labeled_author = ""
                if aut_dict is None:
                    aut_dict = {}

                # ── Step 2: XSLT fallback ─────────────────────────────────
                if not author_labeled:
                    print("[BreakDown] NER pipeline returned no authors, "
                          "falling back to XSLT labeling")
                    try:
                        author_labeled, labeled_author, aut_dict = self.label_docx_authors(
                            docxfile
                        )
                        if labeled_author is None:
                            labeled_author = ""
                        if aut_dict is None:
                            aut_dict = {}
                    except Exception as xslt_exc:
                        print(f"[BreakDown] XSLT fallback also failed: {xslt_exc}")
                        traceback.print_exc()
                        author_labeled = False
                        labeled_author = ""
                        aut_dict = {}

                # ── Step 3: Log final outcome ─────────────────────────────
                if not author_labeled:
                    print(f"[BreakDown] WARNING: All author labeling methods failed "
                          f"for {docxfile}. Continuing with empty author data.")

        self.update_page_setup(docxfile)
        self.add_dummy_lines(docxfile)
        list_macros = ["Apply_Label"]
        process_doc = OpenDocFile()
        process_error, error_log, doc = process_doc.processDocFile(docxfile, True, True, True, list_macros)
        print(f"Applying Label Styles: {docxfile}")
        self.apply_label_styles(docxfile)
        self.remove_dummy_lines(docxfile)

        article_json = os.path.split(docxfile)[0] + "/" + jid + "_" + aid + ".json"
        with open(article_json) as file:
            article_data = json.load(file)

        author_count = len(article_data['authors_info'])
        author_data = article_data['authors_info']
        funder = article_data['funder_info']
        jrn_acronyms = article_data['journal_info']['journal-accronym']

        oa_query_found = False
        oa_query_count = 0
        funding_query_found = False
        funding_query_count = 0
        orcid_query_found = False
        orcid_query_count = 0
        gen_auq_found = False
        gen_auq_count = 0

        if funder is False:
            funder_found = False
        else:
            funder_found = True

        bjid = jid
        if jid not in self.journal_data:
            if jrn_acronyms in self.journal_data:
                bjid = jrn_acronyms
            else:
                ctypes.windll.user32.MessageBoxW(0, f"JID: {jid} - Details Not available in BreakDown.xlsx\n"
                                                    f"Please add details in XLS and run again...",
                                                 "BreakDown [C&M]", 0)
                exit(0)

        if self.journal_data[bjid]["BreakDown"] is True:
            document = Document(docxfile)
            paragraphs = document.paragraphs
            author_label = None
            styles = document.styles
            breakdown_keys = self.breakdown_styles.keys()
            breakdown_values = self.breakdown_styles.values()

            for br_list in breakdown_keys:
                if br_list not in styles:
                    document = new_styles.create_style(document, br_list)
            for br_list in breakdown_values:
                if br_list not in styles:
                    document = new_styles.create_style(document, br_list)

            if self.journal_data[bjid]["FM_Sequence"] is not None:
                fm_text = self.journal_data[bjid]["FM_Sequence"]
            else:
                fm_text = self.default_data["fm_default"]
            fm_sequence = fm_text.split("\n")

            if self.journal_data[bjid]["BM_Sequence"] is not None:
                bm_text = self.journal_data[bjid]["BM_Sequence"]
            else:
                bm_text = self.default_data["bm_default"]
            bm_sequence = bm_text.split("\n")

            bio_text = self.journal_data[bjid]["Bio"]

            if 'Duplicate' not in styles:
                document = new_styles.create_style(document, "Duplicate")
            if 'CL' not in styles:
                document = new_styles.create_style(document, "CL")
            if 'General_Query' not in styles:
                document = new_styles.create_style(document, "General_Query")

            general_query = self.journal_data[bjid]["General_Queries"]
            article_type = article_data['article_info'].get('article-type') or article_data['article_info'].get('issue-section')
            doi_line = article_data['article_info']['article-doi']
            rrh_line = self.journal_data[bjid]["rrh_format"]
            LRH = self.journal_data[bjid]["LRH"]
            funder_head = self.journal_data[bjid]['funder_text']['funder_head']
            decl_head = self.journal_data[bjid]['declaration_text']['declaration_head']

            paragraphs, decl_found, decl_text = self.check_declaration_section(paragraphs)
            corresp_found = self.check_corresponding_author(article_data['authors_info'])

            if article_data['article_info']['openacess'] is True:
                open_access_found = True
                open_access_query = self.journal_data[bjid]['open_access_query']
            else:
                open_access_found = False

            if corresp_found is True:
                corresp_query = "[AQ: PLEASE CONFIRM CORRESPONDING AUTHOR]"

            if article_data['article_info']['supplemental-material'] is True:
                suppl_found = True
            else:
                suppl_found = self.check_supplementary_material(paragraphs)

            if suppl_found is True:
                supp_text = self.journal_data[bjid]['supplementry_text']

            funding_found = False
            if funder_found is True:
                paragraphs, funding_found, funder_para = self.check_funding_section(paragraphs)
                if funding_found is False:
                    paragraphs, funding_found, funder_para = self.check_funder_paragraph(paragraphs)
                if author_count == 1:
                    if funding_found is True:
                        funder_text = self.journal_data[bjid]['funder_text']["AU"] + " " + funder_para
                    else:
                        funder_text = self.journal_data[bjid]['funder_text']["AU"]
                else:
                    if funding_found is True:
                        funder_text = self.journal_data[bjid]['funder_text']["AUs"] + " " + funder_para
                    else:
                        funder_text = self.journal_data[bjid]['funder_text']["AUs"]
            else:
                if author_count == 1:
                    funder_text = self.journal_data[bjid]['funder_text']["Without_AU"]
                else:
                    funder_text = self.journal_data[bjid]['funder_text']["Without_AUs"]
                if re.search(r"([a-z])", funder_text, re.I):
                    funding_found = True

            if decl_found is True:
                if author_count == 1:
                    if decl_found is True:
                        decl_text = self.journal_data[bjid]['declaration_text']["AU"] + " " + decl_text
                    else:
                        decl_text = self.journal_data[bjid]['declaration_text']["AU"]
                else:
                    if decl_found is True:
                        decl_text = self.journal_data[bjid]['declaration_text']["AUs"] + " " + decl_text
                    else:
                        decl_text = self.journal_data[bjid]['declaration_text']["AUs"]
            else:
                if author_count == 1:
                    decl_text = self.journal_data[bjid]['declaration_text']["Without_AU"]
                else:
                    decl_text = self.journal_data[bjid]['declaration_text']["Without_AUs"]
                if re.search(r"([a-z])", decl_text, re.I):
                    decl_found = True

            orcid_details = []
            orcid_dic = {}
            orcid_labels = []
            orcid_found = False
            authors_info = article_data['authors_info']

            new_authors_info, aut_labeled = self.cross_reference_authors(authors_info, labeled_author, aut_dict)

            if len(aut_dict) > 0:
                if len(new_authors_info) == len(aut_labeled):
                    updated_author_info = new_authors_info
                    for author_key in updated_author_info:
                        author = updated_author_info[author_key]
                        if re.search("([a-z0-9])", author['orcid']):
                            orcid_found = True
                            aut_key = list(aut_labeled.keys())[0]
                            if type(aut_key) is int:
                                author_key = int(author_key)
                            orcid_line = aut_labeled[author_key]
                            orcid_line = orcid_line + " https://orcid.org/" + author['orcid']
                            author_id = author['first-name'] + author['middle-name'] + author['last-name']
                            author_id = author_id.translate(str.maketrans('', '', string.punctuation))
                            author_id = author_id.replace(" ", "")
                            author_id = author_id.lower()
                            author_id = re.sub(r"[^a-zA-z0-9\[\]]", "", author_id)
                            orcid_link = " https://orcid.org/" + author['orcid']
                            orcid_dic[author_id] = orcid_link
                            orcid_labels.append(f"[AU{author_key}]")
                            orcid_details.append(orcid_line)
                else:
                    updated_author_info = authors_info
                    for author_key in updated_author_info:
                        author = updated_author_info[author_key]
                        if re.search("([a-z0-9])", author['orcid']):
                            orcid_found = True
                            orcid_line = None
                            if re.search("([a-z0-9])", author['first-name'], re.IGNORECASE):
                                orcid_line = author['first-name']
                            if re.search("([a-z0-9])", author['middle-name'], re.IGNORECASE):
                                orcid_line = orcid_line + " " + author['middle-name']
                            if re.search("([a-z0-9])", author['last-name'], re.IGNORECASE):
                                orcid_line = orcid_line + " " + author['last-name']
                            orcid_line = orcid_line + " https://orcid.org/" + author['orcid']
                            author_id = author['first-name'] + author['middle-name'] + author['last-name']
                            author_id = author_id.translate(str.maketrans('', '', string.punctuation))
                            author_id = author_id.replace(" ", "")
                            author_id = author_id.lower()
                            author_id = re.sub(r"[^a-zA-z0-9\[\]]", "", author_id)
                            orcid_link = " https://orcid.org/" + author['orcid']
                            orcid_dic[author_id] = orcid_link
                            orcid_labels.append(f"[AU{author_key}]")
                            orcid_details.append(orcid_line)
            else:
                updated_author_info = authors_info
                for author_key in updated_author_info:
                    author = new_authors_info[author_key]
                    if re.search("([a-z0-9])", author['orcid']):
                        orcid_found = True
                        if re.search("([a-z0-9])", author['first-name'], re.IGNORECASE):
                            orcid_line = author['first-name']
                        if re.search("([a-z0-9])", author['middle-name'], re.IGNORECASE):
                            orcid_line = orcid_line + " " + author['middle-name']
                        if re.search("([a-z0-9])", author['last-name'], re.IGNORECASE):
                            orcid_line = orcid_line + " " + author['last-name']
                        orcid_line = orcid_line + " https://orcid.org/" + author['orcid']
                        author_id = author['first-name'] + author['middle-name'] + author['last-name']
                        author_id = author_id.translate(str.maketrans('', '', string.punctuation))
                        author_id = author_id.replace(" ", "")
                        author_id = author_id.lower()
                        author_id = re.sub(r"[^a-zA-z0-9\[\]]", "", author_id)
                        orcid_link = " https://orcid.org/" + author['orcid']
                        orcid_dic[author_id] = orcid_link
                        orcid_labels.append(f"[AU{author_key}]")
                        orcid_details.append(orcid_line)

            history_found = False
            history_dic = {}
            if self.journal_data[bjid]['History_Details'] is not None:
                history_found = True
                history_text = self.journal_data[bjid]['History_Details']
                submitted_date = article_data['article_info']['submitted-date']
                revised_date = article_data['article_info']['revised-date']
                accepted_date = article_data['article_info']['accepted-date']
                if re.search("([0-9]+)", submitted_date):
                    history_dic = self.process_date(submitted_date, 'REC', history_dic)
                if re.search("([0-9]+)", revised_date):
                    history_dic = self.process_date(revised_date, 'REV', history_dic)
                if re.search("([0-9]+)", accepted_date):
                    history_dic = self.process_date(accepted_date, 'ACC', history_dic)

            fm_queries = []
            query_count = 1
            if open_access_found is True:
                oa_query_count = query_count
                oa_query_found, open_access_query, query_count = self.process_query_count(open_access_query,
                                                                                          query_count)
                fm_queries.append(open_access_query)

            general_queries = general_query.split("\n")
            if len(general_queries) > 1:
                for gen_query in general_queries:
                    current_query_cnt = query_count
                    xquery_found, gen_query, query_count = self.process_query_count(gen_query, query_count)
                    if re.search("AUTHOR INFORMATION", gen_query, re.IGNORECASE) and au_found is True:
                        gen_auq_found = True
                        gen_auq_count = current_query_cnt
                    if re.search("FUNDING AND CONFLICT OF INTEREST", gen_query, re.IGNORECASE):
                        funding_query_found = True
                        funding_query_count = current_query_cnt
                    fm_queries.append(gen_query)
            else:
                current_query_cnt = query_count
                xquery_found, gen_query, query_count = self.process_query_count(general_query, query_count)
                if re.search("AUTHOR INFORMATION", gen_query, re.IGNORECASE) and au_found is True:
                    gen_auq_found = True
                    gen_auq_count = current_query_cnt
                if re.search("FUNDING AND CONFLICT OF INTEREST", gen_query, re.IGNORECASE):
                    funding_query_found = True
                    funding_query_count = current_query_cnt
                fm_queries.append(gen_query)

            if funding_found is True and decl_found is True:
                funding_query = self.journal_data[bjid]['funding_conflict_query']
                funding_dec_query_count = query_count
                funding_query_count = query_count
                funding_dec_query_found, funding_query, query_count = self.process_query_count(funding_query,
                                                                                               query_count)
                if funding_dec_query_found is False:
                    funding_query_type = 0
                fm_queries.append(funding_query)
            elif funding_found is True:
                funding_query = self.journal_data[bjid]['funding_query']
                funding_query_count = query_count
                funding_query_found, funding_query, query_count = self.process_query_count(funding_query, query_count)
                funding_query_type = 2
                if funding_query_found is False:
                    funding_query_type = 0
                fm_queries.append(funding_query)
            funding_query_type = 1

            if len(orcid_details) > 0 and au_found is True:
                orcid_query = self.journal_data[bjid]['orcid_query']
                orcid_query_count = query_count
                orcid_query_found, orcid_query, query_count = self.process_query_count(orcid_query, query_count)
                fm_queries.append(orcid_query)

            self.insert_paragraph_before(paragraphs[0], "", "Normal")
            paragraphs = document.paragraphs
            fm_queries.reverse()
            first_paragraph = paragraphs[0]
            line_count = 0

            for fm_query in fm_queries:
                self.insert_paragraph_after(first_paragraph, fm_query, "General_Query")
                line_count = line_count + 1

            paragraphs = document.paragraphs

            if doi_line:
                if 'DOI' not in styles:
                    document = new_styles.create_style(document, "DOI")
                doi_line = "DOI: " + doi_line
                self.insert_paragraph_after(paragraphs[line_count], doi_line, "DOI")
                line_count = line_count + 1
                paragraphs = document.paragraphs

            if LRH:
                if "LRH" not in styles:
                    document = new_styles.create_style(document, "LRH")
                LRH = LRH.strip()
                self.insert_paragraph_after(paragraphs[line_count], LRH, "LRH")
                line_count = line_count + 1
                paragraphs = document.paragraphs

            if "RRH" not in styles:
                document = new_styles.create_style(document, "RRH")

            if rrh_line:
                auth_rhead = self.get_author_rhead(author_data)
                if re.search("Surname", rrh_line, re.I):
                    rrh_line = re.sub("Surname", auth_rhead, rrh_line, re.I)
                self.insert_paragraph_after(paragraphs[line_count], rrh_line, "RRH")
                line_count = line_count + 1
                paragraphs = document.paragraphs
            else:
                rrh_line = self.get_running_head_authors(author_data)
                self.insert_paragraph_after(paragraphs[line_count], rrh_line, "RRH")
                line_count = line_count + 1
                paragraphs = document.paragraphs

            breakdown_dic = {}
            if corresp_found is True:
                corresp_text = self.journal_data[bjid]['Corresponding_Author']
                if "Corresponding" not in styles:
                    document = new_styles.create_style(document, "Corresponding")
                breakdown_dic['Corresponding_Query'] = corresp_query
                if "Corresponding_Title" not in styles:
                    document = new_styles.create_style(document, "Corresponding_Title")
                breakdown_dic['Corresponding_Title'] = corresp_text
                breakdown_dic['Corresponding_Query'] = corresp_query
                breakdown_dic['Corresponding_Para'] = "[INSERT CORRESPONDING TEXT HERE]"

            if article_type is not None:
                if "TY" not in styles:
                    document = new_styles.create_style(document, "TY")
                breakdown_dic['TY'] = article_type

            if history_found is True:
                if "DR" not in styles:
                    document = new_styles.create_style(document, "DR")
                if re.search("REC-DD", history_text):
                    history_text = re.sub("REC-DD", history_dic['REC-DD'], history_text)
                if re.search("REC-MM", history_text):
                    history_text = re.sub("REC-MM", history_dic['REC-Month'], history_text)
                if re.search("REC-Month", history_text):
                    history_text = re.sub("REC-Month", history_dic['REC-Month'], history_text)
                if re.search("REC-YYYY", history_text):
                    history_text = re.sub("REC-YYYY", history_dic['REC-YYYY'], history_text)
                if re.search("REV-DD", history_text):
                    history_text = re.sub("REV-DD", history_dic['REV-DD'], history_text)
                if re.search("REV-MM", history_text):
                    history_text = re.sub("REV-MM", history_dic['REV-Month'], history_text)
                if re.search("REV-Month", history_text):
                    history_text = re.sub("REV-Month", history_dic['REV-Month'], history_text)
                if re.search("REV-YYYY", history_text):
                    history_text = re.sub("REV-YYYY", history_dic['REV-YYYY'], history_text)
                if re.search("ACC-DD", history_text):
                    history_text = re.sub("ACC-DD", history_dic['ACC-DD'], history_text)
                if re.search("ACC-MM", history_text):
                    history_text = re.sub("ACC-MM", history_dic['ACC-Month'], history_text)
                if re.search("ACC-Month", history_text):
                    history_text = re.sub("ACC-Month", history_dic['ACC-Month'], history_text)
                if re.search("ACC-YYYY", history_text):
                    history_text = re.sub("ACC-YYYY", history_dic['ACC-YYYY'], history_text)
                breakdown_dic['DR'] = history_text

            if "Funder_Title" not in styles:
                document = new_styles.create_style(document, "Funder_Title")
            if "Funder" not in styles:
                document = new_styles.create_style(document, "Funder")

            breakdown_dic['Funder_Title'] = funder_head

            if "Funder_Para" not in styles:
                document = new_styles.create_style(document, "Funder_Para")
            breakdown_dic['Funder_Para'] = funder_text

            if "Declaration_Title" not in styles:
                document = new_styles.create_style(document, "Declaration_Title")
            breakdown_dic['Declaration_Title'] = decl_head

            if "Declaration_Para" not in styles:
                document = new_styles.create_style(document, "Declaration_Para")
            if "Declaration" not in styles:
                document = new_styles.create_style(document, "Declaration")
            breakdown_dic['Declaration_Para'] = decl_text

            if orcid_found is True:
                orcid_count = len(orcid_details)
                if orcid_count > 1:
                    orcid_title = "ORCID iDs"
                else:
                    orcid_title = "ORCID iD"
                if "Orcid_Title" not in styles:
                    document = new_styles.create_style(document, "Orcid_Title")
                breakdown_dic['Orcid_Title'] = orcid_title
                if "Orcid_Para" not in styles:
                    document = new_styles.create_style(document, "Orcid_Para")
                orcid_details.reverse()
                for orcid_line in orcid_details:
                    if "Orcid_Para" in breakdown_dic:
                        breakdown_dic['Orcid_Para'].append(orcid_line)
                    else:
                        breakdown_dic['Orcid_Para'] = [orcid_line]

            if suppl_found is True:
                if "Supp_Title" not in styles:
                    document = new_styles.create_style(document, "Supp_Title")
                breakdown_dic['Supp_Title'] = "Supplemental material"
                if "Supp_Para" not in styles:
                    document = new_styles.create_style(document, "Supp_Para")
                breakdown_dic['Supp_Para'] = supp_text

            key_para_found, key_para, para_index = self.find_keyword_paragraph(paragraphs)

            for fm_seq in fm_sequence:
                document, para_index = self.process_fmbm_sequence(document, breakdown_dic, fm_seq, para_index)

            paragraphs = document.paragraphs
            ref_head_found, ref_para_found, ref_para, ref_index = self.find_reference_paragraph(paragraphs)

            if ref_para_found is False and ref_head_found is False:
                ref_index = len(paragraphs)
            ref_index = ref_index - 1

            for bm_seq in bm_sequence:
                document, ref_index = self.process_fmbm_sequence(document, breakdown_dic, bm_seq, ref_index)

            document.save(docxfile)
            doc = Document(docxfile)
            orcid_list = []
            present_styles = []

            for paragraph in document.paragraphs:
                present_styles.append(paragraph.style.name)

            for paragraph in doc.paragraphs:
                if paragraph.style.name in ["AU", "AU0"]:
                    author_label = self.extract_authors_from_xml(paragraph._p.xml)

            if au_found is True:
                author_label, matched_authors, unmatched_authors = self.find_matched_authors(orcid_dic, author_label)
                if len(orcid_dic) == 1 and len(author_label) == 1:
                    orcid_list = orcid_labels
                elif len(unmatched_authors) > 0:
                    orcid_list = orcid_labels
                else:
                    for matched_author in matched_authors:
                        orcid_list.append(author_label[matched_author])

            update_declaration = False
            update_funder = False
            if ["Declaration_Title", "Funder_Title"] in present_styles:
                update_declaration = True
            elif "Declaration_Title" in present_styles:
                update_declaration = True
            elif "Funder_Title" in present_styles:
                update_funder = True

            self.insert_query_callouts(docxfile,
                                       oa_query_found=oa_query_found,
                                       oa_query_count=oa_query_count,
                                       funding_query_found=funding_query_found,
                                       funding_query_count=funding_query_count,
                                       orcid_query_found=orcid_query_found,
                                       gen_auq_found=gen_auq_found,
                                       gen_auq_count=gen_auq_count,
                                       orcid_query_count=orcid_query_count,
                                       update_declaration=update_declaration,
                                       update_funder=update_funder,
                                       orcid_list=orcid_list)

            document = Document(docxfile)
            document = self.rearrange_fmbm_sequence(document, fm_sequence, bm_sequence)
            document = self.map_breakdown_styles(document)

            # Convert any remaining *_Title → EH  and  *_Para → AN
            # that were inserted/rearranged by breakDownProcess AFTER
            # applyStyles.apply_styles() had already run.
            apply_sty = ApplyStyles()
            document = apply_sty.bm_named_styles(document)

            para_num = 0

            for paragraph in document.paragraphs:
                if paragraph.style.name == "TY":
                    self.insert_paragraph_before(paragraph, "", "Normal")
                    break
                para_num = para_num + 1

            paragraphs = document.paragraphs
            paragraph = paragraphs[para_num]
            run = paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
            self.update_title_cases(document)
            paragraphs = document.paragraphs
            document.save(docxfile)

    # ====================
    # Front/Back Matter Processing
    # ====================

    def process_fmbm_sequence(self, document, breakdown_dic, seq, para_index):
        """Process front/back matter sequence items."""
        if bool(re.search("Abstract_Title: Abstract_Para", seq, re.I)) is True:
            document = self.process_abstract_section(document)
        elif bool(re.search("Keyword", seq, re.I)) is True:
            document, para_index = self.process_keywords_section(document, seq, para_index)
        else:
            paragraphs = document.paragraphs
            if bool(re.search("HISTORY", seq, re.I)) is True and "DR" in breakdown_dic:
                self.insert_paragraph_after(paragraphs[para_index], breakdown_dic['DR'], "DR")
                para_index = para_index + 1
                paragraphs = document.paragraphs
            elif bool(re.search("Article_Type", seq, re.I)) is True and "TY" in breakdown_dic:
                self.insert_paragraph_after(paragraphs[para_index], breakdown_dic['TY'], "TY")
                para_index = para_index + 1
                paragraphs = document.paragraphs
            elif bool(re.search("Funder", seq, re.I)) is True and (
                    "Funder_Title" in breakdown_dic or "Funder_Para" in breakdown_dic or "Funder" in breakdown_dic):
                if bool(re.search("Funder_Title: Funder_Para", seq, re.I)) is True:
                    fun_para = breakdown_dic['Funder_Title'] + "<br>:<br> " + breakdown_dic['Funder_Para']
                    self.insert_paragraph_after(paragraphs[para_index], fun_para, "Funder")
                    para_index = para_index + 1
                    paragraphs = document.paragraphs
                elif bool(re.search("^Funder_Title$", seq, re.I)) is True:
                    self.insert_paragraph_after(paragraphs[para_index], breakdown_dic["Funder_Title"], "Funder_Title")
                    para_index = para_index + 1
                    paragraphs = document.paragraphs
                elif bool(re.search("^Funder_Para$", seq, re.I)) is True:
                    self.insert_paragraph_after(paragraphs[para_index], breakdown_dic["Funder_Para"], "Funder_Para")
                    para_index = para_index + 1
                    paragraphs = document.paragraphs
            elif bool(re.search("Declaration", seq, re.I)) is True and (
                    "Declaration_Title" in breakdown_dic or "Declaration_Para" in breakdown_dic or "Declaration" in breakdown_dic):
                if bool(re.search("Declaration_Title: Declaration_Para", seq, re.I)) is True:
                    dec_para = breakdown_dic['Declaration_Title'] + "<br>:<br> " + breakdown_dic['Declaration_Para']
                    self.insert_paragraph_after(paragraphs[para_index], dec_para, "Declaration")
                    para_index = para_index + 1
                    paragraphs = document.paragraphs
                elif bool(re.search("^Declaration_Title$", seq, re.I)) is True:
                    self.insert_paragraph_after(paragraphs[para_index], breakdown_dic["Declaration_Title"],
                                                "Declaration_Title")
                    para_index = para_index + 1
                    paragraphs = document.paragraphs
                elif bool(re.search("^Declaration_Para$", seq, re.I)) is True:
                    self.insert_paragraph_after(paragraphs[para_index], breakdown_dic["Declaration_Para"],
                                                "Declaration_Para")
                    para_index = para_index + 1
                    paragraphs = document.paragraphs
            elif bool(re.search("Corresponding", seq, re.I)) is True and (
                    'Corresponding_Title' in breakdown_dic or 'Corresponding_Para' in breakdown_dic or 'Corresponding' in breakdown_dic):
                if bool(re.search("Corresponding_Title: Corresponding_Para", seq, re.I)) is True:
                    dec_para = breakdown_dic['Corresponding_Title'] + "<br>" + breakdown_dic['Corresponding_Para']
                    self.insert_paragraph_after(paragraphs[para_index], dec_para, "Corresponding_Title")
                    para_index = para_index + 1
                    paragraphs = document.paragraphs
                elif bool(re.search(
                        "(Corresponding_Title: |Corresponding_Title:|Corresponding_Title |Corresponding_Title)", seq,
                        re.I)) is True:
                    self.insert_paragraph_after(paragraphs[para_index], breakdown_dic["Corresponding_Title"],
                                                "Corresponding_Title")
                    para_index = para_index + 1
                    paragraphs = document.paragraphs
                elif bool(re.search("^Correspoinding_Para$", seq, re.I)) is True:
                    self.insert_paragraph_after(paragraphs[para_index], breakdown_dic["Correspoinding_Para"],
                                                "Correspoinding_Para")
                    para_index = para_index + 1
                    paragraphs = document.paragraphs
            elif bool(re.search("Orcid", seq, re.I)) is True and 'Orcid_Title' in breakdown_dic:
                for paragraph in paragraphs:
                    para_style = paragraph.style.name
                    if re.match("^(Orcid ID|Orcid IDs)$", paragraph.text, re.I) and para_style != "Orcid_Title":
                        paragraph.clear()
                    for orcid in breakdown_dic['Orcid_Para']:
                        if re.match(orcid, paragraph.text, re.I) and para_style != "Orcid_Para":
                            paragraph.clear()
                if bool(re.search("Orcid_Title: Orcid_Para", seq, re.I)) is True:
                    dec_para = breakdown_dic['Orcid_Title'] + "<br>" + "\n".join(breakdown_dic['Orcid_Para'])
                    self.insert_paragraph_after(paragraphs[para_index], dec_para, "Orcid")
                    para_index = para_index + 1
                    paragraphs = document.paragraphs
                elif bool(re.search("^Orcid_Title$", seq, re.I)) is True:
                    self.insert_paragraph_after(paragraphs[para_index], breakdown_dic["Orcid_Title"],
                                                "Orcid_Title")
                    para_index = para_index + 1
                    paragraphs = document.paragraphs
                elif bool(re.search("^Orcid_Para$", seq, re.I)) is True:
                    orcid_lines = breakdown_dic["Orcid_Para"]
                    for orcid_ln in breakdown_dic["Orcid_Para"]:
                        self.insert_paragraph_after(paragraphs[para_index], orcid_ln, "Orcid_Para")
                    para_index = para_index + len(breakdown_dic["Orcid_Para"])
                    paragraphs = document.paragraphs
            elif bool(re.search("Supp", seq, re.I)) is True and 'Supp_Title' in breakdown_dic:
                if bool(re.search("Supp_Title: Supp_Para", seq, re.I)) is True:
                    dec_para = breakdown_dic['Supp_Title'] + "<br>" + breakdown_dic['Supp_Para']
                    self.insert_paragraph_after(paragraphs[para_index], dec_para, "Supp")
                    para_index = para_index + 1
                    paragraphs = document.paragraphs
                elif bool(re.search("^Supp_Title$", seq, re.I)) is True:
                    self.insert_paragraph_after(paragraphs[para_index], breakdown_dic["Supp_Title"],
                                                "Supp_Title")
                    para_index = para_index + 1
                    paragraphs = document.paragraphs
                elif bool(re.search("^Supp_Para$", seq, re.I)) is True:
                    self.insert_paragraph_after(paragraphs[para_index], breakdown_dic["Supp_Para"],
                                                "Supp_Para")
                    para_index = para_index + 1
                    paragraphs = document.paragraphs
        return document, para_index

    def process_abstract_section(self, document):
        """Process abstract section formatting."""
        paragraphs = document.paragraphs
        abkh_count = 0
        for paragraph in paragraphs:
            if paragraph.style.name.lower() == "abkwh":
                abkh_count += 1
        if abkh_count == 1:
            for paragraph in paragraphs:
                para_text = paragraph.text
                para_style = paragraph.style.name
                if para_style == "ABKWH" and re.search("^Abstract", para_text, re.I):
                    paragraph.clear()
                    paragraph.style = "ABKW"
                    if re.search(":$", para_text):
                        abs_para = paragraph.add_run(f"{para_text} ")
                    elif re.search(": $", para_text):
                        abs_para = paragraph.add_run(f"{para_text}")
                    else:
                        abs_para = paragraph.add_run(f"{para_text}: ")
                    abs_para.font.bold = True
                    break
        return document

    def process_keywords_section(self, document, seq, para_index):
        """Process keywords section formatting."""
        paragraphs = document.paragraphs
        available_styles = []
        for paragraph in document.paragraphs:
            available_styles.append(paragraph.style.name)
        for paragraph in paragraphs:
            para_text = paragraph.text
            para_style = paragraph.style.name
            if re.search(r"^((Keyword|Key word)(s|)( |)(\.|\:|\;|)( |))", para_text, re.IGNORECASE):
                doc_style = document.styles
                if "Keyword_Title" in available_styles and "Keyword_Para" in available_styles:
                    pass
                if "Keyword_Para" in available_styles and "Keyword_Title" not in available_styles:
                    num_remove = False
                    if re.search(r"^((Keyword|Key word)(s|)( |)(\.|\:|\;|)( |))", para_text, re.IGNORECASE):
                        x = re.findall(r"^((Keyword|Key word)(s|)( |)(\.|\:|\;|)( |))", para_text, re.IGNORECASE)
                        cp_label = x[0][0]
                        if para_text.startswith(cp_label) is True:
                            inline = paragraph.runs
                            eq_txt = ""
                            i_value = None
                            for i in range(len(inline)):
                                eq_txt += inline[i].text
                                pattern = re.escape(eq_txt)
                                if re.match(f"^{pattern}$", cp_label):
                                    i_value = i
                                    break
                            if i_value is not None:
                                for i in reversed(range(len(inline))):
                                    if i <= i_value:
                                        run = paragraph.runs[i]
                                        try:
                                            paragraph._p.remove(run._r)
                                            num_remove = True
                                        except:
                                            pass
                    if bool(re.search("Keyword_Title: Keyword_Para", seq, re.I)) is True:
                        if num_remove is True:
                            elem = paragraph._element
                            r = OxmlElement('w:r')
                            elem.insert(1, r)
                            new_run = Run(r, paragraph)
                            new_run.text = "Keywords: "
                            new_run.bold = True
                    elif bool(re.search("(Keyword_Title: |Keyword_Title:|Keyword_Title |Keyword_Title)", seq,
                                        re.I)) is True:
                        if num_remove is True:
                            self.insert_paragraph_before(paragraph, "Keywords", "Keyword_Title")
                            para_index = para_index + 1
        return document, para_index

    # ====================
    # Text Formatting
    # ====================

    def _format_and_add_text(self, new_para, text, style):
        """Format and add text to paragraph based on style."""
        if style == "General_Query":
            new_para.add_run(text)
        elif style == "Correponding_Title":
            corresp_head = new_para.add_run(text)
            corresp_head.font.bold = True
        elif style == "Funder":
            funder_text = re.sub("<br>", "\t", text, 9, re.IGNORECASE)
            funder_list = funder_text.split("\t")
            for hist in funder_list:
                hist_p = new_para.add_run(hist)
                if bool(re.search("^Funding$", hist, re.I)) is True:
                    hist_p.font.bold = True
                    hist_p.font.all_caps = True
                elif bool(re.search("GQ", hist, re.I)) is True:
                    hist_p.font.bold = True
                    hist_p.font.superscript = True
                    hist_p.font.color.rgb = RGBColor(255, 0, 0)
                elif bool(re.search("^:$", hist, re.I)) is True:
                    hist_p.font.bold = True
        elif style == "Declaration":
            decl_text = re.sub("<br>", "\t", text, 9, re.IGNORECASE)
            decl_list = decl_text.split("\t")
            for hist in decl_list:
                hist_p = new_para.add_run(hist)
                if bool(re.search("^Declaration", hist, re.I)) is True:
                    hist_p.font.bold = True
                    hist_p.font.all_caps = True
                elif bool(re.search("GQ", hist, re.I)) is True:
                    hist_p.font.bold = True
                    hist_p.font.superscript = True
                    hist_p.font.color.rgb = RGBColor(255, 0, 0)
                elif bool(re.search("^:$", hist, re.I)) is True:
                    hist_p.font.bold = True
        elif style == "Corresponding_Title":
            corr_text = re.sub("<br>", "\t", text, 9, re.IGNORECASE)
            corr_list = corr_text.split("\t")
            for hist in corr_list:
                hist_p = new_para.add_run(hist)
                if bool(re.search("^Corresponding", hist, re.I)) is True:
                    hist_p.font.bold = True
                else:
                    hist_p.font.color.rgb = RGBColor(255, 0, 0)
                    hist_p.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif style == "Funder_Title":
            funder_text = re.sub("<br>", "\t", text, 9, re.IGNORECASE)
            funder_list = funder_text.split("\t")
            for hist in funder_list:
                hist_p = new_para.add_run(hist)
                if bool(re.search("GQ", hist, re.I)) is True:
                    hist_p.font.bold = True
                    hist_p.font.superscript = True
                    hist_p.font.color.rgb = RGBColor(255, 0, 0)
        elif style == "Declaration_Title":
            decl_text = re.sub("<br>", "\t", text, 9, re.IGNORECASE)
            decl_list = decl_text.split("\t")
            for hist in decl_list:
                hist_p = new_para.add_run(hist)
                if bool(re.search("GQ", hist, re.I)) is True:
                    hist_p.font.bold = True
                    hist_p.font.superscript = True
                    hist_p.font.color.rgb = RGBColor(255, 0, 0)
        elif style == "Orcid_Title":
            para_text = re.sub("<br>", "\t", text, 9, re.IGNORECASE)
            para_list = para_text.split("\t")
            for para in para_list:
                para_p = new_para.add_run(para)
                if bool(re.search("GQ", para, re.I)) is True:
                    para_p.font.bold = True
                    para_p.font.superscript = True
                    para_p.font.color.rgb = RGBColor(255, 0, 0)
        else:
            new_para.add_run(text)

    def update_funding_paragraph(self, document, funding_head, funding_text):
        """Update funding paragraph with specific formatting."""
        paras = document.paragraphs
        funding_index = 0
        for para in paras:
            para_style = para.style.name
            para_text = para.text
            if re.search("^Funding", para_text, re.IGNORECASE) and re.search("(H1|H2|H3|H4)", para_style):
                para.text = funding_head
                next_para = paras[funding_index + 1]
                next_style = next_para.style.name
                if re.search("^TEXT", next_style, re.IGNORECASE):
                    next_para.text = funding_text
            funding_index = funding_index + 1

    # ====================
    # Date Processing
    # ====================

    def process_date(self, date_str, date_type, history_dict):
        """Process date string and add to history dictionary."""
        if not date_str or not re.search(r"([0-9]+)", date_str):
            return history_dict

        try:
            date_parts = date_str.split("-")
            if len(date_parts) < 3:
                return history_dict

            month_abbr = date_parts[1]

            # Convert month abbreviation
            month_num = self.abbr_to_num.get(month_abbr)
            if month_num is None:
                return history_dict

            month_full = datetime.datetime.strptime(month_abbr, '%b').strftime('%B')

            # Add to dictionary
            history_dict.update({
                f"{date_type}-DD": date_parts[0],
                f"{date_type}-MM": month_num,
                f"{date_type}-Month": month_full,
                f"{date_type}-YYYY": date_parts[2]
            })

        except (ValueError, IndexError, KeyError) as e:
            print(f"Error processing date {date_str}: {e}")

        return history_dict

    # ====================
    # Sequence Rearrangement
    # ====================

    def rearrange_fmbm_sequence(self, document, fm_seq, bm_seq):
        """Rearrange front/back matter sequences."""
        paragraphs = document.paragraphs
        abkh_count = 0
        for paragraph in paragraphs:
            if paragraph.style.name.lower() == "abkwh":
                abkh_count += 1
        keyword_paray_found, key_para, para_index = self.find_keyword_paragraph(paragraphs)
        ref_head_found, ref_para_found, ref_para, ref_index = self.find_reference_paragraph(paragraphs)
        fm_seq.insert(0, "General_Query")
        sep_triples = ()
        fm_dic = {}
        fm_cnt = 1
        fm_new_seq = fm_seq
        for fm_s in fm_seq:
            if re.search("Funder_Title: Funder_Para", fm_s, re.I):
                i = fm_seq.index(fm_s)
                fm_seq[i] = "Funder"
                continue
            if re.search("Declaration_Title: Declaration_Para", fm_s, re.I):
                i = fm_seq.index(fm_s)
                fm_seq[i] = "Declaration"
                continue
            if re.search(": ", fm_s, re.I):
                i = fm_seq.index(fm_s)
                fm_sq = fm_s.split(": ")
                fm_seq[i:i + 1] = fm_sq
        for fm in fm_seq:
            fm = re.sub("(:|: | )$", "", fm, re.I)
            if fm in self.breakdown_styles:
                fm_dic[fm_cnt] = self.breakdown_styles[fm]
            fm_cnt = fm_cnt + 1
        fm_paras = {}
        for i, paragraph in enumerate(document.paragraphs):
            style_name = paragraph.style.name
            if bool(re.match("(H1|H2|H3|H4|H5|H6|TEXT IND|TEXT|nomenclature)", style_name, re.I)):
                break
            if abkh_count > 1 and bool(re.match("ABKWH", style_name, re.I)):
                break
            dic_count = 1
            if style_name not in fm_paras:
                fm_paras[style_name] = {dic_count: paragraph._p}
            else:
                dic_count = len(fm_paras[style_name])
                dic_count = dic_count + 1
                new_dic = {dic_count: paragraph._p}
                fm_paras[style_name].update(new_dic)
        final_fmparas = {}
        for cnt in fm_dic:
            fm_ord = fm_dic[cnt]
            if fm_ord in fm_paras:
                final_fmparas[fm_ord] = fm_paras[fm_ord]
        first_p = paragraphs[0]._p
        para_count = 0
        final_dic = {}
        for style in final_fmparas:
            style_item = fm_paras[style]
            for sty in style_item:
                p = style_item[sty]
                if style not in final_dic:
                    final_dic[style] = {}
                    final_dic[style][sty] = p
                else:
                    final_dic[style][sty] = p
        aff_count = 0
        au_count = 0
        if "AU" in final_dic:
            au_dic = final_dic["AU"]
            au_count = len(au_dic)
        if "AF" in final_dic:
            aff_dic = final_dic['AF']
            aff_count = len(aff_dic)
        aff_au_dic = {}
        if au_count > 1 and aff_count > 1:
            au_aff_count = 1
            for au in au_dic:
                author = au_dic[au]
                if au in aff_dic:
                    aff = aff_dic[au]
                if "AU_AFF" not in aff_au_dic:
                    aff_au_dic["AU_AFF"] = {}
                au_aff_dic = aff_au_dic["AU_AFF"]
                if au not in au_aff_dic:
                    aff_au_dic["AU_AFF"][au] = {}
                if author is not None and aff is not None:
                    aff_au_dic["AU_AFF"][au] = {author: aff}
                elif author is not None and aff is None:
                    aff_au_dic["AU_AFF"][au] = {author: ""}
                elif aff is not None and author is None:
                    aff_au_dic["AU_AFF"][au] = {"", aff}
                au_aff_count = au_aff_count + 1
            final_dic["AU"] = aff_au_dic["AU_AFF"]
            del final_dic["AF"]
        paragraphs = document.paragraphs
        first_p = paragraphs[0]._p
        pr_count = 0
        for dic_value in final_dic:
            br_dic = final_dic[dic_value]
            br_count = len(br_dic)
            if br_count > 1 and re.match("ABKWH|AU", dic_value):
                continue
            for br_para in br_dic:
                new_para = br_dic[br_para]
                para_type = type(new_para)
                if type(new_para) is dict:
                    for para_key in new_para:
                        para_value = new_para[para_key]
                        first_p.addnext(para_key)
                        paragraphs = document.paragraphs
                        para_key.addnext(para_value)
                        first_p = para_value
                        pr_count = pr_count + 2
                        paragraphs = document.paragraphs
                else:
                    first_p.addnext(new_para)
                    paragraphs = document.paragraphs
                    pr_count = pr_count + 1
                    first_p = paragraphs[pr_count]._p
        return document

    # ====================
    # File Operations
    # ====================

    def add_dummy_lines(self, docxfile):
        """Add dummy lines for processing."""
        document = Document(docxfile)
        if "remove" not in document.styles:
            document.styles.add_style("remove", WD_STYLE_TYPE.PARAGRAPH)
        for paragraph in document.paragraphs:
            para_style = paragraph.style.name
            if re.match(r"^(CP|AWKH|Keyword_Para)$", para_style):
                self.insert_paragraph_before(paragraph, "", "remove")
        document.save(docxfile)

    def remove_dummy_lines(self, docxfile):
        """Remove dummy lines after processing."""
        document = Document(docxfile)
        for paragraph in document.paragraphs:
            para_style = paragraph.style.name
            if re.match(r"^(remove)$", para_style):
                self.delete_paragraph(paragraph)
        document.save(docxfile)

    def apply_label_styles(self, file_name):
        """Apply label styles to document elements."""
        document = Document(file_name)
        styles = document.styles
        for paragraph in document.paragraphs:
            para_xml = paragraph._p.xml
            para_xml = re.sub(r"<w:t xml:space=\"preserve\">\s+</w:t>", "<w:t xml:space=\"preserve\"> </w:t>", para_xml)
        for paragraph in document.paragraphs:
            style_name = paragraph.style.name
            if style_name == "EQ":
                para_text = paragraph.text
                p = paragraph._p
                num_remove = False
                if re.search(r"(( |)(\(|\[)([A-Z0-9]+|[0-9A-Z]+|[0-9]+)(\)|\]))", para_text, re.IGNORECASE):
                    x = re.findall(r"(( |)(\(|\[)([A-Z0-9]+|[0-9A-Z]+|[0-9]+)(\)|\]))", para_text, re.IGNORECASE)
                    eq_label = x[0][0]
                    if para_text.endswith(eq_label) is True:
                        p = paragraph._p
                        inline = paragraph.runs
                        eq_txt = ""
                        i_value = None
                        eq_range = reversed(range(len(inline)))
                        for i in eq_range:
                            eq_txt = inline[i].text + eq_txt
                            pattern = re.escape(eq_txt)
                            if re.match(pattern, eq_label):
                                i_value = i
                                break
                        if i_value is not None:
                            new_range = reversed(range(len(inline)))
                            for i in new_range:
                                run = paragraph.runs[i]
                                if i >= i_value:
                                    try:
                                        p_xml = run._r.xml
                                        if re.search(r"ProgID\=\"Equation", p_xml) or re.search("m:oMath", p_xml):
                                            pass
                                        else:
                                            paragraph._p.remove(run._r)
                                            num_remove = True
                                    except:
                                        print("error\n")
                                        pass
                                else:
                                    p_xml = run._r.xml
                        if num_remove is True:
                            eq_label = re.sub("^ ", "", eq_label)
                            self.insert_paragraph_after(paragraph, eq_label, "EN")
            if style_name == "ABKW":
                para_text = paragraph.text
                p = paragraph._p
                num_remove = False
                if re.search(r"^(Abstract( |)(\.|\:|\;|)( |))", para_text, re.IGNORECASE):
                    x = re.findall(r"^(Abstract( |)(\.|\:|\;|)( |))", para_text, re.IGNORECASE)
                    cp_label = x[0][0]
                    if para_text.startswith(cp_label) is True:
                        p = paragraph._p
                        inline = paragraph.runs
                        eq_txt = ""
                        i_value = None
                        for i in range(len(inline)):
                            eq_txt += inline[i].text
                            pattern = re.escape(eq_txt)
                            if re.match(f"^{pattern}$", cp_label):
                                i_value = i
                                break
                        if i_value is not None:
                            for i in reversed(range(len(inline))):
                                if i <= i_value:
                                    run = paragraph.runs[i]
                                    try:
                                        paragraph._p.remove(run._r)
                                        num_remove = True
                                    except:
                                        pass
                            if num_remove is True:
                                self.insert_paragraph_before(paragraph, "Abstract", "ABKWH")
            if style_name == "CP":
                para_text = paragraph.text
                p = paragraph._p
                # ------------------------------------------------------------------
                # CP label-split: handles both plain <w:r> runs AND <w:fldSimple>
                # SEQ fields (e.g. "Figure <SEQ> 1 </SEQ>. Caption text").
                # The original run-only walk missed fldSimple nodes, so figure/table
                # numbers stored in SEQ fields were never matched and the split was
                # silently skipped.  We now walk all inline children via lxml.
                # ------------------------------------------------------------------
                _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                _INLINE_TAGS = (
                    f'{{{_W}}}r',
                    f'{{{_W}}}fldSimple',
                    f'{{{_W}}}hyperlink',
                    f'{{{_W}}}ins',
                    f'{{{_W}}}del',
                )

                def _inline_text(node):
                    """Return concatenated text of all <w:t> descendants."""
                    return ''.join((t.text or '') for t in node.iter(f'{{{_W}}}t'))

                def _split_cp_label(paragraph, cp_label, cp_label_pattern):
                    """
                    Walk inline children of *paragraph* via lxml, find the child
                    that completes *cp_label*, remove all children up-to-and-
                    including that child (trimming its text if needed), and return
                    the extracted label string.  Returns None if not matched.
                    """
                    p_elem = paragraph._p
                    inline_children = [c for c in p_elem if c.tag in _INLINE_TAGS]
                    accumulated = ''
                    split_child = None
                    split_accumulated_before = ''

                    for child in inline_children:
                        child_text = _inline_text(child)
                        new_accumulated = accumulated + child_text
                        if re.match(fr'^{re.escape(cp_label_pattern.rstrip())}( |)$',
                                    new_accumulated.rstrip(), re.IGNORECASE):
                            split_child = child
                            split_accumulated_before = accumulated
                            break
                        # label ends mid-child (e.g. ". Caption" run)
                        if len(new_accumulated) >= len(cp_label.rstrip()) and split_child is None:
                            split_child = child
                            split_accumulated_before = accumulated
                            break
                        accumulated = new_accumulated

                    if split_child is None:
                        return None

                    # Trim the split_child so only the label portion remains in it,
                    # then strip the rest of that child's text (the caption start).
                    remaining_label = cp_label[len(split_accumulated_before):]
                    for t_elem in split_child.iter(f'{{{_W}}}t'):
                        t_text = t_elem.text or ''
                        keep = t_text[:len(remaining_label)]
                        leftover = t_text[len(remaining_label):]
                        remaining_label = remaining_label[len(keep):]
                        if leftover:
                            # Put leftover back — it is the start of the caption
                            t_elem.text = leftover.lstrip()
                            break
                        else:
                            # Label consumed this entire <w:t>; clear it so
                            # the run can be removed (avoids duplicate label
                            # in the CP paragraph).
                            t_elem.text = ''

                    # Remove every inline child that comes strictly before split_child
                    reached = False
                    for child in list(p_elem):
                        if child is split_child:
                            reached = True
                            # Remove split_child only if it is now empty
                            if not _inline_text(split_child).strip():
                                p_elem.remove(split_child)
                            break
                        if child.tag in _INLINE_TAGS:
                            p_elem.remove(child)

                    return cp_label

                # ---- Figure / Fig. ----
                if re.search(r"^((Figure|Fig\.)( |)([0-9]+|[0-9a-z]+)(\.|\:|\;|)( |))", para_text, re.IGNORECASE):
                    x = re.findall(r"^((Figure|Fig\.)( |)([0-9]+|[0-9a-z]+)(\.|\:|\;|)( |))", para_text, re.IGNORECASE)
                    cp_label = x[0][0]
                    fig_no = x[0][3]
                    # Normalize: ensure space between keyword and number,
                    # include any separator (. : ;), strip trailing space.
                    cpb_label = f"{x[0][1]} {fig_no}{x[0][4]}"
                    if para_text.startswith(cp_label):
                        extracted = _split_cp_label(paragraph, cp_label, cp_label)
                        if extracted is not None:
                            fig_cl = f"[FIGURE {fig_no} ABOUT HERE]"
                            self.insert_paragraph_before(paragraph, fig_cl, "CL")
                            self.insert_paragraph_before(paragraph, cpb_label, "CPB")

                # ---- Box ----
                if re.search(r"^((Box)( |)([0-9]+|[0-9a-z]+)(\.|\:|\;|)( |))", para_text, re.IGNORECASE):
                    x = re.findall(r"^((Box)( |)([0-9]+|[0-9a-z]+)(\.|\:|\;|)( |))", para_text, re.IGNORECASE)
                    cp_label = x[0][0]
                    cpb_label = f"{x[0][1]} {x[0][3]}{x[0][4]}"
                    if para_text.startswith(cp_label):
                        extracted = _split_cp_label(paragraph, cp_label, cp_label)
                        if extracted is not None:
                            self.insert_paragraph_before(paragraph, cpb_label, "CPB")

                # ---- Table / Tab. ----
                if re.search(r"^((Table|Tab\.)( |)([0-9]+|[0-9a-z]+)(\.|\:|\;|)( |))", para_text, re.IGNORECASE):
                    x = re.findall(r"^((Table|Tab\.)( |)([0-9]+|[0-9a-z]+)(\.|\:|\;|)( |))", para_text, re.IGNORECASE)
                    cp_label = x[0][0]
                    cpb_label = f"{x[0][1]} {x[0][3]}{x[0][4]}"
                    if para_text.startswith(cp_label):
                        extracted = _split_cp_label(paragraph, cp_label, cp_label)
                        if extracted is not None:
                            self.insert_paragraph_before(paragraph, cpb_label, "CPB")
        document.save(file_name)

    def update_page_setup(self, docxfile):
        """Update document page setup to standard format."""
        document = Document(docxfile)
        section = document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.header_distance = Mm(36)
        section.footer_distance = Mm(36)
        document.save(docxfile)

    def update_title_cases(self, document):
        """Update title cases for specific styles."""
        for paragraph in document.paragraphs:
            styleName = paragraph.style.name
            if re.match("AT|H1|H2|H3|H4|EH", styleName, re.IGNORECASE):
                paraText = paragraph.text
                non_ascii_found = False
                for c in paraText:
                    if 0 <= ord(c) <= 127:
                        pass
                    else:
                        non_ascii_found = True
                if not re.search("[a-z]", paraText) and non_ascii_found is False:
                    title_text = titlecase(paraText)
                    if re.match(paraText, title_text, re.IGNORECASE):
                        paragraph.text = title_text

    def map_breakdown_styles(self, document):
        """Map breakdown styles to their corresponding styles."""
        paragraphs = document.paragraphs
        mapping_list = self.breakdown_mapping_tags
        for paragraph in paragraphs:
            style = paragraph.style.name
            if style in mapping_list:
                map_style = mapping_list[style]
                paragraph.style = map_style
        return document

    # ====================
    # Author Processing
    # ====================

    def extract_authors_from_xml(self, input_string):
        """Extract author names and labels from paragraph XML."""
        root = etree.fromstring(input_string)
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        text_elements = root.xpath(
            '//w:t[not(ancestor::w:rPr/w:vertAlign[@w:val="superscript"]) or not(ancestor::w:r/w:rPr/w:vertAlign[@w:val="superscript"])]/text()',
            namespaces=namespaces)
        input_string = ''.join(text_elements)
        input_string = re.sub(r" and | \& ", " ", input_string, re.I)
        pattern = r'([A-Za-z0-9]+)\[(AU\d+)\]'
        input_string = re.sub(r"([^a-zA-Z0-9\[\]]+)", "", input_string)
        authors = {}
        matches = re.findall(pattern, input_string)
        for match in matches:
            author_name = match[0]
            author_index = match[1]
            author_name = author_name.translate(str.maketrans('', '', string.punctuation))
            author_name = author_name.replace(" ", "")
            author_name = author_name.lower()
            author_name = re.sub(r"([0-9]+)", "", author_name)
            authors[author_name] = f"[{author_index}]"
        return authors

    def label_docx_authors(self, docx_file_path):
        """Label authors in DOCX file using XSLT transformation."""
        author_para = ""
        aut_dict = {}
        author_found = False
        tmp_file = docx_file_path + '.tmp'

        try:
            with zipfile.ZipFile(docx_file_path, 'r') as zip_ref:
                zip_ref.extract('word/document.xml', 'temp')
            xslt_file_path = "xsl/label_author_names.xsl"
            subprocess.run(
                ['java', '-jar', 'ParaStyler/saxon9pe.jar', '-s:temp/word/document.xml', '-xsl:' + xslt_file_path,
                 '-o:temp/word/document.xml'])
            with zipfile.ZipFile(docx_file_path, 'r') as zip_read:
                with zipfile.ZipFile(tmp_file, 'w') as zip_write:
                    for item in zip_read.infolist():
                        if item.filename != 'word/document.xml':
                            buffer = zip_read.read(item.filename)
                            zip_write.writestr(item, buffer)
            with zipfile.ZipFile(tmp_file, 'a') as zip_ref:
                zip_ref.write('temp/word/document.xml', 'word/document.xml')
            author_found, author_line = self.get_author_paragraphs("temp/word/document.xml")
            if author_found:
                with open("temp/author_line.xml", "w", encoding="utf-8") as file:
                    file.write(author_line)
                trns_xml = XmlTransform()
                trans_msg = trns_xml.jar_transform("temp/author_line.xml", "xsl/docxXmlToHtml.xsl", "temp/author_line.htm")
                if trans_msg.returncode == 0:
                    with open("temp/author_line.htm", "r", encoding="utf-8") as file:
                        html_content = file.read()

                    html_content = re.sub(", ,", ", ", html_content, re.I)
                    html_content = re.sub(",,", ",", html_content, re.I)
                    html_content = re.sub(r"([a-z]) (\*)", r"\1\2", html_content, re.I)
                    html_content = re.sub(r"<sup>(\,)<\/sup> ([A-Z])", r"\1 \2", html_content)
                    html_content = re.sub(r"<sup>(\,)<\/sup>([A-Z])", r"\1 \2", html_content)
                    html_content = re.sub(r"<sup> (\,)<\/sup>([A-Z])", r"\1 \2", html_content)
                    html_content = re.sub("</sup><sup>", "", html_content)
                    html_content = re.sub(r'(<span([^>+]|)>(((?!</?span>).)*)</span>)</sup>', r"</sup>\1", html_content,
                                          flags=re.I | re.S | re.DOTALL)
                    match = re.search(r"<p>((?:[^<]*|<(?!/p))+)</p>", html_content, re.I)
                    if match:
                        author_para = match.group(1)
                    else:
                        print(f"[label_docx_authors] XSLT produced HTML but no <p> tag found")
                        author_found = False
                else:
                    print(f"[label_docx_authors] XSLT transform failed with return code: {trans_msg.returncode}")
                    author_found = False
            else:
                print("[label_docx_authors] No AU/AU0 paragraph found in document XML")
        except Exception as exc:
            print(f"[label_docx_authors] XSLT pipeline error: {exc}")
            traceback.print_exc()
            author_found = False
        finally:
            # Clean up temp files
            if os.path.exists('temp/word/document.xml'):
                os.remove('temp/word/document.xml')
            if os.path.exists(tmp_file):
                os.replace(tmp_file, docx_file_path)

        return author_found, author_para, aut_dict

    def get_docx_authors(self, docx_file_path, json_path=None):
        """
        Extract and label authors from a DOCX file.

        NEW IMPLEMENTATION — HTML token-stream pipeline (labelAuthorsNER).

        Uses a retry mechanism (up to 2 attempts) to handle transient failures
        like file locks or temporary I/O errors before falling through.

        Returns:
            (docx_labeled: bool, labeled_html: str|None, aut_dict: dict)
            — same signature as the original method so all callers work unchanged.
        """
        max_retries = 2
        last_error = None

        for attempt in range(1, max_retries + 1):
            try:
                docx_labeled, labeled_html, aut_dict = label_au_paragraph_in_docx(
                    docx_file_path,
                    json_path=json_path,
                    output_path=docx_file_path,   # in-place edit
                )
                if docx_labeled:
                    print(f"[get_docx_authors] Labeled {len(aut_dict)} authors "
                          f"via NER pipeline (attempt {attempt})")
                    return docx_labeled, labeled_html, aut_dict
                else:
                    print(f"[get_docx_authors] NER pipeline returned no authors "
                          f"(attempt {attempt})")
                    return False, None, {}

            except Exception as exc:
                last_error = exc
                print(f"[get_docx_authors] NER pipeline error (attempt {attempt}/{max_retries}): {exc}")
                traceback.print_exc()
                if attempt < max_retries:
                    print(f"[get_docx_authors] Retrying in 1 second...")
                    time.sleep(1)

        print(f"[get_docx_authors] NER pipeline failed after {max_retries} attempts: {last_error}")
        return False, None, {}

    def get_author_paragraphs(self, document_xml_path):
        """Extract author paragraphs from XML."""
        tree = etree.parse(document_xml_path)
        root = tree.getroot()
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        xpath_expr = "//w:p[w:pPr/w:pStyle[@w:val='AU0' or @w:val='AU']]"
        paragraphs = root.xpath(xpath_expr, namespaces=namespaces)
        author_paragraph_found = False
        author_paragraph_xml = None
        if paragraphs:
            author_paragraph_xml = etree.tostring(paragraphs[0], encoding='utf-8').decode('utf-8')
            author_paragraph_found = True
        return author_paragraph_found, author_paragraph_xml

    # ====================
    # Query Callouts
    # ====================

    def insert_query_callouts(self, docxfile,
                              oa_query_found, oa_query_count,
                              funding_query_found, funding_query_count,
                              orcid_query_found, orcid_query_count,
                              gen_auq_found, gen_auq_count,
                              update_declaration, update_funder,
                              orcid_list):
        """Insert query callouts — always uses Java Saxon jar (saxonc DLL crashes)."""
        self.insert_query_callouts_jar(docxfile,
                                       oa_query_found, oa_query_count,
                                       funding_query_found, funding_query_count,
                                       orcid_query_found, orcid_query_count,
                                       gen_auq_found, gen_auq_count,
                                       update_declaration, update_funder,
                                       orcid_list)

    def _insert_query_callouts_saxonc(self, docxfile,
                                      oa_query_found, oa_query_count,
                                      funding_query_found, funding_query_count,
                                      orcid_query_found, orcid_query_count,
                                      gen_auq_found, gen_auq_count,
                                      update_declaration, update_funder,
                                      orcid_list):
        """Insert query callouts using Python saxonc module."""
        with zipfile.ZipFile(docxfile, 'r') as zip_ref:
            zip_ref.extract('word/document.xml', 'temp')

        from TransformXml import _run_xslt, _get_proc
        proc = _get_proc()

        # Build parameters dict — _run_xslt creates a fresh Xslt30Processor
        # and calls exception_clear() afterwards, preventing the stale-pointer
        # EXCEPTION_ACCESS_VIOLATION (0xC0000005) seen in libsaxonhec.dll
        params = {}
        if oa_query_found:
            params['oa_query_found']    = proc.make_boolean_value(True)
            params['oa_query_count']    = proc.make_integer_value(oa_query_count)
        if funding_query_found:
            params['funding_query_found']  = proc.make_boolean_value(True)
            params['funding_query_count']  = proc.make_integer_value(funding_query_count)
        if gen_auq_found:
            params['gen_auq_found']  = proc.make_boolean_value(True)
            params['gen_auq_count']  = proc.make_integer_value(gen_auq_count)
        if update_declaration:
            params['update_declaration']  = proc.make_boolean_value(True)
            params['funding_query_count'] = proc.make_integer_value(funding_query_count)
        if update_funder:
            params['update_funder']       = proc.make_boolean_value(True)
            params['funding_query_count'] = proc.make_integer_value(funding_query_count)
        if orcid_query_found:
            params['orcid_query_found'] = proc.make_boolean_value(True)
            params['orcid_query_count'] = proc.make_integer_value(orcid_query_count)
        params['myList'] = proc.make_string_value(",".join(orcid_list))

        if not os.path.exists("D:/mProjects/test"):
            os.mkdir("D:/mProjects/test")
        if os.path.exists('D:/mProjects/test/document.xml'):
            os.remove('D:/mProjects/test/document.xml')
        result_file = 'D:/mProjects/test/document.xml'

        _run_xslt(
            stylesheet_file='xsl/query_callouts.xsl',
            source_file='temp/word/document.xml',
            output_file=result_file,
            parameters=params,
        )

        os.remove('temp/word/document.xml')
        shutil.move(result_file, "temp/word/document.xml")
        with zipfile.ZipFile(docxfile, 'r') as zip_read:
            with zipfile.ZipFile(docxfile + '.tmp', 'w') as zip_write:
                for item in zip_read.infolist():
                    if item.filename != 'word/document.xml':
                        buffer = zip_read.read(item.filename)
                        zip_write.writestr(item, buffer)
        with zipfile.ZipFile(docxfile + '.tmp', 'a') as zip_ref:
            zip_ref.write('temp/word/document.xml', 'word/document.xml')
        os.remove('temp/word/document.xml')
        os.replace(docxfile + '.tmp', docxfile)

    def insert_query_callouts_jar(self, docxfile,
                                  oa_query_found, oa_query_count,
                                  funding_query_found, funding_query_count,
                                  orcid_query_found, orcid_query_count,
                                  gen_auq_found, gen_auq_count,
                                  update_declaration, update_funder,
                                  orcid_list):
        """Insert query callouts using Java Saxon jar."""
        with zipfile.ZipFile(docxfile, 'r') as zip_ref:
            zip_ref.extract('word/document.xml', 'temp')

        params = [
            'oa_query_found=' + str(oa_query_found).lower(),
            'oa_query_count=' + str(oa_query_count),
            'funding_query_found=' + str(funding_query_found).lower(),
            'funding_query_count=' + str(funding_query_count),
            'orcid_query_found=' + str(orcid_query_found).lower(),
            'orcid_query_count=' + str(orcid_query_count),
            'gen_auq_found=' + str(gen_auq_found).lower(),
            'gen_auq_count=' + str(gen_auq_count),
            'update_declaration=' + str(update_declaration).lower(),
            'update_funder=' + str(update_funder).lower(),
            'myList=' + ','.join(orcid_list)
        ]

        if not os.path.exists("D:/mProjects/test"):
            os.mkdir("D:/mProjects/test")
        if os.path.exists('D:/mProjects/test/document.xml'):
            os.remove('D:/mProjects/test/document.xml')

        result_file = 'D:/mProjects/test/document.xml'

        try:
            command = [
                'java', '-jar', "ParaStyler/saxon9pe.jar",
                '-s:temp/word/document.xml',
                '-xsl:xsl/query_callouts.xsl',
                '-o:' + result_file,
            ]

            for param in params:
                command.append(param)

            result = subprocess.run(command, capture_output=True, text=True)
            if result.returncode != 0:
                print(result.stderr)
                return

            os.remove('temp/word/document.xml')
            shutil.move(result_file, "temp/word/document.xml")
            with zipfile.ZipFile(docxfile, 'r') as zip_read:
                with zipfile.ZipFile(docxfile + '.tmp', 'w') as zip_write:
                    for item in zip_read.infolist():
                        if item.filename != 'word/document.xml':
                            buffer = zip_read.read(item.filename)
                            zip_write.writestr(item, buffer)
            with zipfile.ZipFile(docxfile + '.tmp', 'a') as zip_ref:
                zip_ref.write('temp/word/document.xml', 'word/document.xml')
            os.remove('temp/word/document.xml')
            os.replace(docxfile + '.tmp', docxfile)
        except Exception as e:
            print(e)

    # ====================
    # Author Matching
    # ====================

    def find_matched_authors(self, orcid_dic, author_label):
        """Find matches between ORCID dictionary and author labels."""
        matched_authors = []
        unmatched_authors = []

        def is_partial_match(name1, name2):
            return name1 in name2 or name2 in name1

        for orcid_author in orcid_dic.keys():
            matched = False
            for author_name in author_label.keys():
                if re.match(f"^{orcid_author}$", f"{author_name}", re.I):
                    matched_authors.append(orcid_author)
                elif is_partial_match(orcid_author, author_name):
                    matched_authors.append(orcid_author)
                    author_label[f"{orcid_author}"] = author_label[f"{author_name}"]
                    matched = True
                    break
            if not matched:
                unmatched_authors.append(orcid_author)
        return author_label, matched_authors, unmatched_authors

    def transform_caption(self, fig_line):
        """Transform caption XML."""
        pattern = re.compile(
            r'(<w:p[^>]*><w:pPr>.*?<w:pStyle w:val="CP"/>)'
            r'(.*?</w:pPr>)'
            r'(<w:r[^>]*>.*?<w:t>)'
            r'((Figure|Fig\.|Box|Table|Tab\.)( |)([0-9]+|[0-9a-z]+)(\.|\:|\;|)( |))'
            r'(.*?)</w:t></w:r>'
            r'(.*)',
            re.DOTALL
        )

        def replacement(match):
            initial_p = match.group(1)
            ppr_close = match.group(2)
            first_r = match.group(3)
            label = match.group(4)
            rest_content = match.group(8) + match.group(9)
            new_text = (
                f'{initial_p}<w:pStyle w:val="CPB"/>{ppr_close}{first_r}{label}</w:t></w:r>'
                f'<w:p{match.group(2)}><w:pPr><w:pStyle w:val="CP"/></w:pPr>{rest_content}'
            )
            return new_text

        return re.sub(pattern, replacement, fig_line)

    def bd_preclean(self, docx_file_path):
        """Pre-clean document before breakdown processing."""
        with zipfile.ZipFile(docx_file_path, 'r') as zip_ref:
            zip_ref.extract('word/document.xml', 'temp')
        with open("temp/word/document.xml", "r", encoding="utf-8") as file:
            document_xml = file.read()
        document_xml = re.sub(
            r'(<w:p([^>]+|)>(((?!</?w:p>).)*)(<w:pStyle w:val="(CP)"/>)(((?!</?w:p>).)*)<w:r([^>]+|)>(((?!</?w:p>).)*))<w:t([^>]+|)>((Figure|Fig\.|Box|Box\.|Table|Tab\.)( |)([0-9]+|[0-9a-z]+)(\.|\:|\;|))( |)([a-z0-9])',
            r'\1<w:t xml:space="preserve">\13\18</w:t></w:r><w:r><w:t xml:space="preserve">\19',
            document_xml,
            flags=re.I)
        with open("temp/word/document.xml", "w", encoding="utf-8") as file:
            file.write(document_xml)
        with zipfile.ZipFile(docx_file_path, 'r') as zip_read:
            with zipfile.ZipFile(docx_file_path + '.tmp', 'w') as zip_write:
                for item in zip_read.infolist():
                    if item.filename != 'word/document.xml':
                        buffer = zip_read.read(item.filename)
                        zip_write.writestr(item, buffer)
        with zipfile.ZipFile(docx_file_path + '.tmp', 'a') as zip_ref:
            zip_ref.write('temp/word/document.xml', 'word/document.xml')
        os.remove('temp/word/document.xml')
        os.replace(docx_file_path + '.tmp', docx_file_path)

    def label_floats(self, docx_file_path):
        """Label floats in document using XSLT."""
        with zipfile.ZipFile(docx_file_path, 'r') as zip_ref:
            zip_ref.extract('word/document.xml', 'temp')
        xslt_file_path = "xsl/label_floats.xsl"
        subprocess.run(
            ['java', '-jar', 'ParaStyler/saxon9pe.jar', '-s:temp/word/document.xml', '-xsl:' + xslt_file_path,
             '-o:temp/word/document.xml'])
        with zipfile.ZipFile(docx_file_path, 'r') as zip_read:
            with zipfile.ZipFile(docx_file_path + '.tmp', 'w') as zip_write:
                for item in zip_read.infolist():
                    if item.filename != 'word/document.xml':
                        buffer = zip_read.read(item.filename)
                        zip_write.writestr(item, buffer)
        with zipfile.ZipFile(docx_file_path + '.tmp', 'a') as zip_ref:
            zip_ref.write('temp/word/document.xml', 'word/document.xml')
        os.remove('temp/word/document.xml')
        os.replace(docx_file_path + '.tmp', docx_file_path)

    def process_author_line(self, labeled_author):
        """Process labeled author line to extract author dictionary."""
        if not labeled_author:
            return {}

        labeled_author = re.sub(r'<(sup|deg|prefix|suffix|a)>(((?!</?\1>).)*)</\1>', "", labeled_author,
                                flags=re.I | re.S | re.DOTALL)
        labeled_author = re.sub(r'((.+?)(<span([^>]+)>(\[AU(\d+)\])</span>))', r"<au id='\6'>\2</au>", labeled_author,
                                flags=re.I | re.S | re.DOTALL)
        labeled_author = re.sub(r'<au([^>]+|)>(\,|\.|\;| and |and | \&|\&|)(\s+|)', r"<au\1>",
                                labeled_author, re.I | re.S | re.DOTALL)
        labeled_author = re.sub(r'<au([^>]+|)>and ', r"<au\1>", labeled_author, re.I | re.S | re.DOTALL)
        labeled_author = re.sub(r'<au([^>]+|)>\, and ', r"<au\1>", labeled_author, re.I | re.S | re.DOTALL)
        labeled_author = re.sub(r'<au([^>]+|)>\. and ', r"<au\1>", labeled_author, re.I | re.S | re.DOTALL)
        labeled_author = re.sub(r'<au([^>]+|)>\; and ', r"<au\1>", labeled_author, re.I | re.S | re.DOTALL)
        labeled_author = re.sub(r'<au([^>]+|)>\: and ', r"<au\1>", labeled_author, re.I | re.S | re.DOTALL)
        labeled_author = re.sub(r'<au([^>]+|)>\, \& ', r"<au\1>", labeled_author, re.I | re.S | re.DOTALL)
        labeled_author = re.sub(r'<au([^>]+|)>\. \& ', r"<au\1>", labeled_author, re.I | re.S | re.DOTALL)
        labeled_author = re.sub(r'<au([^>]+|)>\; \& ', r"<au\1>", labeled_author, re.I | re.S | re.DOTALL)
        labeled_author = re.sub(r'<au([^>]+|)>\: \& ', r"<au\1>", labeled_author, re.I | re.S | re.DOTALL)

        pattern = re.compile(r"<au id=\'(\d+)\'>(((?!<\/au>).)+)<\/au>")
        matches = pattern.findall(labeled_author)
        author_dict = {}
        for match in matches:
            id = match[0]
            name = match[1]
            author_dict[id] = name.strip()
        return author_dict

    def cross_reference_authors(self, aut_info, aut_line, aut_dict):
        """Cross-reference authors between SMART data and document."""
        dict_count = len(aut_dict)
        if dict_count > 0:
            aut_labeled = aut_dict
        else:
            # Handle case where aut_line might be None or empty
            if aut_line:
                aut_labeled = self.process_author_line(aut_line)
            else:
                aut_labeled = {}

        def format_name(first_name='', middle_name='', last_name=''):
            parts = [first_name, middle_name, last_name]
            return ' '.join(p for p in parts if p)

        name_to_id = {}
        for k, name in aut_labeled.items():
            name_parts = name.split(maxsplit=2)
            formatted_name = format_name(*name_parts)
            name_to_id[formatted_name] = k

        def find_best_match(name, name_to_id):
            if not name_to_id:  # If dictionary is empty
                return None
            result = process.extractOne(name, name_to_id.keys(), scorer=fuzz.token_sort_ratio)
            if result is None:  # Handle case where no match found
                return None
            match, score = result
            return name_to_id.get(match) if score >= 80 else None

        updated_dict2 = {}

        for old_id, info in aut_info.items():
            full_name = format_name(info.get('first-name', ''), info.get('middle-name', ''), info.get('last-name', ''))
            new_id = find_best_match(full_name, name_to_id)
            if new_id:
                updated_dict2[new_id] = info
            else:
                updated_dict2[old_id] = info
        return updated_dict2, aut_labeled


# ====================
# Legacy Compatibility
# ====================

class BreakDownProcess(BreakDownProcessor):
    """Legacy class name for backward compatibility."""

    def __init__(self):
        super().__init__()

    # Map all old method names to new ones
    def insert_para_after(self, paragraph, text, style=None):
        return self.insert_paragraph_after(paragraph, text, style)

    def insert_para_before(self, paragraph, text, pstyle):
        return self.insert_paragraph_before(paragraph, text, pstyle)

    def ins_para_before(self, paragraph, text, style=None):
        return self.insert_paragraph_before(paragraph, text, style)

    def check_delaration(self, paragraphs):
        return self.check_declaration_section(paragraphs)

    def check_supplementary(self, paragraphs):
        return self.check_supplementary_material(paragraphs)

    def check_corresp(self, authors):
        return self.check_corresponding_author(authors)

    def get_funder_para(self, paragraphs):
        return self.check_funding_section(paragraphs)

    def check_funder_para(self, paragraphs):
        return self.check_funder_paragraph(paragraphs)

    def check_query_count(self, query_content, query_count):
        return self.process_query_count(query_content, query_count)

    def get_keyword_para(self, paragraphs):
        return self.find_keyword_paragraph(paragraphs)

    def get_ref_para(self, paragraphs):
        return self.find_reference_paragraph(paragraphs)

    def get_keyword(self, paragraphs):
        return self.check_keywords_section(paragraphs)

    def get_rrh_authors(self, authors_data):
        return self.get_running_head_authors(authors_data)

    # def get_author_rhead(self, authors_data):
    #     return self.get_author_rhead(authors_data)

    def get_author_rhead(self, authors_data):
        return super().get_author_rhead(authors_data)

    def create_breakdown_docx(self, jid, aid, docxfile):
        return self.create_breakdown_document(jid, aid, docxfile)

    def format_and_add(self, new_para, text, style):
        return self._format_and_add_text(new_para, text, style)

    # def process_date(self, date, date_type, history_dic):
    #     return self.process_date(date, date_type, history_dic)

    def process_date(self, date, date_type, history_dic):
        return super().process_date(date, date_type, history_dic)

    def process_fmbm(self, document, breakdown_dic, seq, para_index):
        return self.process_fmbm_sequence(document, breakdown_dic, seq, para_index)

    def process_abstract(self, document):
        return self.process_abstract_section(document)

    def process_keywords(self, document, seq, para_index):
        return self.process_keywords_section(document, seq, para_index)

    def update_tiltle_cases(self, document):
        return self.update_title_cases(document)

    def extract_authors(self, input_string):
        return self.extract_authors_from_xml(input_string)

    def label_docx_author(self, docx_file_path):
        return self.label_docx_authors(docx_file_path)

    def get_docx_author(self, docx_file_path, json_path=None):
        return self.get_docx_authors(docx_file_path, json_path=json_path)

    def auline_process(self, labeled_author):
        return self.process_author_line(labeled_author)

    def cross_smart_vs_doc(self, aut_info, aut_line, aut_dict):
        return self.cross_reference_authors(aut_info, aut_line, aut_dict)

    def rearrange_fmbm(self, document, fm_seq, bm_seq):
        return self.rearrange_fmbm_sequence(document, fm_seq, bm_seq)

    def map_breakdownStyles(self, document):
        return self.map_breakdown_styles(document)

    def process_date(self, date, date_type, history_dic):
        return super().process_date(date, date_type, history_dic)



# breakdown_process = BreakDownProcess()
# breakdown_process.create_breakdown_docx("TAB", "1412846", "V:\\FOR_BREAKDOWN\\BreakDown_DONE\\SAGE\\TAB_1412846\\TAB_1412846_CLN.docx")
# authors_info = {'1': {'salutation': 'Miss', 'first-name': 'Lingli', 'middle-name': '', 'last-name': 'Qiu', 'suffix': '', 'degrees': 'Qiu', 'orcid': '0009-0003-6010-6114', 'orcid-validated': 'true', 'mailing-affiliation': 'Sichuan University West China School of Public Health', 'address1': 'No. 16, Section 3, South Renmin Road, Wuhou District', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': '610041', 'article-affiliation': 'Sichuan University West China School of Public Health', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': '230966', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': 'lingliq119@163.com', 'alternate-email': 'lingliq119@163.com', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '2': {'salutation': 'Miss', 'first-name': 'Lingli', 'middle-name': '', 'last-name': 'Qiu', 'suffix': '', 'degrees': 'Qiu', 'orcid': '0009-0003-6010-6114', 'orcid-validated': 'true', 'mailing-affiliation': 'Sichuan University West China School of Public Health', 'address1': 'No. 16, Section 3, South Renmin Road, Wuhou District', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': '610041', 'article-affiliation': 'Sichuan University West China School of Public Health', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': '230966', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': 'zwqscu@126.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '3': {'salutation': 'Dr', 'first-name': 'Wenqiang', 'middle-name': '', 'last-name': 'Zhang', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'Sichuan University West China School of Public Health', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': '230966', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': '2097724418@qq.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '4': {'salutation': 'Miss', 'first-name': 'Zhixin', 'middle-name': '', 'last-name': 'Tan', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'Sichuan University West China School of Public Health', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': '230966', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': 'W_UXuan@163.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '5': {'salutation': 'Miss', 'first-name': 'Xuan', 'middle-name': '', 'last-name': 'Wu', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'Sichuan University West China School of Public Health', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': '230966', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': 'wangyutong719@stu.scu.edu.cn', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '6': {'salutation': 'Miss', 'first-name': 'Yutong', 'middle-name': '', 'last-name': 'Wang', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'Sichuan University West China School of Public Health', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': '230966', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': 'mstang13@163.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '7': {'salutation': 'Dr', 'first-name': 'Mingshuang', 'middle-name': '', 'last-name': 'Tang', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'Sichuan University West China School of Public Health', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': '230966', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': None, 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '8': {'salutation': 'Dr', 'first-name': 'Lin', 'middle-name': '', 'last-name': 'Chen', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'Sichuan University West China School of Public Health', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': '230966', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': '358730942@qq.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '9': {'salutation': 'Dr', 'first-name': 'Yunjie', 'middle-name': '', 'last-name': 'Liu', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'Sichuan University West China School of Public Health', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': '230966', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': 'liuyunjie@scu.edu.cn', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '10': {'salutation': 'Dr', 'first-name': 'Yunjie', 'middle-name': '', 'last-name': 'Liu', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'Sichuan University West China School of Public Health', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': '230966', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': None, 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '11': {'salutation': 'Dr', 'first-name': 'Bowen', 'middle-name': '', 'last-name': 'Lei', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'Sichuan University West China School of Public Health', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': '230966', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': '15123766508@163.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '12': {'salutation': 'Dr', 'first-name': 'Xiaofeng', 'middle-name': '', 'last-name': 'Ma', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'Sichuan University West China School of Public Health', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': None, 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': '2327398489@qq.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '13': {'salutation': 'Dr', 'first-name': 'Di', 'middle-name': '', 'last-name': 'Zhang', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'Sichuan University West China School of Public Health', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': None, 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': '1124989091@qq.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '14': {'salutation': 'Dr', 'first-name': 'Wenzhi', 'middle-name': '', 'last-name': 'Wang', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'West China Fourth Hospital Sichuan University', 'address2': '', 'city': 'Chengdu', 'state-province': 'Sichuan', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Ultrasound', 'ringgold-institution': 'West China Fourth Hospital Sichuan University', 'ringgold-id': '618726', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': '24504745@qq.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '15': {'salutation': 'Dr', 'first-name': 'Yiping', 'middle-name': '', 'last-name': 'Jia', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'West China Fourth Hospital Sichuan University', 'address2': None, 'city': 'Chengdu', 'state-province': 'Sichuan', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Clinical Laboratory', 'ringgold-institution': 'West China Fourth Hospital Sichuan University', 'ringgold-id': '618726', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': 'heqiurong2009@163.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '16': {'salutation': 'Dr', 'first-name': 'Qiurong', 'middle-name': '', 'last-name': 'He', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'West China Fourth Hospital Sichuan University', 'address2': '', 'city': 'Chengdu', 'state-province': 'Sichuan', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Osteoporosis', 'ringgold-institution': 'West China Fourth Hospital Sichuan University', 'ringgold-id': '618726', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': 'sunlei_hxsy@163.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '17': {'salutation': 'Dr', 'first-name': 'Lei', 'middle-name': '', 'last-name': 'Sun', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'West China Fourth Hospital Sichuan University', 'address2': '', 'city': 'Chengdu', 'state-province': 'Sichuan', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Osteoporosis', 'ringgold-institution': 'West China Fourth Hospital Sichuan University', 'ringgold-id': '618726', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': '78994168@qq.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '18': {'salutation': 'Dr', 'first-name': 'Lu', 'middle-name': '', 'last-name': 'Wang', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'West China Fourth Hospital Sichuan University', 'address2': '', 'city': 'Chengdu', 'state-province': 'Sichuan', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Osteoporosis', 'ringgold-institution': 'West China Fourth Hospital Sichuan University', 'ringgold-id': '618726', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': '78994168@qq.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '19': {'salutation': 'Dr', 'first-name': 'Jian', 'middle-name': '', 'last-name': 'Xu', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'West China Fourth Hospital Sichuan University', 'address2': '', 'city': 'Chengdu', 'state-province': 'Sichuan', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Osteoporosis', 'ringgold-institution': 'West China Fourth Hospital Sichuan University', 'ringgold-id': '618726', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': 'chenyao66@126.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '20': {'salutation': 'Dr', 'first-name': 'Yao', 'middle-name': '', 'last-name': 'Chen', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'West China Fourth Hospital Sichuan University', 'address2': '', 'city': 'Chengdu', 'state-province': 'Sichuan', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Osteoporosis', 'ringgold-institution': 'West China Fourth Hospital Sichuan University', 'ringgold-id': '618726', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': 'fanmengyu@scu.edu.cn', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '21': {'salutation': 'Dr', 'first-name': 'Mengyu', 'middle-name': '', 'last-name': 'Fan', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'Sichuan University West China School of Public Health', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': None, 'article-affiliation': '', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': '230966', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': 'lijiayuan@scu.edu.cn', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '22': {'salutation': 'Professor', 'first-name': 'Jiayuan', 'middle-name': '', 'last-name': 'Li', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'Sichuan University West China School of Public Health', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': '230966', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': 'wcsphwcfh@vip.163.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '23': {'salutation': 'Professor', 'first-name': 'Ben', 'middle-name': '', 'last-name': 'Zhang', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'Sichuan University West China School of Public Health', 'address2': '', 'city': 'Chengdu', 'state-province': '', 'country': 'China', 'postal-code': '', 'article-affiliation': '', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Sichuan University West China School of Public Health', 'ringgold-id': '618726', 'validated-by-peer-review': True, 'validated-by-user': False, 'independent-researcher': True, 'ringgold-id-not-found': False, 'primary-email': '843416455@qq.com', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': False}, '24': {'salutation': 'Dr', 'first-name': 'Xiaoling', 'middle-name': '', 'last-name': 'Wen', 'suffix': '', 'degrees': '', 'orcid': '', 'orcid-validated': None, 'mailing-affiliation': '', 'address1': 'West China Fourth Hospital Sichuan University', 'address2': None, 'city': 'Stockholm', 'state-province': '', 'country': 'Sweden', 'postal-code': '171 77', 'article-affiliation': '', 'article-dept': 'Department of Epidemiology and Biostatistics', 'ringgold-institution': 'Karolinska Institutet', 'ringgold-id': '27106', 'validated-by-peer-review': False, 'validated-by-user': False, 'independent-researcher': False, 'ringgold-id-not-found': False, 'primary-email': 'xiajiang@scu.edu.cn', 'alternate-email': '', 'comp-copies': '0', 'market-via-email': True, 'agreement-received': False, 'corresponding-author': True}}
# labeled_author = '<p class="authors">Lingli Qiu<sup>1,2</sup><span style="color: blue;">[AU1]</span>, Wenqiang Zhang<sup>1</sup><span style="color: blue;">[AU2]</span>, Zhixin Tan<sup>1</sup><span style="color: blue;">[AU3]</span>, Xuan Wu<sup>1</sup><span style="color: blue;">[AU4]</span>, Yutong Wang<sup>1</sup><span style="color: blue;">[AU5]</span>, Mingshuang Tang<sup>1</sup><span style="color: blue;">[AU6]</span>, Lin Chen<sup>1</sup><span style="color: blue;">[AU7]</span>, Yanqiu Zou<sup>1</sup><span style="color: blue;">[AU8]</span>, Yunjie Liu<sup>1</sup><span style="color: blue;">[AU9]</span>, Bowen Lei<sup>1</sup><span style="color: blue;">[AU10]</span>, Xiaofeng Ma<sup>1</sup><span style="color: blue;">[AU11]</span>, Di Zhang<sup>1</sup><span style="color: blue;">[AU12]</span>, Wenzhi Wang<sup>3</sup><span style="color: blue;">[AU13]</span>, Yiping Jia<sup>4</sup><span style="color: blue;">[AU14]</span>, Qiurong He<sup>5</sup><span style="color: blue;">[AU15]</span>, Lei Sun<sup>3</sup><span style="color: blue;">[AU16]</span>, Lu Wang<sup>3</sup><span style="color: blue;">[AU17]</span>, Jian Xu<sup>3</sup><span style="color: blue;">[AU18]</span>, Yao Chen<sup>3</sup><span style="color: blue;">[AU19]</span>, Mengyu Fan<sup>1</sup><span style="color: blue;">[AU20]</span>, Jiayuan Li<sup>1</sup><span style="color: blue;">[AU21]</span>, Ben Zhang<sup>1,6</sup><span style="color: blue;">[AU22]</span>, Xiaoling Wen<sup>7,*</sup><span style="color: blue;">[AU23]</span>, and Xia Jiang<sup>1,8,9,*</sup><span style="color: blue;">[AU24]</span></p>'
# aut_dict = {1: 'Lingli Qiu', 2: 'Wenqiang Zhang', 3: 'Zhixin Tan', 4: 'Xuan Wu', 5: 'Yutong Wang', 6: 'Mingshuang Tang', 7: 'Lin Chen', 8: 'Yanqiu Zou', 9: 'Yunjie Liu', 10: 'Bowen Lei', 11: 'Xiaofeng Ma', 12: 'Di Zhang', 13: 'Wenzhi Wang', 14: 'Yiping Jia', 15: 'Qiurong He', 16: 'Lei Sun', 17: 'Lu Wang', 18: 'Jian Xu', 19: 'Yao Chen', 20: 'Mengyu Fan', 21: 'Jiayuan Li', 22: 'Ben Zhang', 23: 'Xiaoling Wen', 24: 'Xia Jiang'}
# new_authors_info, aut_labeled = breakdown_process.cross_reference_authors(authors_info, labeled_author, aut_dict)
# # # breakdown_process.get_docx_authors("V:\\FOR_BREAKDOWN\\BreakDown_DONE\\SAGE\\TAB_1412846\\docs\\TAB_1412846_CLN_AS.docx")
# breakdown_process.create_breakdown_docx("TAB", "1412846", "V:\\FOR_BREAKDOWN\\BreakDown_DONE\\SAGE\\TAB_1412846\\TAB_1412846_CLN_AS.docx")
# # # breakdown_process.create_breakdown_docx("HPQ", "1278169", "V:\\FOR_BREAKDOWN\\BreakDown_DONE\\SAGE\\HPQ_1278169\\HPQ_1278169_CLN.docx")
