import re
import yaml
import glob
import os
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6 import QtCore, QtGui, QtWidgets
from PyQt6.QtWidgets import QMessageBox, QApplication, QMainWindow, QPushButton, QLabel, QVBoxLayout, QWidget, QDialog, QTextBrowser, QDialogButtonBox
import sys
import getAppPath
from mMerger import DocxMerger
import shutil
from docx2pdf import convert
from PyQt6.QtCore import QUrl
from PyQt6.QtWebEngineWidgets import QWebEngineView
from subprocess import *
from applyStyles import ApplyStyles
from breakDownProcess import BreakDownProcess
import win32com.client as win32
from dbprocess import DataBase
from docx import Document
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from openDocFile_new import OpenDocFile
from CreateArticleInfo import GetArticleId
from loadconfig import getconfig
import subprocess
from docxManipulator import DocxManipulator
from TransformXml import XmlTransform
from createQuoteInfo import CreateParaInfo
from docx.oxml import parse_xml
from lxml import etree
from italic_bookmark import ItalicBookmarkProcessor
import pythoncom
from com_manager import COMManager
from version import __version__


# ---------------------------------------------------------------------------
#  MarkdownViewer  —  lightweight dialog to display .md files
# ---------------------------------------------------------------------------
class MarkdownViewer(QDialog):
    """Renders a Markdown file as formatted HTML inside a scrollable dialog."""

    # Minimal MD→HTML conversion (no extra dependencies)
    @staticmethod
    def _md_to_html(text: str) -> str:
        import re as _re
        lines = text.split("\n")
        html_lines = []
        in_table = False
        in_code = False
        skip_separator = False

        for line in lines:
            # Fenced code blocks
            if line.strip().startswith("```"):
                if in_code:
                    html_lines.append("</code></pre>")
                    in_code = False
                else:
                    html_lines.append("<pre><code>")
                    in_code = True
                continue
            if in_code:
                html_lines.append(line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
                continue

            # Table rows
            if line.strip().startswith("|") and line.strip().endswith("|"):
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                if not in_table:
                    html_lines.append('<table>')
                    tag = "th"
                    in_table = True
                    skip_separator = True
                elif skip_separator and _re.match(r'^[\s|:\-]+$', line):
                    skip_separator = False
                    continue
                else:
                    tag = "td"
                row = "".join(f"<{tag}>{c}</{tag}>" for c in cells)
                html_lines.append(f"<tr>{row}</tr>")
                continue
            elif in_table:
                html_lines.append("</table>")
                in_table = False
                skip_separator = False

            # Headings
            if line.startswith("#### "):
                html_lines.append(f"<h4>{line[5:]}</h4>"); continue
            if line.startswith("### "):
                html_lines.append(f"<h3>{line[4:]}</h3>"); continue
            if line.startswith("## "):
                html_lines.append(f"<h2>{line[3:]}</h2>"); continue
            if line.startswith("# "):
                html_lines.append(f"<h1>{line[2:]}</h1>"); continue

            # HR
            if _re.match(r'^---+$', line.strip()):
                html_lines.append("<hr/>"); continue

            # List items
            if _re.match(r'^\s*[-*+] ', line):
                content = line.strip()[2:]
                html_lines.append(f"<li>{content}</li>"); continue
            if _re.match(r'^\s*\d+\. ', line):
                content = _re.sub(r'^\s*\d+\. ', '', line)
                html_lines.append(f"<li>{content}</li>"); continue

            # Blank line
            if line.strip() == "":
                html_lines.append('<div style="margin:3px 0"></div>'); continue

            html_lines.append(f"<p>{line}</p>")

        if in_table:
            html_lines.append("</table>")

        body = "\n".join(html_lines)
        # Inline: bold, italic, code, links
        body = _re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', body)
        body = _re.sub(r'\*(.+?)\*',     r'<em>\1</em>',         body)
        body = _re.sub(r'`(.+?)`',       r'<code>\1</code>',     body)
        body = _re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2">\1</a>', body)
        body = _re.sub(r'\\([\\`*_{}()\[\]#+\-.!|])', r'\1', body)   # unescape

        return f"""<!DOCTYPE html><html><head><meta charset="utf-8"/><style>
* {{ box-sizing: border-box; }}
body  {{ font-family: Calibri, sans-serif; font-size: 13px;
         background:#2E3440; color:#ECEFF4; margin:12px 16px; line-height:1.5; }}
h1   {{ font-size:18px; color:#88C0D0; font-weight:700;
         margin:0 0 6px 0; padding-bottom:6px;
         border-bottom:2px solid #5E81AC; letter-spacing:0.3px; }}
h2   {{ font-size:14px; color:#88C0D0; font-weight:700;
         margin:14px 0 4px 0; padding:5px 10px;
         background:#3B4252; border-left:3px solid #5E81AC;
         border-radius:0 4px 4px 0; }}
h3   {{ font-size:12px; color:#ECEFF4; font-weight:700;
         margin:8px 0 3px 0; padding:2px 8px;
         display:inline-block; border-radius:3px;
         background:#4C566A; letter-spacing:0.5px; text-transform:uppercase; }}
h4   {{ font-size:12px; color:#81A1C1; margin:6px 0 2px 0; }}
hr   {{ border:none; border-top:1px solid #434C5E; margin:10px 0; }}
code {{ background:#3B4252; padding:1px 5px; border-radius:3px;
        font-family:Consolas,monospace; font-size:11px; color:#A3BE8C; }}
pre  {{ background:#3B4252; padding:10px; border-radius:4px; overflow-x:auto; margin:6px 0; }}
pre code {{ background:none; padding:0; }}
table {{ border-collapse:collapse; width:auto; margin:6px 0 10px 0; }}
th   {{ background:#4C566A; color:#ECEFF4; padding:5px 14px;
        border:1px solid #5E81AC; text-align:left; font-size:12px; }}
td   {{ padding:4px 14px; border:1px solid #3B4252; color:#ECEFF4; font-size:12px; }}
td:first-child {{ color:#81A1C1; font-weight:700; white-space:nowrap; }}
tr:nth-child(even) td {{ background:#3B4252; }}
a    {{ color:#88C0D0; text-decoration:none; }}
a:hover {{ text-decoration:underline; }}
li   {{ margin:2px 0; padding-left:2px; }}
p    {{ margin:3px 0; }}
br   {{ display:block; content:""; margin:2px 0; }}
</style></head><body>{body}</body></html>"""

    def __init__(self, md_path: str, title: str, parent=None):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.resize(780, 560)
        self.setStyleSheet("QDialog { background-color: #2E3440; }")

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 8)
        layout.setSpacing(0)

        self.browser = QTextBrowser()
        self.browser.setOpenExternalLinks(True)
        self.browser.setStyleSheet("""
            QTextBrowser {
                background-color: #2E3440;
                border: none;
            }
            QScrollBar:vertical {
                background: #3B4252; width: 10px; border-radius: 5px;
            }
            QScrollBar::handle:vertical {
                background: #5E81AC; border-radius: 5px; min-height: 20px;
            }
        """)
        layout.addWidget(self.browser)

        btn_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        btn_box.setStyleSheet("""
            QPushButton {
                background-color: #5E81AC; color: #ECEFF4;
                font-weight: bold; border: none; border-radius: 3px;
                padding: 5px 18px;
            }
            QPushButton:hover { background-color: #88C0D0; color: #2E3440; }
        """)
        btn_box.rejected.connect(self.accept)
        layout.addWidget(btn_box, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)

        try:
            with open(md_path, "r", encoding="utf-8") as f:
                content = f.read()
            self.browser.setHtml(self._md_to_html(content))
        except FileNotFoundError:
            self.browser.setHtml(
                f"<body style='background:#2E3440;color:#BF616A;font-family:Calibri;padding:20px'>"
                f"<h3>File not found</h3><p>{md_path}</p></body>"
            )



# ---------------------------------------------------------------------------
#  CreateArticleInfoWorker  — fetches article info from SMART & creates JSON
# ---------------------------------------------------------------------------
class CreateArticleInfoWorker(QThread):
    """
    Runs GetArticleId.smart_login() off the main thread so the UI stays
    responsive while the browser automation runs.

    Signals
    -------
    progress(int)       – 0-100 progress-bar value
    status(str)         – one-line status text
    finished(bool, str) – (success, json_path_or_error_message)
    """
    progress = pyqtSignal(int)
    status   = pyqtSignal(str)
    finished = pyqtSignal(bool, str)

    def __init__(self, art_id: str, jrn_id: str, file_folder: str):
        super().__init__()
        self.art_id      = art_id
        self.jrn_id      = jrn_id
        self.file_folder = file_folder

    # ------------------------------------------------------------------
    # Internet connectivity check
    # ------------------------------------------------------------------
    @staticmethod
    def _check_internet(host: str = "journals.sageapps.com", port: int = 443, timeout: int = 5) -> bool:
        """Return True if the SMART host is reachable, False otherwise."""
        import socket
        try:
            socket.setdefaulttimeout(timeout)
            with socket.create_connection((host, port)):
                return True
        except (socket.timeout, OSError):
            return False

    def run(self):
        import pythoncom
        pythoncom.CoInitialize()
        try:
            self._execute()
        except Exception as exc:
            self.finished.emit(False, f"ERROR: {exc}")
        finally:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    def _execute(self):
        self.status.emit("Checking internet connectivity...")
        self.progress.emit(5)

        if not self._check_internet():
            self.finished.emit(
                False,
                "ERROR: No internet connection.\n\n"
                "Unable to reach the SMART portal (journals.sageapps.com).\n"
                "Please check your network and try again."
            )
            return

        self.status.emit("Connecting to SMART portal...")
        self.progress.emit(10)

        from CreateArticleInfo import GetArticleId

        self.status.emit("Logging in to SMART...")
        self.progress.emit(25)

        creator = GetArticleId()

        self.status.emit("Searching article in SMART...")
        self.progress.emit(50)

        info_found, article_id, jrn_tla = creator.smart_login(
            self.art_id, None, self.jrn_id, self.file_folder
        )

        self.progress.emit(85)

        if not info_found:
            self.finished.emit(
                False,
                f"Article not found in SMART for ID: {self.art_id}"
            )
            return

        json_name = self.jrn_id + "_" + self.art_id + ".json"
        json_path = os.path.join(self.file_folder, json_name)
        if not os.path.exists(json_path):
            self.finished.emit(
                False,
                f"JSON file was not created at expected path:\n{json_path}"
            )
            return

        self.status.emit("JSON created successfully.")
        self.progress.emit(100)
        self.finished.emit(True, json_path)


# ---------------------------------------------------------------------------
#  BreakDownWorker  —  runs mBreakDown pipeline off the main thread
# ---------------------------------------------------------------------------
class BreakDownWorker(QThread):
    """
    Runs the mBreakDown pipeline off the main thread so the UI stays
    responsive and a progress bar can be shown.

    Signals
    -------
    progress(int)       – 0-100 progress-bar value
    status(str)         – one-line status label text
    finished(bool, str) – (close_app, message)
    error(str)          – fatal / unexpected error
    """
    progress = pyqtSignal(int)
    status   = pyqtSignal(str)
    finished = pyqtSignal(bool, str)
    error    = pyqtSignal(str)

    def __init__(self, params: dict):
        super().__init__()
        self.params = params

    # ------------------------------------------------------------------
    # Internet connectivity check (re-used from CreateArticleInfoWorker)
    # ------------------------------------------------------------------
    @staticmethod
    def _check_internet(host: str = "journals.sageapps.com", port: int = 443, timeout: int = 5) -> bool:
        import socket
        try:
            socket.setdefaulttimeout(timeout)
            with socket.create_connection((host, port)):
                return True
        except (socket.timeout, OSError):
            return False

    def run(self):
        import pythoncom
        pythoncom.CoInitialize()
        try:
            self._execute()
        except Exception as exc:
            self.error.emit(str(exc))
        finally:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    def _execute(self):
        p                = self.params
        currentItem      = p["currentItem"]
        article_input    = p["article_input"]
        article_done     = p["article_done"]
        file_full_path   = p["file_full_path"]
        back_up_full_path= p["back_up_full_path"]
        done_file_path   = p["done_file_path"]
        jid              = p["jid"]
        aid              = p["aid"]
        json_path        = p["json_path"]
        uniq_id          = p["uniq_id"]

        import shutil, os
        from breakDownProcess import BreakDownProcess
        from CreateArticleInfo import GetArticleId
        from dbprocess import DataBase

        mydb = DataBase()

        # ── Step 1: Backup ──────────────────────────────────────────────
        self.status.emit("Backing up file...")
        self.progress.emit(5)
        doc_folder = os.path.join(article_input, "docs")
        if not os.path.exists(doc_folder):
            os.mkdir(doc_folder)
        shutil.copy(file_full_path, back_up_full_path)

        # ── Step 2: Fetch JSON if missing ──────────────────────────────
        if not os.path.exists(json_path):
            self.status.emit("Checking internet connectivity...")
            self.progress.emit(10)
            if not self._check_internet():
                self.finished.emit(
                    False,
                    "ERROR: No internet connection.\n\n"
                    "Unable to reach the SMART portal (journals.sageapps.com).\n"
                    "Please check your network and try again."
                )
                return
            self.status.emit("JSON not found — fetching from SMART portal...")
            self.progress.emit(20)
            creat_info = GetArticleId()
            info_generated, article_id, jrn_tla = creat_info.smart_login(
                aid, None, jid, article_input
            )
            self.progress.emit(45)
        else:
            self.progress.emit(20)

        # ── Step 3: Run BreakDown ───────────────────────────────────────
        self.status.emit("Running BreakDown process...")
        self.progress.emit(55)

        # Validate JID exists in BreakDown.json before calling create_breakdown_docx.
        # A missing JID causes a bare KeyError that surfaces as "Unexpected error: 'JRS'"
        # in the UI.  Catch it here with a clear, actionable message.
        try:
            from loadconfig import getconfig as _getconfig
            import yaml as _yaml
            _cfg_path, _ = _getconfig()
            _jrn_json = os.path.join(_cfg_path, "SupportingFiles", "BreakDown.json")
            if os.path.exists(_jrn_json):
                with open(_jrn_json, "r", encoding="utf-8") as _jf:
                    _jrn_data = _yaml.safe_load(_jf)
                if jid not in _jrn_data.get("journal_details", {}):
                    self.finished.emit(
                        False,
                        f"ERROR: Journal ID '{jid}' details not found in BreakDown.json.\n\n"
                        f"Please update BreakDown.xls with the '{jid}' journal details\n"
                        "and run the journal info creator before proceeding."
                    )
                    return
        except Exception as _jid_chk_err:
            # Non-fatal — let create_breakdown_docx surface its own error if needed
            print(f"[WARN] JID pre-check failed: {_jid_chk_err}")

        breakdown_process = BreakDownProcess()
        breakdown_process.create_breakdown_docx(jid, aid, file_full_path)
        self.progress.emit(80)

        # ── Step 4: Move folder to Done ────────────────────────────────
        self.status.emit("Moving to Done folder...")
        self.progress.emit(85)
        process_completed = True
        if os.path.exists(article_done):
            try:
                shutil.rmtree(article_done)
            except Exception:
                process_completed = False

        if process_completed:
            try:
                shutil.move(article_input, article_done)
            except Exception:
                process_completed = False

        self.progress.emit(95)

        if not process_completed:
            self.finished.emit(
                False,
                f"ERROR: BreakDown completed but unable to move folder.\n"
                f"Please move manually:\n  {article_input}\n→ {article_done}"
            )
            return

        # ── Step 5: Update DB ───────────────────────────────────────────
        mydb.update_db(uniq_id, "mBreakDown", "COMPLETED", "", "")
        self.progress.emit(100)
        self.status.emit("BreakDown completed.")
        self.finished.emit(True, f"{currentItem} BreakDown Process Completed|{done_file_path}")


class ParaStylerWorker(QThread):
    """
    Runs the mParaStyler (and optionally mBreakDown) pipeline off the main thread
    so the UI stays responsive.

    Signals
    -------
    progress(int)        – 0-100 progress bar value
    status(str)          – one-line status label text
    finished(bool, str)  – (close_app, message)  delivered on the UI thread
    error(str)           – fatal error message
    json_missing(dict)   – emitted when <jrn>_<aid>.json is absent;
                           payload: {"json_name", "json_source", "inbuild"}
    """
    progress     = pyqtSignal(int)
    status       = pyqtSignal(str)
    finished     = pyqtSignal(bool, str)
    error        = pyqtSignal(str)
    json_missing = pyqtSignal(dict)   # NEW — fires before watcher step

    def __init__(self, params: dict):
        super().__init__()
        self.params = params          # everything the worker needs, no Qt widget refs
        import threading
        self._json_response_event = threading.Event()
        self._json_response = None    # "yes" | "no"  — set by the UI

    # ------------------------------------------------------------------
    # Called from the UI thread to resume the worker after Yes / No
    # ------------------------------------------------------------------
    def provide_json_response(self, answer: str):
        """answer must be 'yes' or 'no'."""
        self._json_response = answer
        self._json_response_event.set()

    # ------------------------------------------------------------------
    def run(self):
        import pythoncom
        pythoncom.CoInitialize()
        try:
            self._execute()
        except Exception as exc:
            self.error.emit(str(exc))
        finally:
            pythoncom.CoUninitialize()

    # ------------------------------------------------------------------
    def _execute(self):
        p          = self.params
        file_full_path   = p["file_full_path"]
        file_folder      = p["file_folder"]
        file_name        = p["file_name"]
        as_file_expected = p["as_file_expected"]
        breakdown_folder = p["breakdown_folder"]
        breakdown_done   = p["breakdown_done"]
        article_folder   = p["article_folder"]
        currentItem      = p["currentItem"]
        jrn_id           = p["jrn_id"]
        art_id           = p["art_id"]
        uniq_id          = p["uniq_id"]
        run_breakdown    = p["run_breakdown"]
        configFolder     = p["configFolder"]
        app_path         = p["app_path"]
        list_macros      = p["list_macros"]

        from italic_bookmark    import ItalicBookmarkProcessor
        from applyStyles        import ApplyStyles
        from openDocFile_new    import OpenDocFile
        from dbprocess          import DataBase
        from TransformXml       import XmlTransform
        from breakDownProcess   import BreakDownProcess
        from CreateArticleInfo  import GetArticleId
        from docx               import Document
        from docx.enum.text     import WD_COLOR_INDEX
        from docx.oxml          import parse_xml
        from lxml               import etree
        from subprocess         import Popen, PIPE
        import re, shutil, os, time

        mydb        = DataBase()
        process_doc = OpenDocFile()

        # ── Step 1: apply_bookmark ──────────────────────────────────────
        self.status.emit("Applying bookmarks...")
        self.progress.emit(5)
        ItalicBookmarkProcessor(file_full_path, mode="apply_bookmark").process()

        # ── Step 2: Run Para Styler (mode decided by watcher_config.yml) ──
        self.status.emit("Reading Para Styler configuration...")
        self.progress.emit(10)

        # Always load watcher_config.yml — it holds the mode flag and
        # (when inbuild_parastyler: false) the hot-folder paths too.
        watcher_config_path = os.path.join(app_path, "watcher_config.yml")
        if not os.path.exists(watcher_config_path):
            self.finished.emit(False, "ERROR: watcher_config.yml not found at: " + watcher_config_path)
            return

        import yaml as _yaml
        with open(watcher_config_path, "r") as _wf:
            watcher_cfg = _yaml.safe_load(_wf)

        # inbuild_parastyler: true  → Direct/JAR mode  (high-spec / outside-network)
        # inbuild_parastyler: false → Watcher hot-folder mode (low-RAM / on-network)
        # Default is false (watcher) so existing installs are not affected.
        inbuild_parastyler = bool(watcher_cfg.get("inbuild_parastyler", False))

        # Common derived values used in both branches
        docx_basename    = os.path.basename(file_full_path)
        prefix           = docx_basename[:-5] if docx_basename.lower().endswith(".docx") else docx_basename
        if prefix.upper().endswith("_CLN"):
            prefix = prefix[:-4]
        expected_as_name = prefix + "_CLN_AS.docx"
        as_file          = os.path.join(file_folder, expected_as_name)

        # ── JSON pre-check (inbuild mode) ──────────────────────────────
        # The JSON is needed later by BreakDown.  Ask the user now
        # rather than silently failing mid-pipeline.
        if inbuild_parastyler:
            _json_name   = jrn_id + "_" + art_id + ".json"
            _json_source = os.path.join(file_folder, _json_name)
            if not os.path.exists(_json_source):
                self._json_response_event.clear()
                self._json_response = None
                self.json_missing.emit({
                    "json_name":   _json_name,
                    "json_source": _json_source,
                    "inbuild":     True,
                })
                self._json_response_event.wait()

                answer = self._json_response
                if answer == "yes":
                    if not os.path.exists(_json_source):
                        self.finished.emit(
                            False,
                            "ERROR: JSON file still not found after creation attempt.\n"
                            + _json_source
                        )
                        return
                else:
                    # "No" in local mode — just untick BreakDown and proceed
                    self.params["_untick_breakdown"] = True
                    run_breakdown = False

        # ── 2A: INBUILD / DIRECT mode ──────────────────────────────────
        if inbuild_parastyler:
            self.status.emit("Running Para Styler (inbuild)... please wait")
            self.progress.emit(12)

            from docxManipulator import DocxManipulator
            docxprocess   = DocxManipulator()
            processresult, as_file = docxprocess.docx_processor(file_full_path)

            # Fallback to run.bat if DocxManipulator could not handle the file
            if processresult is False:
                self.status.emit("Falling back to run.bat...")
                self.progress.emit(20)
                styler_path = os.path.join(configFolder, "ParaStyler\\run.bat")
                p_proc = Popen([styler_path, file_full_path], stdout=PIPE, stderr=PIPE)
                as_file = re.sub(".docx", "_AS.docx", file_full_path)
                p_proc.communicate()
                p_proc.wait()

            self.progress.emit(45)

            if not os.path.exists(as_file):
                self.finished.emit(False, "ERROR: AS file not created. Please proceed manually.")
                return

        # ── 2B: WATCHER / HOT-FOLDER mode ──────────────────────────────
        else:
            self.status.emit("Sending to Para Styler watcher...")
            self.progress.emit(12)

            watcher_input     = watcher_cfg.get("input", "")
            watcher_completed = watcher_cfg.get("completed", "")
            watcher_error     = watcher_cfg.get("error", "")
            watcher_process   = watcher_cfg.get("process", "")

            # ── JSON pre-check (watcher needs the JSON copied to its input) ──
            json_name   = jrn_id + "_" + art_id + ".json"
            json_source = os.path.join(file_folder, json_name)

            if not os.path.exists(json_source):
                # Pause the worker — let the UI show Yes / No dialog
                self._json_response_event.clear()
                self._json_response = None
                self.json_missing.emit({
                    "json_name":   json_name,
                    "json_source": json_source,
                    "inbuild":     False,  # watcher branch
                })
                self._json_response_event.wait()      # blocks until UI responds

                answer = self._json_response          # "yes" | "no"

                if answer == "yes":
                    # UI ran CreateArticleInfoWorker — JSON should now exist
                    if not os.path.exists(json_source):
                        self.finished.emit(
                            False,
                            "ERROR: JSON file still not found after creation attempt.\n"
                            + json_source
                        )
                        return
                    # Fall through — proceed with real JSON in place

                else:
                    # "No" — write a minimal dummy JSON to watcher input so the
                    # watcher does not stall.  Never overwrite a real JSON.
                    dummy_watcher_dest = os.path.join(watcher_input, json_name)
                    if not os.path.exists(dummy_watcher_dest):
                        try:
                            import json as _json
                            _dummy = {
                                "article_info":  {},
                                "journal_info":  {},
                                "authors_info":  {},
                                "funder_info":   {},
                                "_dummy":        True,
                            }
                            with open(dummy_watcher_dest, "w", encoding="utf-8") as _jf:
                                _json.dump(_dummy, _jf)
                            # Remember path so we can clean it up after the
                            # watcher finishes (if it wasn't consumed yet)
                            self.params["_dummy_json_watcher"] = dummy_watcher_dest
                        except Exception:
                            pass
                    # Signal UI to untick "Run with BreakDown"
                    self.params["_untick_breakdown"] = True
                    run_breakdown = False
                    # json_source is still absent — do NOT copy it below;
                    # the dummy copy is already in watcher_input directly
                    json_source = None   # sentinel: skip the copy step

            completed_file = os.path.join(watcher_completed, expected_as_name)
            error_file     = os.path.join(watcher_error, docx_basename)
            error_log_file = os.path.join(watcher_error, docx_basename + ".log")

            # Clean up any stale files from a previous run
            for stale_file in [completed_file, error_file, error_log_file]:
                if os.path.exists(stale_file):
                    try:
                        os.remove(stale_file)
                    except Exception:
                        pass

            # Copy JSON first (if we have a real one), then DOCX
            # (watcher triggers on DOCX arrival)
            self.status.emit("Copying files to watcher input...")
            self.progress.emit(15)
            try:
                if json_source is not None:
                    # Real JSON — copy it to watcher input
                    shutil.copy2(json_source, os.path.join(watcher_input, json_name))
                # else: dummy was written directly to watcher_input — nothing to copy
                shutil.copy2(file_full_path, os.path.join(watcher_input, docx_basename))
            except Exception as copy_ex:
                self.finished.emit(False, "ERROR: Failed to copy files to watcher input:\n" + str(copy_ex))
                return

            # Poll for result — 3-minute timeout
            timeout_seconds = 180
            poll_interval   = 2
            elapsed         = 0
            input_file_path = os.path.join(watcher_input, docx_basename)
            input_drop_time = time.time()

            self.status.emit("Waiting for Para Styler to process...")
            self.progress.emit(20)

            result_found = None   # "completed" | "error" | None (timeout)
            while elapsed < timeout_seconds:
                if os.path.exists(completed_file):
                    result_found = "completed"
                    break
                if os.path.exists(error_file):
                    result_found = "error"
                    break

                # Stall detection: file still in input after 60 s and process folder empty
                if os.path.exists(input_file_path) and (time.time() - input_drop_time) > 60:
                    process_files = []
                    if watcher_process and os.path.isdir(watcher_process):
                        process_files = [
                            f for f in os.listdir(watcher_process)
                            if os.path.isfile(os.path.join(watcher_process, f))
                        ]
                    if not process_files:
                        for _stale in [input_file_path, os.path.join(watcher_input, json_name)]:
                            try:
                                if os.path.exists(_stale):
                                    os.remove(_stale)
                            except Exception:
                                pass
                        # Also remove dummy JSON if we placed one
                        _dummy_path = self.params.get("_dummy_json_watcher")
                        if _dummy_path and os.path.exists(_dummy_path):
                            try:
                                os.remove(_dummy_path)
                            except Exception:
                                pass
                        self.finished.emit(
                            False,
                            "ERROR: Please Check Break Down System, "
                            "ParaStyler seems not running.."
                        )
                        return

                time.sleep(poll_interval)
                elapsed += poll_interval
                pct = min(40, 20 + int(20 * elapsed / timeout_seconds))
                self.progress.emit(pct)
                minutes_left = max(0, (timeout_seconds - elapsed)) // 60
                secs_left    = max(0, (timeout_seconds - elapsed)) % 60
                self.status.emit("Waiting for Para Styler... (%d:%02d remaining)" % (minutes_left, secs_left))

            # Handle poll result
            if result_found == "error":
                error_details = ""
                if os.path.exists(error_log_file):
                    try:
                        with open(error_log_file, "r", encoding="utf-8") as _ef:
                            error_details = _ef.read()
                    except Exception:
                        error_details = "(Could not read error log)"
                try:
                    if os.path.exists(error_file):     os.remove(error_file)
                    if os.path.exists(error_log_file): os.remove(error_log_file)
                except Exception:
                    pass
                self.finished.emit(False,
                    "ERROR: Para Styler failed for " + docx_basename + "\n\n" + error_details)
                return

            if result_found is None:
                self.finished.emit(False,
                    "ERROR: Para Styler timed out (3 minutes) for " + docx_basename
                    + "\nPlease check the watcher process and try again.")
                return

            # result_found == "completed" — wait for file to stabilise before moving
            self.status.emit("Para Styler completed. Copying result back...")
            self.progress.emit(42)
            stable_count = 0
            prev_size    = -1
            for _i in range(30):
                if not os.path.exists(completed_file):
                    break
                cur_size = os.path.getsize(completed_file)
                if cur_size == prev_size and cur_size > 0:
                    stable_count += 1
                    if stable_count >= 3:
                        break
                else:
                    stable_count = 0
                prev_size = cur_size
                time.sleep(1)

            try:
                shutil.move(completed_file, as_file)
            except Exception as move_ex:
                try:
                    shutil.copy2(completed_file, as_file)
                    os.remove(completed_file)
                except Exception as copy_ex2:
                    self.finished.emit(False,
                        "ERROR: Could not move completed file back:\n" + str(copy_ex2))
                    return

            self.progress.emit(45)

            if not os.path.exists(as_file):
                self.finished.emit(False, "ERROR: AS file not created. Please proceed manually.")
                return

        # ── Step 3: apply_italic ────────────────────────────────────────
        self.status.emit("Applying italic styles...")
        self.progress.emit(50)
        ItalicBookmarkProcessor(as_file, mode="apply_italic").process()

        # ── Step 4: post-clean + document fixes ─────────────────────────
        self.status.emit("Running post-clean...")
        self.progress.emit(55)
        applySty = ApplyStyles()
        applySty.as_post_clean(as_file)

        document = Document(as_file)
        # add_math_style inline
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                 "m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
        for prefix, uri in nsmap.items():
            etree.register_namespace(prefix, uri)
        for paragraph in document.paragraphs:
            math_elements = paragraph._element.xpath(".//m:oMathPara")
            if math_elements:
                pg = paragraph._element
                ppr = pg.find(".//w:pPr", namespaces=nsmap)
                if ppr is None:
                    ppr = parse_xml("<w:pPr xmlns:w=\"{}\"/>".format(nsmap["w"]))
                    pg.insert(0, ppr)
                pstyle_elem = ppr.find(".//w:pStyle", namespaces=nsmap)
                if pstyle_elem is not None and pstyle_elem.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") != "displaymath":
                    ppr.remove(pstyle_elem)
                pstyle = parse_xml("<w:pStyle w:val=\"displaymath\" xmlns:w=\"{}\"/>".format(nsmap["w"]))
                ppr.append(pstyle)
        # remove_box inline
        try:
            style = document.styles["abstract"]
            prValue = style._element.pPr
            style._element.remove(prValue)
        except Exception:
            pass
        # articletitle fix
        style_count = 0
        for paragraph in document.paragraphs:
            if paragraph.style.name == "articletitle":
                style_count += 1
                if style_count > 1:
                    paragraph.style = "sectiona"
        if style_count == 0 and document.paragraphs[0].style.name == "sectiona":
            document.paragraphs[0].style = "articletitle"
        # dummy highlight
        for style in document.styles:
            if style.name == "dummy":
                style.font.highlight_color = WD_COLOR_INDEX.YELLOW
                break
        document.save(as_file)

        # ── Step 5: Word macros ─────────────────────────────────────────
        self.status.emit("Running Word macros...")
        self.progress.emit(65)

        _macro_success = False
        _macro_last_err = None
        for _attempt in range(1, 4):          # up to 3 attempts
            try:
                import pythoncom as _pc
                _pc.CoInitialize()
                process_doc.processDocFile(as_file, True, True, True, list_macros)
                _macro_success = True
                break
            except Exception as _macro_err:
                _macro_last_err = _macro_err
                print(f"[WARN] Word macro attempt {_attempt}/3 failed: {_macro_err}")
                try:
                    import subprocess as _sp
                    _sp.call(
                        ["taskkill", "/F", "/IM", "WINWORD.EXE"],
                        stdout=_sp.DEVNULL, stderr=_sp.DEVNULL
                    )
                except Exception:
                    pass
                time.sleep(5 * _attempt)
                try:
                    from openDocFile_new import OpenDocFile as _ODF
                    process_doc = _ODF()
                except Exception:
                    pass

        if not _macro_success:
            self.finished.emit(
                False,
                f"ERROR: Word macros failed after 3 attempts.\n{_macro_last_err}"
            )
            return

        # ── Step 6: apply_styles, check_normal, cross_check ─────────────
        self.status.emit("Applying styles and checks...")
        self.progress.emit(75)
        applyStyles = ApplyStyles()
        applyStyles.apply_styles(as_file)
        if os.path.exists(as_file) and os.path.exists(file_full_path):
            os.remove(file_full_path)
        os.rename(as_file, file_full_path)
        applyStyles.check_normal_styles(file_full_path)
        applyStyles.cross_check_styles(file_full_path)
        XmlTransform().udpate_table_cells(file_full_path)

        # ── Step 7: Move to BreakDown_INPUT ─────────────────────────────
        self.status.emit("Moving files to BreakDown input...")
        self.progress.emit(82)
        breakdown_input = os.path.join(breakdown_folder, article_folder)
        if os.path.exists(breakdown_input):
            try:
                os.remove(breakdown_input)
            except Exception:
                self.finished.emit(False,
                    f"Unable to remove old folder {breakdown_input}\nFolder may be in use. Please move it manually.")
                return
        try:
            shutil.move(file_folder, breakdown_input)
        except Exception as e:
            self.finished.emit(False,
                f"Unable to move folder to {breakdown_input}\nFolder may be in use. Move it manually.")
            return

        mydb.update_db(uniq_id, "mParaStyler", "COMPLETED", "", "")

        # ── Step 8 (optional): mBreakDown ───────────────────────────────
        if not run_breakdown:
            self.progress.emit(100)
            self.status.emit("Para Styler completed.")
            self.finished.emit(False, "Para Styler Completed")
            return

        self.status.emit("Running BreakDown...")
        self.progress.emit(86)

        article_input  = breakdown_input
        article_done   = os.path.join(breakdown_done, article_folder)
        bd_file_path   = os.path.join(article_input, currentItem)

        # backup AS file
        doc_folder = os.path.join(article_input, "docs")
        if not os.path.exists(doc_folder):
            os.mkdir(doc_folder)
        backup_file = re.sub("_CLN", "_CLN_AS", currentItem)
        shutil.copy(bd_file_path, os.path.join(doc_folder, backup_file))

        json_name = jrn_id + "_" + art_id + ".json"
        json_path = os.path.join(article_input, json_name)

        # Validate JID in BreakDown.json before create_breakdown_docx.
        try:
            from loadconfig import getconfig as _getconfig2
            import yaml as _yaml2
            _cfg_path2, _ = _getconfig2()
            _jrn_json2 = os.path.join(_cfg_path2, "SupportingFiles", "BreakDown.json")
            if os.path.exists(_jrn_json2):
                with open(_jrn_json2, "r", encoding="utf-8") as _jf2:
                    _jrn_data2 = _yaml2.safe_load(_jf2)
                if jrn_id not in _jrn_data2.get("journal_details", {}):
                    self.finished.emit(
                        False,
                        f"ERROR: Journal ID '{jrn_id}' details not found in BreakDown.json.\n\n"
                        f"Please update BreakDown.xls with the '{jrn_id}' journal details\n"
                        "and run the journal info creator before proceeding."
                    )
                    return
        except Exception as _jid_chk_err2:
            print(f"[WARN] JID pre-check (ParaStyler) failed: {_jid_chk_err2}")

        breakdown_process = BreakDownProcess()
        if os.path.exists(json_path):
            breakdown_process.create_breakdown_docx(jrn_id, art_id, bd_file_path)
        else:
            creat_info = GetArticleId()
            creat_info.smart_login(art_id, None, jrn_id, article_input)
            breakdown_process.create_breakdown_docx(jrn_id, art_id, bd_file_path)

        self.progress.emit(94)
        self.status.emit("Moving to BreakDown done...")

        if os.path.exists(article_done):
            try:
                shutil.rmtree(article_done)
            except Exception:
                self.finished.emit(False,
                    f"BreakDown done but unable to clear {article_done}. Please move manually.")
                return
        try:
            shutil.move(article_input, article_done)
        except Exception:
            self.finished.emit(False,
                f"BreakDown done but unable to move to {article_done}. Please move manually.")
            return

        mydb.update_db(uniq_id, "mBreakDown", "COMPLETED", "", "")

        done_file_path = os.path.join(article_done, currentItem)
        self.progress.emit(100)
        self.status.emit("BreakDown completed.")
        self.finished.emit(True, f"Para Styler + BreakDown Completed|{done_file_path}")

class ChildWindow(QMainWindow):
    def __init__(self):
        self.processDoc = OpenDocFile()
        self.app_path = getAppPath.getapppath()
        self.configFolder, self.breakDownConfig = getconfig()
        self.db_process = DataBase()
        super(ChildWindow, self).__init__()
        self.setWindowTitle("File Viewer [C&M]")
        self.setGeometry(200, 50, 800, 600)
        self.webView = QWebEngineView()
        self.webView.settings().setAttribute(self.webView.settings().WebAttribute.PluginsEnabled, True)
        self.webView.settings().setAttribute(self.webView.settings().WebAttribute.PdfViewerEnabled, True)
        self.setCentralWidget(self.webView)

    def closeEvent(self, event):
        widgetList = QApplication.topLevelWidgets()
        numWindows = len(widgetList)
        os.remove(pdf_file)
        if numWindows > 1:
            event.accept()
        else:
            event.ignore()

    def url_changed(self):
        self.setWindowTitle(self.webView.title())

    def go_back(self):
        self.webView.back()


class Ui_Dialog(QWidget):
    def __init__(self):
        super(Ui_Dialog, self).__init__()
        self.normalizer_input = None
        self.removed_list = None
        self.selected_list = None
        self.file_names = None
        self.processCombo = None
        self.file_path = None
        self.return_list = ['Select']
        self.input_list = []
        self.listofjournals = []
        self.folder_details = {}
        self.file_details = {}
        self.app_path = getAppPath.getapppath()
        self.configFolder, self.breakDownConfig = getconfig()
        dialog_yml = os.path.join(self.configFolder, 'config\\dialogConfig.yaml')
        with open(dialog_yml, "r") as stream:
            self.dialog_config = yaml.safe_load(stream)
        breakdown_yaml = os.path.join(self.configFolder, 'config\\breakDown.yaml')
        with open(breakdown_yaml, "r") as stream:
            self.breakDownConfig = yaml.safe_load(stream)
        self.mergerInput = self.breakDownConfig['FOLDERS']['MERGER_INPUT']
        self.mergerError = self.breakDownConfig['FOLDERS']['MERGER_ERROR']
        self.ParaStylerInput = self.breakDownConfig['FOLDERS']['ParaStyler_INPUT']
        self.ParaStylerError = self.breakDownConfig['FOLDERS']['ParaStyler_ERROR']
        self.BreakDownInput = self.breakDownConfig['FOLDERS']['BreakDown_INPUT']
        self.BreakDownError = self.breakDownConfig['FOLDERS']['BreakDown_ERROR']
        self.BreakDownDone = self.breakDownConfig['FOLDERS']['BreakDown_DONE']
        self.firstprocess = self.dialog_config['PROCESS'][0]
        self.firstcustomer = self.dialog_config['CUSTOMERS'][0]
        self.jid_aid_list, self.folder_details = self.get_aid_jid_list(self.firstprocess, self.firstcustomer,
                                                                       self.dialog_config)

    def setupUi(self, Dialog):
        Dialog.setObjectName("Dialog")
        Dialog.resize(700, 460)

        # Enable minimize and maximize buttons
        Dialog.setWindowFlags(
            Qt.WindowType.Window |
            Qt.WindowType.WindowMinimizeButtonHint |
            Qt.WindowType.WindowMaximizeButtonHint |
            Qt.WindowType.WindowCloseButtonHint
        )
        Dialog.setMinimumSize(QtCore.QSize(700, 460))

        # Force dark background via palette — covers QMainWindow central widget
        from PyQt6.QtGui import QPalette, QColor
        dark_palette = Dialog.palette()
        dark_palette.setColor(QPalette.ColorRole.Window,      QColor("#2E3440"))
        dark_palette.setColor(QPalette.ColorRole.WindowText,  QColor("#ECEFF4"))
        dark_palette.setColor(QPalette.ColorRole.Base,        QColor("#2E3440"))
        dark_palette.setColor(QPalette.ColorRole.AlternateBase, QColor("#3B4252"))
        dark_palette.setColor(QPalette.ColorRole.Text,        QColor("#ECEFF4"))
        dark_palette.setColor(QPalette.ColorRole.Button,      QColor("#5E81AC"))
        dark_palette.setColor(QPalette.ColorRole.ButtonText,  QColor("#ECEFF4"))
        Dialog.setPalette(dark_palette)
        Dialog.setAutoFillBackground(True)

        # Set the overall stylesheet for the dialog
        Dialog.setStyleSheet("""
            QDialog, QMainWindow, QWidget#Dialog {
                background-color: #2E3440; /* Dark gray background */
                color: #ECEFF4; /* Light gray text */
            }
            QFrame {
                background-color: #3B4252; /* Slightly lighter gray for frames */
                border-radius: 5px;
            }
            QLabel {
                color: #ECEFF4; /* Light gray text */
            }
            QComboBox {
                background-color: #4C566A; /* Dark gray background */
                color: #ECEFF4; /* Light gray text */
                border: 1px solid #5E81AC; /* Blue border */
                border-radius: 3px;
                padding: 5px;
            }
            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 15px;
                border-left-width: 1px;
                border-left-color: #5E81AC; /* Blue border */
                border-left-style: solid;
                border-top-right-radius: 3px;
                border-bottom-right-radius: 3px;
            }
            QComboBox::down-arrow {
                width: 0;
                height: 0;
                border-style: solid;
                border-width: 6px 5px 0 3px; /* Triangle shape */
                border-color: #ECEFF4 transparent transparent transparent; /* White triangle */
                margin-right: 4px; /* Align it neatly within the drop-down */
                margin-left: 6px; /* Align it neatly within the drop-down */
            }
            QComboBox QAbstractItemView {
                background-color: #4C566A; /* Dark gray background */
                color: #ECEFF4; /* Light gray text */
                selection-background-color: #5E81AC; /* Blue selection */
            }
            QPushButton {
                background-color: #5E81AC; /* Blue background */
                color: #ECEFF4; /* Light gray text */
                font-weight: bold;
                border: none;
                border-radius: 3px;
                padding: 5px 10px;
            }
            QPushButton:hover {
                background-color: #88C0D0; /* Lighter blue on hover */
                color: #2E3440; /* Dark text for improved readability */
            }
            QPushButton:pressed {
                background-color: #BF616A; /* Dark gray when pressed */
                color: #ECEFF4; /* Keeping text light for contrast */
            }
            QListWidget {
                background-color: #4C566A; /* Dark gray background */
                color: #ECEFF4; /* Light gray text */
                border: 1px solid #5E81AC; /* Blue border */
                border-radius: 3px;
            }
            QListWidget::item:selected {
                background-color: #FFFFFF;
                color: #000000;
                font-weight: bold;
                }
            QListWidget::item:hover {
                background-color: #434C5E;
                }
            QProgressBar {
                background-color: #4C566A; /* Dark gray background */
                color: #ECEFF4; /* Light gray text */
                border: 1px solid #5E81AC; /* Blue border */
                border-radius: 3px;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #5E81AC; /* Blue progress */
                border-radius: 3px;
            }
            QCheckBox {
                color: #ECEFF4;
                spacing: 6px;
                font-weight: bold;
            }
            QCheckBox::indicator {
                width: 16px;
                height: 16px;
                border: 2px solid #5E81AC;
                border-radius: 3px;
                background-color: #4C566A;
            }
            QCheckBox::indicator:checked {
                background-color: #5E81AC;
                border: 2px solid #88C0D0;
            }
            QCheckBox::indicator:unchecked {
                background-color: #4C566A;
                border: 2px solid #5E81AC;
            }
            QCheckBox::indicator:disabled {
                border: 2px solid #434C5E;
                background-color: #3B4252;
            }
            QCheckBox:disabled {
                color: #4C566A;
                font-weight: normal;
            }
        """)

        font = QtGui.QFont()
        font.setFamily("Calibri")
        font.setPointSize(9)
        font.setWeight(75)
        font1 = QtGui.QFont()
        font1.setFamily("Calibri")
        font1.setPointSize(10)
        font1.setWeight(600)
        Dialog.setFont(font)

        # ── Help & Changelog icon buttons (top-right corner) ────────────
        _icon_style = """
            QPushButton {{
                background-color: {bg};
                color: #ECEFF4;
                font-weight: bold;
                font-size: 11px;
                border: none;
                border-radius: 14px;
                padding: 0px;
            }}
            QPushButton:hover  {{ background-color: #88C0D0; color: #2E3440; }}
            QPushButton:pressed {{ background-color: #BF616A; }}
        """

        self.btnHelp = QtWidgets.QPushButton("?", Dialog)
        self.btnHelp.setGeometry(QtCore.QRect(664, 21, 28, 28))
        self.btnHelp.setToolTip("User Guide")
        self.btnHelp.setStyleSheet(_icon_style.format(bg="#5E81AC"))
        self.btnHelp.setObjectName("btnHelp")
        self.btnHelp.setCursor(QtGui.QCursor(QtCore.Qt.CursorShape.PointingHandCursor))

        self.btnChangelog = QtWidgets.QPushButton("V", Dialog)
        self.btnChangelog.setGeometry(QtCore.QRect(632, 21, 28, 28))
        self.btnChangelog.setToolTip("Changelog")
        self.btnChangelog.setStyleSheet(_icon_style.format(bg="#4C566A"))
        self.btnChangelog.setObjectName("btnChangelog")
        self.btnChangelog.setCursor(QtGui.QCursor(QtCore.Qt.CursorShape.PointingHandCursor))
        # ─────────────────────────────────────────────────────────────────

        self.frame = QtWidgets.QFrame(Dialog)
        self.frame.setGeometry(QtCore.QRect(20, 10, 601, 50))
        self.frame.setFrameShape(QtWidgets.QFrame.Shape.StyledPanel)
        self.frame.setFrameShadow(QtWidgets.QFrame.Shadow.Raised)
        self.frame.setObjectName("frame")

        self.processLabel = QtWidgets.QLabel(self.frame)
        self.processLabel.setGeometry(QtCore.QRect(15, 10, 60, 28))
        self.processLabel.setFont(font)
        self.processLabel.setObjectName("label")

        self.processCombo = QtWidgets.QComboBox(self.frame)
        self.processCombo.setGeometry(QtCore.QRect(75, 10, 125, 28))
        self.processCombo.setObjectName("comboBox")
        self.processCombo.addItems(self.dialog_config['PROCESS'])
        self.processCombo.currentTextChanged.connect(self.process_combobox_changed)

        self.customerCombo = QtWidgets.QComboBox(self.frame)
        self.customerCombo.setGeometry(QtCore.QRect(310, 10, 80, 28))
        self.customerCombo.setObjectName("comboBox_2")
        self.customerCombo.addItems(self.dialog_config['CUSTOMERS'])
        self.customerCombo.currentTextChanged.connect(self.customer_combobox_changed)

        self.jidAidCombo = QtWidgets.QComboBox(self.frame)
        self.jidAidCombo.setGeometry(QtCore.QRect(480, 10, 111, 28))
        self.jidAidCombo.setObjectName("comboBox_3")
        self.jidAidCombo.addItems(self.jid_aid_list)
        self.jidAidCombo.currentTextChanged.connect(self.get_file_list)

        self.customerLabel = QtWidgets.QLabel(self.frame)
        self.customerLabel.setGeometry(QtCore.QRect(235, 10, 65, 28))
        self.customerLabel.setFont(font)
        self.customerLabel.setObjectName("label_2")

        self.jidAidLabel = QtWidgets.QLabel(self.frame)
        self.jidAidLabel.setGeometry(QtCore.QRect(420, 10, 50, 28))
        self.jidAidLabel.setFont(font)
        self.jidAidLabel.setObjectName("label_4")

        self.frame_2 = QtWidgets.QFrame(Dialog)
        self.frame_2.setGeometry(QtCore.QRect(20, 70, 670, 340))
        self.frame_2.setFrameShape(QtWidgets.QFrame.Shape.StyledPanel)
        self.frame_2.setFrameShadow(QtWidgets.QFrame.Shadow.Raised)
        self.frame_2.setObjectName("frame_2")

        self.selectedList = QtWidgets.QListWidget(self.frame_2)
        self.selectedList.setGeometry(QtCore.QRect(20, 30, 530, 145))
        self.selectedList.setObjectName("listWidget")

        self.removeList = QtWidgets.QListWidget(self.frame_2)
        self.removeList.setGeometry(QtCore.QRect(20, 200, 530, 95))
        self.removeList.setObjectName("listWidget_2")

        self.btnIns = QtWidgets.QPushButton(self.frame_2)
        self.btnIns.setGeometry(QtCore.QRect(575, 30, 85, 20))
        self.btnIns.setFont(font)
        self.btnIns.setObjectName("btnIns")

        self.btnOpen = QtWidgets.QPushButton(self.frame_2)
        self.btnOpen.setGeometry(QtCore.QRect(575, 60, 85, 20))
        self.btnOpen.setFont(font)
        self.btnOpen.setObjectName("btnOpen")

        self.btnMoveUp = QtWidgets.QPushButton(self.frame_2)
        self.btnMoveUp.setGeometry(QtCore.QRect(575, 90, 85, 20))
        self.btnMoveUp.setFont(font)
        self.btnMoveUp.setObjectName("btnMoveUp")

        self.btnMoveDown = QtWidgets.QPushButton(self.frame_2)
        self.btnMoveDown.setGeometry(QtCore.QRect(575, 120, 85, 20))
        self.btnMoveDown.setFont(font)
        self.btnMoveDown.setObjectName("btnMoveDown")

        self.btnRemove = QtWidgets.QPushButton(self.frame_2)
        self.btnRemove.setGeometry(QtCore.QRect(575, 150, 85, 20))
        self.btnRemove.setFont(font)
        self.btnRemove.setObjectName("btnRemove")

        self.btnAdd = QtWidgets.QPushButton(self.frame_2)
        self.btnAdd.setGeometry(QtCore.QRect(575, 230, 85, 20))
        self.btnAdd.setFont(font)
        self.btnAdd.setObjectName("btnAdd")

        self.selFilesLabel = QtWidgets.QLabel(self.frame_2)
        self.selFilesLabel.setGeometry(QtCore.QRect(20, 10, 130, 20))
        self.selFilesLabel.setFont(font)
        self.selFilesLabel.setObjectName("label_5")

        self.rmvFilesLabel = QtWidgets.QLabel(self.frame_2)
        self.rmvFilesLabel.setGeometry(QtCore.QRect(20, 180, 110, 20))
        self.rmvFilesLabel.setFont(font)
        self.rmvFilesLabel.setObjectName("label_6")

        # --- Run with BreakDown checkbox (placed below removeList, inside frame_2) ---
        self.chkRunWithBreakDown = QtWidgets.QCheckBox(self.frame_2)
        self.chkRunWithBreakDown.setGeometry(QtCore.QRect(20, 300, 200, 22))
        self.chkRunWithBreakDown.setFont(font)
        self.chkRunWithBreakDown.setObjectName("chkRunWithBreakDown")
        self.chkRunWithBreakDown.setChecked(False)
        self.chkRunWithBreakDown.setEnabled(False)   # enabled only when mParaStyler is selected

        self.btnOK = QtWidgets.QPushButton(Dialog)
        self.btnOK.setGeometry(QtCore.QRect(160, 420, 80, 25))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Fixed, QtWidgets.QSizePolicy.Policy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.btnOK.sizePolicy().hasHeightForWidth())
        self.btnOK.setSizePolicy(sizePolicy)
        self.btnOK.setSizeIncrement(QtCore.QSize(0, 0))
        self.btnOK.setFont(font1)
        self.btnOK.setObjectName("pushButton")

        self.btnCancel = QtWidgets.QPushButton(Dialog)
        self.btnCancel.setGeometry(QtCore.QRect(290, 420, 80, 25))
        self.btnCancel.setFont(font1)
        self.btnCancel.setObjectName("btnCancel")

        self.btnLog = QtWidgets.QPushButton(Dialog)
        self.btnLog.setGeometry(QtCore.QRect(412, 420, 80, 25))
        self.btnLog.setFont(font1)
        self.btnLog.setObjectName("btnLog")

        self.pbar = QtWidgets.QProgressBar(self.frame_2)
        self.pbar.setGeometry(20, 258, 530, 18)
        self.pbar.setHidden(True)
        self.pbar.setTextVisible(True)
        self.pbar.setFormat("%p%")

        self.statusLabel = QtWidgets.QLabel(self.frame_2)
        self.statusLabel.setGeometry(QtCore.QRect(20, 276, 530, 18))
        self.statusLabel.setObjectName("statusLabel")
        self.statusLabel.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.statusLabel.setStyleSheet("color: #88C0D0; font-style: italic;")
        self.statusLabel.setHidden(True)

        self.retranslateUi(Dialog)
        self.update_buttons_status()
        self.connections()
        QtCore.QMetaObject.connectSlotsByName(Dialog)

    def retranslateUi(self, Dialog):
        _translate = QtCore.QCoreApplication.translate
        Dialog.setWindowTitle(_translate("Dialog", "BreakDown [C&M]"))
        breakdown_logo = os.path.join(self.configFolder, 'SupportingFiles', 'BD.ico')
        if os.path.exists(breakdown_logo):
            Dialog.setWindowIcon(QtGui.QIcon(breakdown_logo))
            if sys.platform == "win32":
                import ctypes
                myappid = 'com.companyname.BreakDown'  # Unique ID for your application
                ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
        self.processLabel.setText(_translate("Dialog",
                                             "<html><head/><body><p><span style=\"  font-weight:600; font-size:10pt;\">Process:</span></p></body></html>"))
        self.customerLabel.setText(_translate("Dialog",
                                              "<html><head/><body><p><span style=\" font-weight:600; font-size:10pt;\">Customer:</span></p></body></html>"))
        self.jidAidLabel.setText(_translate("Dialog",
                                            "<html><head/><body><p><span style=\"  font-weight:600; font-size:10pt;\">JID_AID:</span></p></body></html>"))
        self.btnIns.setText(_translate("Dialog", "Instruction(s)"))
        self.btnOpen.setText(_translate("Dialog", "Open"))
        self.btnMoveUp.setText(_translate("Dialog", "Move UP"))
        self.btnMoveDown.setText(_translate("Dialog", "Move Down"))
        self.btnRemove.setText(_translate("Dialog", "Remove"))
        self.btnAdd.setText(_translate("Dialog", "Add"))
        self.selFilesLabel.setText(_translate("Dialog",
                                              "<html><head/><body><p><span style=\" font-weight:600; color:#FFFFFF;\">Selected Files:</span></p></body></html>"))
        self.rmvFilesLabel.setText(_translate("Dialog",
                                              "<html><head/><body><p><span style=\" font-weight:600; color:#FFFFFF;\">Removed Files:</span></p></body></html>"))
        self.chkRunWithBreakDown.setText(_translate("Dialog", "Run with BreakDown"))
        self.btnOK.setText(_translate("Dialog", "OK"))
        self.btnCancel.setText(_translate("Dialog", "Cancel"))
        self.btnLog.setText(_translate("Dialog", "Log"))

    def process_combobox_changed(self):
        self.customerCombo.clear()
        self.selectedList.clear()
        self.removeList.clear()
        self.customerCombo.addItems(self.dialog_config['CUSTOMERS'])
        # Enable checkbox only for mParaStyler; reset to unchecked when switching away
        is_para_styler = self.processCombo.currentText() == "mParaStyler"
        self.chkRunWithBreakDown.setEnabled(is_para_styler)
        if not is_para_styler:
            self.chkRunWithBreakDown.setChecked(False)
            self.chkRunWithBreakDown.setText("Run with BreakDown")
            self.chkRunWithBreakDown.setStyleSheet("color: #4C566A; font-weight: normal;")

    def customer_combobox_changed(self):
        self.jidAidCombo.clear()
        self.selectedList.clear()
        self.removeList.clear()
        self.jid_aid_list = []
        self.jid_aid_list, self.folder_details = self.aid_jid_list(self.processCombo.currentText(),
                                                                   self.customerCombo.currentText(), self.dialog_config)
        self.jidAidCombo.addItems(self.jid_aid_list)
        return self.jid_aid_list

    def aid_jid_list(self, process, customer, config_details):
        input_path = config_details['FOLDERS'][process]['INPUT_PATH']
        input_path = re.sub(r"\[CUSTOMER\]", customer, input_path, re.IGNORECASE)
        self.input_list = glob.glob(input_path + r"\*")
        self.return_list = ['Select']
        self.selectedList.clear()
        self.removeList.clear()
        if len(self.input_list) == 0:
            self.return_list = ['None']
            self.folder_details = {}
        for path_name in self.input_list:
            if os.path.isdir(path_name):
                folder_name = os.path.split(path_name)[1]
                self.folder_details[folder_name] = path_name
                self.return_list.append(folder_name)
        return self.return_list, self.folder_details

    def get_aid_jid_list(self, process, customer, config_details):
        input_path = config_details['FOLDERS'][process]['INPUT_PATH']
        input_path = re.sub(r"\[CUSTOMER\]", customer, input_path, re.IGNORECASE)
        self.input_list = glob.glob(input_path + r"\*")
        self.return_list = ['Select']
        if len(self.input_list) == 0:
            self.return_list = ['None']
            self.folder_details = {}
        for path_name in self.input_list:
            if os.path.isdir(path_name):
                folder_name = os.path.split(path_name)[1]
                self.folder_details[folder_name] = path_name
                self.return_list.append(folder_name)
        return self.return_list, self.folder_details

    def get_file_list(self, value):
        self.file_details = {}
        self.file_names = []
        self.selectedList.clear()
        self.removeList.clear()
        if value == "Select" or value == "" or value == "None":
            self.selectedList.clear()
            self.btnOK.setDisabled(True)
        else:
            self.file_path = self.folder_details[value]
            file_fullpath = glob.glob(self.file_path + u"/*.doc")
            file_fullpath.extend(glob.glob(self.file_path + u"/*.docx"))
            file_fullpath.extend(glob.glob(self.file_path + u"/*.xls"))
            file_fullpath.extend(glob.glob(self.file_path + u"/*.xlsx"))
            for file in file_fullpath:
                file_name = os.path.split(file)[1]
                self.file_details[file_name] = file
                self.file_names.append(file_name)
            if len(self.file_names) > 0:
                self.btnOK.setEnabled(True)
            self.selectedList.addItems(self.file_names)
        return self.file_details, self.file_names

    @QtCore.pyqtSlot()
    def update_buttons_status(self):
        self.btnMoveUp.setDisabled(not bool(self.selectedList.selectedItems()) or self.selectedList.currentRow() == 0)
        self.btnRemove.setDisabled(not bool(self.selectedList.selectedItems()))
        self.btnIns.setDisabled(not bool(self.selectedList.currentRow() == 0))
        self.btnOpen.setDisabled(not bool(self.selectedList.selectedItems()))
        self.btnOpen.setDisabled(not bool(self.selectedList.selectedItems()))
        self.btnAdd.setDisabled(not bool(self.selectedList.selectedItems()))
        self.btnOK.setDisabled(self.selectedList.count() == 0)
        self.btnMoveDown.setDisabled(
            not bool(self.selectedList.selectedItems()) or self.selectedList.currentRow() == (
                        self.selectedList.count() - 1))

    def connections(self):
        self.selectedList.itemSelectionChanged.connect(self.update_buttons_status)
        self.removeList.itemSelectionChanged.connect(self.update_buttons_status)
        self.btnOpen.clicked.connect(self.on_btnOpen_clicked)
        self.btnIns.clicked.connect(self.on_btnInstruction_clicked)
        self.btnMoveUp.clicked.connect(self.on_btnMoveUP_clicked)
        self.btnMoveDown.clicked.connect(self.on_btnMoveDown_clicked)
        self.btnRemove.clicked.connect(self.on_btnRemove_clicked)
        self.btnAdd.clicked.connect(self.on_btnAdd_clicked)
        self.btnCancel.clicked.connect(self.on_btnCancel_clicked)
        self.btnOK.clicked.connect(self.on_btnOK_clicked)
        self.chkRunWithBreakDown.stateChanged.connect(self.on_chkRunWithBreakDown_changed)
        self.btnHelp.clicked.connect(self.on_btnHelp_clicked)
        self.btnChangelog.clicked.connect(self.on_btnChangelog_clicked)

    @QtCore.pyqtSlot()
    def on_btnHelp_clicked(self):
        """Open the User Guide markdown viewer."""
        guide_path = os.path.join(self.app_path, "USER_GUIDE.md")
        viewer = MarkdownViewer(guide_path, f"BreakDown User Guide  —  v{__version__}", self)
        viewer.exec()

    @QtCore.pyqtSlot()
    def on_btnChangelog_clicked(self):
        """Open the Changelog markdown viewer."""
        log_path = os.path.join(self.app_path, "CHANGELOG.md")
        viewer = MarkdownViewer(log_path, f"BreakDown Changelog  —  v{__version__}", self)
        viewer.exec()

    @QtCore.pyqtSlot()
    def on_btnOpen_clicked(self):
        row = self.selectedList.currentRow()
        currentItem = self.selectedList.item(row).text()
        selectedFile = self.file_details[currentItem]
        try:
            os.startfile(selectedFile, 'open')
        except Exception as e:
            QMessageBox.information(self, "mSelect [C&M]",
                                    f"File: {selectedFile}\nUnable to Open file.\n"
                                    f"Please open, Check, and Save manually..")
            os.startfile(selectedFile)

    @QtCore.pyqtSlot()
    def on_btnView_clicked(self):
        row = self.selectedList.currentRow()
        currentItem = self.selectedList.item(row).text()
        selectedFile = self.file_details[currentItem]
        self.open_pdf_file(selectedFile)

    # ------------------------------------------------------------------
    # JID validation helper
    # ------------------------------------------------------------------
    def _validate_jid_in_breakdown_json(self, jid: str) -> tuple:
        """
        Check whether *jid* exists in BreakDown.json → journal_details.

        Returns
        -------
        (True,  "")           – JID found, all good
        (False, error_msg)    – JID missing; caller should show the message
                                and abort the process.
        """
        try:
            configPath, _ = getconfig()
            jrn_json_path = os.path.join(configPath, "SupportingFiles", "BreakDown.json")
            if not os.path.exists(jrn_json_path):
                return False, (
                    f"BreakDown.json not found at:\n{jrn_json_path}\n\n"
                    "Please verify the configuration path and try again."
                )
            with open(jrn_json_path, "r", encoding="utf-8") as f:
                jrn_data = yaml.safe_load(f)
            journal_details = jrn_data.get("journal_details", {})
            if jid not in journal_details:
                return False, (
                    f"Journal ID '{jid}' details not found in BreakDown.json.\n\n"
                    f"Please update BreakDown.xls with the '{jid}' journal details\n"
                    "and run the journal info creator before proceeding."
                )
            return True, ""
        except Exception as exc:
            return False, f"Error reading BreakDown.json:\n{exc}"

    @QtCore.pyqtSlot()
    def on_btnInstruction_clicked(self):
        configPath, breakDownConfig = getconfig()
        jrnJson = os.path.join(configPath, "SupportingFiles/BreakDown.json")
        with open(jrnJson, 'r') as file:
            jsonDetails = yaml.safe_load(file)
        jid_aid = self.jidAidCombo.currentText()
        jidAid = jid_aid.split("_")
        jid = jidAid[0]
        jrnInstructions = jsonDetails['journal_details'][jid]['Instruction']
        jrnIns = jid + ": Instructions\n--------------------\n\n" + jrnInstructions
        self.open_instructions(jrnIns)

    @QtCore.pyqtSlot()
    def on_btnMoveUP_clicked(self):
        row = self.selectedList.currentRow()
        currentItem = self.selectedList.takeItem(row)
        self.selectedList.insertItem(row - 1, currentItem)
        self.selectedList.setCurrentRow(row - 1)

    @QtCore.pyqtSlot()
    def on_btnMoveDown_clicked(self):
        row = self.selectedList.currentRow()
        currentItem = self.selectedList.takeItem(row)
        self.selectedList.insertItem(row + 1, currentItem)
        self.selectedList.setCurrentRow(row + 1)

    @QtCore.pyqtSlot()
    def on_btnRemove_clicked(self):
        row = self.selectedList.currentRow()
        currentItem = self.selectedList.takeItem(row)
        self.removeList.insertItem(0, currentItem)

    @QtCore.pyqtSlot()
    def on_btnAdd_clicked(self):
        row = self.removeList.currentRow()
        currentItem = self.removeList.takeItem(row)
        self.selectedList.insertItem(0, currentItem)

    @QtCore.pyqtSlot(int)
    def on_chkRunWithBreakDown_changed(self, state):
        if state == 2:  # Qt.CheckState.Checked
            self.chkRunWithBreakDown.setText("✔  Run with BreakDown")
            self.chkRunWithBreakDown.setStyleSheet("color: #88C0D0; font-weight: bold;")
        else:
            self.chkRunWithBreakDown.setText("Run with BreakDown")
            self.chkRunWithBreakDown.setStyleSheet("color: #ECEFF4; font-weight: bold;")

    @QtCore.pyqtSlot()
    def on_btnCancel_clicked(self):
        sys.exit(0)

    def open_doc_file(self, file):
        file_name = os.path.split(file)[1]
        QMessageBox.information(self, "mSelect [C&M]",
                                f"File: {file_name}\nIf any updates, Please save and close...")
        if file.lower().endswith(".doc") or file.lower().endswith(".docx"):
            word = win32.gencache.EnsureDispatch('Word.Application')
            word.Visible = True
            try:
                doc = word.Documents.Open(file)
            except Exception as e:
                QMessageBox.information(self, "mSelect [C&M]",
                                        f"File: {file}\nUnable to Open file.\n"
                                        f"Please open, Check, and Save manually..")
                os.startfile(file)

    def open_pdf_file(self, file):
        global pdf_file
        pdf_file = re.sub(".docx", ".pdf", file)
        convert(file, pdf_file)
        self.childWin = ChildWindow()
        if self.childWin.isVisible():
            self.childWin.close()
        else:
            self.childWin.show()
            url = QUrl.fromLocalFile(pdf_file)
            self.childWin.webView.load(url)

    def open_instructions(self, instructions):
        temp_file = "temp.txt"
        if os.path.exists("temp.txt"):
            subprocess.run(["del", temp_file], shell=True)
        with open(temp_file, "w") as file:
            file.write(instructions)
        subprocess.run(["notepad.exe", temp_file])

    def _run_breakdown_after_parastyler(self, jid, aid, article_folder, customer,
                                        breakdown_folder, breakdown_done,
                                        currentItem, mydb, uniq_id, process_doc):
        """
        Runs the mBreakDown step immediately after a successful mParaStyler.
        The file is expected to already be in breakdown_folder/article_folder/currentItem
        (moved there at the end of the mParaStyler step).
        Returns True if breakdown completed and app should exit, False otherwise.
        """
        article_input = os.path.join(breakdown_folder, article_folder)
        article_done  = os.path.join(breakdown_done,   article_folder)
        file_full_path = os.path.join(article_input, currentItem)

        # Backup the AS file before breakdown
        doc_folder = os.path.join(article_input, "docs")
        if not os.path.exists(doc_folder):
            os.mkdir(doc_folder)
        backup_file = re.sub("_CLN", "_CLN_AS", currentItem)
        back_up_full_path = os.path.join(doc_folder, backup_file)
        shutil.copy(file_full_path, back_up_full_path)

        # Create article JSON if missing, then run breakdown
        json_name = jid + "_" + aid + ".json"
        json_path = os.path.join(article_input, json_name)
        breakdown_process = BreakDownProcess()
        if os.path.exists(json_path):
            breakdown_process.create_breakdown_docx(jid, aid, file_full_path)
        else:
            creat_info = GetArticleId()
            creat_info.smart_login(aid, None, jid, article_input)
            breakdown_process.create_breakdown_docx(jid, aid, file_full_path)

        # Move article_input → article_done
        process_completed = True
        if os.path.exists(article_done):
            try:
                shutil.rmtree(article_done)
            except Exception as e:
                process_completed = False

        if process_completed:
            try:
                shutil.move(article_input, article_done)
            except Exception as e:
                process_completed = False

        if not process_completed:
            QMessageBox.information(self, "mSelect [C&M]",
                                    f"BreakDown completed but unable to move folder.\n"
                                    f"Please move {article_input} to {article_done} manually.")
            return False

        done_file_path = os.path.join(article_done, currentItem)
        mydb.update_db(uniq_id, "mBreakDown", "COMPLETED", "", "")
        QMessageBox.information(self, "mBreakDown [C&M]",
                                f"{currentItem} BreakDown Process Completed")

        word_instance, doc_instance = process_doc.openWordDocumentVisible(done_file_path)
        if word_instance is not None and doc_instance is not None:
            print("Document opened successfully and should be visible.")
            print(f"Opened: {doc_instance.Name}")
            return True
        else:
            print("Failed to open document.")
            return False

    # ------------------------------------------------------------------
    # Bookmark validation
    # ------------------------------------------------------------------
    def _check_body_back_bookmarks(self, docx_path: str):
        import zipfile
        import re as _re

        missing = []
        try:
            with zipfile.ZipFile(docx_path, 'r') as z:
                xml = z.read('word/document.xml').decode('utf-8', errors='replace')
            for bm_name in ('body', 'back'):
                # search for w:bookmarkStart ... w:name="body" (double or single quote)
                pattern = 'w:bookmarkStart[^>]*w:name=' + '"' + bm_name + '"'
                if not _re.search(pattern, xml):
                    # also try single-quoted variant
                    pattern2 = "w:bookmarkStart[^>]*w:name='" + bm_name + "'"
                    if not _re.search(pattern2, xml):
                        missing.append('"' + bm_name + '"')
        except Exception as exc:
            return False, "Could not read bookmarks from file:\n" + str(exc)

        if missing:
            names = " and ".join(missing)
            msg = (
                "Bookmark(s) " + names + " not found in the selected document.\n\n"
                "Please insert the body and back bookmarks using the\n"
                "Bookmark → body-bookmark / back-bookmark buttons\n"
                "before running Para Styler."
            )
            return False, msg
        return True, ""
        return True, ""

    @QtCore.pyqtSlot()
    def on_btnOK_clicked(self):
        items = []
        ritems = []
        self.selected_list = {}
        self.removed_list = {}
        self.normalizer_input = {}
        close_dialog = False
        for x in range(self.selectedList.count()):
            items.append(self.selectedList.item(x).text().encode("utf8").decode("utf8"))
        for x in range(self.removeList.count()):
            ritems.append(self.removeList.item(x).text().encode("utf8").decode("utf8"))
        for ritem in ritems:
            self.removed_list[ritem] = self.file_details[ritem]
        for item in items:
            self.selected_list[item] = self.file_details[item]
        if len(self.selected_list) > 0:
            self.normalizer_input['selected'] = self.selected_list
            self.normalizer_input['customer'] = self.customerCombo.currentText()
            self.normalizer_input['folder'] = self.jidAidCombo.currentText()
            self.normalizer_input['process'] = self.processCombo.currentText()
        if len(self.removed_list) > 0:
            self.normalizer_input['removed'] = self.removed_list
        if len(self.selected_list) == 0:
            QMessageBox.information(self, "mSelect [C&M]",
                                    "Please Select atleast one document to proceed...")
        else:
            process_doc = OpenDocFile()
            selected_process = self.processCombo.currentText()
            self.btnOK.setDisabled(True)
            self.btnCancel.setDisabled(True)
            self.jidAidCombo.setDisabled(True)
            jid_aid = self.jidAidCombo.currentText()
            jidAid = jid_aid.split("_")
            jrn_id = jidAid[0]
            art_id = jidAid[1]
            mydb = DataBase()
            uniq_id = mydb.get_uniqueid(jrn_id, art_id)
            if selected_process == "mMerger":
                self.pbar.setHidden(False)
                self.pbar.setValue(50)
                merger = DocxMerger()
                merger_result = merger.merge_docx_robust(self.normalizer_input, uniq_id)
                if merger_result:
                    merger.move_files_to_docs(self.normalizer_input, uniq_id)
                    mydb.update_db(uniq_id, "mMerger", "COMPLETED", "", "")
                else:
                    QMessageBox.warning(self, "mSelect [C&M]",
                                        "Merging failed. Please check logs and try manual merge.")
                    mydb.update_db(uniq_id, "mMerger", "ERROR", "",
                                   "All merge strategies failed")
                self.pbar.setValue(100)
                if merger_result is True:
                    if merger_result is True:
                        merger_folder = self.normalizer_input['folder']
                        customer = self.normalizer_input['customer']
                        source_folder = re.sub(r"\[CUSTOMER\]", customer, self.mergerInput)
                        dest_folder = re.sub(r"\[CUSTOMER\]", customer, self.ParaStylerInput)
                        source_folder = os.path.join(source_folder, merger_folder)
                        dest_folder = os.path.join(dest_folder, merger_folder)
                        try:
                            shutil.move(source_folder, dest_folder)
                        except Exception as e:
                            print(e)
                    QMessageBox.information(self, "mSelect [C&M]",
                                            "Merging Process Completed...")
                    self.jidAidCombo.removeItem(self.jidAidCombo.currentIndex())
                    mydb.update_db(uniq_id, "mMerger", "COMPLETED", "", "")
                elif merger_result is False:
                    QMessageBox.information(self, "mSelect [C&M]",
                                            "Some of the documents may contain shapes please merger manually")
                    mydb.update_db(uniq_id, "mMerger", "ERROR", "", f"Some of the documents may contain shapes please merger manually")
                self.pbar.setHidden(True)

            elif selected_process == "mParaStyler":
                if len(self.selected_list) > 1:
                    QMessageBox.information(self, "mSelect [C&M]", "For Para Styler input must be single file\nPlease remove unwanted files and run again.")
                elif len(self.selected_list) == 1:
                    if self.selectedList.currentItem() is None:
                        first_item = self.selectedList.item(0)
                        self.selectedList.setCurrentItem(first_item)
                    currentItem = self.selectedList.currentItem().text()
                    article_folder = self.normalizer_input['folder']
                    customer = self.normalizer_input['customer']
                    parastyler_folder = re.sub(r"\[CUSTOMER\]", customer, self.ParaStylerInput)
                    breakdown_folder  = re.sub(r"\[CUSTOMER\]", customer, self.BreakDownInput)
                    breakdown_done    = re.sub(r"\[CUSTOMER\]", customer, self.BreakDownDone)
                    file_full_path = os.path.join(parastyler_folder, article_folder, currentItem)
                    file_folder    = os.path.split(file_full_path)[0]
                    file_name      = os.path.split(file_full_path)[1]
                    # Backup original before anything starts
                    backup_folder = os.path.join(file_folder, "docs")
                    if not os.path.exists(backup_folder):
                        os.makedirs(backup_folder)
                    shutil.copy(file_full_path, os.path.join(backup_folder, file_name))
                    info_path = os.path.join(file_folder, "para_info.xml")
                    if os.path.exists(info_path):
                        os.remove(info_path)
                    # Build worker params — no Qt widget references cross thread
                    worker_params = {
                        "file_full_path":   file_full_path,
                        "file_folder":      file_folder,
                        "file_name":        file_name,
                        "as_file_expected": re.sub(".docx", "_AS.docx", file_full_path),
                        "breakdown_folder": breakdown_folder,
                        "breakdown_done":   breakdown_done,
                        "article_folder":   article_folder,
                        "currentItem":      currentItem,
                        "jrn_id":           jrn_id,
                        "art_id":           art_id,
                        "uniq_id":          uniq_id,
                        "run_breakdown":    self.chkRunWithBreakDown.isChecked(),
                        "configFolder":     self.configFolder,
                        "app_path":         self.app_path,
                        "list_macros":      ["AutoFitTable", "LoadSageStyles", "RemoveLineNumbers"],
                    }
                    # ── Bookmark guard ─────────────────────────────────
                    bm_ok, bm_msg = self._check_body_back_bookmarks(file_full_path)
                    if not bm_ok:
                        self.btnOK.setEnabled(True)
                        self.btnCancel.setEnabled(True)
                        self.jidAidCombo.setEnabled(True)
                        QMessageBox.warning(self, "mSelect [C&M] — Bookmarks Required", bm_msg)
                        return
                    # ────────────────────────────────────────────────────

                    self._start_parastyler_worker(worker_params)
                    return   # UI returns immediately; worker signals handle the rest

            elif selected_process == "mBreakDown":
                if len(self.selected_list) > 1:
                    QMessageBox.information(self, "mSelect [C&M]", "For Breakdown Process input must be Single file\nPlease remove unwanted files and run again.")
                elif len(self.selected_list) == 1:
                    if self.selectedList.currentItem() is None:
                        first_item = self.selectedList.item(0)
                        self.selectedList.setCurrentItem(first_item)
                    currentItem      = self.selectedList.currentItem().text()
                    article_folder   = self.normalizer_input['folder']
                    customer         = self.normalizer_input['customer']
                    breakdown_folder = re.sub(r"\[CUSTOMER\]", customer, self.BreakDownInput)
                    breakdown_done   = re.sub(r"\[CUSTOMER\]", customer, self.BreakDownDone)
                    article_input    = os.path.join(breakdown_folder, article_folder)
                    article_done     = os.path.join(breakdown_done, article_folder)
                    file_full_path   = os.path.join(breakdown_folder, article_folder, currentItem)
                    backup_file      = re.sub("_CLN", "_CLN_AS", currentItem)
                    back_up_full_path = os.path.join(breakdown_folder, article_folder, "docs", backup_file)
                    done_file_path   = os.path.join(article_done, currentItem)
                    art_parts        = article_folder.split("_")
                    jid              = art_parts[0]
                    aid              = art_parts[1]
                    json_name        = jid + "_" + aid + ".json"
                    json_path        = os.path.join(article_input, json_name)

                    # ── JID pre-check: must exist in BreakDown.json ────────
                    jid_ok, jid_msg = self._validate_jid_in_breakdown_json(jid)
                    if not jid_ok:
                        self.btnOK.setEnabled(True)
                        self.btnCancel.setEnabled(True)
                        self.jidAidCombo.setEnabled(True)
                        QMessageBox.warning(self, "mBreakDown [C&M] — Journal Not Found", jid_msg)
                        return
                    # ──────────────────────────────────────────────────────

                    QMessageBox.information(self, "mBreakDown [C&M]", f"{currentItem} BreakDown Process")

                    bd_params = {
                        "currentItem":        currentItem,
                        "article_input":      article_input,
                        "article_done":       article_done,
                        "file_full_path":     file_full_path,
                        "back_up_full_path":  back_up_full_path,
                        "done_file_path":     done_file_path,
                        "jid":                jid,
                        "aid":                aid,
                        "json_path":          json_path,
                        "uniq_id":            uniq_id,
                    }
                    self._start_breakdown_worker(bd_params)
                    return   # UI returns immediately; worker signals handle the rest

            self.selectedList.clear()
            self.removeList.clear()
            self.btnOK.setEnabled(True)
            self.btnCancel.setEnabled(True)
            self.jidAidCombo.setEnabled(True)
            if close_dialog is True:
                sys.exit(0)

    # ------------------------------------------------------------------
    # Worker launcher + signal slots
    # ------------------------------------------------------------------
    def _start_parastyler_worker(self, params: dict):
        """Kick off ParaStylerWorker and show the progress bar."""
        self.pbar.setValue(0)
        self.pbar.setHidden(False)
        self.statusLabel.setText("Starting...")
        self.statusLabel.setHidden(False)
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)

        self._worker = ParaStylerWorker(params)
        self._worker.progress.connect(self._on_worker_progress)
        self._worker.status.connect(self._on_worker_status)
        self._worker.finished.connect(self._on_worker_finished)
        self._worker.error.connect(self._on_worker_error)
        self._worker.json_missing.connect(self._on_worker_json_missing)   # NEW
        self._worker.start()

    @QtCore.pyqtSlot(int)
    def _on_worker_progress(self, value: int):
        self.pbar.setValue(value)

    @QtCore.pyqtSlot(str)
    def _on_worker_status(self, msg: str):
        self.statusLabel.setText(msg)

    @QtCore.pyqtSlot(bool, str)
    def _on_worker_finished(self, close_app: bool, message: str):
        QApplication.restoreOverrideCursor()
        self.pbar.setHidden(True)
        self.statusLabel.setHidden(True)
        self.selectedList.clear()
        self.removeList.clear()
        self.btnOK.setEnabled(True)
        self.btnCancel.setEnabled(True)
        self.jidAidCombo.setEnabled(True)

        # If the worker asked us to untick "Run with BreakDown" (No-path), do it now
        if hasattr(self, "_worker") and self._worker is not None:
            if self._worker.params.get("_untick_breakdown"):
                self.chkRunWithBreakDown.setChecked(False)
                self.chkRunWithBreakDown.setText("Run with BreakDown")
                self.chkRunWithBreakDown.setStyleSheet("color: #ECEFF4; font-weight: bold;")

        # message format for breakdown result: "title|done_file_path"
        if "|" in message:
            display_msg, done_file_path = message.split("|", 1)
        else:
            display_msg   = message
            done_file_path = None

        if message.startswith("ERROR:"):
            QMessageBox.warning(self, "mSelect [C&M]", display_msg)
        else:
            QMessageBox.information(self, "mSelect [C&M]", display_msg)

        if close_app and done_file_path:
            process_doc = OpenDocFile()
            word_instance, doc_instance = process_doc.openWordDocumentVisible(done_file_path)
            if word_instance is not None and doc_instance is not None:
                print(f"Opened: {doc_instance.Name}")
            sys.exit(0)

    @QtCore.pyqtSlot(str)
    def _on_worker_error(self, err: str):
        QApplication.restoreOverrideCursor()
        self.pbar.setHidden(True)
        self.statusLabel.setHidden(True)
        self.btnOK.setEnabled(True)
        self.btnCancel.setEnabled(True)
        self.jidAidCombo.setEnabled(True)
        QMessageBox.critical(self, "mSelect [C&M]", f"Unexpected error:\n{err}")

    # ------------------------------------------------------------------
    # JSON-missing dialog + CreateArticleInfoWorker slots
    # ------------------------------------------------------------------
    @QtCore.pyqtSlot(dict)
    def _on_worker_json_missing(self, info: dict):
        """
        Called on the UI thread when ParaStylerWorker discovers the JSON is absent.

        Shows a dialog:
          • Yes  → launch CreateArticleInfoWorker to fetch & write the JSON,
                   then resume the ParaStylerWorker with 'yes'
          • No   → resume immediately with 'no' (dummy JSON / skip BreakDown logic
                   is handled inside the worker)
        """
        QApplication.restoreOverrideCursor()

        json_name  = info["json_name"]
        is_inbuild = info["inbuild"]

        msg = (
            f"JSON file not available:\n  {json_name}\n\n"
            "Would you like to create it now by searching SMART?\n\n"
            "  Yes  – Search SMART and create the JSON, then continue.\n"
            "  No   – Skip JSON creation and proceed without it\n"
            "         (BreakDown will be unticked automatically)."
        )
        reply = QMessageBox.question(
            self,
            "mSelect [C&M] — JSON Not Found",
            msg,
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.Yes,
        )

        if reply == QMessageBox.StandardButton.Yes:
            # ── Internet connectivity guard ─────────────────────────────
            if not CreateArticleInfoWorker._check_internet():
                QMessageBox.critical(
                    self,
                    "mSelect [C&M] — No Internet Connection",
                    "Cannot reach the SMART portal (journals.sageapps.com).\n\n"
                    "Please check your network connection and try again.\n\n"
                    "Process has been stopped.",
                )
                self._worker.provide_json_response("no")
                self.chkRunWithBreakDown.setChecked(False)
                self.chkRunWithBreakDown.setText("Run with BreakDown")
                self.chkRunWithBreakDown.setStyleSheet("color: #ECEFF4; font-weight: bold;")
                return
            # ───────────────────────────────────────────────────────────

            # Launch CreateArticleInfoWorker — progress bar re-used
            jid_aid = self.jidAidCombo.currentText()
            parts   = jid_aid.split("_")
            jrn_id  = parts[0]
            art_id  = parts[1]

            # Derive file_folder from the worker's own params
            file_folder = self._worker.params["file_folder"]

            self.statusLabel.setText("Creating JSON from SMART...")
            self.statusLabel.setHidden(False)
            self.pbar.setValue(0)
            self.pbar.setHidden(False)
            QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)

            self._cai_worker = CreateArticleInfoWorker(art_id, jrn_id, file_folder)
            self._cai_worker.progress.connect(self._on_cai_progress)
            self._cai_worker.status.connect(self._on_cai_status)
            self._cai_worker.finished.connect(self._on_cai_finished)
            self._cai_worker.start()
            # ParaStylerWorker remains paused on its threading.Event

        else:
            # No — untick checkbox on UI immediately, then resume worker
            self.chkRunWithBreakDown.setChecked(False)
            self.chkRunWithBreakDown.setText("Run with BreakDown")
            self.chkRunWithBreakDown.setStyleSheet("color: #ECEFF4; font-weight: bold;")
            QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
            self._worker.provide_json_response("no")

    @QtCore.pyqtSlot(int)
    def _on_cai_progress(self, value: int):
        self.pbar.setValue(value)

    @QtCore.pyqtSlot(str)
    def _on_cai_status(self, msg: str):
        self.statusLabel.setText(msg)

    @QtCore.pyqtSlot(bool, str)
    def _on_cai_finished(self, success: bool, path_or_err: str):
        """Called when CreateArticleInfoWorker finishes."""
        QApplication.restoreOverrideCursor()

        if success:
            # Show completed JSON path then resume ParaStylerWorker
            QMessageBox.information(
                self,
                "mSelect [C&M] — JSON Created",
                f"Article JSON created successfully:\n{path_or_err}\n\n"
                "Resuming Para Styler...",
            )
            self.statusLabel.setText("JSON created. Resuming Para Styler...")
            self.pbar.setValue(0)
            QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
            self._worker.provide_json_response("yes")
        else:
            # Creation failed — stop everything and report
            self.pbar.setHidden(True)
            self.statusLabel.setHidden(True)
            self.btnOK.setEnabled(True)
            self.btnCancel.setEnabled(True)
            self.jidAidCombo.setEnabled(True)
            QMessageBox.critical(
                self,
                "mSelect [C&M] — JSON Creation Failed",
                f"Could not create the JSON file:\n\n{path_or_err}\n\n"
                "Process ended. Please create the JSON manually and try again.",
            )
            # Tell the worker to stop (send "no" so it unblocks, then we
            # rely on the worker finishing cleanly — it will emit finished
            # with an ERROR message which _on_worker_finished will handle)
            self._worker.provide_json_response("no")

    # ------------------------------------------------------------------
    # BreakDown Worker launcher + signal slots
    # ------------------------------------------------------------------
    def _start_breakdown_worker(self, params: dict):
        """Kick off BreakDownWorker and show the progress bar."""
        self.pbar.setValue(0)
        self.pbar.setHidden(False)
        self.statusLabel.setText("Starting BreakDown...")
        self.statusLabel.setHidden(False)
        self.btnOK.setDisabled(True)
        self.btnCancel.setDisabled(True)
        self.jidAidCombo.setDisabled(True)
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)

        self._bd_worker = BreakDownWorker(params)
        self._bd_worker.progress.connect(self._on_bd_progress)
        self._bd_worker.status.connect(self._on_bd_status)
        self._bd_worker.finished.connect(self._on_bd_finished)
        self._bd_worker.error.connect(self._on_bd_error)
        self._bd_worker.start()

    @QtCore.pyqtSlot(int)
    def _on_bd_progress(self, value: int):
        self.pbar.setValue(value)

    @QtCore.pyqtSlot(str)
    def _on_bd_status(self, msg: str):
        self.statusLabel.setText(msg)

    @QtCore.pyqtSlot(bool, str)
    def _on_bd_finished(self, close_app: bool, message: str):
        QApplication.restoreOverrideCursor()
        self.pbar.setHidden(True)
        self.statusLabel.setHidden(True)
        self.selectedList.clear()
        self.removeList.clear()
        self.btnOK.setEnabled(True)
        self.btnCancel.setEnabled(True)
        self.jidAidCombo.setEnabled(True)

        # message format on success: "display text|done_file_path"
        if "|" in message:
            display_msg, done_file_path = message.split("|", 1)
        else:
            display_msg    = message
            done_file_path = None

        if message.startswith("ERROR:"):
            QMessageBox.warning(self, "mBreakDown [C&M]", display_msg)
        else:
            QMessageBox.information(self, "mBreakDown [C&M]", display_msg)

        if close_app and done_file_path:
            process_doc = OpenDocFile()
            word_instance, doc_instance = process_doc.openWordDocumentVisible(done_file_path)
            if word_instance is not None and doc_instance is not None:
                print(f"Opened: {doc_instance.Name}")
            sys.exit(0)

    @QtCore.pyqtSlot(str)
    def _on_bd_error(self, err: str):
        QApplication.restoreOverrideCursor()
        self.pbar.setHidden(True)
        self.statusLabel.setHidden(True)
        self.btnOK.setEnabled(True)
        self.btnCancel.setEnabled(True)
        self.jidAidCombo.setEnabled(True)
        QMessageBox.critical(self, "mBreakDown [C&M]", f"Unexpected error:\n{err}")

    def run_para_styler(self, doc_file):
        """Legacy fallback — kept for mBreakDown-only path."""
        styler_path = os.path.join(self.configFolder, "ParaStyler\\run.bat")
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        p = Popen([styler_path, doc_file], stdout=PIPE, stderr=PIPE)
        as_file = re.sub(".docx", "_AS.docx", doc_file)
        output, errors = p.communicate()
        QApplication.restoreOverrideCursor()
        p.wait()
        return as_file

    def remove_box(self, document):
        style = document.styles['abstract']
        prValue = style._element.pPr
        style._element.remove(prValue)
        return document

    def header_count(self, table):
        span_count = 1
        for row in table.rows:
            row_xml = row._tr.xml
            if re.search(r"(w\:gridSpan)", row_xml, re.IGNORECASE):
                span_count = span_count + 1
            else:
                break
        return span_count

    def update_table_styles(self, document):
        styles = document.styles
        if "tablehead" not in styles:
            style = document.styles.add_style("tablehead", WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = document.styles["Normal"]
        if "tablebody" not in styles:
            style = document.styles.add_style("tablebody", WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = document.styles["Normal"]
        styles = document.styles
        for table in document.tables:
            thead_count = self.header_count(table)
            row_count = 1
            for row in table.rows:
                if row_count <= thead_count:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.style = styles['tablehead']
                else:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.style = styles['tablebody']
                    row_count = row_count + 1
                row_count = row_count + 1
        return document

    def add_math_style(self, doc):
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                 'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
        for prefix, uri in nsmap.items():
            etree.register_namespace(prefix, uri)
        for paragraph in doc.paragraphs:
            math_elements = paragraph._element.xpath('.//m:oMathPara')
            if math_elements:
                p = paragraph._element
                ppr = p.find('.//w:pPr', namespaces=nsmap)
                if ppr is None:
                    ppr = parse_xml('<w:pPr xmlns:w="{}"/>'.format(nsmap['w']))
                    p.insert(0, ppr)
                pstyle_elem = ppr.find('.//w:pStyle', namespaces=nsmap)
                if pstyle_elem is not None and pstyle_elem.get(
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') != "displaymath":
                    ppr.remove(pstyle_elem)
                xml_str = '<w:pStyle w:val="displaymath" xmlns:w="{}"/>'.format(nsmap['w'])
                pstyle = parse_xml(xml_str)
                ppr.append(pstyle)
        return doc

    def clean_document_xml(self, document):
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        p_style = paragraph.style
                        para_style = paragraph.style.name
                        stId = document.styles[para_style]._element.styleId
                        pprElement = paragraph._p.pPr
                        new_ppr = OxmlElement("w:pPr")
                        new_ppr.style = stId
                        paragraph._p.remove(pprElement)
                        paragraph._p.insert(0, new_ppr)
                        for run in paragraph.runs:
                            rprValue = run._r.rPr
                            if rprValue is not None:
                                szVal = run._r.rPr.sz
                                colVal = run._r.rPr.color
                                rFontsValue = run._element.rPr.rFonts
                                asciiValue = run._element.rPr.rFonts_ascii
                                hAnsiValue = run._element.rPr.rFonts_hAnsi
                                if asciiValue is not None and asciiValue == "Arial":
                                    run._r.rPr.remove(rFontsValue)
                                if szVal is not None:
                                    run._r.rPr.remove(szVal)
                                if colVal is not None:
                                    colXml = run._r.rPr.color.xml
                                    if bool(re.search('w:val="000000"', colXml)) is True:
                                        run._r.rPr.remove(colVal)
        for paragraph in document.paragraphs:
            p_style = paragraph.style
            para_style = paragraph.style.name
            para_text = paragraph.text
            if para_style != "Normal":
                stId = document.styles[para_style]._element.styleId
                pprElement = paragraph._p.pPr
                new_ppr = OxmlElement("w:pPr")
                new_ppr.style = stId
                paragraph._p.remove(pprElement)
                paragraph._p.insert(0, new_ppr)
                for run in paragraph.runs:
                    rprValue = run._r.rPr
                    if rprValue is not None:
                        szVal = run._r.rPr.sz
                        colVal = run._r.rPr.color
                        boldVal = run._r.rPr.b
                        rFontsValue = run._element.rPr.rFonts
                        asciiValue = run._element.rPr.rFonts_ascii
                        hAnsiValue = run._element.rPr.rFonts_hAnsi
                        if asciiValue is not None and asciiValue == "Arial":
                            run._r.rPr.remove(rFontsValue)
                        if asciiValue is not None and asciiValue == "Calibri":
                            run._r.rPr.remove(rFontsValue)
                        if asciiValue is not None and asciiValue == "majorHAnsi":
                            run._r.rPr.remove(rFontsValue)
                        if boldVal is not None:
                            if re.search("articletitle|abstracttitle|sectiona|sectionb|sectionc|sectiond|conflictofinterest|acknowledgementstitle|references", para_style, re.IGNORECASE):
                                run._r.rPr.remove(boldVal)
                        if szVal is not None:
                            run._r.rPr.remove(szVal)
                        if colVal is not None:
                            colXml = run._r.rPr.color.xml
                            if bool(re.search('w:val="000000"', colXml)) is True:
                                run._r.rPr.remove(colVal)
        return document
