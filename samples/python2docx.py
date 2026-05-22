import sys

import docx
from docx2python import docx2python as dx2py


def ns_tag_name(node, name):
    if node.nsmap and node.prefix:
        return "{{{:s}}}{:s}".format(node.nsmap[node.prefix], name)
    return name


def descendants(node, desc_strs):
    if node is None:
        return []
    if not desc_strs:
        return [node]
    ret = {}
    for child_str in desc_strs[0]:
        for child in node.iterchildren(ns_tag_name(node, child_str)):
            descs = descendants(child, desc_strs[1:])
            if not descs:
                continue
            cd = ret.setdefault(child_str, [])
            if isinstance(descs, list):
                cd.extend(descs)
            else:
                cd.append(descs)
    return ret


def simplified_descendants(desc_dict):
    ret = []
    for vs in desc_dict.values():
        for v in vs:
            if isinstance(v, dict):
                ret.extend(simplified_descendants(v))
            else:
                ret.append(v)
    return ret


def process_list_data(attrs, dx2py_elem):
    #print(simplified_descendants(attrs))
    desc = simplified_descendants(attrs)[0]
    level = int(desc.attrib[ns_tag_name(desc, "val")])
    elem = [i for i in dx2py_elem[0].split("\t") if i][0]#.rstrip(")")
    return "    " * level + elem + " "


def main(*argv):
    fname = r"./sample.docx"
    docd = docx.Document(fname)
    docdpy = dx2py(fname)
    dr = docdpy.docx_reader
    #print(dr.files)  # !!! Check word/numbering.xml !!!
    docdpy_runs = docdpy.document_runs[0][0][0]
    if len(docd.paragraphs) != len(docdpy_runs):
        print("Lengths don't match. Abort")
        return -1
    subnode_tags = (("pPr",), ("numPr",), ("ilvl",))  # (("pPr",), ("numPr",), ("ilvl", "numId"))  # numId is for matching elements from word/numbering.xml
    for idx, (par, l) in enumerate(zip(docd.paragraphs, docdpy_runs)):
        #print(par.text, l)
        numbered_attrs = descendants(par._element, subnode_tags)
        #print(numbered_attrs)
        if numbered_attrs:
            print(process_list_data(numbered_attrs, l) + par.text)
        else:
            print(par.text)


if __name__ == "__main__":
    print("Python {:s} {:03d}bit on {:s}\n".format(" ".join(elem.strip() for elem in sys.version.split("\n")),
                                                   64 if sys.maxsize > 0x100000000 else 32, sys.platform))
    rc = main(*sys.argv[1:])
    print("\nDone.")
    sys.exit(rc)