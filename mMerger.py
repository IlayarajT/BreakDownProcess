import os
import re
import shutil
import time

import win32com.client as win32
import yaml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import (
    WD_COLOR_INDEX,
    WD_PARAGRAPH_ALIGNMENT,
    WD_LINE_SPACING,
)
from docx.shared import Pt, RGBColor
from docxcompose.composer import Composer
from win32com.client import constants
import zipfile
import tempfile
import getAppPath
from dbprocess import DataBase
from loadconfig import getconfig
from com_manager import WordApplicationManager, COMManager
import threading

class DocxMerger:
    # ==============================================================
    # INIT
    # ==============================================================

    def __init__(self):
        getAppPath.getapppath()  # preserved side-effect
        self.configFolder, self.breakDownConfig = getconfig()
        self.db_process = DataBase()

        merger_yaml = os.path.join(
            self.configFolder, "config", "mMerger.yaml"
        )
        with open(merger_yaml, "r") as stream:
            self.merger_config = yaml.safe_load(stream)

    # ==============================================================
    # FILE MOVE / CLEANUP
    # ==============================================================

    def move_files_to_docs(self, input_list, unique_id):
        source_files_option = self.merger_config["SOURCE"]
        selected_docs = input_list["selected"]
        removed_docs = input_list.get("removed", {})

        files = list(selected_docs.values())
        file_path = os.path.dirname(files[0])
        backup_path = os.path.join(file_path, "docs")

        os.makedirs(backup_path, exist_ok=True)

        def move_to_backup(doc_map):
            for name, path in doc_map.items():
                shutil.move(path, os.path.join(backup_path, name))

        def remove_files(doc_map):
            for path in doc_map.values():
                os.remove(path)

        if source_files_option == "Retain":
            move_to_backup(selected_docs)
            move_to_backup(removed_docs)
        elif source_files_option == "Remove":
            remove_files(selected_docs)
            remove_files(removed_docs)

    # ==============================================================
    # MERGE USING WORD (DOC)
    # ==============================================================

    def merge_in_doc_old(self, input_list, unique_id):
        merger_result = False
        customer = input_list["customer"]
        folder = input_list["folder"]

        cln_doc_name = self.merger_config["FILENAME"][customer]
        cln_doc_name = re.sub("FOLDER", folder, cln_doc_name)

        files = list(input_list["selected"].values())
        file_path = os.path.dirname(files[0])
        composed_file = os.path.join(file_path, cln_doc_name)

        file1 = files.pop(0)
        macro_body = ""

        for file in files:
            file = os.path.abspath(file)
            macro_body += (
                f'\t\t.InsertFile FileName := "{file}", '
                "ConfirmConversions := False, "
                "Link := False, Attachment := False\n"
            )

        word = win32.gencache.EnsureDispatch("Word.Application")
        word.Visible = True
        doc = word.Documents.Open(file1)
        time.sleep(3)

        try:
            macro = doc.VBProject.VBComponents.Add(1)
        except Exception as e:
            print(e)
            return False

        code = f"""
        sub FileMerger()
            Selection.EndKey Unit:=wdStory
            With Selection
            {macro_body}
            End With
        end sub
        """

        macro.CodeModule.AddFromString(code)

        try:
            doc.Application.Run("FileMerger")
            merger_result = True
        except Exception as e:
            print(e)

        if merger_result:
            try:
                word.ActiveDocument.SaveAs(
                    composed_file,
                    FileFormat=constants.wdFormatXMLDocument,
                )
                word.ActiveDocument.Close()
                self.find_duplicates_new(composed_file)
            except Exception as e:
                print(e)
                merger_result = False

        time.sleep(3)
        return merger_result

    def merge_in_doc(self, input_list, unique_id):
        """Merge documents using Word automation with proper COM handling"""
        merger_result = False
        customer = input_list["customer"]
        folder = input_list["folder"]

        cln_doc_name = self.merger_config["FILENAME"][customer]
        cln_doc_name = re.sub("FOLDER", folder, cln_doc_name)

        files = list(input_list["selected"].values())
        file_path = os.path.dirname(files[0])
        composed_file = os.path.join(file_path, cln_doc_name)

        file1 = files.pop(0)
        macro_body = ""

        for file in files:
            file = os.path.abspath(file)
            macro_body += (
                f'\t\t.InsertFile FileName := "{file}", '
                "ConfirmConversions := False, "
                "Link := False, Attachment := False\n"
            )

        try:
            # Use WordApplicationManager for proper COM handling
            with WordApplicationManager(visible=True) as word:
                # Open the first document
                doc = word.Documents.Open(file1)
                time.sleep(2)  # Give Word time to open

                # Try to add macro
                try:
                    macro = doc.VBProject.VBComponents.Add(1)
                except Exception as e:
                    print(f"Could not add VBProject macro: {e}")
                    # Fallback: use Selection.InsertFile directly
                    word.Selection.EndKey(Unit=constants.wdStory)

                    for file in files:
                        try:
                            word.Selection.InsertFile(
                                FileName=file,
                                ConfirmConversions=False,
                                Link=False,
                                Attachment=False
                            )
                        except Exception as insert_error:
                            print(f"Error inserting file {file}: {insert_error}")
                            continue

                    merger_result = True
                else:
                    # Add macro code
                    code = f"""
                    sub FileMerger()
                        Selection.EndKey Unit:=wdStory
                        With Selection
                        {macro_body}
                        End With
                    end sub
                    """

                    try:
                        macro.CodeModule.AddFromString(code)
                        doc.Application.Run("FileMerger")
                        merger_result = True
                    except Exception as macro_error:
                        print(f"Macro execution failed: {macro_error}")
                        # Try alternative method
                        word.Selection.EndKey(Unit=constants.wdStory)
                        for file in files:
                            word.Selection.InsertFile(
                                FileName=file,
                                ConfirmConversions=False,
                                Link=False,
                                Attachment=False
                            )
                        merger_result = True

                # Save the merged document
                if merger_result:
                    try:
                        doc.SaveAs(
                            FileName=composed_file,
                            FileFormat=constants.wdFormatXMLDocument
                        )
                        doc.Close(SaveChanges=0)

                        # Process duplicates
                        self.find_duplicates_new(composed_file)

                    except Exception as save_error:
                        print(f"Error saving document: {save_error}")
                        merger_result = False

        except Exception as e:
            print(f"Word automation failed: {e}")
            merger_result = False

        return merger_result

    # ==============================================================
    # THREAD-SAFE WORD OPERATIONS
    # ==============================================================

    def merge_in_doc_threadsafe(self, input_list, unique_id):
        """Thread-safe version for use in worker threads"""
        result = [False]  # Use list to capture result from nested function

        def merge_worker():
            try:
                # Use com_context for worker threads
                from com_manager import COMManager
                with COMManager.com_context():
                    result[0] = self.merge_in_doc(input_list, unique_id)
            except Exception as e:
                print(f"Thread-safe merge failed: {e}")
                result[0] = False

        # Run in a separate thread
        thread = threading.Thread(target=merge_worker)
        thread.start()
        thread.join(timeout=300)  # 5 minute timeout

        if thread.is_alive():
            print("Word automation timeout")
            return False

        return result[0]

    # ==============================================================
    # ENHANCED MERGE WITH FALLBACK STRATEGY
    # ==============================================================

    def merge_docx_robust(self, input_list, unique_id):
        """Robust merging with multiple fallback strategies"""
        customer = input_list["customer"]
        folder = input_list["folder"]

        cln_doc_name = self.merger_config["FILENAME"][customer]
        cln_doc_name = re.sub("FOLDER", folder, cln_doc_name)

        files = list(input_list["selected"].values())
        file_path = os.path.dirname(files[0])
        composed_file = os.path.join(file_path, cln_doc_name)

        # Strategy 1: Try python-docx for simple documents
        if not self._has_diagrams_or_complex_structure(files):
            try:
                if len(files) == 1:
                    shutil.copy(files[0], composed_file)
                    self._fix_orphan_notes(composed_file, source_files=files)
                    self.find_duplicates_new(composed_file)
                    return True

                result = Document(files[0])
                composer = Composer(result)

                for i, fname in enumerate(files[1:], start=1):
                    doc = Document(fname)
                    if i != len(files) - 1:
                        doc.add_page_break()
                    composer.append(doc)

                composer.save(composed_file)
                self._fix_orphan_notes(composed_file, source_files=files)
                self.find_duplicates_new(composed_file)
                return True

            except Exception as e:
                print(f"Python-docx merge failed: {e}, falling back to Word...")

        # Strategy 2: Try Word automation
        try:
            result = self.merge_in_doc(input_list, unique_id)
            if result:
                try:
                    self._fix_orphan_notes(composed_file, source_files=files)
                except Exception:
                    pass
                return True
        except Exception as e:
            print(f"Word automation failed: {e}")

        # Strategy 3: Thread-safe Word automation as last resort
        try:
            print("Attempting thread-safe Word merge...")
            result = self.merge_in_doc_threadsafe(input_list, unique_id)
            if result:
                try:
                    self._fix_orphan_notes(composed_file, source_files=files)
                except Exception:
                    pass
            return result
        except Exception as e:
            print(f"All merge strategies failed: {e}")

        return False

    # ==============================================================
    # ERROR HANDLING AND LOGGING
    # ==============================================================

    def log_merge_attempt(self, files, method, success, error=None):
        """Log merge attempts for debugging"""
        log_entry = {
            'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
            'files': [os.path.basename(f) for f in files],
            'method': method,
            'success': success,
            'error': str(error) if error else None
        }

        # Log to file
        log_file = os.path.join(self.configFolder, 'merge_log.json')
        try:
            import json
            logs = []
            if os.path.exists(log_file):
                with open(log_file, 'r') as f:
                    logs = json.load(f)
            logs.append(log_entry)
            with open(log_file, 'w') as f:
                json.dump(logs, f, indent=2)
        except:
            pass

        # Update database if available
        if not success and hasattr(self, 'db_process'):
            error_msg = f"{method}: {error}" if error else method
            self.db_process.update_db(
                unique_id, "mMerger", "ERROR", "", error_msg
            )

    # ==============================================================
    # MERGE DOCX (COMPOSER)
    # ==============================================================
    def merge_docx(self, input_list, unique_id):
        customer = input_list["customer"]
        folder = input_list["folder"]

        cln_doc_name = self.merger_config["FILENAME"][customer]
        cln_doc_name = re.sub("FOLDER", folder, cln_doc_name)

        files = list(input_list["selected"].values())
        file_path = os.path.dirname(files[0])
        composed_file = os.path.join(file_path, cln_doc_name)

        # For documents with diagrams or complex structure, use Word
        if self._has_diagrams_or_complex_structure(files):
            print("Document has diagrams, using Word automation...")
            return self.merge_in_doc(input_list, unique_id)

        # Simple documents can use python-docx
        try:
            if len(files) > 1:
                result = Document(files[0])
                composer = Composer(result)

                for i, fname in enumerate(files[1:], start=1):
                    doc = Document(fname)
                    if i != len(files) - 1:
                        doc.add_page_break()
                    composer.append(doc)

                composer.save(composed_file)
                self.find_duplicates_new(composed_file)
                return True
            else:
                shutil.copy(files[0], composed_file)
                self.find_duplicates_new(composed_file)
                return True

        except Exception as e:
            print(f"Python-docx merge failed: {e}, falling back to Word...")
            return self.merge_in_doc(input_list, unique_id)

    def _has_diagrams_or_complex_structure(self, files):
        """Check if any document has diagrams or might cause merge issues"""
        for file in files:
            try:
                with zipfile.ZipFile(file, 'r') as doc_zip:
                    file_list = doc_zip.namelist()
                    # Check for diagrams or other complex elements
                    if any('diagram' in f.lower() for f in file_list):
                        return True
                    if any('chart' in f.lower() for f in file_list):
                        return True
                    if any('drawing' in f.lower() for f in file_list):
                        return True
            except:
                pass
        return False

    def merge_docx_old(self, input_list, unique_id):
        customer = input_list["customer"]
        folder = input_list["folder"]

        cln_doc_name = self.merger_config["FILENAME"][customer]
        cln_doc_name = re.sub("FOLDER", folder, cln_doc_name)

        files = list(input_list["selected"].values())
        file_path = os.path.dirname(files[0])
        composed_file = os.path.join(file_path, cln_doc_name)

        if len(files) > 1:
            result = Document(files[0])
            result.add_page_break()
            composer = Composer(result)

            for i, fname in enumerate(files[1:], start=1):
                doc = Document(fname)
                if i != len(files) - 1:
                    doc.add_page_break()
                try:
                    composer.append(doc)
                except Exception as e:
                    self.db_process.update_db(
                        unique_id,
                        "mMerger",
                        "ERROR",
                        "",
                        f"Unable to append document {fname}",
                    )
                    print(e)
                    return False

            composer.save(composed_file)
            self.find_duplicates_new(composed_file)
            return True

        # Single file case
        shutil.copy(files[0], composed_file)
        self.find_duplicates_new(composed_file)
        return True

    # ==============================================================
    # DUPLICATE DETECTION (ACTIVE VERSION)
    # ==============================================================

    def _fix_orphan_notes(self, docxfile, source_files=None):
        """Merge endnote/footnote definitions from source docx files into the
        merged docx, so that references in document.xml have matching
        definitions. Also merge the corresponding rels files so that
        hyperlinks inside endnotes/footnotes resolve correctly.
        Prevents "Show Repairs" and "Word found unreadable content" dialogs.

        When docxcompose.Composer merges files:
          - It keeps only the FIRST file's endnotes.xml / footnotes.xml,
            so references from later files have no matching definition.
          - Even when an endnote block IS preserved (because it was in the
            first file or by coincidence), the rels file for that notes
            part (word/_rels/endnotes.xml.rels) is often dropped, leaving
            hyperlinks inside the endnote with unresolved rIds. Word shows
            "Word found unreadable content" in this case.

        This method:
          1. Finds orphan references in document.xml and merges the missing
             endnote/footnote blocks from source files.
          2. For any rId reference inside merged note blocks, also merges
             the matching <Relationship> entry from the source rels file
             (renumbering to avoid collisions with the merged rels).
          3. As a final safety net, also restores the rels file for notes
             that already exist in the merged output but have lost their
             rels (the case where docxcompose copied the block but not the
             rels).
          4. Drops only the <w:endnoteReference/> / <w:footnoteReference/>
             element (NOT the wrapping <w:r>) for any reference whose
             definition truly cannot be recovered.

        Args:
            docxfile: path to the merged docx to fix.
            source_files: list of source docx paths used in the merge.
        """
        try:
            import tempfile

            # ── 1. Read merged file state ──────────────────────────────────
            with zipfile.ZipFile(docxfile, "r") as zin:
                names = zin.namelist()
                doc_xml = zin.read("word/document.xml").decode("utf-8")
                endnotes_xml = (zin.read("word/endnotes.xml").decode("utf-8")
                                if "word/endnotes.xml" in names else "")
                footnotes_xml = (zin.read("word/footnotes.xml").decode("utf-8")
                                 if "word/footnotes.xml" in names else "")
                endnotes_rels = (zin.read("word/_rels/endnotes.xml.rels")
                                 .decode("utf-8")
                                 if "word/_rels/endnotes.xml.rels" in names
                                 else "")
                footnotes_rels = (zin.read("word/_rels/footnotes.xml.rels")
                                  .decode("utf-8")
                                  if "word/_rels/footnotes.xml.rels" in names
                                  else "")

            original_doc = doc_xml
            original_endnotes_xml = endnotes_xml
            original_footnotes_xml = footnotes_xml
            original_endnotes_rels = endnotes_rels
            original_footnotes_rels = footnotes_rels

            existing_endnote_ids = set(re.findall(
                r'<w:endnote\s[^>]*w:id="(-?\d+)"', endnotes_xml
            ))
            existing_footnote_ids = set(re.findall(
                r'<w:footnote\s[^>]*w:id="(-?\d+)"', footnotes_xml
            ))

            doc_endnote_refs = re.findall(
                r'<w:endnoteReference\s[^>]*w:id="(-?\d+)"', doc_xml
            )
            doc_footnote_refs = re.findall(
                r'<w:footnoteReference\s[^>]*w:id="(-?\d+)"', doc_xml
            )
            orphan_endnote_ids = set(doc_endnote_refs) - existing_endnote_ids
            orphan_footnote_ids = set(doc_footnote_refs) - existing_footnote_ids

            # ── 2. Helpers for rels handling ───────────────────────────────
            def _next_rid(rels_xml):
                """Return the next available rId (e.g. 'rId9') for a rels file."""
                rids = re.findall(r'\bId="rId(\d+)"', rels_xml or "")
                nums = [int(x) for x in rids] if rids else []
                return max(nums) + 1 if nums else 1

            def _ensure_rels_skeleton(rels_xml):
                """If rels_xml is empty, return a fresh skeleton."""
                if rels_xml:
                    return rels_xml
                return ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                        '<Relationships xmlns="http://schemas.openxmlformats.org/'
                        'package/2006/relationships"></Relationships>')

            def _extract_rel(rels_xml, rid):
                """Extract a single <Relationship Id="rId"> element."""
                if not rels_xml:
                    return None
                m = re.search(
                    r'<Relationship\s[^>]*?Id="' + re.escape(rid) + r'"[^>]*?/>',
                    rels_xml
                )
                return m.group(0) if m else None

            def _insert_rel(rels_xml, rel_elem):
                """Insert a relationship element before </Relationships>."""
                rels_xml = _ensure_rels_skeleton(rels_xml)
                return re.sub(
                    r'</Relationships>',
                    rel_elem + '</Relationships>',
                    rels_xml,
                    count=1
                )

            def _renumber_rids_in_block(block_xml, rid_map):
                """Replace r:id="rIdOld" / r:embed="rIdOld" etc. with new rId."""
                def _replace_all(xml):
                    for old, new in rid_map.items():
                        pat = re.compile(
                            r'(r:(?:id|embed|link)=")' + re.escape(old) + r'(")'
                        )
                        xml = pat.sub(
                            lambda m: f'{m.group(1)}{new}{m.group(2)}',
                            xml
                        )
                    return xml
                return _replace_all(block_xml)

            def _merge_rids_for_block(block_xml, src_rels_xml,
                                     target_rels_xml):
                """For all rIds referenced in block_xml:
                   - Look up matching <Relationship> in src_rels_xml
                   - Allocate a fresh rId in target_rels_xml
                   - Insert into target_rels_xml with new ID
                   - Return (renumbered_block_xml, updated_target_rels_xml)
                """
                referenced = set(re.findall(
                    r'r:(?:id|embed|link)="(rId\d+)"', block_xml
                ))
                if not referenced:
                    return block_xml, target_rels_xml

                rid_map = {}
                for old_rid in referenced:
                    src_rel = _extract_rel(src_rels_xml, old_rid)
                    if not src_rel:
                        continue
                    # If target already has the same Target (e.g. same URL),
                    # we could de-dup. Simpler: always allocate a fresh ID.
                    new_n = _next_rid(target_rels_xml)
                    new_rid = f'rId{new_n}'
                    new_rel = re.sub(
                        r'(Id=")rId\d+(")',
                        lambda m: f'{m.group(1)}{new_rid}{m.group(2)}',
                        src_rel,
                        count=1
                    )
                    target_rels_xml = _insert_rel(target_rels_xml, new_rel)
                    rid_map[old_rid] = new_rid

                if rid_map:
                    block_xml = _renumber_rids_in_block(block_xml, rid_map)
                return block_xml, target_rels_xml

            # ── 3. Merge orphan endnote/footnote definitions ──────────────
            if source_files and (orphan_endnote_ids or orphan_footnote_ids):
                _existing_ids_int = [int(x) for x in existing_endnote_ids
                                     if x.lstrip('-').isdigit()]
                next_endnote_id = (max(_existing_ids_int) + 1
                                   if _existing_ids_int else 1)
                next_endnote_id = max(next_endnote_id, 1)

                _existing_fids_int = [int(x) for x in existing_footnote_ids
                                      if x.lstrip('-').isdigit()]
                next_footnote_id = (max(_existing_fids_int) + 1
                                    if _existing_fids_int else 1)
                next_footnote_id = max(next_footnote_id, 1)

                endnote_block_re = re.compile(
                    r'<w:endnote\s[^>]*w:id="(-?\d+)"[^>]*>.*?</w:endnote>',
                    re.DOTALL
                )
                footnote_block_re = re.compile(
                    r'<w:footnote\s[^>]*w:id="(-?\d+)"[^>]*>.*?</w:footnote>',
                    re.DOTALL
                )

                for src_path in source_files:
                    try:
                        with zipfile.ZipFile(src_path, "r") as zsrc:
                            snames = zsrc.namelist()
                            src_endnotes = (zsrc.read("word/endnotes.xml")
                                            .decode("utf-8")
                                            if "word/endnotes.xml" in snames
                                            else "")
                            src_footnotes = (zsrc.read("word/footnotes.xml")
                                             .decode("utf-8")
                                             if "word/footnotes.xml" in snames
                                             else "")
                            src_endnotes_rels = (zsrc.read(
                                "word/_rels/endnotes.xml.rels"
                            ).decode("utf-8")
                                if "word/_rels/endnotes.xml.rels" in snames
                                else "")
                            src_footnotes_rels = (zsrc.read(
                                "word/_rels/footnotes.xml.rels"
                            ).decode("utf-8")
                                if "word/_rels/footnotes.xml.rels" in snames
                                else "")
                    except Exception:
                        continue

                    for m in endnote_block_re.finditer(src_endnotes):
                        block = m.group(0)
                        src_id = m.group(1)
                        if src_id in ('-1', '0'):
                            continue
                        if src_id not in orphan_endnote_ids:
                            continue

                        # Re-number the endnote ID itself
                        new_block = re.sub(
                            r'(<w:endnote\s[^>]*w:id=")(-?\d+)(")',
                            lambda mm: f'{mm.group(1)}{next_endnote_id}{mm.group(3)}',
                            block, count=1
                        )
                        # Merge rels for hyperlinks inside the block
                        new_block, endnotes_rels = _merge_rids_for_block(
                            new_block, src_endnotes_rels, endnotes_rels
                        )

                        if endnotes_xml:
                            endnotes_xml = re.sub(
                                r'</w:endnotes>',
                                new_block + '</w:endnotes>',
                                endnotes_xml,
                                count=1
                            )
                            doc_xml = re.sub(
                                r'(<w:endnoteReference\s[^>]*w:id=")'
                                + re.escape(src_id) + r'(")',
                                lambda mm: f'{mm.group(1)}{next_endnote_id}{mm.group(2)}',
                                doc_xml
                            )
                            existing_endnote_ids.add(str(next_endnote_id))
                            orphan_endnote_ids.discard(src_id)
                            next_endnote_id += 1

                    for m in footnote_block_re.finditer(src_footnotes):
                        block = m.group(0)
                        src_id = m.group(1)
                        if src_id in ('-1', '0'):
                            continue
                        if src_id not in orphan_footnote_ids:
                            continue

                        new_block = re.sub(
                            r'(<w:footnote\s[^>]*w:id=")(-?\d+)(")',
                            lambda mm: f'{mm.group(1)}{next_footnote_id}{mm.group(3)}',
                            block, count=1
                        )
                        new_block, footnotes_rels = _merge_rids_for_block(
                            new_block, src_footnotes_rels, footnotes_rels
                        )

                        if footnotes_xml:
                            footnotes_xml = re.sub(
                                r'</w:footnotes>',
                                new_block + '</w:footnotes>',
                                footnotes_xml,
                                count=1
                            )
                            doc_xml = re.sub(
                                r'(<w:footnoteReference\s[^>]*w:id=")'
                                + re.escape(src_id) + r'(")',
                                lambda mm: f'{mm.group(1)}{next_footnote_id}{mm.group(2)}',
                                doc_xml
                            )
                            existing_footnote_ids.add(str(next_footnote_id))
                            orphan_footnote_ids.discard(src_id)
                            next_footnote_id += 1

            # ── 4. Fix lost-rels case: notes XML has rIds but no rels file ─
            #     This happens when docxcompose copied the endnote block
            #     (because it was in the first file or matched by id) but
            #     dropped the rels file. We need to restore rels from any
            #     source whose own endnote block produced the same rIds.
            def _fix_lost_rels(notes_xml, notes_rels, source_files,
                               source_xml_key, source_rels_key):
                referenced = set(re.findall(
                    r'r:(?:id|embed|link)="(rId\d+)"', notes_xml or ""
                ))
                if not referenced:
                    return notes_xml, notes_rels
                # Which rIds are already resolved in the existing rels file?
                resolved_rids = set(re.findall(
                    r'\bId="(rId\d+)"', notes_rels or ""
                ))
                missing = referenced - resolved_rids
                if not missing:
                    return notes_xml, notes_rels

                # Try to find missing rIds in source files' rels
                rid_map = {}
                for src_path in source_files or []:
                    try:
                        with zipfile.ZipFile(src_path, "r") as zsrc:
                            snames = zsrc.namelist()
                            if source_rels_key not in snames:
                                continue
                            src_rels = zsrc.read(source_rels_key).decode("utf-8")
                    except Exception:
                        continue

                    for rid in list(missing):
                        if rid in rid_map:
                            continue
                        src_rel = _extract_rel(src_rels, rid)
                        if not src_rel:
                            continue
                        # Insert with the SAME id (no collision because
                        # it was missing from target). If id collides for
                        # some reason, allocate fresh.
                        existing_ids = set(re.findall(
                            r'\bId="(rId\d+)"', notes_rels or ""
                        ))
                        if rid in existing_ids:
                            new_n = _next_rid(notes_rels)
                            new_rid = f'rId{new_n}'
                            new_rel = re.sub(
                                r'(Id=")rId\d+(")',
                                lambda m: f'{m.group(1)}{new_rid}{m.group(2)}',
                                src_rel, count=1
                            )
                            notes_rels = _insert_rel(notes_rels, new_rel)
                            rid_map[rid] = new_rid
                        else:
                            notes_rels = _insert_rel(notes_rels, src_rel)
                            rid_map[rid] = rid
                        missing.discard(rid)

                # Renumber any rIds that needed to change
                actual_renames = {k: v for k, v in rid_map.items() if k != v}
                if actual_renames:
                    notes_xml = _renumber_rids_in_block(
                        notes_xml, actual_renames
                    )
                return notes_xml, notes_rels

            if source_files:
                endnotes_xml, endnotes_rels = _fix_lost_rels(
                    endnotes_xml, endnotes_rels, source_files,
                    "word/endnotes.xml", "word/_rels/endnotes.xml.rels"
                )
                footnotes_xml, footnotes_rels = _fix_lost_rels(
                    footnotes_xml, footnotes_rels, source_files,
                    "word/footnotes.xml", "word/_rels/footnotes.xml.rels"
                )

            # ── 5. Final safety net: drop unresolved reference elements only
            doc_endnote_refs_after = re.findall(
                r'<w:endnoteReference\s[^>]*w:id="(-?\d+)"', doc_xml
            )
            doc_footnote_refs_after = re.findall(
                r'<w:footnoteReference\s[^>]*w:id="(-?\d+)"', doc_xml
            )
            still_orphan_e = set(doc_endnote_refs_after) - existing_endnote_ids
            still_orphan_f = set(doc_footnote_refs_after) - existing_footnote_ids

            def _drop_ref_only(xml, tag, bad_ids):
                if not bad_ids:
                    return xml
                pat = re.compile(
                    r'<w:' + tag + r'\b[^/]*?w:id="(-?\d+)"\s*/>'
                )
                return pat.sub(
                    lambda m: "" if m.group(1) in bad_ids else m.group(0),
                    xml
                )

            doc_xml = _drop_ref_only(doc_xml, "endnoteReference", still_orphan_e)
            doc_xml = _drop_ref_only(doc_xml, "footnoteReference", still_orphan_f)

            # ── 6. Detect change and write back ───────────────────────────
            unchanged = (doc_xml == original_doc
                         and endnotes_xml == original_endnotes_xml
                         and footnotes_xml == original_footnotes_xml
                         and endnotes_rels == original_endnotes_rels
                         and footnotes_rels == original_footnotes_rels)
            if unchanged:
                return

            tmp_fd, tmp_path = tempfile.mkstemp(
                suffix=".docx", dir=os.path.dirname(docxfile)
            )
            os.close(tmp_fd)
            try:
                with zipfile.ZipFile(docxfile, "r") as zin:
                    existing_names = set(zin.namelist())
                    with zipfile.ZipFile(tmp_path, "w",
                                         zipfile.ZIP_DEFLATED) as zout:
                        # Write existing entries, overriding the parts we
                        # changed and skipping rels files that we are about
                        # to re-add (so we don't duplicate them).
                        skip = set()
                        if endnotes_rels:
                            skip.add("word/_rels/endnotes.xml.rels")
                        if footnotes_rels:
                            skip.add("word/_rels/footnotes.xml.rels")

                        for item in zin.infolist():
                            if item.filename == "word/document.xml":
                                zout.writestr(item, doc_xml.encode("utf-8"))
                            elif item.filename == "word/endnotes.xml":
                                zout.writestr(item, endnotes_xml.encode("utf-8"))
                            elif item.filename == "word/footnotes.xml":
                                zout.writestr(item,
                                              footnotes_xml.encode("utf-8"))
                            elif item.filename in skip:
                                # We'll re-add below
                                continue
                            else:
                                zout.writestr(item, zin.read(item.filename))

                        # Add the rels files (new or updated)
                        if endnotes_rels and \
                                "word/_rels/endnotes.xml.rels" not in existing_names:
                            zout.writestr(
                                "word/_rels/endnotes.xml.rels",
                                endnotes_rels.encode("utf-8")
                            )
                        elif endnotes_rels:
                            zout.writestr(
                                "word/_rels/endnotes.xml.rels",
                                endnotes_rels.encode("utf-8")
                            )
                        if footnotes_rels and \
                                "word/_rels/footnotes.xml.rels" not in existing_names:
                            zout.writestr(
                                "word/_rels/footnotes.xml.rels",
                                footnotes_rels.encode("utf-8")
                            )
                        elif footnotes_rels:
                            zout.writestr(
                                "word/_rels/footnotes.xml.rels",
                                footnotes_rels.encode("utf-8")
                            )

                os.replace(tmp_path, docxfile)
                print(f"[merger] Fixed endnote/footnote references in "
                      f"{os.path.basename(docxfile)}")
            except Exception:
                try:
                    if os.path.exists(tmp_path):
                        os.remove(tmp_path)
                except Exception:
                    pass
                raise
        except Exception as e:
            print(f"[merger] _fix_orphan_notes warning: {e}")

    def find_duplicates_new(self, docxfile):
        document = Document(docxfile)

        try:
            document.styles["Duplicate"].delete()
        except KeyError:
            pass

        dup_style = document.styles.add_style(
            "Duplicate", WD_STYLE_TYPE.PARAGRAPH
        )
        dup_style.base_style = document.styles["Normal"]

        font = dup_style.font
        font.name = "Times New Roman"
        font.size = Pt(12)
        font.color.rgb = RGBColor(255, 0, 0)
        font.highlight_color = WD_COLOR_INDEX.YELLOW

        para_fmt = dup_style.paragraph_format
        para_fmt.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para_fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para_fmt.left_indent = Pt(0)
        para_fmt.right_indent = Pt(0)
        para_fmt.space_before = Pt(0)
        para_fmt.space_after = Pt(0)

        para_map = {}
        for idx, para in enumerate(document.paragraphs):
            if para.text.strip():
                para_map.setdefault(para.text, []).append(idx)

        for indexes in para_map.values():
            if len(indexes) > 1:
                for dup_idx in indexes[1:]:
                    document.paragraphs[dup_idx].style = dup_style

        document.save(docxfile)

    # ==============================================================
    # PARAGRAPH DELETE (UTILITY)
    # ==============================================================

    def delete_paragraph(self, paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
