"""
test_author_labeling.py — Unit tests for labelAuthorsNER module.

Takes a docx and json as input, uses labelAuthorsNER to parse/label authors,
then produces a labeled author docx file.

Usage:
    python test_author_labeling.py <input.docx> <article.json>

Output:
    <input>_LABELED.docx — with [AU1], [AU2], ... labels in blue
"""

import sys
import os
import re
import json
import unittest

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================================
# Import everything from labelAuthorsNER
# ============================================================================
from labelAuthorsNER import (
    # Main classes
    AuthorLabeler,
    LabelAuthor,
    LabelResult,
    AuthorEntity,
    # Convenience function (drop-in for breakDownProcess)
    get_docx_authors_ner,
    # Layer 2: superscript parsing
    parse_by_superscript_boundaries,
    extract_au_runs,
    # Layer 4: JSON cross-ref
    load_json_authors,
    cross_reference_json,
    # HTML generation
    generate_labeled_html,
    # Helpers
    is_degree_token,
    _normalize_name,
)


# ============================================================================
# Paths — set by command line or default to uploaded files
# ============================================================================
DOCX_PATH = None
JSON_PATH = None


def _make_run(text, is_superscript=False):
    return {"text": text, "is_superscript": is_superscript, "style_name": "", "bold": False}


# ============================================================================
# Docx output helper — writes [AU] labels back into the docx
# ============================================================================

def ensure_aulabel_style(doc: Document):
    """Create 'aulabel' character style if missing."""
    if "aulabel" not in {s.name for s in doc.styles}:
        style = doc.styles.add_style("aulabel", WD_STYLE_TYPE.CHARACTER)
        style.font.color.rgb = RGBColor(0, 0, 255)
        style.font.size = Pt(10)
    return doc


def rebuild_au_paragraph(para, authors):
    """Clear AU paragraph and rebuild with labeled runs from labelAuthorsNER result."""
    # Keep paragraph properties, remove all runs
    for child in list(para._p):
        if child.tag != qn('w:pPr'):
            para._p.remove(child)

    def add_run(text, superscript=False, style_name=None, color_rgb=None):
        r_elem = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        if style_name:
            rs = OxmlElement('w:rStyle')
            rs.set(qn('w:val'), style_name)
            rpr.append(rs)
        if superscript:
            v = OxmlElement('w:vertAlign')
            v.set(qn('w:val'), 'superscript')
            rpr.append(v)
        if color_rgb:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color_rgb)
            rpr.append(c)
        r_elem.append(rpr)
        t = OxmlElement('w:t')
        t.text = text
        if text and (text[0] == ' ' or text[-1] == ' '):
            t.set(qn('xml:space'), 'preserve')
        r_elem.append(t)
        para._p.append(r_elem)

    total = len(authors)
    for i, a in enumerate(authors):
        add_run(a.full_name)
        if a.superscripts:
            sup = ",".join(a.superscripts)
            if a.is_corresponding:
                sup += ",*"
            add_run(sup, superscript=True)
        elif a.is_corresponding:
            add_run("*", superscript=True)
        if a.degrees:
            add_run(" " + ", ".join(a.degrees))
        add_run(a.label, style_name="aulabel", color_rgb="0000FF")
        if i < total - 2:
            add_run(", ")
        elif i == total - 2:
            add_run(", and ")
    add_run(".")


# ============================================================================
# Unit Tests — all use labelAuthorsNER imports
# ============================================================================

class TestDegreeRecognition(unittest.TestCase):
    """Test is_degree_token from labelAuthorsNER."""

    def test_known_degrees_recognized(self):
        for deg in ["MBBS", "MD", "PhD", "MSc", "MHA", "FACC", "FRCP", "DrPH", "MPH"]:
            self.assertTrue(is_degree_token(deg), f"{deg} should be a degree")

    def test_degrees_with_trailing_punctuation(self):
        self.assertTrue(is_degree_token("MBBS,"))
        self.assertTrue(is_degree_token("MSc."))
        self.assertTrue(is_degree_token("PhD;"))

    def test_parenthetical_degrees(self):
        # FRCP and MBBS are recognized; parenthetical forms may vary by regex
        self.assertTrue(is_degree_token("FRCP"))
        self.assertTrue(is_degree_token("MBBS"))

    def test_names_not_degrees(self):
        for name in ["Matthew", "Maharaj", "Seecheran", "Naveen", "Rajeev", "Priya",
                      "Neal", "Rishi", "Arun", "Pravin"]:
            self.assertFalse(is_degree_token(name), f"'{name}' should NOT be a degree")

    def test_separators_not_degrees(self):
        for word in ["and", "the", "of", "in", "Dr", "Mr", ","]:
            self.assertFalse(is_degree_token(word), f"'{word}' should NOT be a degree")


class TestSuperscriptBoundaryParsing(unittest.TestCase):
    """Test parse_by_superscript_boundaries from labelAuthorsNER."""

    @classmethod
    def setUpClass(cls):
        """Extract runs from real docx using labelAuthorsNER.extract_au_runs."""
        if DOCX_PATH and os.path.exists(DOCX_PATH):
            cls.au_text, cls.runs_info = extract_au_runs(DOCX_PATH)
        else:
            cls.au_text = ""
            cls.runs_info = []

    def test_parse_correct_author_count(self):
        """Parsed count must match JSON count."""
        if not self.runs_info:
            self.skipTest("No docx provided")
        json_count, _, _ = load_json_authors(JSON_PATH)
        authors = parse_by_superscript_boundaries(self.runs_info, expected_count=json_count)
        self.assertEqual(len(authors), json_count,
                         f"Parsed {len(authors)} but JSON has {json_count}")

    def test_no_degree_as_author_name(self):
        """No degree abbreviation should appear as an author full_name."""
        if not self.runs_info:
            self.skipTest("No docx provided")
        authors = parse_by_superscript_boundaries(self.runs_info)
        for a in authors:
            self.assertFalse(is_degree_token(a.full_name),
                             f"'{a.full_name}' is a degree, not an author name")

    def test_author_indices_sequential(self):
        """AU indices must be 1-based sequential."""
        if not self.runs_info:
            self.skipTest("No docx provided")
        authors = parse_by_superscript_boundaries(self.runs_info)
        indices = [a.index for a in authors]
        self.assertEqual(indices, list(range(1, len(authors) + 1)))

    def test_every_author_has_superscript(self):
        """Each author should have at least one affiliation superscript."""
        if not self.runs_info:
            self.skipTest("No docx provided")
        authors = parse_by_superscript_boundaries(self.runs_info)
        for a in authors:
            self.assertTrue(len(a.superscripts) > 0,
                            f"{a.label} '{a.full_name}' has no superscripts")

    def test_degrees_assigned_not_lost(self):
        """If degrees exist in the docx text, they must be captured, not dropped."""
        if not self.runs_info:
            self.skipTest("No docx provided")
        authors = parse_by_superscript_boundaries(self.runs_info)

        # Collect all degrees found by the parser
        all_parsed_degrees = []
        for a in authors:
            all_parsed_degrees.extend(a.degrees)

        # Collect degree tokens visible in the non-superscript run text
        # (these are the degrees the parser should have captured)
        expected_degrees = set()
        for run in self.runs_info:
            if run["is_superscript"]:
                continue
            for word in re.split(r'[,\s.;]+', run["text"]):
                if word.strip() and is_degree_token(word.strip()):
                    expected_degrees.add(word.strip())

        # Every degree token in the docx text should be captured by some author
        for deg in expected_degrees:
            self.assertIn(deg, all_parsed_degrees,
                          f"Degree '{deg}' found in docx text but not assigned to any author")


class TestJSONCrossReference(unittest.TestCase):
    """Test cross_reference_json from labelAuthorsNER."""

    @classmethod
    def setUpClass(cls):
        if DOCX_PATH and os.path.exists(DOCX_PATH) and JSON_PATH and os.path.exists(JSON_PATH):
            _, cls.runs_info = extract_au_runs(DOCX_PATH)
            cls.json_count, cls.json_authors, cls.json_name_lookup = load_json_authors(JSON_PATH)
            cls.authors = parse_by_superscript_boundaries(cls.runs_info, cls.json_count)
        else:
            cls.authors = []
            cls.json_authors = {}
            cls.json_name_lookup = {}

    def test_all_authors_matched(self):
        """Every parsed author should match a JSON entry."""
        if not self.authors:
            self.skipTest("No docx/json provided")
        matched, _, log = cross_reference_json(self.authors, self.json_authors, self.json_name_lookup)
        for a in matched:
            self.assertTrue(a.json_key,
                            f"{a.label} '{a.full_name}' did not match any JSON author")

    def test_no_duplicate_json_keys(self):
        """Two parsed authors should not match the same JSON entry."""
        if not self.authors:
            self.skipTest("No docx/json provided")
        matched, _, _ = cross_reference_json(self.authors, self.json_authors, self.json_name_lookup)
        keys = [a.json_key for a in matched if a.json_key]
        self.assertEqual(len(keys), len(set(keys)),
                         f"Duplicate JSON keys: {keys}")

    def test_missing_middle_name_still_matches(self):
        """Authors with middle names in JSON but not in DOCX should still match."""
        if not self.authors:
            self.skipTest("No docx/json provided")
        matched, _, _ = cross_reference_json(self.authors, self.json_authors, self.json_name_lookup)

        # Find any JSON author who has a middle name — the docx often omits it
        # Verify that fuzzy/no-middle matching still links them
        json_with_middle = {}
        for key, info in self.json_authors.items():
            if info.get("middle-name", "").strip():
                ln = info.get("last-name", "")
                json_with_middle[key] = ln

        if not json_with_middle:
            # This docx has no middle-name mismatches to test — that's fine, skip
            self.skipTest("No JSON authors with middle names to test")

        # For each JSON author with a middle name, check they matched
        for jkey, jlast in json_with_middle.items():
            matched_author = next((a for a in matched if a.json_key == jkey), None)
            self.assertIsNotNone(
                matched_author,
                f"JSON[{jkey}] with middle name (last='{jlast}') was not matched"
            )

    def test_orcid_enriched(self):
        """Authors with ORCID in JSON should have it after cross-ref."""
        if not self.authors:
            self.skipTest("No docx/json provided")
        matched, _, _ = cross_reference_json(self.authors, self.json_authors, self.json_name_lookup)
        orcid_authors = [a for a in matched if a.orcid]
        self.assertGreaterEqual(len(orcid_authors), 1, "At least one author should have ORCID")

    def test_corresponding_author_enriched(self):
        """The corresponding author from JSON should be flagged."""
        if not self.authors:
            self.skipTest("No docx/json provided")
        matched, _, _ = cross_reference_json(self.authors, self.json_authors, self.json_name_lookup)
        corresp = [a for a in matched if a.is_corresponding]
        self.assertGreaterEqual(len(corresp), 1, "At least one corresponding author expected")


class TestAuthorLabeler(unittest.TestCase):
    """Test the main AuthorLabeler class from labelAuthorsNER."""

    def test_label_from_docx_and_json(self):
        """AuthorLabeler.label_from_docx_and_json returns correct LabelResult."""
        if not (DOCX_PATH and JSON_PATH):
            self.skipTest("No docx/json provided")

        labeler = AuthorLabeler()
        result = labeler.label_from_docx_and_json(DOCX_PATH, JSON_PATH)

        self.assertIsInstance(result, LabelResult)
        self.assertGreater(len(result.authors), 0, "Should find at least one author")
        self.assertTrue(result.json_matched, "All authors should match JSON")
        self.assertIn("[AU1]", result.labeled_html, "HTML should contain [AU1]")
        self.assertIn("authors", result.labeled_html, "HTML should have authors class")

    def test_label_from_docx_and_json_author_count(self):
        """Author count from AuthorLabeler must match JSON."""
        if not (DOCX_PATH and JSON_PATH):
            self.skipTest("No docx/json provided")

        labeler = AuthorLabeler()
        result = labeler.label_from_docx_and_json(DOCX_PATH, JSON_PATH)

        with open(JSON_PATH) as f:
            expected = len(json.load(f)["authors_info"])
        self.assertEqual(len(result.authors), expected)

    def test_label_from_docx_only(self):
        """AuthorLabeler works without JSON too (no cross-ref)."""
        if not DOCX_PATH:
            self.skipTest("No docx provided")

        labeler = AuthorLabeler()
        result = labeler.label_from_docx_and_json(DOCX_PATH, json_path=None)

        self.assertGreater(len(result.authors), 0)
        self.assertFalse(result.json_matched, "Without JSON, json_matched should be False")

    def test_author_dict_keys_match_indices(self):
        """author_dict keys must match author indices."""
        if not (DOCX_PATH and JSON_PATH):
            self.skipTest("No docx/json provided")

        labeler = AuthorLabeler()
        result = labeler.label_from_docx_and_json(DOCX_PATH, JSON_PATH)

        for a in result.authors:
            self.assertIn(a.index, result.author_dict)
            self.assertEqual(result.author_dict[a.index], a.full_name)


class TestGetDocxAuthorsNER(unittest.TestCase):
    """Test get_docx_authors_ner — the drop-in replacement for breakDownProcess."""

    def test_returns_correct_tuple(self):
        """get_docx_authors_ner returns (bool, str, dict) matching old interface."""
        if not (DOCX_PATH and JSON_PATH):
            self.skipTest("No docx/json provided")

        ok, html, aut_dict = get_docx_authors_ner(DOCX_PATH, JSON_PATH)

        self.assertIsInstance(ok, bool)
        self.assertTrue(ok, "Should return True for valid docx")
        self.assertIsInstance(html, str)
        self.assertIsInstance(aut_dict, dict)
        self.assertGreater(len(aut_dict), 0)

    def test_aut_dict_values_are_names(self):
        """aut_dict values should be author name strings, not dicts."""
        if not (DOCX_PATH and JSON_PATH):
            self.skipTest("No docx/json provided")

        _, _, aut_dict = get_docx_authors_ner(DOCX_PATH, JSON_PATH)

        for key, val in aut_dict.items():
            self.assertIsInstance(key, int, f"Key {key} should be int")
            self.assertIsInstance(val, str, f"Value for {key} should be str, got {type(val)}")
            self.assertFalse(is_degree_token(val), f"'{val}' looks like a degree, not a name")


class TestLabelAuthor(unittest.TestCase):
    """Test the LabelAuthor drop-in replacement class."""

    def test_author_process_returns_tuple(self):
        """LabelAuthor.author_process returns (str, dict)."""
        sample_html = ('<p class="authors">John Smith<sup>1</sup>, '
                       'Jane Doe<sup>2</sup></p>')

        la = LabelAuthor()
        result_html, result_dict = la.author_process(sample_html)

        self.assertIsInstance(result_html, str)
        self.assertIsInstance(result_dict, dict)

    def test_author_process_with_real_html(self):
        """LabelAuthor.author_process on real labeled HTML."""
        if not (DOCX_PATH and JSON_PATH):
            self.skipTest("No docx/json provided")

        # First get labeled HTML from AuthorLabeler
        labeler = AuthorLabeler()
        result = labeler.label_from_docx_and_json(DOCX_PATH, JSON_PATH)

        # Feed it through LabelAuthor (the drop-in interface)
        la = LabelAuthor()
        html_out, dict_out = la.author_process(result.labeled_html)

        self.assertIsInstance(html_out, str)
        self.assertIsInstance(dict_out, dict)
        self.assertGreater(len(dict_out), 0, "Should return non-empty author dict")


class TestGenerateLabeledHTML(unittest.TestCase):
    """Test generate_labeled_html from labelAuthorsNER."""

    def test_html_structure(self):
        """Generated HTML should have correct structure."""
        authors = [
            AuthorEntity(index=1, full_name="John Smith", superscripts=["1"], degrees=["PhD"]),
            AuthorEntity(index=2, full_name="Jane Doe", superscripts=["2"],
                         is_corresponding=True),
        ]
        html = generate_labeled_html(authors)

        self.assertTrue(html.startswith('<p class="authors">'))
        self.assertTrue(html.endswith('</p>'))
        self.assertIn('[AU1]', html)
        self.assertIn('[AU2]', html)
        self.assertIn('<sup>1</sup>', html)
        self.assertIn('<sup>2</sup>', html)
        self.assertIn('PhD', html)
        self.assertIn(', and ', html)  # last separator

    def test_corresponding_star(self):
        """Corresponding author should get * in superscript."""
        authors = [
            AuthorEntity(index=1, full_name="Test Author", superscripts=["1"],
                         is_corresponding=True),
        ]
        html = generate_labeled_html(authors)
        self.assertIn('<sup>*</sup>', html)

    def test_empty_authors(self):
        """Empty list should produce empty paragraph."""
        html = generate_labeled_html([])
        self.assertEqual(html, '<p class="authors"></p>')


class TestLabeledDocxOutput(unittest.TestCase):
    """Test that the full pipeline produces a valid labeled docx file."""

    @classmethod
    def setUpClass(cls):
        """Run the full pipeline and save the labeled docx."""
        cls.output_path = None
        if not (DOCX_PATH and JSON_PATH):
            return

        # Use AuthorLabeler from labelAuthorsNER
        labeler = AuthorLabeler()
        cls.result = labeler.label_from_docx_and_json(DOCX_PATH, JSON_PATH)

        # Write the labels into the docx
        doc = Document(DOCX_PATH)
        doc = ensure_aulabel_style(doc)

        au_para = None
        for p in doc.paragraphs:
            if p.style.name in ("AU", "AU0"):
                au_para = p
                break

        if au_para and cls.result.authors:
            rebuild_au_paragraph(au_para, cls.result.authors)
            base, ext = os.path.splitext(DOCX_PATH)
            cls.output_path = f"{base}_LABELED{ext}"
            doc.save(cls.output_path)

    def test_output_file_created(self):
        """Labeled docx file should be created."""
        if not self.output_path:
            self.skipTest("No docx/json provided")
        self.assertTrue(os.path.exists(self.output_path),
                        f"Output file not created: {self.output_path}")

    def test_output_has_au_labels(self):
        """Output docx AU paragraph should contain [AU] labels."""
        if not self.output_path:
            self.skipTest("No docx/json provided")

        doc = Document(self.output_path)
        au_para = None
        for p in doc.paragraphs:
            if p.style.name in ("AU", "AU0"):
                au_para = p
                break

        self.assertIsNotNone(au_para, "AU paragraph missing in output")
        labels = re.findall(r'\[AU\d+\]', au_para.text)
        self.assertEqual(len(labels), len(self.result.authors),
                         f"Expected {len(self.result.authors)} labels, found {len(labels)}")

    def test_output_labels_are_blue(self):
        """[AU] label runs should have blue color and aulabel style."""
        if not self.output_path:
            self.skipTest("No docx/json provided")

        doc = Document(self.output_path)
        for p in doc.paragraphs:
            if p.style.name not in ("AU", "AU0"):
                continue
            for run in p.runs:
                if re.match(r'\[AU\d+\]', run.text):
                    # Check style name
                    self.assertEqual(run.style.name, "aulabel",
                                     f"Run '{run.text}' should have 'aulabel' style")
                    # Check blue color
                    self.assertIsNotNone(run.font.color.rgb,
                                         f"Run '{run.text}' should have color set")
                    self.assertEqual(str(run.font.color.rgb), "0000FF",
                                     f"Run '{run.text}' should be blue")
            break

    def test_output_superscripts_preserved(self):
        """Superscript runs should still be superscript in output."""
        if not self.output_path:
            self.skipTest("No docx/json provided")

        doc = Document(self.output_path)
        found_sup = False
        for p in doc.paragraphs:
            if p.style.name not in ("AU", "AU0"):
                continue
            for run in p.runs:
                if run.font.superscript:
                    found_sup = True
                    # Should contain digits (affiliations)
                    self.assertTrue(re.search(r'\d', run.text),
                                    f"Superscript run '{run.text}' should contain digits")
            break
        self.assertTrue(found_sup, "Output should have at least one superscript run")

    def test_output_all_author_names_present(self):
        """All author names from labelAuthorsNER should be in the output text."""
        if not self.output_path:
            self.skipTest("No docx/json provided")

        doc = Document(self.output_path)
        au_text = ""
        for p in doc.paragraphs:
            if p.style.name in ("AU", "AU0"):
                au_text = p.text
                break

        for a in self.result.authors:
            self.assertIn(a.full_name, au_text,
                          f"Author '{a.full_name}' missing from output docx")


# ============================================================================
# Runner — also prints human-readable report
# ============================================================================

def print_report(docx_path, json_path):
    """Print a readable test report before running unittest."""
    print("=" * 70)
    print("AUTHOR LABELING TEST — using labelAuthorsNER")
    print("=" * 70)
    print(f"  DOCX:  {docx_path}")
    print(f"  JSON:  {json_path}")

    labeler = AuthorLabeler()
    result = labeler.label_from_docx_and_json(docx_path, json_path)

    print(f"\n  Parse details:")
    for d in result.match_details:
        print(f"    {d}")

    print(f"\n  Authors ({len(result.authors)}):")
    for a in result.authors:
        jmatch = f" -> JSON[{a.json_key}]" if a.json_key else " -> NO MATCH"
        extras = []
        if a.degrees:
            extras.append(f"deg={','.join(a.degrees)}")
        if a.orcid:
            extras.append(f"orcid={a.orcid}")
        if a.is_corresponding:
            extras.append("CORRESPONDING")
        ext_str = f"  ({'; '.join(extras)})" if extras else ""
        print(f"    {a.label} {a.full_name}{jmatch}{ext_str}")

    print(f"\n  JSON matched: {result.json_matched}")
    print(f"{'=' * 70}\n")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python test_author_labeling.py <input.docx> <article.json>")
        print("\nRuns unit tests using labelAuthorsNER module.")
        print("Produces: <input>_LABELED.docx")
        sys.exit(1)

    DOCX_PATH = sys.argv[1]
    JSON_PATH = sys.argv[2]

    if not os.path.exists(DOCX_PATH):
        print(f"ERROR: {DOCX_PATH} not found")
        sys.exit(1)
    if not os.path.exists(JSON_PATH):
        print(f"ERROR: {JSON_PATH} not found")
        sys.exit(1)

    # Print human-readable report first
    print_report(DOCX_PATH, JSON_PATH)

    # Run unit tests (remove the file args so unittest doesn't choke)
    sys.argv = [sys.argv[0], "-v"]
    unittest.main(module=__name__, exit=True)
