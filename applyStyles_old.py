from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from getAppPath import getapppath
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from loadconfig import getconfig
from docx.oxml.table import CT_Row, CT_Tc
import yaml
import re
import os
import json
import docx
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import zipfile
from TransformXmlJar import XmlTransformJar


class ApplyStyles:
    def __init__(self):
        self.configFolder, self.breakDownConfig = getconfig()
        yaml_file = os.path.join(self.configFolder, "config\\paraStyles.yaml")
        with open(yaml_file, "r") as stream:
            self.style_details = yaml.safe_load(stream)
        self.para_styles = self.style_details['Paragraph_Styles']
        self.style_mapping = self.style_details['Parastyle_mapping']
        breakdown_sequence = os.path.join(self.configFolder, 'config\\breakdownSequence.json')
        with open(breakdown_sequence, "r") as breakdown_styles:
            self.breakdownJson = json.loads(breakdown_styles.read())
            self.breakdownStyles = self.breakdownJson['breakdownStyles']
        backmatter_json_file = os.path.join(self.configFolder, 'config\\backMatterTitles.json')
        with open(backmatter_json_file, "r") as backmatter_json:
            self.backmatterTitles = json.loads(backmatter_json.read())
        self.template_path = 'SupportingFiles/SAGE_styles.docx'  # Adjust the path to your .dot file
        self.template = Document(self.template_path)

    def apply_styles(self, docxfile):
        document = Document(docxfile)
        document = self.add_breakdown_styles(document)
        document = self.append_style(document)
        document = self.update_paras(document)
        document = self.bm_styles(document)
        document = self.update_table_styles(document)
        document = self.remove_formatting_from_styles(document)
        document = self.remove_rpr_styles(document)
        document = self.clean_heading_formats(document)
        # document = self.remove_run_properties(document)
        document.save(docxfile)
        # document.save('V:\FOR_BREAKDOWN\ParaStyler_INPUT\SAGE\CPJ_1227087\CPJ_1227087_CLN_AS1.docx')

    def clean_heading_formats(self, document):
        for paragraph in document.paragraphs:
            style_val = paragraph.style.name
            if paragraph._p.pPr is not None:
                if style_val in ["AU0", "AU", "H1", "EH", "EH0", "Reference_Title", "Reference_Title0", "ABKWH", "ABKWH0"]:
                    # Find the <w:rPr> element
                    pPr = paragraph._p.pPr
                    rPr = pPr.xpath('./w:rPr')
                    if rPr is not None:
                        for child in rPr:
                            if child.tag.endswith('b') or child.tag.endswith('bCs'):
                                rPr.remove(child)
                    for run in paragraph.runs:
                        rPr = run._element.rPr
                        if rPr is not None:
                            for child in rPr:
                                if child.tag.endswith('b') or child.tag.endswith('bCs'):
                                    rPr.remove(child)
        return document

    def remove_rpr_styles(self, document):
        for paragraph in document.paragraphs:
            if paragraph._p.pPr is not None:
                pPr = paragraph._p.pPr
                rPr = pPr.xpath('./w:rPr')
                # Find all <w:sz> and <w:szCs> elements
                if rPr is not None:
                    for child in rPr:
                        if child.tag.endswith('sz') or child.tag.endswith('szCs'):
                            rPr.remove(child)
        for paragraph in document.paragraphs:
            pPr = paragraph._element.pPr
            if pPr is not None:
                for child in pPr:
                    # Specify the tags of the child elements you want to remove
                    if child.tag.endswith('rPr'):
                        pPr.remove(child)
            for run in paragraph.runs:
                rPr = run._element.rPr
                if rPr is not None:
                    for child in rPr:
                        if child.tag.endswith('sz') or child.tag.endswith('szCs') or \
                                child.tag.endswith('shd') or \
                                child.tag.endswith('rStyle') or \
                                child.tag.endswith('color') or \
                                child.tag.endswith('u') or \
                                child.tag.endswith('spacing') or \
                                child.tag.endswith('textOutline') or \
                                child.tag.endswith('position'):
                            rPr.remove(child)
        #                                 child.tag.endswith('rFonts') or \
        return document

    def remove_formatting_from_styles(self, document):
        target_styles = ['H1', 'H2', 'H3', 'H4', 'ABKH', 'EH', "TY"]
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        for paragraph in document.paragraphs:
            # Get the paragraph's style
            style = paragraph.style.name.lower()
            # Check if the paragraph's style is in the target_styles list
            if style.lower() in [s.lower() for s in target_styles]:
                p_element = paragraph._element
                # Find all <w:r> elements within the paragraph
                r_elements = p_element.findall('.//w:r', namespaces=nsmap)
                # If there is only one <w:r> element in the paragraph
                if len(r_elements) == 1:
                    r_element = r_elements[0]
                    # Find all <w:rPr> elements within the <w:r> element
                    rPr_elements = r_element.findall('.//w:rPr', namespaces=nsmap)
                    # If there is only one <w:rPr> element in the <w:r> element
                    if len(rPr_elements) == 1:
                        rPr_element = rPr_elements[0]
                        # Find and remove <w:b> and <w:bCs> child elements from <w:rPr>
                        for child in rPr_element.findall('.//w:b', namespaces=nsmap):
                            rPr_element.remove(child)
                        for child in rPr_element.findall('.//w:bCs', namespaces=nsmap):
                            rPr_element.remove(child)
                        for child in rPr_element.findall('.//w:i', namespaces=nsmap):
                            rPr_element.remove(child)
                        for child in rPr_element.findall('.//w:iCs', namespaces=nsmap):
                            rPr_element.remove(child)
        return document

    def update_paras(self, document):
        para_count = 0
        for paragraph in document.paragraphs:
            style_name = paragraph.style.name
            styles = document.styles
            if style_name == "TEXT IND":
                if para_count > 0:
                    pre_count = para_count - 1
                    pre_para = document.paragraphs[pre_count]
                    pre_style = pre_para.style.name
                    if re.search("H1|H2|H3|H4|H5|H6", pre_style, re.IGNORECASE):
                        paragraph.style = styles['TEXT']
            # if style_name == "ABKW":
            #     self.remove_box(paragraph)
            para_count = para_count + 1
        return document

    def apply_styles_old(self, docxfile):
        new_styles = []
        document = Document(docxfile)
        document = self.add_breakdown_styles(document)
        styles = document.styles
        # latent_styles = document.styles.latent_styles
        # for table in document.tables:
        #     self.modifyBorder(table)
        #     table.allow_autofit = True
        para_count = 0
        if "TEXT" not in styles:
            document = self.create_style(document, "TEXT")
        for paragraph in document.paragraphs:
            style_name = paragraph.style.name
            document, new_style = self.append_style(document, style_name)
            styles = document.styles
            if new_style is not None:
                if new_style == "TEXT IND":
                    if para_count > 0:
                        pre_count = para_count - 1
                        pre_para = document.paragraphs[pre_count]
                        pre_style = pre_para.style.name
                        if re.search("H1|H2|H3|H4|H5|H6", pre_style, re.IGNORECASE):
                            paragraph.style = styles['TEXT']
                        else:
                            paragraph.style = styles['TEXT IND']
                    else:
                        paragraph.style = styles[new_style]
                else:
                    paragraph.style = styles[new_style]
                new_styles.append(new_style)
            if new_style == "ABKW":
                self.remove_box(paragraph)
            para_count = para_count + 1
        all_styles = self.para_styles.keys()
        pending_styles = [e for e in all_styles if e not in new_styles]
        for pending_style in pending_styles:
            #self.append_style(document, pending_style)
            exstyles = document.styles
            if pending_style not in exstyles:
                document = self.create_style(document, pending_style)
        #newdocx = re.sub(".docx", "_new.docx", docxfile)
        document = self.bm_styles(document)
        document = self.update_table_styles(document)
        document.save(docxfile)

    def load_styles(self, document):
        para_keys = self.para_styles
        styles = document.styles
        for br_list in para_keys:
            if br_list not in styles:
                document = self.create_style(document, br_list)

    def add_breakdown_styles(self, document):
        breakdown_keys = self.breakdownStyles.keys()
        breakdown_values = self.breakdownStyles.values()
        styles = document.styles
        for br_list in breakdown_keys:
            if br_list not in styles:
                document = self.create_style(document, br_list)
        for br_list in breakdown_values:
            if br_list not in styles:
                document = self.create_style(document, br_list)
        return document

    def append_style(self, document):
        para_mapping = self.style_mapping
        for paragraph in document.paragraphs:
            current_style = paragraph.style.name
            if current_style in para_mapping:
                new_style = para_mapping[current_style]
                if new_style in document.styles:
                    # base_style_name = "Normal_" + current_style
                    # if base_style_name not in document.styles:
                    #     base_st = document.styles.add_style(base_style_name, WD_STYLE_TYPE.PARAGRAPH)
                    #     document.styles[base_style_name].base_style = document.styles["Normal"]
                    #     document.styles[new_style].base_style = document.styles[base_style_name]
                    #     font = base_st.font
                    #     font.name = 'Times New Roman'
                    #     font.size = Pt(12)
                    #     font.color.rgb = RGBColor(0, 0, 0)
                    #     para = base_st.paragraph_format
                    #     para.first_line_indent = None
                    #     para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    #     para.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    #     para.left_indent = Pt(0)
                    #     para.right_indent = Pt(0)
                    #     para.space_after = Pt(0)
                    #     para.space_before = Pt(0)
                    paragraph.style = new_style
            else:
                pass
        return document

    def append_style_old(self, document, current_style):
        para_dic = self.para_styles
        para_mapping = self.style_mapping
        latent_styles = document.styles.latent_styles
        new_style = None
        if current_style in para_mapping:
            new_style = para_mapping[current_style]
            if new_style is not None:
                para_style = para_dic[new_style]
                try:
                    style_exist = document.styles[new_style]
                except Exception as e:
                    styles = document.styles
                    try:
                        styles[new_style].delete()
                    except Exception as e:
                        pass
                    cr_name = "Normal_" + current_style
                    cr_style = document.styles.add_style(cr_name, WD_STYLE_TYPE.PARAGRAPH)
                    cr_style.base_style = document.styles["Normal"]
                    style = document.styles.add_style(new_style, WD_STYLE_TYPE.PARAGRAPH)
                    style.base_style = document.styles[cr_name]
                    #font = style.font
                    font_name = para_style['font']['name']
                    font_size = int(para_style['font']['size'])
                    font_color = para_style['font']['color'].split(",")
                    red = int(font_color[0])
                    green = int(font_color[1])
                    blue = int(font_color[2])
                    style.font.name = font_name
                    style.font.size = Pt(font_size)
                    style.font.color.rgb = RGBColor(red, green, blue)
                    if para_style['font']['bold'] == True:
                        style.font.bold = True
                    if para_style['font']['italic'] == True:
                        style.font.italic = True
                    if para_style['font']['smallcaps'] == True:
                        style.font.small_caps = True
                    para_format = style.paragraph_format
                    if para_style['paragraph']['alignment'] == "LEFT":
                        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif para_style['paragraph']['alignment'] == "RIGHT":
                        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    elif para_style['paragraph']['alignment'] == "CENTER":
                        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    if para_style['paragraph']['line_spacing_rule'] == "ONE_POINT_FIVE":
                        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    elif para_style['paragraph']['line_spacing_rule'] == "SINGLE":
                        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    elif para_style['paragraph']['line_spacing_rule'] == "DOUBLE":
                        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                    below_space = int(para_style['paragraph']['below_space'])
                    above_space = int(para_style['paragraph']['above_space'])
                    line_indent = int(para_style['paragraph']['first_line_indent'])
                    left_indent = int(para_style['paragraph']['left_indent'])
                    right_indent = int(para_style['paragraph']['right_indent'])
                    if line_indent == 0:
                        style.paragraph_format.first_line_indent = None
                    else:
                        style.paragraph_format.first_line_indent = Pt(line_indent)
                    if left_indent == 0:
                        style.paragraph_format.left_indent = None
                    else:
                        style.paragraph_format.left_indent = Pt(left_indent)
                    if right_indent == 0:
                        style.paragraph_format.right_indent = None
                    else:
                        style.paragraph_format.right_indent = Pt(line_indent)
                    if new_style == "ABKW":
                        style.paragraph_format.left_indent = Pt(0)
                        style.paragraph_format.right_indent = Pt(0)
                    style.paragraph_format.space_after = Pt(below_space)
                    style.paragraph_format.space_before = Pt(above_space)
        return document, new_style

    def clear_styles(self, docxfile):
        document = Document(docxfile)
        styles = document.styles
        styles["NormalPara"].delete()
        style = document.styles.add_style('NormalPara', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
        font = style.font
        font.name = 'Times New Roman'
        font.color.rgb = RGBColor(0, 0, 0)
        para = style.paragraph_format
        para.first_line_indent = None
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para.first_line_indent = None
        para.left_indent = Pt(0)
        para.right_indent = Pt(0)
        para.space_after = Pt(0)
        para.space_before = Pt(0)
        for table in document.tables:
            try:
                self.modifyBorder(table)
            except Exception as e:
                pass
        for paragraph in document.paragraphs:
            try:
                paragraph.style = document.styles['Normal']
            except Exception as e:
                pass
        newdocx = re.sub(".docx", "_new.docx", docxfile)
        document.save(newdocx)

    def create_character_style(self, document, style_name):
        my_styles = document.styles
        style = document.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
        style.base_style = my_styles['Default Paragraph Font']
        font = style.font
        font.color.rgb = RGBColor(255, 0, 0)
        font.bold = True
        font.superscript = False
        return document

    def create_style(self, document, style_name):
        temp_style, style_object = self.find_style_in_template(style_name)
        if temp_style:
            # new_style = document.styles.add_style(temp_style, WD_STYLE_TYPE.PARAGRAPH)
            # Copy the formatting properties from the source style to the target style
            # new_style._element.clear()
            # for element in style_object._element:
            #     new_style._element.append(element)
            new_style = document.styles.add_style(style_object.name, style_object.type)
            new_style.base_style = style_object.base_style
            if style_object.font.name is not None:
                new_style.font.name = style_object.font.name
            elif style_object.base_style.font.name is not None:
                new_style.font.name = style_object.base_style.font.name
            new_style.font.size = style_object.font.size
            new_style.font.color.rgb = style_object.font.color.rgb
            temp_paragraph_format = style_object.paragraph_format
            paragraph_format = new_style.paragraph_format
            paragraph_format.alignment = temp_paragraph_format.alignment
            paragraph_format.first_line_indent = temp_paragraph_format.first_line_indent
            paragraph_format.keep_together = temp_paragraph_format.keep_together
            paragraph_format.keep_with_next = temp_paragraph_format.keep_with_next
            paragraph_format.left_indent = temp_paragraph_format.left_indent
            paragraph_format.line_spacing = temp_paragraph_format.line_spacing
            paragraph_format.line_spacing_rule = temp_paragraph_format.line_spacing_rule
            paragraph_format.page_break_before = temp_paragraph_format.page_break_before
            paragraph_format.right_indent = temp_paragraph_format.right_indent
            paragraph_format.space_after = temp_paragraph_format.space_after
            paragraph_format.space_before = temp_paragraph_format.space_before
            paragraph_format.widow_control = temp_paragraph_format.widow_control
        else:
            if style_name in self.para_styles:
                para_style = self.para_styles[style_name]
            else:
                para_style = self.para_styles['default']
            style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = document.styles["Normal"]
            font = style.font
            font_name = para_style['font']['name']
            font_size = int(para_style['font']['size'])
            font_color = para_style['font']['color'].split(",")
            red = int(font_color[0])
            green = int(font_color[1])
            blue = int(font_color[2])
            font.name = font_name
            font.size = Pt(font_size)
            font.color.rgb = RGBColor(red, green, blue)
            if para_style['font']['bold'] == True:
                font.bold = True
            if para_style['font']['italic'] == True:
                font.italic = True
            if para_style['font']['smallcaps'] == True:
                font.small_caps = True
            if para_style['font']['allcaps'] == True:
                font.all_caps = True
            try:
                h_color = para_style['font']['highlight_color']
            except Exception as e:
                print(e)
            if para_style['font']['highlight_color'] is not False:
                hl_color = para_style['font']['highlight_color']
                if hl_color == "YELLOW":
                    font.highlight_color = WD_COLOR_INDEX.YELLOW
            paragraph_format = style.paragraph_format
            if para_style['paragraph']['alignment'] == "LEFT":
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif para_style['paragraph']['alignment'] == "RIGHT":
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif para_style['paragraph']['alignment'] == "CENTER":
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if para_style['paragraph']['line_spacing_rule'] == "ONE_POINT_FIVE":
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            elif para_style['paragraph']['line_spacing_rule'] == "SINGLE":
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            elif para_style['paragraph']['line_spacing_rule'] == "DOUBLE":
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            below_space = int(para_style['paragraph']['below_space'])
            above_space = int(para_style['paragraph']['above_space'])
            line_indent = int(para_style['paragraph']['first_line_indent'])
            left_indent = int(para_style['paragraph']['left_indent'])
            right_indent = int(para_style['paragraph']['right_indent'])
            if line_indent == 0:
                paragraph_format.first_line_indent = None
            else:
                paragraph_format.first_line_indent = Pt(line_indent)
            if left_indent == 0:
                paragraph_format.left_indent = None
            else:
                paragraph_format.left_indent = Pt(left_indent)
            if right_indent == 0:
                paragraph_format.right_indent = None
            else:
                paragraph_format.right_indent = Pt(line_indent)
            paragraph_format.space_after = Pt(below_space)
            paragraph_format.space_before = Pt(above_space)
        return document

    def modifyBorder(self, table):
        tbl = table._tbl
        for cell in tbl.iter_tcs():
            tcPr = cell.tcPr  # get tcPr element, in which we can define style of borders
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '4')
            top.set(qn('w:space'), '0')
            top.set(qn('w:color'), 'auto')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '4')
            left.set(qn('w:space'), '0')
            left.set(qn('w:color'), 'auto')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), 'auto')
            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'single')
            right.set(qn('w:sz'), '4')
            right.set(qn('w:space'), '0')
            right.set(qn('w:color'), 'auto')
            tcBorders.append(top)
            tcBorders.append(left)
            tcBorders.append(bottom)
            tcBorders.append(right)
            tcPr.append(tcBorders)

    def remove_box(self, paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr,
                                  'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                                  'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                                  'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                                  'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                                  'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                                  'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                                  'w:pPrChange'
                                  )
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'none')
        top.set(qn('w:sz'), '0')
        top.set(qn('w:space'), '0')
        top.set(qn('w:color'), 'auto')
        pBdr.append(top)
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'none')
        left.set(qn('w:sz'), '0')
        left.set(qn('w:space'), '0')
        left.set(qn('w:color'), 'auto')
        pBdr.append(left)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'none')
        bottom.set(qn('w:sz'), '0')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'none')
        right.set(qn('w:sz'), '0')
        right.set(qn('w:space'), '0')
        right.set(qn('w:color'), 'auto')
        pBdr.append(right)

    def find_duplicates(self, docxfile):
        document = Document(docxfile)
        styles = document.styles
        try:
            styles["Duplicate"].delete()
        except Exception as e:
            pass
        style = document.styles.add_style('Duplicate', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.color.rgb = RGBColor(255, 0, 0)
        font.highlight_color = WD_COLOR_INDEX.YELLOW
        para = style.paragraph_format
        para.first_line_indent = None
        para.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para.first_line_indent = None
        para.left_indent = Pt(0)
        para.right_indent = Pt(0)
        para.space_after = Pt(0)
        para.space_before = Pt(0)
        para_count = 0
        para_dic = {}
        lines = document.paragraphs
        # for line in lines:
        #     if len(line.text) == 0:
        #         para_object = len(line._p.xpath("w:r"))
        #         if para_object == 0:
        #             self.delete_paragraph(line)
        for paragraph in document.paragraphs:
            align = str(paragraph.alignment)
            if re.search("CENTER", align, re.IGNORECASE):
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif re.search("JUSTIFY", align, re.IGNORECASE):
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif re.search("RIGHT", align, re.IGNORECASE):
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif re.search("None", align, re.IGNORECASE):
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            empty_paras = []
            try:
                text_length = len(paragraph.text)
                para_object = len(paragraph._p.xpath("w:r"))
                if text_length == 0 and para_object == 0:
                   empty_paras.append(para_count)
                else:
                    para_dic[para_count] = paragraph.text
                para_count = para_count + 1
            except Exception as e:
                pass
        flipped = {}
        for key, value in para_dic.items():
            if value not in flipped:
                flipped[value] = [key]
            else:
                flipped[value].append(key)
        for flip_key in flipped.keys():
            flip_val = flipped[flip_key]
            if len(flip_val) > 1:
                flip_val.pop(0)
                for flip_value in flip_val:
                    duplicate_para = document.paragraphs[flip_value]
                    duplicate_para.style = document.styles["Duplicate"]
        document.save(docxfile)

    def reapply_styles(self, docxfile):
        document = Document(docxfile)
        for paragraph in document.paragraphs:
            style_name = paragraph.style.name
            paragraph.style = document.styles[style_name]
        document.save(docxfile)

    def delete_paragraph(self, paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    def bm_styles(self, document):
        styles = document.styles
        if "EH" not in styles:
            document = self.create_style(document, "EH")
        if "AN" not in styles:
            document = self.create_style(document, "AN")
        paragraphs = document.paragraphs
        para_count = 0
        for paragraph in paragraphs:
            paratext = paragraph.text
            paratext = re.sub(r"(\; |\;|\: |\:| |\. |\.| )$", "", paratext)
            if paratext in self.backmatterTitles:
                paragraph.style = "EH"
                next_para_count = para_count + 1
                next_para = paragraphs[next_para_count]
                next_para_style = next_para.style.name
                if next_para_style == "TEXT" or next_para_style == "TEXT IND":
                    next_para.style = "AN"
            para_count = para_count + 1
        return document

    def check_normal_styles(self, docxfile):
        document = Document(docxfile)
        for paragraph in document.paragraphs:
            style_name = paragraph.style.name
            pattern = "Normal"
            if re.match(pattern, style_name):
                para_text = paragraph.text
                para_xml = paragraph._p.xml
                if len(para_text) > 0:
                    cp_pattern = r"^(Table|Figure|Fig|TABLE|FIGURE|FIG)(\.| |)([0-9]+)(\.|\:|) [A-Z]"
                    text_ind_pattern = "^[A-Z].*"
                    text_pattern = "^[a-z0-9].*"
                    if re.match(cp_pattern, para_text):
                        paragraph.style = "CP"
                    elif re.match(text_ind_pattern, para_text):
                        paragraph.style = "TEXT IND"
                    else:
                        if re.search(r"ProgID\=\"Equation", para_xml):
                            if not re.match(text_pattern, para_text):
                                paragraph.style = "EQ"
                        else:
                            paragraph.style = "TEXT"
                elif len(para_text) == 0:
                    if re.search(r"ProgID\=\"Equation", para_xml):
                        paragraph.style = "EQ"
        document.save(docxfile)

    def cross_check_styles(self, docxfile):
        document = Document(docxfile)
        for paragraph in document.paragraphs:
            para_text = paragraph.text
            if len(para_text) > 0:
                regex_pattern = r"^(Table|Figure|Fig|TABLE|FIGURE|FIG)(\.| |)([0-9]+)(\.|\:|) [A-Z]"
                para_style = paragraph.style.name
                if re.match(regex_pattern, para_text):
                    if not re.match("CP", para_style):
                        paragraph.style = "CP"
        document.save(docxfile)

    def insert_para_after(self, paragraph, text, style=None):
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        if style is not None:
            new_para.style = style
        return new_para

    def insert_para_before(self, paragraph, text, pstyle):
        paragraph.insert_paragraph_before(text, style=pstyle)

    def update_table_styles_old(self, document):
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_style = paragraph.style.name
                        if paragraph_style == "tablehead":
                            paragraph.style = document.styles['TCH']
                        elif paragraph_style == "tablebody":
                            paragraph.style = document.styles['TT']
        return document

    from docx.text.paragraph import Paragraph

    def update_table_styles(self, document):
        """
        Table style updater with belt-and-suspenders protection.

        - Fast path: python-docx row.cells (for simple tables)
        - Safe path: XML traversal (for merged / complex tables)
        """

        # ---------- FAST PATH ----------
        try:
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph_style = paragraph.style.name
                            if paragraph_style == "tablehead":
                                paragraph.style = document.styles['TCH']
                            elif paragraph_style == "tablebody":
                                paragraph.style = document.styles['TT']
            return document

        except ValueError:
            # Triggers on: merged cells, gridSpan, header rows, etc.
            pass

        for table in document.tables:
            tbl = table._tbl
            nsmap = tbl.nsmap

            for tr in tbl.findall('.//w:tr', namespaces=nsmap):
                for tc in tr.findall('.//w:tc', namespaces=nsmap):
                    for p in tc.findall('.//w:p', namespaces=nsmap):
                        paragraph = Paragraph(p, table._parent)

                        style_name = paragraph.style.name
                        if style_name == "tablehead":
                            paragraph.style = document.styles['TCH']
                        elif style_name == "tablebody":
                            paragraph.style = document.styles['TT']

        return document

    def get_table_header_old(self, document):
        for table in document.tables:
            span_count = 0
            for row in table.rows:
                row_xml = row._tr.xml
                if re.search(r"(w\:gridSpan)", row_xml, re.IGNORECASE):
                    span_count = span_count + 1
        return span_count

    def update_table_styles_old(self, document, header_count):
        styles = document.styles
        if "TCH" not in styles:
            document = self.create_style(document, "TCH")
        if "TT" not in styles:
            document = self.create_style(document, "TT")
        styles = document.styles
        for table in document.tables:
            row_count = 0
            for row in table.rows:
                if row_count <= header_count:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.style = styles['TCH']
                else:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.style = styles['TT']
                    row_count = row_count + 1
                row_count = row_count + 1
        return document

    def find_style_in_template(self, style_name):
        for style in self.template.styles:
            if style.name == style_name:
                return style_name, style
        return None, None

    def as_post_clean(self, docx_file_path):
        with zipfile.ZipFile(docx_file_path, 'r') as zip_ref:
            zip_ref.extract('word/document.xml', 'temp')
        trans_xml = XmlTransformJar()
        trans_xml.trans_xml("temp/word/document.xml", "xsl/postAs.xsl", "temp/word/document.xml")
        with open("temp/word/document.xml", "r", encoding="utf-8") as file:
            document_xml = file.read()
        document_xml = re.sub("（", "(", document_xml)
        document_xml = re.sub("）", ")", document_xml)
        document_xml = re.sub(
            r'(<m:t( xml:space="preserve"|)> (\(|\[)([0-9]+|[0-9]+[a-z]|[a-z][0-9]+)(\)|\])</m:t></m:r></m:oMath>(</m:oMathPara>|))',
            r'</m:r></m:oMath>\6<w:r><w:t xml:space="preserve"> \3\4\5</w:t></w:r>',
            document_xml,
            flags=re.I)
        with open("temp/word/document.xml", "w", encoding="utf-8") as file:
            file.write(document_xml)
        with zipfile.ZipFile(docx_file_path, 'r') as zip_read:
            with zipfile.ZipFile(docx_file_path + '.tmp', 'w') as zip_write:
                for item in zip_read.infolist():
                    if item.filename != 'word/document.xml':
                        buffer = zip_read.read(item.filename)
                        zip_write.writestr(item, buffer)
        with zipfile.ZipFile(docx_file_path + '.tmp', 'a') as zip_ref:
            zip_ref.write('temp/word/document.xml', 'word/document.xml')
        os.remove('temp/word/document.xml')
        os.replace(docx_file_path + '.tmp', docx_file_path)


# applyStyles = ApplyStyles()
# applyStyles.apply_styles("V:\\FOR_BREAKDOWN\\ParaStyler_INPUT\SAGE\\JIV_1395658\\JIV_1395658_CLN_AS.docx")
# applyStyles.as_post_clean("V:\FOR_BREAKDOWN\BreakDown_INPUT\SAGE\PIC_1262898\docs\PIC_1262898_CLN_AS.docx")
# document = Document('V:\FOR_BREAKDOWN\BreakDown_INPUT\SAGE\PHR_1251982\PHR_1251982_CLN.docx')
# applyStyles.clean_heading_formats(document)
# document.save('V:\FOR_BREAKDOWN\BreakDown_INPUT\SAGE\PHR_1251982\PHR_1251982.docx')
# applyStyles.remove_bolding_from_styles(document)
# # # # header_count = applyStyles.get_table_header(document)
# # # # document = applyStyles.update_table_styles(document, header_count)
# document.save('V:\FOR_BREAKDOWN\BreakDown_INPUT\SAGE\FAS_1251903\FAS_1251903.docx')
# # # # document.save('V:\FOR_BREAKDOWN\BreakDown_INPUT\SAGE\BRQ_1162872\BRQ_1162872_CLN_AS1.docx')