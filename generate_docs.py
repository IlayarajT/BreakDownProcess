"""
generate_docs.py  —  BreakDown Documentation Generator
=======================================================
Produces:  BreakDown_Documentation.docx

Sections
    1. Title Page
    2. Table of Contents
    3. System Overview
    4. End-to-End Pipeline Workflow Diagram  (visual table-based flowchart)
    5. Component Architecture                (5-layer description + ASCII)
    6. Deployment Architecture               (topology + folder tree)
    7. Installation Guide
    8. Configuration Reference
    9. Batch File Reference
   10. Troubleshooting
   11. Appendix  (JAR inventory, file inventory)

Run:
    python generate_docs.py
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.table import CT_Tbl
import copy

# ──────────────────────────────────────────────────────────────────────────────
#  Colour palette  (R, G, B)
# ──────────────────────────────────────────────────────────────────────────────
CLR = {
    "title_bg":    (0,   70,  127),   # dark navy
    "title_fg":    (255, 255, 255),
    "stage_blue":  (0,   112, 192),   # process boxes
    "stage_green": (0,   176, 80),    # success / output
    "stage_red":   (192, 0,   0),     # error / warning
    "stage_orange":(255, 153, 0),     # supporting
    "stage_purple":(112, 48,  160),   # java / external
    "stage_teal":  (0,   176, 170),   # config
    "arrow_bg":    (166, 166, 166),   # arrow cells
    "head_bg":     (68,  114, 196),   # section heading bg
    "head_fg":     (255, 255, 255),
    "row_alt":     (217, 225, 242),   # alternate table row
    "row_hdr":     (68,  114, 196),
    "border":      (89,  89,  89),
    "code_bg":     (242, 242, 242),
    "white":       (255, 255, 255),
    "black":       (0,   0,   0),
    "light_blue":  (189, 215, 238),
    "light_green": (198, 239, 206),
    "light_orange":(255, 235, 156),
    "light_gray":  (247, 247, 247),
}


# ══════════════════════════════════════════════════════════════════════════════
#  LOW-LEVEL XML HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _rgb_hex(r, g, b):
    return f"{r:02X}{g:02X}{b:02X}"


def set_cell_bg(cell, rgb_tuple):
    r, g, b = rgb_tuple
    hex_color = _rgb_hex(r, g, b)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    # remove existing shd if any
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 6)))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(el)


def set_table_no_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBorders)


def set_table_width(table, width_pct=100):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(width_pct * 50))
    tblW.set(qn("w:type"), "pct")


def cell_para(cell, text, bold=False, italic=False, font_size=9,
              color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
              space_before=0, space_after=0):
    """Replace all paragraphs in cell with a single formatted paragraph."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p


def merge_cells_horizontal(table, row_idx, start_col, end_col):
    """Merge cells in a row from start_col to end_col (inclusive)."""
    row = table.rows[row_idx]
    cell = row.cells[start_col]
    for col in range(start_col + 1, end_col + 1):
        cell = cell.merge(row.cells[col])
    return cell


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    return p


def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")

    class _FakeBr:
        _r = br
        def __init__(self):
            pass

    # python-docx run.add_break() signature
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*CLR["title_bg"])
    return p


def add_body(doc, text, size=10, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_code_block(doc, text):
    """Add a monospace code / ASCII-art block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # light gray shading via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _rgb_hex(*CLR["code_bg"]))
    pPr.append(shd)

    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_info_box(doc, text, bg=CLR["light_blue"], font_size=9):
    """Single-cell table used as a coloured info/note box."""
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, 100)
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    doc.add_paragraph()
    return table


def add_simple_table(doc, headers, rows, col_widths=None):
    """Add a styled data table."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    set_table_width(table, 100)

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        set_cell_bg(c, CLR["row_hdr"])
        cell_para(c, h, bold=True, font_size=9, color=CLR["white"],
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ri, row_data in enumerate(rows):
        bg = CLR["white"] if ri % 2 == 0 else CLR["row_alt"]
        for ci, val in enumerate(row_data):
            c = table.rows[ri + 1].cells[ci]
            set_cell_bg(c, bg)
            cell_para(c, str(val), font_size=9)

    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)

    doc.add_paragraph()
    return table


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

def build_title_page(doc):
    # Large colour band
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, 100)
    cell = table.cell(0, 0)
    set_cell_bg(cell, CLR["title_bg"])
    cell.paragraphs[0].clear()

    for line, size, bold in [
        ("BreakDown", 36, True),
        ("Automated Manuscript Processing Pipeline", 18, False),
        ("", 10, False),
        ("Architecture & Deployment Documentation", 14, False),
        ("", 10, False),
        ("Version 1.2.5  ·  Build 20260513  ·  May 2026", 11, False),
    ]:
        if line == "":
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
        else:
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p2.add_run(line)
            r.font.size = Pt(size)
            r.bold = bold
            r.font.color.rgb = RGBColor(*CLR["title_fg"])
            p2.paragraph_format.space_after = Pt(2)

    cell.paragraphs[0].paragraph_format.space_before = Pt(30)
    cell.paragraphs[-1].paragraph_format.space_after = Pt(30)

    doc.add_paragraph()

    # Subtitle line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SAGE Publications  ·  Internal Technical Reference")
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = RGBColor(*CLR["title_bg"])
    doc.add_paragraph().add_run("").font.size = Pt(6)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

def build_system_overview(doc):
    add_heading(doc, "1. System Overview", 1)

    add_body(doc,
        "BreakDown is a Windows-based automated manuscript processing pipeline for SAGE Publications. "
        "It accepts raw author submission packages (ZIP/TAR/RAR/7Z) deposited into a watched network hotfolder "
        "and transforms them into fully styled Word documents ready for typesetting.",
        size=10)

    add_heading(doc, "Key Facts", 2)
    add_simple_table(doc,
        ["Property", "Value"],
        [
            ["Application Name",   "BreakDown"],
            ["Version",            "1.2.5 (build 20260513)"],
            ["Platform",           "Windows 10/11  64-bit"],
            ["Language",           "Python 3.12  +  Java 17 (JAR components)"],
            ["Word Integration",   "Microsoft Word via Win32 COM automation"],
            ["Config Drive",       "V:\\TOOLS\\BreakDown\\  (network share)"],
            ["Hotfolder",          "V:\\FOR_BREAKDOWN\\INPUT\\SAGE\\"],
            ["Output",             "V:\\FOR_CONVERSION\\SAGE\\[JID]\\[AID]\\*_AS.docx"],
            ["Poll Interval",      "Every 10 seconds"],
            ["Per-file Timeout",   "900 seconds (15 min)"],
            ["Saxon XSLT",         "Saxon 9 PE  (requires licence file)"],
        ],
        col_widths=[2.0, 4.5]
    )

    add_heading(doc, "Input Package Contents", 2)
    add_simple_table(doc,
        ["File Type", "Pattern / Extension", "Purpose"],
        [
            ["Short Metadata",  "SAGE-metadata-*.xml  |  dd-nnn.xml",   "Article & journal IDs (auto-lookup)"],
            ["Long Metadata",   "*-metadata.xml (article_set root)",     "Full merge order + file list"],
            ["Manifest",        "*-manifest.html",                       "Package manifest"],
            ["Word Documents",  ".doc  /  .docx",                        "Manuscript files to be processed"],
            ["Graphics",        ".tif .jpg .png .svg .eps",              "Figures (staged, not processed)"],
            ["Supplementary",   "supp*/supplement* .docx",               "Treated specially (large-table check)"],
            ["Archive",         ".zip .tar .rar .7z",                    "Nested archives (recursive extract)"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — PIPELINE WORKFLOW DIAGRAM  (visual)
# ══════════════════════════════════════════════════════════════════════════════

def _flow_box(table, row, col, text, bg, fg=CLR["white"], bold=True, font_size=9, span=1):
    """Write text into a flow-chart box (coloured cell)."""
    if span > 1:
        cell = merge_cells_horizontal(table, row, col, col + span - 1)
    else:
        cell = table.cell(row, col)
    set_cell_bg(cell, bg)
    cell_para(cell, text, bold=bold, font_size=font_size, color=fg,
              align=WD_ALIGN_PARAGRAPH.CENTER,
              space_before=4, space_after=4)
    return cell


def _arrow(table, row, col, span=1):
    """Downward arrow cell."""
    if span > 1:
        cell = merge_cells_horizontal(table, row, col, col + span - 1)
    else:
        cell = table.cell(row, col)
    set_cell_bg(cell, CLR["white"])
    cell_para(cell, "▼", bold=False, font_size=11,
              color=CLR["arrow_bg"],
              align=WD_ALIGN_PARAGRAPH.CENTER,
              space_before=1, space_after=1)
    return cell


def _empty(table, row, col):
    cell = table.cell(row, col)
    set_cell_bg(cell, CLR["white"])
    cell_para(cell, "")
    return cell


def build_pipeline_workflow(doc):
    add_heading(doc, "2. End-to-End Pipeline Workflow", 1)

    add_body(doc,
        "The diagram below shows the complete data flow from package arrival in the hotfolder "
        "through every pipeline stage to the final styled output document.",
        size=10)

    doc.add_paragraph()

    # ── Layout: 7 columns
    # col 0,1,2  left lane   | col 3  centre divider | col 4,5,6  right lane
    # We use a 11-column grid: 0..10
    # Simple linear layout with a decision branch

    COLS = 11
    rows_data = []

    # Row definitions: (row_type, ...)
    # Row types: 'box', 'arrow', 'split_arrow', 'decision', 'blank', 'merge_arrow'

    # We'll build a table with enough rows
    NROWS = 36

    table = doc.add_table(rows=NROWS, cols=COLS)
    set_table_no_border(table)
    set_table_width(table, 100)

    # initialise all cells as white empty
    for r in range(NROWS):
        for c in range(COLS):
            _empty(table, r, c)

    # Column widths (approx)
    col_w = [0.55, 0.55, 0.55,  0.55, 0.55, 0.55,  0.3,  0.55, 0.55, 0.55,  0.55]
    for row in table.rows:
        for ci, w in enumerate(col_w):
            row.cells[ci].width = Inches(w)

    def box(r, c, text, bg, fg=CLR["white"], bold=True, sz=9, span=1):
        _flow_box(table, r, c, text, bg, fg, bold, sz, span)

    def arr(r, c, span=1):
        _arrow(table, r, c, span)

    def h_arrow(r, c, direction="→"):
        cell = table.cell(r, c)
        set_cell_bg(cell, CLR["white"])
        cell_para(cell, direction, bold=False, font_size=11,
                  color=CLR["arrow_bg"],
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=1, space_after=1)

    # ── Row 0: UPSTREAM (input)
    box(0, 0, "UPSTREAM", CLR["stage_teal"], span=5)
    box(0, 6, "", CLR["white"])
    box(0, 7, "WATCHER MACHINE", CLR["stage_blue"], span=4)

    # Row 1: arrow + arrow
    arr(1, 0, span=5)
    arr(1, 7, span=4)

    # Row 2: Input drop  |  watcher process
    box(2, 0, "[ DROP ]  ZIP / TAR / RAR / 7Z\nto V:\\FOR_BREAKDOWN\\INPUT\\SAGE\\",
        CLR["light_blue"], CLR["black"], bold=False, sz=8, span=5)
    box(2, 7, "[WATCH]  watcher.exe\nPolls every 10 seconds",
        CLR["stage_blue"], span=4)

    # Row 3: arrow connecting
    arr(3, 0, span=5)
    arr(3, 7, span=4)

    # Row 4: mAnalyzer header
    box(4, 0, "STAGE 1 — mAnalyzer  (BreakDown.exe -p=mAnalyzer)", CLR["stage_blue"], span=11)

    # Row 5: arrow
    arr(5, 0, span=11)

    # Row 6: mAnalyzer sub-steps
    box(6, 0, "① Extract\nZIP/TAR/RAR/7Z\n(recursive)", CLR["light_blue"], CLR["black"], bold=False, sz=8, span=3)
    box(6, 3, "② Classify\nFiles\n(metadata/doc/graphics)", CLR["light_blue"], CLR["black"], bold=False, sz=8, span=3)
    h_arrow(6, 6)
    box(6, 7, "③ Read Metadata\nShort/Long XML\n→ JID + AID", CLR["light_blue"], CLR["black"], bold=False, sz=8, span=2)
    box(6, 9, "④ GetArticleId\nSage login\n(smart_login)", CLR["light_blue"], CLR["black"], bold=False, sz=8, span=2)

    # Row 7: arrow
    arr(7, 0, span=11)

    # Row 8: Decision diamond
    box(8, 3, "[?] Single .docx\n+ Short Metadata?", CLR["stage_orange"], span=5)

    # Row 9: YES left / NO right branches
    box(9, 0, "YES →  skip_merger", CLR["light_green"], CLR["black"], bold=False, sz=8, span=4)
    box(9, 7, "NO  →  prepare MERGER_INPUT", CLR["light_orange"], CLR["black"], bold=False, sz=8, span=4)

    # Row 10: arrows in both lanes
    arr(10, 0, span=4)
    arr(10, 7, span=4)

    # Row 11: Stage 2a vs Stage 2b
    box(11, 0, "Direct to\nParaStyler_INPUT\n(rename → JID_AID_CLN.docx)",
        CLR["stage_green"], span=4)
    box(11, 7, "STAGE 2 — mNormalizer\n(BreakDown.exe -p=mNormalizer)",
        CLR["stage_blue"], span=4)

    # Row 12: skip (nothing in left lane) ; step details in right
    arr(12, 7, span=4)

    # Row 13: Normalizer sub-steps
    box(13, 7, "① Kill WINWORD / EXCEL", CLR["light_blue"], CLR["black"], bold=False, sz=8, span=4)
    arr(13, 0, span=4)

    # Row 14
    box(14, 7, "② DocxPreClean  +  ImageCleaner", CLR["light_blue"], CLR["black"], bold=False, sz=8, span=4)

    # Row 15
    arr(14, 0, span=4)
    arr(15, 7, span=4)
    box(15, 0, "▶  Output already in\nParaStyler_INPUT", CLR["light_green"], span=4)

    # Row 16
    box(16, 7, "③ .doc → .docx  (COM SaveAs)", CLR["light_blue"], CLR["black"], bold=False, sz=8, span=4)
    arr(17, 7, span=4)
    box(17, 0, "", CLR["white"])

    # Row 18
    box(18, 7, "④ Open in Word (COM)\nWordSessionController\n(restart every 3 files)", CLR["light_blue"], CLR["black"], bold=False, sz=8, span=4)
    arr(19, 7, span=4)

    # Row 20
    box(20, 7, "⑤ Run 15+ Word Macros\n(AcceptTrackChange, TotalCleanUp,\nRemoveLineNumbers …)", CLR["light_blue"], CLR["black"], bold=False, sz=8, span=4)
    arr(21, 7, span=4)

    # Row 22
    box(22, 7, "⑥ Save .docx  +  PreClean JAR\n(sage-auto-styler.jar -pre)", CLR["light_blue"], CLR["black"], bold=False, sz=8, span=4)
    arr(23, 7, span=4)

    # Row 24: mMerger
    box(24, 7, "STAGE 3 — mMerger\nMerge all .docx → JID_AID_CLN.docx\n(docxcompose Composer)", CLR["stage_blue"], span=4)
    arr(25, 7, span=4)

    # Row 26: both lanes converge to ParaStyler
    box(26, 3, "V:\\FOR_BREAKDOWN\\ParaStyler_INPUT\\SAGE\\[JID_AID]\\", CLR["stage_teal"], span=5)
    arr(27, 3, span=5)

    # Row 28: ParaStyler
    box(28, 0, "STAGE 4 — ParaStyler  (External Java Process)", CLR["stage_purple"], span=11)
    arr(29, 0, span=11)

    # Row 30: ParaStyler sub-steps
    box(30, 0, "Saxon XSLT\n(saxon9pe.jar)\npara_info.xsl\nauthor_label.xsl",
        CLR["light_blue"], CLR["black"], bold=False, sz=8, span=5)
    h_arrow(30, 5)
    box(30, 6, "+", CLR["white"], CLR["black"], bold=True, sz=12)
    h_arrow(30, 7)
    box(30, 8, "DocxManipulator JAR\n(sage-auto-styler.jar -ipas)\nApply SAGE paragraph styles",
        CLR["light_blue"], CLR["black"], bold=False, sz=8, span=3)

    arr(31, 0, span=11)

    # Row 32: Output
    box(32, 0,
        "[OUTPUT]  V:\\FOR_CONVERSION\\SAGE\\[JID]\\[AID]\\[JID]_[AID]_AS.docx",
        CLR["stage_green"], span=11)

    # Row 33-34: error path note
    doc.add_paragraph()
    add_info_box(doc,
        "ERROR PATH:  If mAnalyzer/mNormalizer/mMerger fails or times out (900 s), "
        "the package is moved to V:\\FOR_BREAKDOWN\\ERROR\\SAGE\\ and logged in error_log.html.",
        bg=CLR["light_orange"])

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — COMPONENT ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

def build_component_architecture(doc):
    add_heading(doc, "3. Component Architecture", 1)
    add_body(doc, "Five architecture layers, from entry-point orchestration down to Java components.", size=10)

    # ── Layer diagram as table ────────────────────────────────────────────────
    layers = [
        ("LAYER 1 — Entry Point / Orchestration",   CLR["title_bg"],   CLR["white"],
         "watcher.exe  →  polls INPUT\\SAGE every 10 s, invokes BreakDown.exe per file\n"
         "BreakDown.exe  →  CLI dispatcher; reads startupConfig.yaml → config folder; "
         "routes -p argument to sub-module (mAnalyzer / mSelect / mNormalizer / mMerger)"),

        ("LAYER 2 — Pipeline Stages",               CLR["stage_blue"],  CLR["white"],
         "mAnalyzer (mAnalyser.py)  —  Extract archives, classify files, read XML metadata, "
         "call GetArticleId.smart_login(), stage .docx files to MERGER_INPUT or ParaStyler_INPUT\n"
         "mNormalizer (mNormalizer.py)  —  Kill orphan processes, pre-clean DOCX, .doc→.docx "
         "conversion via COM, open in Word, run 15+ macros, save, invoke pre-clean JAR\n"
         "mMerger (mMerger.py)  —  Merge multiple .docx via docxcompose Composer in MergeOrder, "
         "produce JID_AID_CLN.docx, move to ParaStyler_INPUT\n"
         "ParaStyler  —  External Java: Saxon XSLT (para_info.xsl, author_label.xsl) + "
         "sage-auto-styler.jar; produces *_AS.docx in FOR_CONVERSION"),

        ("LAYER 3 — Supporting Modules",            CLR["stage_teal"],  CLR["white"],
         "DataBase (dbprocess.py)  —  Optional MySQL tracking; db_system:false by default\n"
         "TransformXml.py  —  Lazy-init Saxon JVM (one JVM/process); JET dump cleanup\n"
         "CreateArticleInfo.py  —  Selenium + sageJournalInfo.json lookup for JID→TLA\n"
         "com_manager.py / com_utils.py  —  COM context manager, gen_py cache repair\n"
         "DocxPreClean.py  —  XML-level clean before Word opens the file\n"
         "DocxImageCleaner.py  —  Strip oversized embedded images\n"
         "utils/  —  file_utils, process_runner, resource_monitor, error_logger, progress, retry"),

        ("LAYER 4 — Configuration",                CLR["stage_orange"], CLR["white"],
         "startupConfig.yaml  →  bootstrap only (CONFIG.BreakDown = V:\\TOOLS\\BreakDown)\n"
         "config\\breakDown.yaml  →  folder paths, logger, timeout\n"
         "config\\watcher.yaml  →  exe path, customers, INPUT/ERROR folders per customer\n"
         "config\\mAnalyser.yaml  →  DocTypes, MergeOrder, folder overrides\n"
         "config\\mNormalizer.yaml  →  KillProcess list, RunMacros list, PreClean settings\n"
         "config\\mMerger.yaml  →  output filename pattern, source-file retention\n"
         "config\\paraStyles.yaml  →  60+ SAGE paragraph style definitions + XML→style map\n"
         "config\\breakdownSequence.json  →  style-name ↔ mapping-tag table"),

        ("LAYER 5 — Java Components",              CLR["stage_purple"], CLR["white"],
         "DocxManipulator\\sage-auto-styler.jar  —  Pre-clean (-pre) and style-apply (-ipas) modes\n"
         "DocxManipulator\\jar\\  —  Aspose-Words 22.10, Jackson, jsoup, log4j, slf4j, …\n"
         "ParaStyler\\saxon9pe.jar  —  Saxon XSLT PE (requires saxon-license.lic)\n"
         "ParaStyler\\weka-stable-3.6.6.jar  —  Random-committee ML classifier\n"
         "ParaStyler\\asprop30x.arff.randomCommitee_50.model  —  Trained classifier model\n"
         "ParaStyler\\*.xsl  —  XSLT stylesheets (normalize, paraStyle, author_label, …)"),
    ]

    for title, bg, fg, desc in layers:
        t = doc.add_table(rows=2, cols=1)
        set_table_width(t, 100)
        set_table_no_border(t)

        # Header cell
        hc = t.cell(0, 0)
        set_cell_bg(hc, bg)
        cell_para(hc, title, bold=True, font_size=10, color=fg,
                  space_before=4, space_after=4)

        # Body cell
        bc = t.cell(1, 0)
        set_cell_bg(bc, CLR["light_gray"])
        p = bc.paragraphs[0]
        p.clear()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        for line in desc.split("\n"):
            if line.strip():
                run = p.add_run(line + "\n")
                run.font.size = Pt(9)
                if line.strip().endswith("  —") or "  —  " in line or "  →  " in line or "  ↔  " in line:
                    parts = re.split(r"(  —  |  →  |  ↔  )", line, 1)
                    if len(parts) == 3:
                        p.clear()
                        r1 = p.add_run(parts[0])
                        r1.bold = True
                        r1.font.size = Pt(9)
                        r2 = p.add_run(parts[1] + parts[2])
                        r2.font.size = Pt(9)
                        p.add_run("\n").font.size = Pt(9)

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — DEPLOYMENT ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

def build_deployment_architecture(doc):
    add_heading(doc, "4. Deployment Architecture", 1)
    add_body(doc, "Three Windows machines sharing a single network drive (V:).", size=10)

    # ── Machine summary table ─────────────────────────────────────────────────
    add_heading(doc, "4.1  Machines & Roles", 2)
    add_simple_table(doc,
        ["Machine", "Role", "Key Process", "Batch File to Open"],
        [
            ["Watcher Machine",     "Monitors hotfolder continuously",
             "watcher.exe  (-l=HOTFOLDER)",    "start_watcher.bat  ◄ keep open"],
            ["Processing Machine",  "Runs BreakDown pipeline per file",
             "BreakDown.exe  (all -p modes)",   "Invoked automatically by Watcher"],
            ["ParaStyler Machine",  "Applies SAGE paragraph styles",
             "ParaStyler Java process",         "ParaStyler\\run.bat"],
        ],
        col_widths=[1.5, 1.7, 2.2, 2.1]
    )

    # ── Network share layout ──────────────────────────────────────────────────
    add_heading(doc, "4.2  Network Share Layout  (V: drive)", 2)
    add_info_box(doc,
        "V:  =  \\\\192.168.0.102\\d$\\REPOSITORY   (mapped on every machine via  net use  or  map_drive.bat)",
        bg=CLR["light_blue"])

    add_simple_table(doc,
        ["V:\\ Path", "Purpose", "Owner"],
        [
            ["TOOLS\\BreakDown\\",               "Tool installation (EXEs, JARs, configs)",   "Admin"],
            ["FOR_BREAKDOWN\\INPUT\\SAGE\\",       "Hotfolder  —  drop packages HERE",          "Upstream/Watcher"],
            ["FOR_BREAKDOWN\\PROCESS\\",           "Extracted in-flight packages",              "mAnalyzer"],
            ["FOR_BREAKDOWN\\ERROR\\",             "Failed packages",                           "Pipeline"],
            ["FOR_BREAKDOWN\\LOG\\",               "break_down.log + per-article logs",         "Pipeline"],
            ["FOR_BREAKDOWN\\MERGER_INPUT\\SAGE\\","Pre-merge staged .docx files",             "mAnalyzer"],
            ["FOR_BREAKDOWN\\ParaStyler_INPUT\\",  "CLN.docx awaiting ParaStyler",             "mMerger"],
            ["FOR_BREAKDOWN\\BreakDown_DONE\\",    "Completed packages",                        "Pipeline"],
            ["FOR_CONVERSION\\SAGE\\[JID]\\[AID]\\","Final *_AS.docx output",                  "ParaStyler"],
        ],
        col_widths=[2.6, 2.4, 1.5]
    )

    # ── Install folder tree ───────────────────────────────────────────────────
    add_heading(doc, "4.3  Install Folder Tree  (V:\\TOOLS\\BreakDown\\)", 2)
    add_code_block(doc, """\
V:\\TOOLS\\BreakDown\\
│
├── BreakDown.exe               main CLI tool
├── watcher.exe                 watcher daemon
├── start_watcher.bat  ◄─────── OPEN THIS on Watcher Machine (HOTFOLDER)
├── start_watcher_s3.bat        (optional S3 mode)
├── startupConfig.yaml          bootstrap: CONFIG.BreakDown path
│
├── _internal\\                  PyInstaller runtime (DLLs + packages)
│   ├── DocxManipulator\\
│   │   ├── sage-auto-styler.jar         pre-clean + style-apply JAR
│   │   └── jar\\  (Aspose, Jackson, jsoup, log4j, …)
│   └── ParaStyler\\
│       ├── saxon9pe.jar                 Saxon XSLT processor
│       ├── saxon-license.lic  ◄── REQUIRED (PE licence)
│       ├── weka-stable-3.6.6.jar
│       ├── asprop30x.arff.randomCommitee_50.model
│       └── *.xsl
│
├── config\\
│   ├── breakDown.yaml           folder paths, logger, timeout
│   ├── watcher.yaml             exe path, customers, hotfolder
│   ├── mAnalyser.yaml           DocTypes, MergeOrder
│   ├── mNormalizer.yaml         RunMacros, KillProcess, PreClean
│   ├── mMerger.yaml             filename pattern, source retention
│   ├── dbConfig.yaml            db_system: false (MySQL optional)
│   ├── paraStyles.yaml          60+ SAGE paragraph style definitions
│   ├── breakdownSequence.json   style ↔ mapping-tag table
│   ├── backMatterTitles.json    back-matter section patterns
│   └── saxon-license.lic        Saxon licence backup copy
│
└── SupportingFiles\\
    ├── SAGE_styles.docx  ◄── REQUIRED (Word style template)
    ├── SAGESTYLES.dotx
    ├── CMSTYLES.dotx
    ├── checkDocRunning.exe      COM-safe file-lock status helper
    ├── checkDocRunning.yaml     info_path config
    ├── sageJournalInfo.json  ◄── REQUIRED (JID → TLA lookup)
    └── defaultValue.json""")

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — INSTALLATION GUIDE
# ══════════════════════════════════════════════════════════════════════════════

def build_installation_guide(doc):
    add_heading(doc, "5. Installation Guide", 1)

    # --- Pre-requisites ---
    add_heading(doc, "5.1  Pre-requisites", 2)
    add_simple_table(doc,
        ["Software", "Version", "Required On", "Notes"],
        [
            ["Java JRE/JDK",    "17+",      "All machines",     "Must be in PATH.  adoptium.net"],
            ["V: drive mapping","—",         "All machines",     "net use V: \\\\192.168.0.102\\d$\\REPOSITORY"],
            ["Microsoft Word",  "2016/2019/365","Processing Mach.","COM automation for mNormalizer"],
            ["Python",          "3.12",     "Dev machine only", "Only if running from source"],
            ["PyInstaller",     "latest",   "Dev machine only", "For  build.bat"],
        ],
        col_widths=[1.5, 0.9, 1.5, 2.6]
    )

    # --- Quick install ---
    add_heading(doc, "5.2  Quick Install  (install_breakdown.bat)", 2)
    add_body(doc,
        "Run the installer from the project root or release folder. "
        "It completes all 14 steps automatically.", size=10)
    add_code_block(doc, """\
:: From Command Prompt (run as Administrator recommended)
install_breakdown.bat

:: Install to a custom path
install_breakdown.bat "D:\\MyTools\\BreakDown"
""")

    add_heading(doc, "Installer Steps", 3)
    add_simple_table(doc,
        ["Step", "Action"],
        [
            ["1",  "Validate: V: drive mapped + Java in PATH"],
            ["2",  "Locate BreakDown.exe  (current dir or dist\\BreakDown_v*\\)"],
            ["3",  "Create V:\\TOOLS\\BreakDown\\ sub-folders + all V:\\FOR_BREAKDOWN\\ working folders"],
            ["4",  "Copy BreakDown.exe + watcher.exe"],
            ["5",  "Copy _internal\\  (PyInstaller runtime)"],
            ["6",  "Copy all JAR files  (DocxManipulator, ParaStyler, Aspose-Words)"],
            ["7",  "Copy Saxon licence files  (saxon-license.lic)"],
            ["8",  "Copy config YAML / JSON files"],
            ["9",  "Copy SupportingFiles  (SAGE_styles.docx, checkDocRunning.exe, …)"],
            ["10", "Copy XSL stylesheets"],
            ["11", "Generate startupConfig.yaml"],
            ["12", "Verify all working folders exist"],
            ["13", "Generate start_watcher.bat + start_watcher_s3.bat in install root"],
            ["14", "Print summary + next-steps"],
        ],
        col_widths=[0.4, 6.1]
    )

    # --- Build ---
    add_heading(doc, "5.3  Build from Source  (build.bat)", 2)
    add_code_block(doc, """\
build.bat              :: patch bump  1.0.0 → 1.0.1  +  PyInstaller
build.bat minor        :: minor bump  1.0.1 → 1.1.0
build.bat major        :: major bump  1.1.0 → 2.0.0
build.bat set 1.5.0    :: force version
build.bat build-only   :: no version bump, update build date only

:: Output:  dist\\BreakDown_v<version>\\
""")

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — CONFIGURATION REFERENCE
# ══════════════════════════════════════════════════════════════════════════════

def build_config_reference(doc):
    add_heading(doc, "6. Configuration Reference", 1)

    configs = [
        ("startupConfig.yaml  —  Bootstrap",
"""MAPPING:
    DRIVE: V
    PATH: \\\\192.168.0.102\\d$\\REPOSITORY
CONFIG:
    BreakDown: V:\\TOOLS\\BreakDown   # ← must match install root""",
         "Read by loadconfig.py from the same folder as BreakDown.exe.  "
         "Everything else is resolved from CONFIG.BreakDown."),

        ("config\\watcher.yaml  —  Watcher",
"""BREAKDOWN_EXE: V:\\TOOLS\\BreakDown\\BreakDown.exe

HOTFOLDER:
  CUSTOMERS: [SAGE]          # add customer names here
  SAGE:
    FOLDERS:
      INPUT: V:\\FOR_BREAKDOWN\\INPUT\\SAGE
      ERROR: V:\\FOR_BREAKDOWN\\ERROR\\SAGE

S3:
  CUSTOMERS: [SAGE]
  SAGE:
    REPOSITORY:    # S3 bucket
    ACCESSID:      # AWS access key ID
    ACCESSKEY:     # AWS secret access key""",
         "Controls which customers the watcher monitors and the hotfolder paths.  "
         "BREAKDOWN_EXE points at the installed executable."),

        ("config\\mNormalizer.yaml  —  Macros & Pre-clean",
"""KillProcess: [WINWORD.exe, EXCEL.exe, POWERPNT.exe]

SAGE:
  RunMacros:               # executed in this order
    - AcceptTrackChange
    - TotalCleanUP
    - UnlinkFieldcodesExceptMath
    - ConvertEndnoteToFootnote
    - EnqTableToText
    - RemoveLineNumbers
    - FlattenTextBoxes
    - RemoveAllFramesInDoc
    - RemoveUnwantedSpaces
    - EliminateMultipleSpaces
    - RemoveDocVar
    - CitationsToStaticText
    - RemoveFirstLineIndent
    - RemoveAllHyperlinks
    - CleanTableCells
  PreClean:
    enabled: false           # true → invoke sage-auto-styler.jar
    replace_macros: false    # true → skip Word macros entirely
    jar_name: sage-auto-styler.jar
    jar_args: ["-pre"]
    timeout: 300""",
         "Set PreClean.enabled: true to activate the Java pre-clean step after macros.  "
         "Set replace_macros: true to skip Word macros and rely solely on the JAR."),

        ("config\\mMerger.yaml  —  Merger",
"""FILENAME:
  SAGE: FOLDER_CLN.docx    # output filename template

SOURCE: Retain             # Retain | Remove""",
         "SOURCE: Retain moves source .docx files to a docs\\ sub-folder.  "
         "SOURCE: Remove deletes them after merging."),

        ("config\\mAnalyser.yaml  —  Analyser",
"""DocTypes: [title, main, ack, author_note, bio, figure, table, other]
MergeOrder: {1: title, 2: main, 3: ack, 4: author_note, 5: bio, 6: figure, 7: table}

FOLDERS:
  MERGER: V:\\FOR_BREAKDOWN\\MERGER_INPUT\\[CUSTOMER]\\[JID]_[AID]
  PROCESS: V:\\FOR_BREAKDOWN\\PROCESS
  ERROR:   V:\\FOR_BREAKDOWN\\ERROR""",
         "DocTypes defines valid file-type values from long metadata.  "
         "MergeOrder controls the order in which document parts are merged."),
    ]

    for title, code, note in configs:
        add_heading(doc, title, 2)
        if note:
            add_body(doc, note, size=9, italic=True, color=(80, 80, 80))
        add_code_block(doc, code)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — BATCH FILES & MANUAL COMMANDS
# ══════════════════════════════════════════════════════════════════════════════

def build_operations(doc):
    add_heading(doc, "7. Batch Files & Manual Commands", 1)

    add_heading(doc, "7.1  Batch Files Reference", 2)
    add_simple_table(doc,
        ["File", "Machine", "Purpose", "Keep Open?"],
        [
            ["start_watcher.bat",     "Watcher",    "Start watcher daemon — HOTFOLDER mode",     "✅ YES"],
            ["start_watcher_s3.bat",  "Watcher",    "Start watcher daemon — S3 mode",            "✅ YES"],
            ["map_drive.bat",         "Any",        "Map V: drive to \\\\192.168.0.102\\…",       "No"],
            ["install_breakdown.bat", "Any",        "Full 14-step installation",                  "No"],
            ["local_styler.bat",      "Any",        "Run sage-auto-styler.jar on one .docx",      "No"],
            ["build.bat",             "Dev only",   "Build EXEs via PyInstaller",                 "No"],
            ["ParaStyler\\run.bat",   "ParaStyler", "Run ParaStyler on a package",               "No"],
        ],
        col_widths=[1.9, 1.2, 3.0, 0.9]
    )

    add_info_box(doc,
        "NOTE:  start_watcher.bat launches watcher.exe as a FOREGROUND process.  "
        "The CMD window MUST remain open.  Close it and the watcher stops.",
        bg=CLR["light_orange"])

    add_heading(doc, "7.2  Manual Processing Commands", 2)

    cmds = [
        ("Run mAnalyzer on a specific package",
         'BreakDown.exe -p="mAnalyzer" -f="V:\\FOR_BREAKDOWN\\INPUT\\SAGE\\Article_Attachments-2026-05-25.zip" -c="SAGE"'),
        ("Run mNormalizer manually",
         'BreakDown.exe -p="mNormalizer" -c="SAGE" -jf="V:\\FOR_BREAKDOWN\\MERGER_INPUT\\SAGE\\SGO_123456\\SGO_123456.json"'),
        ("Run mMerger manually",
         'BreakDown.exe -p="mMerger" -c="SAGE" -jf="V:\\FOR_BREAKDOWN\\MERGER_INPUT\\SAGE\\SGO_123456\\SGO_123456.json"'),
        ("Open mSelect GUI (manual doc selection)",
         'BreakDown.exe -p="mSelect"'),
        ("Update SAGE Journal Info",
         'BreakDown.exe -p="createSageJournalInfo"'),
        ("Run local pre-clean JAR on a .docx",
         'local_styler.bat "V:\\FOR_BREAKDOWN\\ParaStyler_INPUT\\SAGE\\SGO_123456\\SGO_123456_CLN.docx"'),
    ]

    for label, cmd in cmds:
        add_body(doc, label, size=9, bold=True)
        add_code_block(doc, cmd)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════

def build_troubleshooting(doc):
    add_heading(doc, "8. Troubleshooting", 1)

    sections = [
        ("Watcher Issues", [
            ("Files sit in INPUT\\ unprocessed",       "Is start_watcher.bat CMD window still open?"),
            ("Error: watcher.exe not found",           "Run install_breakdown.bat first"),
            ("Error: V: drive not mapped",             "Run map_drive.bat"),
            ("Timeout errors in log",                  "Check if WINWORD.EXE is hung; kill via Task Manager"),
            ("Package moved to ERROR immediately",     "Check ZIP contains *-metadata.xml or SAGE-metadata-*.xml"),
        ]),
        ("mNormalizer / Word COM Issues", [
            ("CLSIDToClassMap error on startup",       "BreakDown auto-clears gen_py cache — re-run"),
            ("COMException when opening document",     "Kill all WINWORD.EXE processes; re-run"),
            ("Document opens as read-only",            "Close all open Word windows before processing"),
            ("Macro not found error",                  "Ensure CMSTYLES.dotx / SAGESTYLES.dotx are loaded in Word"),
            ("Word crashes (0xC0000409)",              "Increase Word restart interval in WordSessionController"),
        ]),
        ("JAR / Java Issues", [
            ("JAR file not found: sage-auto-styler.jar","Check _internal\\DocxManipulator\\ exists in install root"),
            ("Java not found in PATH",                  'setx PATH "%PATH%;C:\\Program Files\\Java\\jdk-17\\bin"'),
            ("Saxon licence error",                    "Verify ParaStyler\\saxon-license.lic exists and is valid"),
            ("Weka model error",                       "Check asprop30x.arff.randomCommitee_50.model present"),
        ]),
        ("Log Locations", [
            ("Main BreakDown log",                     "V:\\FOR_BREAKDOWN\\LOG\\break_down.log"),
            ("Per-article log",                        "V:\\FOR_BREAKDOWN\\LOG\\[CUSTOMER]\\[JID]\\[AID]\\"),
            ("Watcher error log",                      "V:\\FOR_BREAKDOWN\\LOG\\error_log.html"),
            ("Watcher startup log",                    "V:\\TOOLS\\BreakDown\\watcher_startup.log"),
        ]),
    ]

    for heading, rows in sections:
        add_heading(doc, heading, 2)
        add_simple_table(doc,
            ["Symptom", "Solution / Location"],
            rows,
            col_widths=[2.8, 3.7]
        )

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — APPENDIX: JAR + FILE INVENTORY
# ══════════════════════════════════════════════════════════════════════════════

def build_appendix(doc):
    add_heading(doc, "Appendix A — JAR & File Inventory", 1)

    add_heading(doc, "A.1  JAR Files", 2)
    add_simple_table(doc,
        ["JAR File", "Folder", "Purpose"],
        [
            ["sage-auto-styler.jar",              "DocxManipulator\\",      "Pre-clean + SAGE style apply"],
            ["docx-manipulator.jar",              "DocxManipulator\\",      "Alias for above"],
            ["aspose-words-22.10-jdk17.jar",      "DocxManipulator\\jar\\", "Word document manipulation"],
            ["jackson-databind-2.9.8.jar",        "DocxManipulator\\jar\\", "JSON serialisation"],
            ["jackson-core-2.9.8.jar",            "DocxManipulator\\jar\\", "JSON core"],
            ["jackson-annotations-2.9.0.jar",     "DocxManipulator\\jar\\", "JSON annotations"],
            ["jsoup-1.8.3.jar",                   "DocxManipulator\\jar\\", "HTML parsing"],
            ["commons-io-2.4.jar",                "DocxManipulator\\jar\\", "File utilities"],
            ["commons-logging-1.1.1.jar",         "DocxManipulator\\jar\\", "Logging bridge"],
            ["log4j-api-2.16.0.jar",              "DocxManipulator\\jar\\", "Logging API"],
            ["log4j-core-2.16.0.jar",             "DocxManipulator\\jar\\", "Logging core"],
            ["slf4j-api-1.5.6.jar",               "DocxManipulator\\jar\\", "SLF4J facade"],
            ["filters-2.0.235.jar",               "DocxManipulator\\jar\\", "Content filters"],
            ["gluegen-rt-main-2.3.2.jar",         "DocxManipulator\\jar\\", "OpenGL (Aspose)"],
            ["jogl-all-main-2.3.2.jar",           "DocxManipulator\\jar\\", "OpenGL (Aspose)"],
            ["jai-imageio-core-1.3.0.jar",        "DocxManipulator\\jar\\", "Image I/O"],
            ["mime-util-2.1.1.jar",               "DocxManipulator\\jar\\", "MIME type detection"],
            ["saxon9pe.jar",                      "ParaStyler\\",           "Saxon XSLT PE processor"],
            ["weka-stable-3.6.6.jar",             "ParaStyler\\",           "ML classifier runner"],
            ["commons-cli-1.2.jar",               "ParaStyler\\",           "CLI argument parsing"],
            ["commons-io-2.4.jar",                "ParaStyler\\",           "File utilities"],
            ["guava-10.0.jar",                    "ParaStyler\\",           "Google Guava"],
            ["log4j-api-2.0-beta8.jar",           "ParaStyler\\",           "Logging"],
            ["log4j-core-2.0-beta8.jar",          "ParaStyler\\",           "Logging"],
            ["jsr305-1.3.9.jar",                  "ParaStyler\\",           "JSR-305 annotations"],
        ],
        col_widths=[2.5, 2.0, 2.0]
    )

    add_heading(doc, "A.2  Critical Supporting Files", 2)
    add_simple_table(doc,
        ["File", "Folder", "Mandatory?", "Purpose"],
        [
            ["SAGE_styles.docx",        "SupportingFiles\\", "✅ YES", "Word template with all SAGE styles"],
            ["sageJournalInfo.json",     "SupportingFiles\\", "✅ YES", "JID → TLA mapping database"],
            ["checkDocRunning.exe",      "SupportingFiles\\", "YES",    "COM-safe file-lock status helper"],
            ["checkDocRunning.yaml",     "SupportingFiles\\", "YES",    "info_path for status files"],
            ["saxon-license.lic",        "ParaStyler\\",      "✅ YES", "Saxon PE licence (XSLT required)"],
            ["asprop30x.arff…model",     "ParaStyler\\",      "YES",    "Trained para-style ML classifier"],
            ["CMSTYLES.dotx",            "SupportingFiles\\", "Rec.",   "Word template for macros"],
            ["SAGESTYLES.dotx",          "SupportingFiles\\", "Rec.",   "Word template for macros"],
            ["defaultValue.json",        "SupportingFiles\\", "No",     "Default field values"],
            ["BreakDownLogo.png",         "SupportingFiles\\", "No",     "Application logo"],
        ],
        col_widths=[2.2, 1.5, 0.9, 2.9]
    )


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            "BreakDown_Documentation.docx")

    doc = Document()

    # ── Page layout: A4, narrow margins ──────────────────────────────────────
    section = doc.sections[0]
    section.page_width   = Cm(21)
    section.page_height  = Cm(29.7)
    section.left_margin  = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin   = Cm(1.8)
    section.bottom_margin= Cm(1.8)

    # ── Default Normal style ──────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)

    # ── Heading styles ────────────────────────────────────────────────────────
    for lvl, size in [(1, 14), (2, 12), (3, 10)]:
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(*CLR["title_bg"])
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after  = Pt(4)

    print("Building: Title page ...")
    build_title_page(doc)

    print("Building: System overview ...")
    build_system_overview(doc)
    doc.add_page_break()

    print("Building: Pipeline workflow diagram ...")
    build_pipeline_workflow(doc)

    print("Building: Component architecture ...")
    build_component_architecture(doc)

    print("Building: Deployment architecture ...")
    build_deployment_architecture(doc)

    print("Building: Installation guide ...")
    build_installation_guide(doc)

    print("Building: Configuration reference ...")
    build_config_reference(doc)

    print("Building: Operations / batch files ...")
    build_operations(doc)

    print("Building: Troubleshooting ...")
    build_troubleshooting(doc)

    print("Building: Appendix ...")
    build_appendix(doc)

    doc.save(out_path)
    size_kb = os.path.getsize(out_path) // 1024
    print(f"\n[OK]  Document saved: {out_path}")
    print(f"      File size: {size_kb} KB")
    return out_path


if __name__ == "__main__":
    main()
