import os
import html
import re
from docx import Document
from lxml import etree
import xml.etree.ElementTree as ET
import xml.dom.minidom
from docx.enum.style import WD_STYLE_TYPE
import zipfile
from docx.oxml.ns import qn

class CreateParaInfo:
    def __init__(self):
        pass

    def is_quote(self, para_info, curr_para_info, ref_para_num, prev_para_info=None, next_para_info=None):
        ldquo = '\u201C'
        rdquo = '\u201D'
        lsquo = '\u2018'
        rsquo = '\u2019'
        rdquodot = '\u201D' + "."
        rsquodot = '\u2019' + "."
        # if re.search("Sophia", curr_para_info["firstWord"]):
        #     print(curr_para_info["firstWord"])
        # if prev_para_info is not None and prev_para_info['lastWord'] == "labor:":
        #     print("test")
        #     print(int(curr_para_info['leftIndent']))
        #     print(int(prev_para_info['leftIndent']))
        if prev_para_info is None:
            return ""
        elif int(ref_para_num) > 0 and int(para_info["number"]) > ref_para_num:
            return ""
        elif (int(curr_para_info['leftIndent']) != int(prev_para_info['leftIndent']) and int(curr_para_info["wordCount"]) < 7):
            return ""
        elif int(curr_para_info['wordCount']) > 0:
            if curr_para_info['firstWord'].startswith(("Fig", 'Tab', "References")) or \
                    prev_para_info['firstWord'].startswith(("Fig", 'Tab', "References")) or \
                    curr_para_info['paraStyle'].startswith(("List")):
                return ""
            elif curr_para_info['lastWord'].endswith(("?", "?.")):
                return ""
            elif curr_para_info['firstWord'].startswith((lsquo, ldquo)) and curr_para_info['lastWord'].endswith((rsquo, rsquodot, rdquo, rdquodot, ")", ").")) and prev_para_info['lastChar'].endswith(':'):
                return f'<para number="{para_info["number"]}" quote="Left"/>'
            elif curr_para_info['firstWord'].startswith((lsquo, ldquo)) and curr_para_info['firstWord'].endswith((rsquo, rsquodot, rdquo, rdquodot)):
                return f'<para number="{para_info["number"]}" quote="Left"/>'
            elif int(curr_para_info['leftIndent']) > int(prev_para_info['leftIndent']) and \
                    next_para_info['alignment'] == "right":
                return f'<para number="{para_info["number"]}" quote="Left"/>'
            elif int(curr_para_info['leftIndent']) > int(prev_para_info['leftIndent']) and \
                    next_para_info['firstChar'].startswith("(") and \
                    next_para_info['lastWord'].endswith(")", ")."):
                return f'<para number="{para_info["number"]}" quote="Left"/>'
            elif int(prev_para_info['leftIndent']) > 0 and \
                    (int(curr_para_info['leftIndent']) > 0 or int(curr_para_info["wordCount"]) < 10) and \
                    (int(curr_para_info['leftIndent']) == int(prev_para_info['leftIndent']) or int(curr_para_info["wordCount"]) < 10):
                if curr_para_info['alignment'] == "right":
                    return f'<para number="{para_info["number"]}" quote="Right"/>'
                else:
                    return f'<para number="{para_info["number"]}" quote="Left"/>'
            else:
                if int(curr_para_info['leftIndent']) > int(prev_para_info['leftIndent']) and \
                        (next_para_info is None or int(curr_para_info['leftIndent']) != int(next_para_info['leftIndent'])):
                    if curr_para_info['isTitleCase'] == "False" and int(curr_para_info['wordCount']) > 15:
                        if curr_para_info['firstChar'].startswith((lsquo, ldquo)) or \
                                not prev_para_info['lastChar'].endswith('.'):
                            return f'<para number="{para_info["number"]}" quote="Left"/>'
                        elif int(curr_para_info['leftIndent']) > 0 and int(curr_para_info['rightIndent']) > 0:
                            return f'<para number="{para_info["number"]}" quote="Left"/>'
                        elif curr_para_info['firstChar'].startswith((lsquo, ldquo)) and \
                                curr_para_info['lastChar'].endswith((")", ").", rsquo, rdquo)):
                            return f'<para number="{para_info["number"]}" quote="Left"/>'
        return ""

    def is_list(self, para_info, curr_para_info, ref_para_num, prev_para_info=None, next_para_info=None):
        # if int(para_info["number"]) == 38:
        #     print(int(para_info["number"]))
        #     print(curr_para_info["firstChar"].isdigit())
        if int(curr_para_info['wordCount']) > 0:
            if int(ref_para_num) > 0 and int(para_info["number"]) > ref_para_num:
                return ""
            elif curr_para_info['paraStyle'].startswith(("List")):
                bullet = '\u25CF'
                if re.search(curr_para_info["listType"], "None") and curr_para_info["firstChar"].isdigit() is True:
                    return f'<para number="{para_info["number"]}" listType="decimal" level="{curr_para_info["listLvl"]}" list="True"/>'
                elif re.search("None", curr_para_info["listType"]) and re.search(bullet, curr_para_info["firstChar"]):
                        return f'<para number="{para_info["number"]}" listType="bullet" level="{curr_para_info["listLvl"]}" list="True"/>'
                else:
                    return f'<para number="{para_info["number"]}" listType="{curr_para_info["listType"]}" level="{curr_para_info["listLvl"]}" list="True"/>'
            elif re.search("None", curr_para_info["listType"]) and re.search("-", curr_para_info["firstChar"]):
                if re.search("-", curr_para_info["firstChar"]) and re.search("displayquote", curr_para_info["paraStyle"]):
                    return f'<para number="{para_info["number"]}" listType="tabbed" level="{curr_para_info["listLvl"]}" list="True"/>'
                elif re.search("-", curr_para_info["firstChar"]) and re.search("-", prev_para_info["firstChar"]):
                    return f'<para number="{para_info["number"]}" listType="tabbed" level="{curr_para_info["listLvl"]}" list="True"/>'
                elif re.search("-", curr_para_info["firstChar"]) and re.search("-", next_para_info["firstChar"]):
                    return f'<para number="{para_info["number"]}" listType="tabbed" level="{curr_para_info["listLvl"]}" list="True"/>'
                else:
                    return ""
            else:
                return ""
        return ""

    def process_xml(self, file_path):
        tree = ET.parse(file_path)
        root = tree.getroot()
        prev_para_info = None
        ref_para_num = 0
        for ref_para in root.findall('./paraInfo[@isRefTitle="True"][@number]'):
            ref_para_num = int(ref_para.attrib['number'])
        result = f"<info><refPara num=\"{ref_para_num}\"/>\n<quot>\n"
        for para_info in root.findall('.//paraInfo'):
            prev_para_info = para_info.find('./previousPara')
            curr_para_info = para_info.find('./para')
            next_para_info = para_info.find('./nextPara')
            quote_result = self.is_quote(para_info.attrib, curr_para_info.attrib, ref_para_num,
                                         prev_para_info.attrib if prev_para_info is not None else None,
                                         next_para_info.attrib if next_para_info is not None else None)
            if quote_result:
                result += quote_result + "\n"

        result += "</quot><list>\n"
        for para_info in root.findall('.//paraInfo'):
            curr_para_info = para_info.find('./para')
            next_para_info = para_info.find('./nextPara')
            prev_para_info = para_info.find('./previousPara')
            list_result = self.is_list(para_info.attrib, curr_para_info.attrib, ref_para_num,
                                         prev_para_info.attrib if prev_para_info is not None else None,
                                         next_para_info.attrib if next_para_info is not None else None)
            if list_result:
                result += list_result + "\n"
        result += "</list></info>\n"
        return result

    def is_title_case(self, text):
        words = text.split()
        capital_words = sum(1 for word in words if word[0].isupper())
        ratio = capital_words / len(words)
        return ratio > 0.7

    def pretty_format_xml(self, xml_string):
        dom = xml.dom.minidom.parseString(xml_string)
        return dom.toprettyxml()

    def get_num_fmt(self, num_id, numbering_content):
        root = ET.fromstring(numbering_content)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        num_element = root.find('.//w:num[@w:numId="{}"]'.format(num_id), namespaces=ns)
        if num_element is not None:
            abstract_num_id = num_element.find('.//w:abstractNumId', namespaces=ns).attrib[
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val']
            abstract_num_element = root.find('.//w:abstractNum[@w:abstractNumId="{}"]'.format(abstract_num_id),
                                             namespaces=ns)
            if abstract_num_element is not None:
                # Extract the w:numFmt value
                num_fmt = abstract_num_element.find('.//w:numFmt', namespaces=ns).attrib[
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val']
                return num_fmt
        # Return None if the numId or abstractNumId is not found
        return None        # for num in numbering_tree.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num'):
        #     if num.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId').attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'] == num_id:
        #         abstract_num_id = num.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId').attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val']
        #         for abstract_num in numbering_tree.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNum'):
        #             if abstract_num.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId'] == abstract_num_id:
        #                 num_fmt = abstract_num.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numFmt').attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val']
        #                 return num_fmt

    def get_paragraph_info(self, docxfile, document, paragraph):
        info = {}
        # Get paragraph properties
        # properties = paragraph.findall('.//w:pPr', namespaces=nsmap)
        numbering_found = False
        with zipfile.ZipFile(docxfile, 'r') as zip_ref:
            if "numbering.xml" in zip_ref.namelist():
                numbering_found = True
            # zip_ref.extract('word/document.xml', 'temp')

        if numbering_found is True:
            numbering = document.part.numbering_part.numbering_definitions._numbering.xml
        para_element = etree.fromstring(paragraph._element.xml)
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        properties = para_element.findall('./w:pPr', namespaces=nsmap)
        info['leftIndent'] = 0
        info['alignment'] = None
        info['rightIndent'] = 0
        info['paraStyle'] = None
        info['listLvl'] = None
        info['listType'] = None
        info['refTitle'] = False
        if properties:
            prop = properties[0]
            # Check for bold
            if prop.find('.//w:numPr', namespaces=nsmap) is not None:
                numLevel = prop.find('.//w:numPr/w:ilvl', namespaces=nsmap)
                if numLevel is not None:
                    numLvl = numLevel.get(
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                numberId = prop.find('.//w:numPr/w:numId', namespaces=nsmap)
                if numberId is not None:
                    numId = numberId.get(
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                #<w:num w:numId="15"><w:abstractNumId w:val="3"
                #<w:abstractNumId w:val="3"  - <w:numFmt w:val="bullet"/>
                #absId = numbering.find(f'.//w:num[@w:numId="{numId}"]', namespaces=nsmap)
                listType = None
                info['listLvl'] = 0
                if numbering_found is True:
                    listType = self.get_num_fmt(numId, numbering)
                info['listLvl'] = numLvl
                info['listType'] = listType
            style = prop.find('.//w:pStyle', namespaces=nsmap)
            styleName = style.get(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if style is not None else 'Normal'
            info['paraStyle'] = styleName
            styleLeftIndent = 0
            styleRightIndent = 0
            if not re.search(r"(Normal|Heading|List|author)", styleName):
                for style in document.styles:
                    docstyle = style.name
                    docstyle = re.sub(r'\s+', '', docstyle)
                    if docstyle == styleName:  # Example: Match style by name
                        styleLeftIndent = style.paragraph_format.left_indent if style.paragraph_format.left_indent is not None else 0
                        styleRightIndent = style.paragraph_format.right_indent if style.paragraph_format.right_indent is not None else 0
                        break

            left_indent = prop.find('.//w:ind[@w:left]', namespaces=nsmap)
            l_indent = left_indent.attrib[
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left'] if left_indent is not None else 0
            # Get rightIndent
            if l_indent == 0 and styleLeftIndent > 0:
                l_indent = styleLeftIndent
            info['leftIndent'] = l_indent
            right_indent = prop.find('.//w:ind[@w:right]', namespaces=nsmap)
            r_indent = right_indent.attrib[
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}right'] if right_indent is not None else 0
            if r_indent == 0 and styleRightIndent > 0:
                r_indent = styleRightIndent
            info['rightIndent'] = r_indent
            alignment = prop.find('.//w:jc', namespaces=nsmap)
            info['alignment'] = alignment.get(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if alignment is not None else 'left'
            # Get tab position
        line = paragraph.text
        line = line.strip()
        words = line.split(" ")
        line = line.strip()
        info['isTitleCase'] = False
        info['wordCount'] = 0
        info['firstChar'] = None
        info['lastChar'] = None
        info['lastWord'] = None
        info['firstWord'] = None
        info['firstWord'] = None
        if len(line) > 0:
            isTitle = self.is_title_case(line)
            info['isTitleCase'] = isTitle if line else False
            words = line.split(" ")
            info['wordCount'] = str(len(words))
            # Get first and last characters
            first_char = words[0][0] if words else ''
            info['firstChar'] = html.escape(first_char, quote=True)
            last_char = words[-1][-1] if words else ''
            info['lastChar'] = html.escape(last_char, quote=True)
            last_word = words[-1] if words else ''
            info['lastWord'] = html.escape(last_word, quote=True)
            first_word = words[0] if words else ''
            info['firstWord'] = html.escape(first_word, quote=True)
            if len(words) == 1:
                if info['firstWord'].startswith("Referen"):
                    info['refTitle'] = True
        return info

    def create_info_xml(self, docfile):
        file_name = os.path.basename(docfile)
        file_path = os.path.dirname(docfile)
        doc = Document(docfile)
        output = []
        prev_info = None
        para_count = 0
        for paragraph in doc.paragraphs:
            para_count = para_count + 1
            para_name = paragraph.style.name
            current_info = self.get_paragraph_info(docfile, doc, paragraph)
            isRefTitle = current_info["refTitle"]
            output.append(f'<paraInfo number="{para_count}" isRefTitle="{isRefTitle}">')
            if prev_info:
                output.append('<previousPara ')
                output.append(' '.join([f'{k}="{v}"' for k, v in prev_info.items()]))
                output.append('/>')
            output.append('<para ')
            output.append(' '.join([f'{k}="{v}"' for k, v in current_info.items()]))
            output.append('/>')
            next_para = self.get_paragraph_info(docfile, doc, doc.paragraphs[para_count]) if para_count < len(doc.paragraphs) else {}
            output.append('<nextPara ')
            output.append(' '.join([f'{k}="{v}"' for k, v in next_para.items()]))
            output.append('/>')
            output.append('</paraInfo>')
            prev_info = current_info
        # Write the output to a file
        output = "<info>" + "".join(output) + "</info>"
        formatted_output = self.pretty_format_xml(output)  # Pretty format the output XML
        info_path = os.path.join(file_path, "para_info.xml")
        with open(info_path, 'w', encoding="utf-8") as f:
            f.write(formatted_output)
        quot_xml = self.process_xml(info_path)
        # os.remove(info_path)
        if re.search("<para", quot_xml):
            formatted_quot = self.pretty_format_xml(quot_xml)
            quot_path = os.path.join(file_path, "quot_info.xml")
            with open(quot_path, 'w', encoding="utf-8") as f:
                f.write(formatted_quot)
            return quot_path
        else:
            return None

    def add_quot_prefix_to_paragraphs(self, docfile, xml_file):
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        doc = Document(docfile)
        tree = ET.parse(xml_file)
        root = tree.getroot()
        style_changed = False
        for para_info in root.findall('.//quot/para'):
            para_number = int(para_info.attrib['number'])
            para_align = para_info.attrib['quote']
            if para_number <= len(doc.paragraphs):
                paragraph = doc.paragraphs[para_number - 1]
                prefix_run = paragraph.add_run()
                if para_align == "Right":
                    prefix_run.text = "[candm-quote-right]"
                else:
                    prefix_run.text = "[candm-quote-left]"
                prefix_run.font.hidden = True
                style_changed = True
        if style_changed is True:
            doc.save(docfile)

    def add_list_prefix_to_paragraphs(self, docfile, xml_file):
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        doc = Document(docfile)
        tree = ET.parse(xml_file)
        root = tree.getroot()
        style_changed = False
        for para_info in root.findall('.//list/para'):
            para_number = int(para_info.attrib['number'])
            para_style = para_info.attrib['listType']
            if para_number <= len(doc.paragraphs):
                paragraph = doc.paragraphs[para_number - 1]
                prefix_run = paragraph.add_run()
                if para_style == "decimal":
                    prefix_run.text = "[candm-nl-text]"
                elif para_style == "bullet":
                    prefix_run.text = "[candm-bl-text]"
                else:
                    prefix_run.text = "[candm-ul-text]"
                prefix_run.font.hidden = True
                style_changed = True
        if style_changed is True:
            doc.save(docfile)

    def ensure_paragraph_style(self, document, style_name, base="Normal"):
        styles = document.styles
        if style_name in styles:
            return styles[style_name]

        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles[base]
        return style

    def apply_styles_to_paragraphs(self, docfile, xml_file):
        ref_para = 0
        if os.path.exists(xml_file):
            tree = ET.parse(xml_file)
            root = tree.getroot()
            for ref in root.findall(".//refPara"):
                ref_para = int(ref.attrib['num'])
        doc = Document(docfile)
        styles = {s.name for s in doc.styles}
        doc_styles = []
        for style in doc.styles:
            style_name = style.name
            doc_styles.append(style_name)
        style_changed = False
        MARKER_MAP = {
            "[candm-quote-left]": ("displayquote", None),
            "[candm-quote-right]": ("displayquote", "right"),
            "[candm-nl-text]": ("numlist", None),
            "[candm-bl-text]": ("bullist", None),
            "[candm-ul-text]": ("tablist", None),
        }
        for paragraph in doc.paragraphs:
            para_text = paragraph.text
            if not para_text:
                continue
            style_changed = True
            for marker, (style_name, align) in MARKER_MAP.items():
                if marker not in para_text:
                    continue
                current_style = paragraph.style.name if paragraph.style else ""
                skip_style = (
                        current_style.lower().startswith("sect")
                        and style_name in {"numlist", "tablist"}
                )
                if not skip_style:
                    styles.add(style_name)
                    self.ensure_paragraph_style(doc, style_name)
                    paragraph.style = style_name
                    if align == "right":
                        pPr = paragraph._p.get_or_add_pPr()
                        jc = pPr.get_or_add_jc()
                        jc.set(qn("w:val"), "right")
                for run in list(paragraph.runs):
                    if marker in run.text:
                        paragraph._p.remove(run._element)
                break

        # para_count = 0
        # for paragraph in doc.paragraphs:
        #     for run in paragraph.runs:
        #         if re.search(r"\[candm-quote-(left|right)\]", run.text):
        #             if "displayquote" not in doc_styles:
        #                 doc.styles.add_style("displayquote", WD_STYLE_TYPE.PARAGRAPH)
        #                 doc_styles.append("displayquote")
        #             if paragraph.style.name in ["paragraph", "dummy", "lista", "listb", "sectiona", "sectionb", "sectionc"]:
        #                 paragraph.style = "displayquote"
        #             elif ref_para < para_count and paragraph.style.name in ["bibitem"]:
        #                 paragraph.style = "displayquote"
        #             if re.search("\[candm-quote-right\]", run.text):
        #                 paragraph.alignment = 2
        #             paragraph._p.remove(run._element)
        #             style_changed = True
        #         elif re.search("\[candm-nl-text\]", run.text):
        #             if "numlist" not in doc_styles:
        #                 doc.styles.add_style("numlist", WD_STYLE_TYPE.PARAGRAPH)
        #                 doc_styles.append("numlist")
        #             if paragraph.style.name in ["paragraph", "dummy", "lista", "listb"]:
        #                 paragraph.style = "numlist"
        #             elif ref_para < para_count and paragraph.style.name in ["bibitem"]:
        #                 paragraph.style = "numlist"
        #             paragraph._p.remove(run._element)
        #             style_changed = True
        #         elif re.search("\[candm-bl-text\]", run.text):
        #             if "bullist" not in doc_styles:
        #                 doc.styles.add_style("bullist", WD_STYLE_TYPE.PARAGRAPH)
        #                 doc_styles.append("bullist")
        #             if paragraph.style.name in ["paragraph", "dummy", "lista", "listb"]:
        #                 paragraph.style = "bullist"
        #             elif ref_para < para_count and paragraph.style.name in ["bibitem"]:
        #                 paragraph.style = "bullist"
        #             paragraph._p.remove(run._element)
        #             style_changed = True
        #         elif re.search("\[candm-ul-text\]", run.text):
        #             if "tablist" not in doc_styles:
        #                 doc.styles.add_style("tablist", WD_STYLE_TYPE.PARAGRAPH)
        #                 doc_styles.append("tablist")
        #             if paragraph.style.name in ["paragraph", "dummy", "lista", "listb", "displayquote"]:
        #                 paragraph.style = "tablist"
        #             elif ref_para < para_count and paragraph.style.name in ["bibitem"]:
        #                 paragraph.style = "tablist"
        #             paragraph._p.remove(run._element)
        #             style_changed = True
        #         para_count += 1
        para_count = 0
        for paragraph in doc.paragraphs:
            if paragraph.style.name == "bibitem" and ref_para > para_count:
                paragraph.style = "paragraph"
                style_changed = True
            para_count += 1
        if style_changed is True:
            doc.save(docfile)

# info = CreateParaInfo()
# docfile = 'V:\FOR_BREAKDOWN\BreakDown_INPUT\SAGE\ASR_1263497\docs\ASR_1263497_CLN.docx'
# xml_file = "V:\FOR_BREAKDOWN\ParaStyler_INPUT\SAGE\JIV_1260004\quot_info.xml"
# info.create_info_xml(docfile)
# info.add_quot_prefix_to_paragraphs(docfile, xml_file)
# info.add_list_prefix_to_paragraphs(docfile, xml_file)
# info.apply_styles_to_paragraphs(docfile, xml_file)
# Extract XML representation of the document
